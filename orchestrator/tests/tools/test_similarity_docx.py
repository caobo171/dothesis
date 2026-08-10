"""The document self-check: what it finds, and what it must never claim.

The rule the whole feature hangs on: "no matches found" and "nobody looked" are
different answers, and a student who reads the second as the first submits on
the strength of it. `corpus_checked` carries that distinction and nothing may
collapse it.
"""
import io

import pytest
from docx import Document

from orchestrator.tools.similarity_docx import scan_docx, similarity_docx

_DUP = ("Nghiên cứu sử dụng phương pháp chọn mẫu phi xác suất có chủ đích kết hợp "
        "thuận tiện, phát bảng câu hỏi trực tuyến đến người tiêu dùng đang sinh sống "
        "và làm việc tại Thành phố Hồ Chí Minh trong khoảng thời gian từ tháng ba "
        "đến tháng năm năm 2025.")
_QUOTE = ('Theo lý thuyết, "độ tin cậy của người phát ngôn quyết định mức độ chấp nhận '
          'thông điệp của người nhận trong bối cảnh truyền thông xã hội" và điều đó '
          'rất quan trọng với nghiên cứu này.')


def _doc(*, duplicate=True, quote=True, refs=True) -> bytes:
    d = Document()
    d.add_heading("CHƯƠNG 3", level=1)
    d.add_paragraph(_DUP)
    d.add_paragraph("Thang đo kế thừa từ nghiên cứu trước và được điều chỉnh (Ohanian, 1990).")
    d.add_heading("CHƯƠNG 4", level=1)
    if duplicate:
        d.add_paragraph(_DUP)
    if quote:
        d.add_paragraph(_QUOTE)
    d.add_paragraph("Kết quả cho thấy cả ba giả thuyết đều được chấp nhận (Nguyen, 2021).")
    if refs:
        d.add_heading("TÀI LIỆU THAM KHẢO", level=1)
        d.add_paragraph("Ohanian, R. (1990). Construction and validation of a scale.")
        d.add_paragraph("Tran, M. (2015). Consumer behaviour in emerging markets.")
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


class _Provider:
    name = "fake"

    def __init__(self, matches=None, boom=False):
        self._matches, self._boom = matches or [], boom

    def check(self, text, *, language="vi"):
        if self._boom:
            raise RuntimeError("provider down")
        return {"score": 0.2, "matches": self._matches, "provider": self.name}


# --- the honesty invariant ----------------------------------------------------

def test_without_a_provider_the_report_says_nobody_looked():
    _out, rep = similarity_docx(_doc())
    assert rep["corpus_checked"] is False
    assert rep["corpus_matches"] == []
    assert rep["provider"] is None


def test_a_provider_failure_is_not_a_clean_result():
    """The one wrong answer: a transport error degrading into "no matches"."""
    _out, rep = similarity_docx(_doc(), provider=_Provider(boom=True))
    assert rep["corpus_checked"] is False
    assert rep["corpus_error"] == "provider_error"


def test_a_provider_that_runs_is_recorded_as_having_run():
    _out, rep = similarity_docx(
        _doc(), provider=_Provider(matches=[{"source": "x", "overlap": 0.3}]))
    assert rep["corpus_checked"] is True
    assert rep["provider"] == "fake" and len(rep["corpus_matches"]) == 1


# --- what it finds ------------------------------------------------------------

def test_a_passage_repeated_inside_the_document_is_found():
    _out, rep = similarity_docx(_doc())
    dups = rep["internal_duplication"]
    assert len(dups) == 1
    assert dups[0]["tokens"] >= 25
    assert "chọn mẫu phi xác suất" in dups[0]["excerpt"]


def test_a_document_that_does_not_repeat_itself_reports_nothing():
    _out, rep = similarity_docx(_doc(duplicate=False))
    assert rep["internal_duplication"] == []


def test_a_quotation_with_no_citation_is_found():
    _out, rep = similarity_docx(_doc())
    assert len(rep["uncited_quotations"]) == 1
    assert "độ tin cậy của người phát ngôn" in rep["uncited_quotations"][0]["excerpt"]


def test_a_citation_missing_from_the_reference_list_is_found():
    _out, rep = similarity_docx(_doc())
    assert rep["cited_not_in_references"] == ["Nguyen (2021)"]


def test_a_reference_never_cited_is_found():
    _out, rep = similarity_docx(_doc())
    assert [e.split(",")[0] for e in rep["references_never_cited"]] == ["Tran"]


def test_a_reference_that_IS_cited_is_not_reported():
    """Ohanian is cited in chapter 3. Keyed per key rather than per entry, the
    surname pattern's other hits (the initial, the first title word) each
    reported the entry uncited on their own."""
    _out, rep = similarity_docx(_doc())
    assert not any("Ohanian" in e for e in rep["references_never_cited"])


def test_the_reference_list_is_never_flagged_as_a_finding():
    """Bibliographic data is not prose. Two entries by the same author in the
    same year would otherwise read as internal duplication."""
    _out, rep = similarity_docx(_doc())
    ref_paras = {d["paragraph_a"] for d in rep["internal_duplication"]}
    ref_paras |= {d["paragraph_b"] for d in rep["internal_duplication"]}
    assert all(p < 8 for p in ref_paras)


# --- the annotated copy -------------------------------------------------------

def test_the_annotated_copy_highlights_only_the_findings():
    from docx.enum.text import WD_COLOR_INDEX

    out, rep = similarity_docx(_doc())
    assert out
    doc = Document(io.BytesIO(out))
    lit = {i for i, p in enumerate(doc.paragraphs)
           if any(r.font.highlight_color == WD_COLOR_INDEX.YELLOW for r in p.runs)}
    assert lit == ({d["paragraph_a"] for d in rep["internal_duplication"]}
                   | {d["paragraph_b"] for d in rep["internal_duplication"]}
                   | {q["paragraph"] for q in rep["uncited_quotations"]})


def test_the_summary_says_it_is_not_turnitin():
    out, _rep = similarity_docx(_doc(), language="vi")
    text = "\n".join(p.text for p in Document(io.BytesIO(out)).paragraphs)
    assert "không phải bản quét Turnitin" in text
    out_en, _ = similarity_docx(_doc(), language="en")
    text_en = "\n".join(p.text for p in Document(io.BytesIO(out_en)).paragraphs)
    assert "not a Turnitin scan" in text_en


def test_the_summary_survives_a_document_that_defines_no_list_or_heading_styles():
    """python-docx resolves "List Bullet" and "Heading 1" through the STUDENT's
    styles.xml. A thesis template that names its styles anything else raised
    KeyError and cost the entire annotated copy — the same dangling-style trap
    the cover page hit."""
    d = Document()
    d.add_paragraph(_DUP)
    d.add_paragraph(_DUP)
    b = io.BytesIO()
    d.save(b)
    for style in list(Document(io.BytesIO(b.getvalue())).styles):
        pass  # (sanity: the fixture opens)
    out, rep = similarity_docx(b.getvalue())
    assert out and not rep.get("annotation_failed")


# --- scan ---------------------------------------------------------------------

def test_the_scan_counts_what_the_run_would_look_at():
    out = scan_docx(_doc())
    assert out["ok"] and out["body_paragraphs"] > 0 and out["words"] > 50
    assert out["reference_entries"] == 2
    assert out["quotations"] == 1


@pytest.mark.parametrize("fn", [scan_docx, similarity_docx])
def test_a_file_that_is_not_a_docx_fails_honestly(fn):
    res = fn(b"not a docx")
    rep = res[1] if isinstance(res, tuple) else res
    assert rep["ok"] is False and rep["error"] == "unreadable"
