"""Humanize a .docx IN PLACE, so the student gets their document back.

Why this exists rather than reusing the text extractor: `_extract_docx_text`
(api/app/routers/uploads.py) was built to feed the AGENT, where layout costs
nothing — it takes `p.text`, so heading levels and character formatting are
gone, and it appends every table at the END as pipe-joined rows, detached from
the caption that introduced it. That is fine for classification and useless as
the input to a rewrite you intend to hand to a supervisor.

So this walks the document object instead and writes the rewrite back into the
same paragraphs. Paragraph styles, heading levels, numbering, tables and the
overall structure survive because they are never touched.

Four things are skipped BY DESIGN, not by omission:

  tables   — they are data. The frozen-token rule already forbids changing a
             number, so re-voicing a results table is pure risk. python-docx's
             `doc.paragraphs` excludes table-cell paragraphs, so this falls out
             of the walk for free.
  headings — structural labels, not prose. Re-voicing "4.3.3 Thang đo Chuyên
             môn" produces a heading that no longer matches the table of
             contents.
  boilerplate — the declaration the student signs, the keyword line, cover-page
             form fields. Supplied verbatim by the university, so a rewrite
             makes the signed statement stop matching the form. See
             _BOILERPLATE_ANYWHERE / _BOILERPLATE_OPENERS.
  references — bibliographic data in the same category as the tables, just
             typed as paragraphs. Re-voicing an entry corrupts a SOURCE — a
             title, an edition, a page range, a publisher — and no gain is
             possible: Turnitin does not flag a bibliography as AI prose. See
             _reference_indices, and the regression note on it, before ever
             re-enabling them.

Known loss, stated plainly: replacing a paragraph's text rebuilds its runs, so
a **bold** phrase mid-sentence does not survive. The first run is kept (which
preserves the paragraph's style and that run's character formatting) and the
rest are blanked.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Callable

logger = logging.getLogger(__name__)

# A paragraph shorter than this is a caption, a label, a figure note or a
# stray fragment — "Bảng 4.7: Chi tiết độ tin cậy thang đo" has nothing to
# re-voice, and rewriting it risks breaking a cross-reference.
_MIN_WORDS = 12

# Characters per batch. Consecutive body paragraphs are rewritten together:
# one call per ~2 pages instead of one per paragraph turns a 60-page thesis
# from ~300 calls into ~20, and gives the rewrite enough context to vary
# sentence rhythm ACROSS paragraphs rather than inside each one in isolation.
_BATCH_CHARS = 3000

_HEADING_STYLES = ("heading", "title", "subtitle", "toc", "caption")

# Front-matter the student SIGNS, or that the university supplies verbatim.
#
# WHY: a real run re-voiced Bolton Business School's plagiarism declaration —
# the form the student signs, including the sentence swearing no AI tool was
# used. That text is the university's template, not the student's prose, so
# editing its wording makes the signed statement stop matching the form it
# claims to be. Turnitin flagged the paragraph anyway: all risk, no benefit.
#
# Matched by PHRASE, not by position: a declaration sits before every heading
# while "Keywords:" sits after the abstract heading, so no single position rule
# covers both. Vietnamese entries carry the same weight as the English ones —
# "Tôi xin cam đoan" is the standard LỜI CAM ĐOAN opener.
#
# Two classes, because a marker that is safe to match ANYWHERE and one that is
# only safe at the START are not the same kind of phrase. From this same
# dissertation: the ethics paragraph reads "an Ethics Approval Form was
# submitted to and approved by the supervising tutor" — ordinary body prose
# carrying a cover-page phrase. Matching "submitted to" anywhere froze it, and
# in a one-paragraph document froze the whole run to `no_prose`. Generic
# form-field labels therefore only count when they OPEN the paragraph.
_BOILERPLATE_ANYWHERE = (
    "i confirm that", "i declare that", "i hereby declare", "i certify that",
    "student signature", "marking tutor", "gai not permitted",
    "tôi xin cam đoan", "lời cam đoan",
)
_BOILERPLATE_OPENERS = (
    "keywords:", "submitted to", "student number", "word count",
    "từ khóa:", "sinh viên thực hiện", "giảng viên hướng dẫn", "mã sinh viên",
)


def _is_boilerplate(text: str) -> bool:
    low = text.strip().lower()
    return (any(m in low for m in _BOILERPLATE_ANYWHERE)
            or low.startswith(_BOILERPLATE_OPENERS))


# Where the reference list starts. This is the ENGINE's single notion of that
# boundary: it lived in cite_docx, which already imports from this module, so
# the definition is LIFTED here and imported back — a second, divergent regex
# in either direction is exactly how the two walks would one day disagree
# about where the bibliography begins. (The API's extract_references keeps its
# own copy because it works on flat pasted text, not paragraph objects, and
# needs re.M — different substrate, same vocabulary.) Matched by TEXT, not by
# style: students paste headings without styling them. Last match wins,
# because the table of contents carries the same words earlier in the file.
_REF_HEADING = re.compile(
    r"^\s*(?:danh\s+m[ụu]c\s+)?(?:t[àa]i\s+li[ệe]u\s+tham\s+kh[ảa]o"
    r"|references?|bibliography|works\s+cited)\s*:?\s*$",
    re.I,
)

# Where the reference region ENDS: the appendix that customarily follows it is
# body prose again. Matched as a fallback by text for the same unstyled-paste
# reason as _REF_HEADING; the word-count cap keeps a body sentence that merely
# starts with "Appendix A shows..." from ending the region early.
_APPENDIX_HEADING = re.compile(r"^(appendix|annexe?|ph[ụu]\s+l[ụu]c)\b", re.I)


def _ref_heading_index(paragraphs: list[Any]) -> int | None:
    """Index of the reference-list heading, or None. Last match wins."""
    found = None
    for i, p in enumerate(paragraphs):
        if _REF_HEADING.match((p.text or "").strip()):
            found = i
    return found


def _reference_indices(paragraphs: list[Any]) -> set[int]:
    """Paragraph indices of the reference list — never offered for rewrite.

    WHY, so nobody "helpfully" re-enables them: reference entries are
    bibliographic DATA. A re-voice can alter a title, an edition, a page
    range, a publisher — corrupting a real source the student will be marked
    against — and buys nothing, because a bibliography is not prose a detector
    scores. Regression that forced this: _batches packed a real bibliography
    into two 16-paragraph batches whose whole-batch rewrites happened to fail,
    so the references survived only BY ACCIDENT — until the per-paragraph
    retry rewrote each entry individually and re-voiced the entire list.

    Positional, not per-entry pattern matching: an entry is indistinguishable
    from a body sentence that cites ("Nguyen (2019) showed that...") one line
    at a time, but the region between the references heading and the next
    section boundary is unambiguous. The region ends at the first following
    heading (styled, or an unstyled short "Appendix"/"Phụ lục" line) so
    appendix prose stays eligible.
    """
    start = _ref_heading_index(paragraphs)
    if start is None:
        return set()
    out: set[int] = set()
    for j in range(start + 1, len(paragraphs)):
        p = paragraphs[j]
        text = (p.text or "").strip()
        if text and (_is_heading(p)
                     or (_APPENDIX_HEADING.match(text) and len(text.split()) <= 8)):
            break
        out.add(j)
    return out


def _is_heading(p: Any) -> bool:
    try:
        name = (p.style.name or "").strip().lower()
    except Exception:  # noqa: BLE001 — a malformed style must not sink the walk
        return False
    return any(name.startswith(h) for h in _HEADING_STYLES)


def _eligible(p: Any) -> bool:
    text = (p.text or "").strip()
    if not text or _is_heading(p) or _is_boilerplate(text):
        return False
    return len(text.split()) >= _MIN_WORDS


def _batches(indexed: list[tuple[int, str]]) -> list[list[int]]:
    """Group consecutive eligible paragraphs into rewrite passages.

    Breaks on a gap in the index (an intervening heading or table), so a batch
    never spans a section boundary — which would ask the model to keep one
    voice across two different topics.
    """
    out: list[list[int]] = []
    cur: list[int] = []
    size = 0
    prev_idx: int | None = None
    for idx, text in indexed:
        gap = prev_idx is not None and idx != prev_idx + 1
        if cur and (gap or size + len(text) > _BATCH_CHARS):
            out.append(cur)
            cur, size = [], 0
        cur.append(idx)
        size += len(text)
        prev_idx = idx
    if cur:
        out.append(cur)
    return out


def _open(body: bytes):
    from docx import Document  # noqa: PLC0415 — heavy, and only needed here
    return Document(io.BytesIO(body))


def scan_docx(body: bytes) -> dict:
    """What a rewrite would touch, without touching it. No LLM, no charge.

    Exists so the student sees the size of the job before paying for it: a
    thesis is hundreds of paragraphs, and "upload and hope" is the wrong way to
    spend someone's credits.
    """
    try:
        doc = _open(body)
    except Exception:  # noqa: BLE001
        logger.exception("humanize_docx: could not open document")
        return {"ok": False, "error": "unreadable"}

    paragraphs = doc.paragraphs
    # Same exclusion as the rewrite walk below, so the quote matches the job:
    # a student must not be priced for reference entries that are never touched.
    refs = _reference_indices(paragraphs)
    eligible = [(i, (p.text or "").strip()) for i, p in enumerate(paragraphs)
                if i not in refs and _eligible(p)]
    headings = sum(1 for p in paragraphs if _is_heading(p) and (p.text or "").strip())
    short = sum(
        1 for p in paragraphs
        if (p.text or "").strip() and not _is_heading(p) and len((p.text or "").split()) < _MIN_WORDS
    )
    return {
        "ok": True,
        "body_paragraphs": len(eligible),
        "headings": headings,
        "short_or_captions": short,
        "tables": len(doc.tables),
        "passages": len(_batches(eligible)),
        "chars": sum(len(t) for _, t in eligible),
        # Whitespace-separated tokens across the SAME eligible prose the walk
        # rewrites — references, headings and captions excluded, so a quote
        # priced on this counts only work that will actually be done.
        #
        # Tokens, not dictionary words, and deliberately: Vietnamese writes one
        # word as several space-separated syllables ("nghiên cứu"), so this
        # over-counts a Vietnamese document relative to an English one — which
        # is right, because the model cost it stands in for scales with
        # syllables too, not with lexical words.
        "words": sum(len(t.split()) for _, t in eligible),
    }


def _set_paragraph_text(p: Any, text: str) -> None:
    """Replace a paragraph's text, keeping its style and first run's format."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


# Failure kinds a smaller passage cannot change: they are properties of the
# DEPLOYMENT (no anchor installed for this language) or of the input being
# empty, not of the batch the model was handed. Retrying them per paragraph
# spends one model call per paragraph to learn the same fact len(batch) times.
# Everything else — frozen_violation, paragraph_count_changed, llm_failed,
# flatter_than_original, empty_rewrite — IS shape- or moment-dependent, and a
# single paragraph has fewer frozen tokens to hold, no boundaries to merge,
# and a fresh provider call, so most recover on the second pass.
_NO_RETRY = frozenset({"no_anchor", "empty_input"})

# Below this HANDLED share of eligible prose — rewritten, or declined because
# it was already good (see _NOT_A_FAILURE) — the run failed at its one job
# even if SOMETHING changed. `ok = rewritten > 0` is how a real 10,950-word
# dissertation shipped with ~70% of its prose byte-identical, reported
# success, charged the student — and still failed Turnitin, because the
# flagged text was exactly the untouched original. Half, not higher: with the
# per-paragraph fallback near-full handling is the healthy norm, so a run
# under 50% means the provider or the document is broken and the caller must
# see a failure, not a file.
_MIN_COVERAGE = 0.5

# Skip kinds that are the tool DECLINING, not the run BREAKING.
#
# flatter_than_original is the burstiness guard doing its job on prose that is
# already good: _is_burstier is a RELATIVE test — a rewrite ships only if it
# is at least as varied as what it was handed — so the better the input, the
# more rewrites get refused. Measured on a re-uploaded, already-humanized
# dissertation: 51 of 86 measurable paragraphs already sat at or above the
# 0.47 CV bar detector.py calibrated from Turnitin-clean text, so refusing to
# rewrite most of it is the EXPECTED outcome, not an outage. Counting those
# refusals against the coverage floor meant "your prose already reads as
# human" was answered with a 422 and a token charge — and re-uploads are a
# first-class flow (/runs/{id}/rerun exists).
#
# Everything else stays on the failure side, INCLUDING unknown kinds: a gate
# added upstream must announce itself here before it can be excused, or it
# silently becomes "fine".
_NOT_A_FAILURE = frozenset({"flatter_than_original"})


# Markdown emphasis the model leaks despite the prompt — a real re-run shipped
# literal asterisks around every italicised title (`*Improving Organizational
# Effectiveness*`) because _set_paragraph_text copies characters verbatim. The
# frozen-token gate cannot catch this (markers are not numbers or citations),
# so the markers are stripped here, at the last step before text can reach the
# document. Anchored the way markdown itself anchors so ONLY paired emphasis
# matches: an opener must hug the following word, a closer the preceding one,
# and `_` additionally requires non-word neighbours. That is what keeps the
# legitimate uses of the same characters intact — a lone significance star
# (p < 0.05*), a footnote marker opening a line, multiplication (3 * 4), and
# the snake_case indicator names this corpus is full of (TL_1, JS_2).
_MD_BOLD_RE = re.compile(r"(?<!\*)\*\*(?![\s*])([^*\n]+?)(?<![\s*])\*\*(?!\*)")
_MD_STAR_RE = re.compile(r"(?<![\w*])\*(?![\s*])([^*\n]+?)(?<![\s*])\*(?![\w*])")
_MD_UBOLD_RE = re.compile(r"(?<!\w)__(?![\s_])([^_\n]+?)(?<![\s_])__(?!\w)")
_MD_UNDER_RE = re.compile(r"(?<!\w)_(?![\s_])([^_\n]+?)(?<![\s_])_(?!\w)")


def _strip_markdown_emphasis(text: str) -> str:
    """Drop emphasis MARKERS, keep the words. Double markers before single,
    so `**bold**` is not half-eaten by the single-star pass."""
    out = _MD_BOLD_RE.sub(r"\1", text)
    out = _MD_STAR_RE.sub(r"\1", out)
    out = _MD_UBOLD_RE.sub(r"\1", out)
    return _MD_UNDER_RE.sub(r"\1", out)


def _attempt(humanize_fn: Callable[..., dict], texts: list[str], *,
             language: str | None, user_anchor: str | None,
             usage: list[dict]) -> tuple[list[str] | None, str | None]:
    """One shot at rewriting `texts` as a single passage.

    Returns (parts, None) on success — one rewritten part per input paragraph,
    positionally aligned — or (None, error_kind) on any failure. The
    paragraph-count check lives HERE so the whole-batch pass and the
    per-paragraph retry get the identical guarantee: reassembly is positional,
    and a reply with the wrong number of paragraphs must never reach the
    document, on the first try or the second.

    Usage is appended on every path including failures: a rewrite that burned
    tokens and then failed verification still cost real money, and the caller
    bills usage, not successes.
    """
    r = humanize_fn("\n\n".join(texts), language=language, user_anchor=user_anchor)
    usage.extend(r.get("usage") or [])
    if not r.get("ok"):
        return None, str(r.get("error") or "unknown")
    parts = [s.strip() for s in (r.get("text") or "").split("\n\n") if s.strip()]
    if len(parts) != len(texts):
        return None, "paragraph_count_changed"
    # Stripped AFTER the count check (markers never span paragraphs, so the
    # alignment is unaffected) and inside _attempt so both the whole-batch pass
    # and the per-paragraph retry write clean text — see _MD_*_RE above.
    return [_strip_markdown_emphasis(p) for p in parts], None


def humanize_docx(
    body: bytes,
    *,
    language: str | None = None,
    user_anchor: str | None = None,
    humanize_fn: Callable[..., dict] | None = None,
    on_progress: Callable[[int, int], None] | None = None,
) -> tuple[bytes | None, dict]:
    """Rewrite every eligible body paragraph and return the rebuilt .docx.

    `humanize_fn` defaults to humanize_prose and is injectable so tests can
    drive the walk without a model.

    Returns (docx_bytes, report). A batch whose rewrite fails verification — or
    whose paragraph count comes back different from what was sent — is retried
    one paragraph at a time (see _NO_RETRY for the exceptions), and only the
    paragraphs that fail ALONE keep their originals. The count check is
    load-bearing at both levels: reassembly is positional, so a model that
    merges two paragraphs would otherwise shift every later paragraph's text
    into the wrong slot and silently scramble the document.
    """
    if humanize_fn is None:
        from .humanize import humanize_prose  # noqa: PLC0415
        humanize_fn = humanize_prose

    try:
        doc = _open(body)
    except Exception:  # noqa: BLE001
        logger.exception("humanize_docx: could not open document")
        return None, {"ok": False, "error": "unreadable"}

    paragraphs = doc.paragraphs
    # The bibliography is carved out BEFORE eligibility, not inside _eligible:
    # membership is positional (everything between the references heading and
    # the next section), which a per-paragraph predicate cannot see. See
    # _reference_indices for why the entries must never be offered for rewrite.
    refs = _reference_indices(paragraphs)
    eligible = [(i, (p.text or "").strip()) for i, p in enumerate(paragraphs)
                if i not in refs and _eligible(p)]
    if not eligible:
        return None, {"ok": False, "error": "no_prose",
                      "detail": "No body paragraphs long enough to rewrite — the "
                                "document looks like headings, tables or captions."}

    by_idx = dict(eligible)
    usage: list[dict] = []
    rewritten = skipped = 0
    failures: list[dict] = []

    # Materialised so the caller can be told the total up front: a progress bar
    # that only learns its denominator on the last batch is not a progress bar.
    # `on_progress` is best-effort — a reporting failure must not abort a walk
    # the student is paying for.
    batches = list(_batches(eligible))
    total = len(batches)
    if on_progress:
        try:
            on_progress(0, total)
        except Exception:  # noqa: BLE001
            logger.exception("humanize_docx: progress callback failed")

    for done, batch in enumerate(batches, start=1):
        source = [by_idx[i] for i in batch]
        parts, err = _attempt(humanize_fn, source, language=language,
                              user_anchor=user_anchor, usage=usage)
        # Reported after the FIRST attempt, not at the end of the loop body:
        # the model call is the minute-long part, and bumping only on success
        # would leave a progress bar that stalls whenever something goes wrong.
        # The per-paragraph retries below ride inside the same batch tick — the
        # denominator is batches, and a retry is this batch taking longer, not
        # more batches.
        if on_progress:
            try:
                on_progress(done, total)
            except Exception:  # noqa: BLE001
                logger.exception("humanize_docx: progress callback failed")

        if err is None:
            for idx, new_text in zip(batch, parts):
                _set_paragraph_text(paragraphs[idx], new_text)
            rewritten += len(batch)
            continue

        # The whole batch failed. All-or-nothing here was the production
        # defect: with 4-16 paragraphs per batch, one bad reply silently kept
        # a page of original text and the run still said ok — measured on a
        # real dissertation, 39 of 70 batches (4,990 words) came back
        # byte-identical. So retry each paragraph on its own and keep
        # whichever succeed; only a paragraph that fails ALONE stays original.
        if len(batch) > 1 and err not in _NO_RETRY:
            logger.warning(
                "humanize_docx: batch %d/%d failed (%s) — retrying its %d "
                "paragraphs individually", done, total, err, len(batch))
            for idx in batch:
                p_parts, p_err = _attempt(
                    humanize_fn, [by_idx[idx]], language=language,
                    user_anchor=user_anchor, usage=usage)
                if p_err is None:
                    _set_paragraph_text(paragraphs[idx], p_parts[0])
                    rewritten += 1
                else:
                    skipped += 1
                    # "retried": the reader of a run report needs to tell "we
                    # gave up after one try" from "this paragraph failed even
                    # alone" — the second is a fact about the paragraph.
                    failures.append({"paragraphs": 1, "error": p_err,
                                     "retried": True})
                    logger.warning(
                        "humanize_docx: paragraph %d kept as original after "
                        "retry (%s) — %d words", idx, p_err,
                        len(by_idx[idx].split()))
            continue

        skipped += len(batch)
        failures.append({"paragraphs": len(batch), "error": err})
        # One line per lost batch, counts only: "why did 39 batches fail" must
        # be answerable from the logs alone, and was not — the reasons were
        # computed, returned, and dropped by every caller.
        logger.warning(
            "humanize_docx: batch %d/%d kept as original (%s) — %d "
            "paragraph(s), %d chars", done, total, err, len(batch),
            sum(len(by_idx[i]) for i in batch))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    attempted = rewritten + skipped
    # Partition the skips before judging the run: a paragraph the tool
    # DECLINED to make worse (_NOT_A_FAILURE) was handled correctly, and only
    # the rest are the run breaking. Derived from the `failures` list rather
    # than counted in the loop, so the verdict can never disagree with its own
    # breakdown.
    failed = sum(int(f.get("paragraphs") or 0) for f in failures
                 if f.get("error") not in _NOT_A_FAILURE)
    declined = skipped - failed
    handled = rewritten + declined
    handled_ratio = (handled / attempted) if attempted else 0.0
    report = {
        # Not `rewritten > 0`: that let 1 rewritten paragraph out of 132 count
        # as success. And not rewritten-only coverage either: judging declines
        # as failures 422'd a document whose prose was already good. Success
        # means most of the prose was HANDLED — rewritten, or already human
        # enough that the guard refused to touch it. A run below the floor
        # still returns its (partial) bytes so the caller can store them, but
        # reports failure.
        "ok": handled > 0 and handled_ratio >= _MIN_COVERAGE,
        "rewritten": rewritten,
        "skipped": skipped,
        "declined": declined,
        # Distinct from `ok` on purpose: "done" and "nothing needed doing"
        # call for different sentences to the student, and the headless
        # surfaces read status codes and headers, not run history. True only
        # when the WHOLE document was declined — any rewrite or any genuine
        # failure makes "your writing already reads as human" a false claim.
        "already_human": bool(declined) and not rewritten and not failed,
        # Share actually rewritten — kept as the descriptive stat it always
        # was; the ok verdict above is what now excludes declines.
        "coverage": round((rewritten / attempted) if attempted else 0.0, 3),
        "failures": failures,
        "usage": usage,
        # Whether the caller is getting the student's own file back. Explicit
        # rather than inferred from `ok`, because the headless surfaces and the
        # billing path both branch on it and neither should re-derive the rule.
        "reverted": False,
    }
    if handled > 0 and handled_ratio < _MIN_COVERAGE:
        report["error"] = "mostly_skipped"
        report["detail"] = (
            f"Only {rewritten} of {attempted} eligible paragraphs could be "
            f"rewritten. A part-rewritten document scores worse than an "
            f"untouched one, so your original file was returned unchanged.")
        # Hand back the ORIGINAL, not the partial rewrite.
        #
        # Measured on a real submission (Turnitin, before/after, same student,
        # one day apart): a document rewritten in patches scored WORSE than the
        # untouched original — 23% -> 30%. Turnitin classifies OVERLAPPING
        # segments of roughly 5-10 sentences, so a segment straddling the join
        # between a rewritten paragraph and an untouched one is scored as one
        # unit. Paragraphs never touched — byte-identical in both files — went
        # from 0.0% flagged to 8.2% purely from what landed next to them, while
        # untouched paragraphs far from any edit moved -0.4.
        #
        # So a partial rewrite is not "some benefit"; it is a net harm, and the
        # student is better off with the file they uploaded.
        report["reverted"] = True
        logger.warning(
            "humanize_docx: coverage %.2f below floor %.2f — returning the "
            "original document instead of a partial rewrite (%d/%d rewritten)",
            handled_ratio, _MIN_COVERAGE, rewritten, attempted)
        return body, report
    return buf.getvalue(), report
