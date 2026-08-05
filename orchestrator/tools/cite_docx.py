"""Cite a .docx IN PLACE — the document comes back with its citations fixed.

Same shape as humanize_docx (walk the document, write back into the paragraphs
that are already there, never rebuild it), and for the same reason: a thesis
that loses its heading levels, tables and numbering is not something a student
can hand to a supervisor.

Two jobs, run in this order:

  A. RESOLVE what is already cited. Every in-text "(Hair và cộng sự, 2010)" is
     looked up in CrossRef and the reference list is rebuilt from the records
     that came back. This is deterministic, costs no model tokens, and is what
     actually gets a thesis past a supervisor: the list finally matches the
     citations in the body.

  B. CITE what is not. Claims that need a source and have none get one — but
     ONLY a source that was found in CrossRef and then confirmed, by a separate
     model call, to actually support that claim. A claim with no confirmed
     source gets a visible "[cần nguồn]" marker instead.

  C. LINK them. Every in-text citation becomes a Word internal hyperlink onto a
     bookmark on its own reference entry, so clicking "(Hair và cộng sự, 2010)"
     jumps to the bottom of the thesis. This is also the only way the student
     can SEE which mentions the tool understood: a citation that renders as
     plain black text is one that did not resolve to any entry.

Phase B is the dangerous half and is built to fail closed. The model never
supplies a reference: it supplies a SEARCH QUERY and a yes/no on candidates that
CrossRef returned. Nothing reaches the reference list that CrossRef did not
return first, so this cannot invent a plausible-looking source — which is the
single most common way an LLM-assisted thesis fails, and the thing the citation
checker in this same product exists to catch.
"""
from __future__ import annotations

import copy
import io
import json
import logging
import os
import re
from concurrent.futures import ThreadPoolExecutor
from typing import Any, Callable

from .citations import (
    InText, clean, dedupe, format_reference, intext_form, is_citable,
    match_reference_line, parse_intext_citations, parse_reference_lines,
    resolve_verbose, uses_vietnamese_convention,
)
from .crossref import search as crossref_search
from .humanize_docx import _batches, _is_heading, _open, _set_paragraph_text

logger = logging.getLogger(__name__)

# Where the reference list starts. Same vocabulary as the API's list checker;
# the last match wins because the table of contents carries the words too.
_REF_HEADING = re.compile(
    r"^\s*(?:danh\s+m[ụu]c\s+)?(?:t[àa]i\s+li[ệe]u\s+tham\s+kh[ảa]o"
    r"|references?|bibliography|works\s+cited)\s*:?\s*$",
    re.I,
)

# A paragraph shorter than this is a caption or a label — nothing to cite.
_MIN_WORDS = 12

# Marker left on a claim we could not source. Visible on purpose: a student can
# search for it, and a supervisor reading it knows the tool made no claim.
MARKER_VI = "[cần nguồn]"
MARKER_EN = "[citation needed]"

# Ceiling on phase B. Each claim costs a CrossRef search plus a confirmation
# call, and a thesis has more unsourced sentences than anyone wants to pay for
# in one run.
_MAX_NEW_CITATIONS = 40

_CLAIM_PROMPT = """You are helping a graduate student find which sentences in \
their thesis need a citation.

Return ONLY sentences that state a claim a reader would expect a source for: an \
established theory, a definition, a threshold or rule of thumb, a prior finding, \
or a statistic about the world.

Do NOT return:
- sentences that already carry a citation
- the author's own results, data, sample or procedure
- transitions, structure sentences, or descriptions of this study's own tables

For each one, give a SHORT English search query (author names, method names and \
key terms) that would find the source paper in an academic index.

Respond with JSON only:
{{"claims": [{{"sentence": "<the sentence, copied exactly>", "query": "<english search query>"}}]}}

If nothing in the passage needs a citation, return {{"claims": []}}.

PASSAGE:
{passage}"""

_VERIFY_PROMPT = """A student wrote this claim in their thesis:

CLAIM: {claim}

Here are papers found in CrossRef. Decide whether any of them ACTUALLY supports \
that specific claim — same concept, same finding, not merely the same field.

{candidates}

Answer with JSON only: {{"index": <0-based index of the supporting paper>}} or \
{{"index": null}} if none of them genuinely supports the claim.

Answering with a paper that does not support the claim would put a fabricated \
citation into a thesis. When unsure, answer null."""


def _get_llm(temperature: float = 0.0):
    from orchestrator.llm import get_orchestrator_llm  # noqa: PLC0415

    # Same override knobs as humanize: this runs on whatever route the operator
    # points the writing tools at.
    return get_orchestrator_llm(
        temperature=temperature,
        route=os.getenv("CITE_LLM_ROUTE") or os.getenv("HUMANIZE_LLM_ROUTE") or None,
        model=os.getenv("CITE_LLM_MODEL") or os.getenv("HUMANIZE_LLM_MODEL") or None,
    )


def _invoke(llm, prompt: str, usage: list) -> str:
    """The single LLM chokepoint, so metering is attached in one place."""
    from orchestrator.message_utils import text_of  # noqa: PLC0415

    resp = llm.invoke(prompt)
    try:
        from orchestrator.token_meter import _usage_from_response  # noqa: PLC0415
        p_tok, c_tok = _usage_from_response(resp)
        usage.append({"model": str(getattr(llm, "model", "unknown")),
                      "prompt_tokens": p_tok, "completion_tokens": c_tok})
    except Exception:  # noqa: BLE001
        # Accounting must never cost a student their document.
        logger.exception("cite_docx: token accounting failed")
    return text_of(resp)


def _json_out(raw: str) -> dict:
    """Parse a model's JSON reply, tolerating code fences and stray prose."""
    s = (raw or "").strip()
    s = re.sub(r"^```(?:json)?|```$", "", s, flags=re.M).strip()
    start, end = s.find("{"), s.rfind("}")
    if start == -1 or end == -1:
        return {}
    try:
        return json.loads(s[start:end + 1])
    except Exception:  # noqa: BLE001
        logger.warning("cite_docx: unparseable model JSON: %r", s[:200])
        return {}


# --- cross-references: bookmarks and internal hyperlinks ---------------------
#
# python-docx has no API for either, so this is raw WordprocessingML. Kept in
# one block, small and boring, because the failure mode of getting it wrong is a
# document Word refuses to open — which for a student is worse than no links.

# Word's own hyperlink blue. Set explicitly rather than via the "Hyperlink"
# character style: a thesis template downloaded from a faculty website often
# does not define that style, and a link that renders as plain black text is a
# link the student cannot see and will never click.
_LINK_COLOR = "1155CC"


def _ox():
    from docx.oxml import OxmlElement  # noqa: PLC0415 — heavy, docx-only path
    from docx.oxml.ns import qn        # noqa: PLC0415
    return OxmlElement, qn


def _run_el(text: str, rpr_proto, *, link: bool):
    """A <w:r>, optionally styled as a link, inheriting the surrounding format.

    The rPr is rebuilt rather than copied-and-appended: CT_RPr is a sequence
    with a fixed child order, and appending w:color after w:lang produces XML
    Word will not open. Only the properties that matter for matching the
    paragraph's look are carried across.
    """
    OxmlElement, qn = _ox()
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    def carry(tag: str) -> None:
        if rpr_proto is None:
            return
        found = rpr_proto.find(qn(tag))
        if found is not None:
            rpr.append(copy.deepcopy(found))

    carry("w:rFonts")
    carry("w:b")
    carry("w:i")
    if link:
        color = OxmlElement("w:color")
        color.set(qn("w:val"), _LINK_COLOR)
        rpr.append(color)
    carry("w:sz")
    carry("w:szCs")
    if link:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    carry("w:lang")
    r.append(rpr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _next_bookmark_id(doc) -> int:
    """One past the highest id already in the document.

    A thesis arrives full of bookmarks — Word writes one per table-of-contents
    entry — and reusing an id silently breaks that TOC.
    """
    _, qn = _ox()
    top = 0
    for el in doc.element.body.iter(qn("w:bookmarkStart")):
        try:
            top = max(top, int(el.get(qn("w:id")) or 0))
        except ValueError:
            continue
    return top + 1


def _bookmark(paragraph, name: str, bid: int) -> None:
    """Make a paragraph the target of an internal link."""
    OxmlElement, qn = _ox()
    p = paragraph._element  # noqa: SLF001 — python-docx exposes no XML API
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    # After w:pPr, never before it: pPr must be the first child of w:p.
    ppr = p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(start)
    else:
        p.insert(0, start)
    p.append(end)


def _segments(paragraph) -> list[tuple[int, int, Any, str]]:
    """(start, end, run_element_or_None, text) for the visible text, in order.

    Two kinds of segment are OPAQUE — they count towards the offsets, and
    nothing is ever written into them:

      hyperlinks — a URL the student pasted, or a citation this pass linked a
        moment ago. `paragraph.runs` does not report their text at all, so
        measuring against runs alone would put every citation after them out of
        position by that length and land a link inside a word.

      runs that are not exactly one <w:t> — a run holding a tab or a line break
        alongside its text. Rewriting one of those through python-docx's
        Run.text replaces ALL of its content, which silently deletes the tab and
        moves the line.
    """
    _, qn = _ox()
    out: list[tuple[int, int, Any, str]] = []
    pos = 0
    for child in paragraph._element:  # noqa: SLF001
        if child.tag == qn("w:r"):
            ts = child.findall(qn("w:t"))
            text = "".join(t.text or "" for t in ts)
            out.append((pos, pos + len(text), child if len(ts) == 1 else None, text))
        elif child.tag == qn("w:hyperlink"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            out.append((pos, pos + len(text), None, text))
        else:
            continue
        pos = out[-1][1]
    return out


def _set_run_text(run_el, text: str) -> None:
    """Write a run's text without disturbing its other children.

    Only ever called on runs `_segments` marked writable, i.e. holding exactly
    one <w:t>.
    """
    _, qn = _ox()
    t = run_el.find(qn("w:t"))
    t.set(qn("xml:space"), "preserve")
    t.text = text


def _visible(segments: list[tuple[int, int, Any, str]]) -> str:
    """The paragraph text as the offsets in `segments` measure it.

    Not `paragraph.text`: that renders <w:tab/> and <w:br/> as characters this
    walk does not count, and one tab in a paragraph would shift every citation
    after it.
    """
    return "".join(t for _, _, _, t in segments)


def _wrap_span_as_link(paragraph, start: int, end: int, anchor: str) -> bool:
    """Turn text[start:end] of a paragraph into a hyperlink to `anchor`.

    Surgical on purpose. Rewriting the whole paragraph (what `_set_paragraph_text`
    does) would flatten every bold and italic in it, and unlike a rewrite, this
    pass touches nearly every paragraph in the thesis — so the formatting loss
    that is acceptable for a sentence being reworded is not acceptable here.
    Only the runs the citation actually overlaps are split.
    """
    segments = _segments(paragraph)
    if not segments:
        return False

    # Zero-length segments are skipped: a tab-only run sitting inside the span
    # is not text to rewrite, and blanking it would delete the tab.
    touched = [seg for seg in segments
               if seg[1] > seg[0] and seg[1] > start and seg[0] < end]
    if not touched or any(seg[2] is None for seg in touched):
        return False  # overlaps something opaque; leave the paragraph alone
    first_s, _, first_r, first_t = touched[0]
    last_s, last_e, last_r, last_t = touched[-1]
    if start < first_s or end > last_e:
        return False  # the span is not fully covered by writable runs

    proto = first_r.find(_ox()[1]("w:rPr"))
    link_el = _hyperlink(anchor, _visible(segments)[start:end], proto)

    _set_run_text(first_r, first_t[: start - first_s])
    if first_r is last_r:
        first_r.addnext(link_el)
        suffix = last_t[end - last_s:]
        if suffix:
            link_el.addnext(_run_el(suffix, proto, link=False))
    else:
        for seg in touched[1:-1]:
            _set_run_text(seg[2], "")
        _set_run_text(last_r, last_t[end - last_s:])
        first_r.addnext(link_el)
    return True


def _hyperlink(anchor: str, text: str, rpr_proto):
    OxmlElement, qn = _ox()
    h = OxmlElement("w:hyperlink")
    h.set(qn("w:anchor"), anchor)
    h.append(_run_el(text, rpr_proto, link=True))
    return h


def _link_citations(paragraphs: list[Any], body_idx: range,
                    anchors: dict[str, str]) -> int:
    """Link every in-text citation that has a reference entry to it."""
    linked = 0
    for i in body_idx:
        p = paragraphs[i]
        text = _visible(_segments(p))
        if "(" not in text:
            continue
        cits = [c for c in parse_intext_citations(text)
                if c.span[0] >= 0 and c.key in anchors]
        # Right to left: linking a span removes it from p.runs, so working
        # backwards keeps every earlier offset valid.
        for c in sorted(cits, key=lambda c: -c.span[0]):
            if _wrap_span_as_link(p, c.span[0], c.span[1], anchors[c.key]):
                linked += 1
    return linked


# --- document structure -----------------------------------------------------

def _ref_heading_index(paragraphs: list[Any]) -> int | None:
    """Index of the reference-list heading, or None. Last match wins."""
    found = None
    for i, p in enumerate(paragraphs):
        if _REF_HEADING.match((p.text or "").strip()):
            found = i
    return found


def _body_range(paragraphs: list[Any]) -> range:
    """Paragraph indices that are prose, i.e. everything before the references.

    Citations are parsed and claims are scanned only here. Running either over
    the reference list itself would re-cite the bibliography.
    """
    end = _ref_heading_index(paragraphs)
    return range(0, end if end is not None else len(paragraphs))


def _existing_reference_lines(paragraphs: list[Any]) -> list[str]:
    """The reference entries the student already typed.

    Read BEFORE the list is rebuilt, for two reasons. They are the best query
    CrossRef will ever get for these citations (a full bibliographic line beats
    "author + year" by a wide margin), and when a line still will not resolve it
    is what gets carried forward — replacing a real entry the student typed with
    "Tác giả (2019). [chưa đối chiếu được]" throws away work and information
    they already had.
    """
    ref_i = _ref_heading_index(paragraphs)
    if ref_i is None:
        return []
    raw = []
    for p in paragraphs[ref_i + 1:]:
        if _is_heading(p):
            break
        raw.append(p.text or "")
    return parse_reference_lines(raw)


def _eligible(p: Any) -> bool:
    text = (p.text or "").strip()
    if not text or _is_heading(p):
        return False
    return len(text.split()) >= _MIN_WORDS


def scan_cite_docx(body: bytes) -> dict:
    """What citing would touch, without touching it. No LLM, no charge.

    The confirm-before-you-spend step, same as the humanize scan: phase A is
    free, but phase B costs tokens per claim and the student should see how many
    claims that is before agreeing to it.
    """
    try:
        doc = _open(body)
    except Exception:  # noqa: BLE001
        logger.exception("cite_docx: could not open document")
        return {"ok": False, "error": "unreadable"}

    paragraphs = doc.paragraphs
    body_idx = _body_range(paragraphs)
    prose = [(i, (paragraphs[i].text or "").strip()) for i in body_idx
             if _eligible(paragraphs[i])]

    all_text = "\n".join((paragraphs[i].text or "") for i in body_idx)
    cits = parse_intext_citations(all_text)
    ref_i = _ref_heading_index(paragraphs)
    existing = 0
    if ref_i is not None:
        for p in paragraphs[ref_i + 1:]:
            if _is_heading(p):
                break
            if (p.text or "").strip():
                existing += 1

    return {
        "ok": True,
        "intext_citations": len(cits),
        "distinct_sources": len(dedupe(cits)),
        "existing_references": existing,
        "has_reference_section": ref_i is not None,
        "body_paragraphs": len(prose),
        "passages": len(_batches(prose)),
        "headings": sum(1 for p in paragraphs if _is_heading(p) and (p.text or "").strip()),
        "tables": len(doc.tables),
    }


def _write_reference_section(doc, paragraphs: list[Any],
                             records: list[dict]) -> None:
    """Replace (or create) the reference list, bookmarking every entry.

    Old entries are removed rather than appended to, because leaving them turns
    a fixed list into a list with two versions of every source. The walk stops
    at the next heading so an appendix after the references survives.

    Each entry gets its `anchor` bookmark on the way in — that is what the
    in-text hyperlinks jump to.
    """
    ref_i = _ref_heading_index(paragraphs)
    bid = _next_bookmark_id(doc)

    def emit(para, rec) -> None:
        nonlocal bid
        _bookmark(para, rec["anchor"], bid)
        bid += 1

    if ref_i is None:
        doc.add_paragraph("")
        doc.add_paragraph("TÀI LIỆU THAM KHẢO")
        for rec in records:
            emit(doc.add_paragraph(rec["text"]), rec)
        return

    for p in paragraphs[ref_i + 1:]:
        if _is_heading(p):
            break
        p._element.getparent().remove(p._element)  # noqa: SLF001 — python-docx has no public API

    cursor = paragraphs[ref_i]._element  # noqa: SLF001
    for rec in records:
        para = doc.add_paragraph(rec["text"])  # lands at the end of the body...
        cursor.addnext(para._element)          # noqa: SLF001 — ...then moves into place
        cursor = para._element                 # noqa: SLF001
        emit(para, rec)


def _insert_citation(text: str, sentence: str, citation: str) -> str | None:
    """Put `citation` at the end of `sentence` inside `text`.

    Exact substring match only. A fuzzy insertion that lands in the wrong place
    corrupts a student's sentence, and there is no way for them to tell which
    sentences were touched — so a claim we cannot locate is simply skipped.
    """
    if sentence not in text:
        return None
    trimmed = sentence.rstrip()
    if trimmed.endswith((".", "!", "?")):
        cited = f"{trimmed[:-1]} {citation}{trimmed[-1]}"
    else:
        cited = f"{trimmed} {citation}"
    return text.replace(sentence, cited, 1)


def _confirm(claim: str, candidates: list[dict], llm, usage: list) -> dict | None:
    """Ask the model which candidate really supports the claim. None is a valid
    and expected answer — it is what keeps a wrong source out of the thesis."""
    # Same gate as phase A: a supplemental file or an authorless stub must not
    # even be offered to the model, because a "yes" on one of those is a
    # reference line no supervisor will accept.
    candidates = [c for c in candidates if is_citable(c)]
    if not candidates:
        return None
    listing = []
    for i, c in enumerate(candidates):
        parts = (c.get("issued") or {}).get("date-parts") or [[]]
        year = parts[0][0] if parts and parts[0] else "n.d."
        authors = ", ".join(
            clean(a.get("family") or a.get("name") or "")
            for a in (c.get("author") or [])[:3])
        abstract = clean(c.get("abstract") or "")[:600]
        listing.append(
            f"[{i}] {clean((c.get('title') or [''])[0])} — {authors} ({year}). "
            f"{clean((c.get('container-title') or [''])[0])}\n    {abstract}".strip())
    raw = _invoke(llm, _VERIFY_PROMPT.format(claim=claim, candidates="\n\n".join(listing)),
                  usage)
    idx = _json_out(raw).get("index")
    if isinstance(idx, int) and 0 <= idx < len(candidates):
        return candidates[idx]
    return None


def cite_docx(
    body: bytes,
    *,
    add_missing: bool = True,
    llm=None,
    search_fn: Callable[[str, int], list[dict]] | None = None,
    resolve_fn: Callable[[InText, str | None], dict | None] | None = None,
) -> tuple[bytes | None, dict]:
    """Resolve the citations a document has, add the ones it needs, return it.

    `llm`, `search_fn` and `resolve_fn` are injectable so the walk can be tested
    without a model or a network.

    Returns (docx_bytes, report). The document is always returned when it could
    be opened, even if nothing resolved — phase A rebuilding a reference list of
    "unresolved" entries is still a truthful, useful answer.
    """
    search_fn = search_fn or crossref_search
    resolve_fn = resolve_fn or resolve_verbose

    try:
        doc = _open(body)
    except Exception:  # noqa: BLE001
        logger.exception("cite_docx: could not open document")
        return None, {"ok": False, "error": "unreadable"}

    paragraphs = doc.paragraphs
    body_idx = _body_range(paragraphs)
    all_text = "\n".join((paragraphs[i].text or "") for i in body_idx)
    vietnamese = uses_vietnamese_convention(all_text)
    usage: list[dict] = []

    # One record per reference line: its text, and every in-text key that should
    # link to it. Two mentions written differently ("Hair và cộng sự, 2010" and
    # "Hair et al., 2010") resolve to the same work and must share one entry —
    # keyed by text so the list never carries the same source twice.
    records: list[dict] = []
    by_text: dict[str, dict] = {}

    def add_entry(text: str, key: str | None) -> None:
        rec = by_text.get(text)
        if rec is None:
            rec = {"text": text, "keys": set()}
            by_text[text] = rec
            records.append(rec)
        if key:
            rec["keys"].add(key)

    # --- phase A: resolve what is already cited -----------------------------
    cits = dedupe(parse_intext_citations(all_text))
    old_lines = _existing_reference_lines(paragraphs)
    hints = {c.key: match_reference_line(c, old_lines) for c in cits}
    label = "chưa đối chiếu được" if vietnamese else "not verified"

    # Three workers, not four: a thesis is fifty-odd lookups fired in a burst,
    # and CrossRef throttles that. Being told a real reference does not exist
    # because we hit a rate limit is worse than taking a few more seconds.
    with ThreadPoolExecutor(max_workers=3) as pool:
        answers = list(pool.map(lambda c: resolve_fn(c, hints.get(c.key)), cits))

    # An injected resolve_fn (tests) answers with the record alone; the real one
    # also says how strong the match was.
    def unpack(answer) -> tuple[dict | None, str]:
        return answer if isinstance(answer, tuple) else (answer, "line")

    # A citation the student never listed can only be matched on a surname and a
    # year, and "Nunnally 1978" is both "Psychometric theory" and "1K Delay Line
    # Digitizer". The entry is real, but it may not be the one they meant, and
    # the student is the only one who can say. Silence here would be presenting
    # a guess as a checked reference.
    weak = "khớp theo tên và năm" if vietnamese else "matched on name and year"

    used_lines: set[str] = set()
    resolved = unresolved = weak_matches = 0
    for cit, answer in zip(cits, answers, strict=True):
        msg, how = unpack(answer)
        line = hints.get(cit.key)
        if line:
            used_lines.add(line)
        if msg:
            entry = format_reference(msg)
            if how == "author-year":
                entry = f"{entry} [{weak}]"
                weak_matches += 1
            add_entry(entry, cit.key)
            resolved += 1
        else:
            # Carried, not dropped and not invented. The student's own line wins
            # when they wrote one: it holds a title, a journal and page numbers
            # that "Hair và cộng sự (2019)" does not, and losing that is a worse
            # outcome than the failed lookup it is being labelled for.
            kept = f"{line} [{label}]" if line else f"{cit.authors} ({cit.year}). [{label}]"
            add_entry(kept, cit.key)
            unresolved += 1

    # Entries with no in-text citation anywhere in the body. Kept, because
    # deleting a line the student typed is not this tool's call — and flagged,
    # because an uncited entry in the list is the other half of what a
    # supervisor checks.
    orphan = "không thấy trích dẫn trong bài" if vietnamese else "not cited in the text"
    orphans = 0
    for line in old_lines:
        if line not in used_lines:
            add_entry(f"{line} [{orphan}]", None)
            orphans += 1

    # --- phase B: cite what is not ------------------------------------------
    added = marked = 0
    if add_missing:
        llm = llm or _get_llm(0.0)
        prose = [(i, (paragraphs[i].text or "").strip()) for i in body_idx
                 if _eligible(paragraphs[i])]
        by_idx = dict(prose)
        marker = MARKER_VI if vietnamese else MARKER_EN

        for batch in _batches(prose):
            if added >= _MAX_NEW_CITATIONS:
                break
            passage = "\n\n".join(by_idx[i] for i in batch)
            try:
                claims = _json_out(
                    _invoke(llm, _CLAIM_PROMPT.format(passage=passage), usage)
                ).get("claims") or []
            except Exception:  # noqa: BLE001
                logger.exception("cite_docx: claim pass failed for a batch")
                continue

            for claim in claims:
                if added >= _MAX_NEW_CITATIONS:
                    break
                sentence = (claim or {}).get("sentence") or ""
                query = (claim or {}).get("query") or sentence
                if not sentence.strip():
                    continue
                # Which paragraph of this batch holds the sentence.
                target = next((i for i in batch if sentence in (paragraphs[i].text or "")), None)
                if target is None:
                    continue

                try:
                    candidates = search_fn(query, 3)
                    match = _confirm(sentence, candidates, llm, usage) if candidates else None
                except Exception:  # noqa: BLE001
                    # One failed lookup or confirmation must not cost the student
                    # the other 200 paragraphs. Failing closed here means the
                    # claim falls through to the marker, never to a guess.
                    logger.exception("cite_docx: sourcing failed for one claim")
                    match = None

                if match:
                    form = intext_form(match, vietnamese=vietnamese)
                    inserted = _insert_citation(
                        paragraphs[target].text or "", sentence, form)
                    if inserted is None:
                        continue
                    _set_paragraph_text(paragraphs[target], inserted)
                    # Key the new entry the same way the linking pass will read
                    # the citation back out of the paragraph, so an inserted
                    # citation is clickable like any the student wrote.
                    parsed = parse_intext_citations(form)
                    add_entry(format_reference(match),
                              parsed[0].key if parsed else None)
                    added += 1
                else:
                    inserted = _insert_citation(
                        paragraphs[target].text or "", sentence, marker)
                    if inserted is None:
                        continue
                    _set_paragraph_text(paragraphs[target], inserted)
                    marked += 1

    # Alphabetical, then numbered — the anchor has to be assigned after the sort
    # or every link points one entry off.
    records.sort(key=lambda r: r["text"].casefold())
    anchors: dict[str, str] = {}
    for n, rec in enumerate(records, 1):
        rec["anchor"] = f"dtref{n}"
        for key in rec["keys"]:
            anchors.setdefault(key, rec["anchor"])

    # Link before writing the list: `_write_reference_section` rebuilds the tail
    # of the document, and the body indices this pass walks were taken from the
    # paragraph list as it stands now.
    linked = _link_citations(paragraphs, body_idx, anchors)
    _write_reference_section(doc, doc.paragraphs, records)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue(), {
        "ok": True,
        "resolved": resolved,
        "unresolved": unresolved,
        "weak": weak_matches,
        "orphans": orphans,
        "added": added,
        "marked": marked,
        "linked": linked,
        "references": len(records),
        "usage": usage,
    }
