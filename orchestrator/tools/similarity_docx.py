"""Document-level similarity & citation self-check — .docx in, .docx out.

Shaped like humanize_docx / cite_docx: the student uploads their thesis, gets a
free scan of what would be looked at, then pays for a run that returns the same
document with the findings highlighted and a summary appended.

WHAT THIS IS, AND WHAT IT IS NOT
--------------------------------
It is NOT a Turnitin scan and must never be presented as one. A corpus-backed
similarity check needs the web, a paper index and the institution's own
submissions; DoThesis has none of those, and orchestrator/tools/plagiarism.py
holds the vendor seam for the day one is bought. When a provider IS configured
this module runs it too and folds the matches in — but `corpus_checked` in the
report says plainly which happened, because "no matches" and "nobody looked"
are the two answers a student must never see confused.

What it does offline, and what is genuinely worth having:

  1. INTERNAL DUPLICATION — the same passage appearing twice in the student's
     own document. Real theses are assembled over months and repeat themselves;
     a supervisor notices, and so does a marker reading two chapters in a row.

  2. QUOTE HYGIENE — the thing that actually causes a similarity problem:
     - quoted text with no citation near it,
     - in-text citations with no matching entry in the reference list,
     - reference entries never cited anywhere in the body.

     None of that needs a corpus. All of it is what a supervisor sends a thesis
     back over, and (2) is the part that would have raised the Turnitin score in
     the first place.

The matching engine is quality/similarity.py — the same winnowing, the same
detection floor. A second implementation of shingling in this repo would be a
second set of thresholds to keep in step.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any, Callable, Optional

logger = logging.getLogger(__name__)

# A duplicated run must be at least this many tokens to report. Higher than
# quality/similarity.py's intra-chapter floor of 20: inside ONE document, shared
# methodology boilerplate ("as shown in the table below, the results indicate
# that…") is normal and flagging it trains the student to ignore the report.
MIN_DUPLICATE_TOKENS = 25

# Cap the report. A thesis that duplicates itself 200 times has one problem, not
# two hundred, and a wall of findings is read as noise.
MAX_DUPLICATES = 25
MAX_QUOTE_FINDINGS = 40
EXCERPT = 180


def _open(body: bytes):
    from docx import Document  # noqa: PLC0415 — heavy, only needed here
    return Document(io.BytesIO(body))


def _helpers():
    """The docx-walking rules, borrowed rather than re-derived.

    "Which paragraphs are the reference list" and "which are headings" already
    have one careful answer in humanize_docx, with the reasoning attached. A
    second copy here would drift the first time either is corrected.
    """
    from .humanize_docx import (  # noqa: PLC0415
        _eligible, _is_heading, _reference_indices,
    )
    return _reference_indices, _is_heading, _eligible


def _engine():
    from quality.similarity import (  # noqa: PLC0415 — lazy: quality imports
        _is_cited, citation_spans, matched_spans,  # orchestrator lazily too
        normalize_tokens, quote_regions,
    )
    return normalize_tokens, matched_spans, quote_regions, _is_cited, citation_spans


def _excerpt(text: str, start: int, end: int) -> str:
    frag = (text or "")[start:end].strip()
    return frag if len(frag) <= EXCERPT else frag[:EXCERPT].rstrip() + "…"


# --- free scan ----------------------------------------------------------------

def scan_docx(body: bytes) -> dict:
    """What a check would look at, without checking. No provider, no charge.

    Same contract as humanize_docx.scan_docx: the student sees the size of the
    job before spending anything on it.
    """
    try:
        doc = _open(body)
    except Exception:  # noqa: BLE001
        logger.exception("similarity_docx: could not open document")
        return {"ok": False, "error": "unreadable"}

    ref_indices, is_heading, _eligible = _helpers()
    _norm, _spans, quote_regions, _cited, citation_spans = _engine()

    paragraphs = doc.paragraphs
    refs = ref_indices(paragraphs)
    body_idx = [i for i, p in enumerate(paragraphs)
                if i not in refs and not is_heading(p) and (p.text or "").strip()]
    text = "\n".join((paragraphs[i].text or "") for i in body_idx)
    return {
        "ok": True,
        "paragraphs": len(paragraphs),
        "body_paragraphs": len(body_idx),
        "words": len(text.split()),
        "quotations": len(quote_regions(text)),
        "in_text_citations": len(citation_spans(text)),
        "reference_entries": len(refs),
    }


# --- reference-list cross-check -----------------------------------------------

# "Nguyen (2019)", "(Nguyen & Tran, 2019)", "(Ohanian, 1990; Wang, 2018)".
_YEAR_RE = re.compile(r"(19|20)\d{2}")
_SURNAME_RE = re.compile(r"[A-ZÀ-Ỹ][\w'’\-]+")


def _citation_keys(text: str) -> set[tuple[str, str]]:
    """(surname, year) pairs mentioned in the body."""
    _n, _m, _q, _c, citation_spans = _engine()
    keys: set[tuple[str, str]] = set()
    for s, e in citation_spans(text or ""):
        frag = text[s:e]
        year = _YEAR_RE.search(frag)
        if not year:
            continue
        for name in _SURNAME_RE.findall(frag):
            low = name.lower()
            if low in ("and", "et", "al", "và", "cộng", "sự"):
                continue
            keys.add((low, year.group(0)))
    return keys


def _reference_keys(entries: list[str]) -> list[tuple[str, set[tuple[str, str]]]]:
    """[(entry, {(surname, year), …})] — every key that entry could be cited by.

    Keyed per ENTRY, not per key. Flattened into one key→entry map, an entry was
    reported "never cited" as soon as ANY of its keys went unmatched — and every
    entry has several, because the surname pattern also picks up the initial and
    the first word of the title. "Ohanian, R. (1990). Construction and…" yielded
    ohanian/r/construction, matched on the first, and was still reported
    uncited on the other two.
    """
    out: list[tuple[str, set[tuple[str, str]]]] = []
    for entry in entries:
        year = _YEAR_RE.search(entry or "")
        if not year:
            continue
        # len > 1 drops the initial in "Ohanian, R." — a one-letter token is
        # never the surname a citation is written under.
        names = [n for n in _SURNAME_RE.findall(entry or "") if len(n) > 1]
        keys = {(n.lower(), year.group(0)) for n in names[:3]}
        if keys:
            out.append((entry.strip(), keys))
    return out


# --- the check ----------------------------------------------------------------

def similarity_docx(
    body: bytes,
    *,
    provider: Any | None = None,
    language: str = "vi",
    on_progress: Optional[Callable[[int, int], None]] = None,
) -> tuple[bytes | None, dict]:
    """Run the self-check and return (annotated .docx, report).

    `provider` is an orchestrator.tools.plagiarism.SimilarityProvider or None.
    None is not a failure and not a clean result — the report says
    `corpus_checked: false` and every caller must surface that distinction.
    """
    try:
        doc = _open(body)
    except Exception:  # noqa: BLE001
        logger.exception("similarity_docx: could not open document")
        return None, {"ok": False, "error": "unreadable"}

    ref_indices, is_heading, _elig = _helpers()
    normalize_tokens, matched_spans, quote_regions, is_cited, _cs = _engine()

    paragraphs = doc.paragraphs
    refs = ref_indices(paragraphs)
    body_idx = [i for i, p in enumerate(paragraphs)
                if i not in refs and not is_heading(p) and len((p.text or "").split()) >= 8]
    total_steps = 3 + (1 if provider is not None else 0)

    def _step(n):
        if on_progress:
            try:
                on_progress(n, total_steps)
            except Exception:  # noqa: BLE001 — progress must not sink the run
                logger.debug("similarity_docx: progress callback failed", exc_info=True)

    _step(0)

    # 1. Internal duplication ---------------------------------------------------
    #
    # Pairwise winnowing over every paragraph pair is O(n²) and a thesis has
    # hundreds. Index the fingerprints first and only compare paragraphs that
    # actually share one — the same trick a real similarity index uses, and the
    # difference between a second and a minute.
    toks = {i: normalize_tokens(paragraphs[i].text or "") for i in body_idx}
    from quality.similarity import shingle_hashes, winnow  # noqa: PLC0415

    fingerprint_owners: dict[int, list[int]] = {}
    for i in body_idx:
        words = [t for t, _s, _e in toks[i]]
        for h, _pos in winnow(shingle_hashes(words)):
            fingerprint_owners.setdefault(h, []).append(i)
    candidates: set[tuple[int, int]] = set()
    for owners in fingerprint_owners.values():
        uniq = sorted(set(owners))
        if len(uniq) < 2 or len(uniq) > 12:   # a phrase in 12 paragraphs is boilerplate
            continue
        for a_i, a in enumerate(uniq):
            for b in uniq[a_i + 1:]:
                candidates.add((a, b))

    duplicates = []
    for a, b in sorted(candidates):
        for m in matched_spans(toks[a], toks[b], MIN_DUPLICATE_TOKENS):
            # a_char, not a_start: a_start is a TOKEN index, and slicing the
            # paragraph text with it would quote the wrong words entirely.
            c0, c1 = m["a_char"]
            duplicates.append({
                "paragraph_a": a, "paragraph_b": b, "tokens": m["tokens"],
                "excerpt": _excerpt(paragraphs[a].text or "", c0, c1),
            })
    duplicates.sort(key=lambda d: -d["tokens"])
    duplicates = duplicates[:MAX_DUPLICATES]
    _step(1)

    # 2. Quote hygiene ----------------------------------------------------------
    uncited_quotations = []
    for i in body_idx:
        text = paragraphs[i].text or ""
        for s, e in quote_regions(text):
            if e - s < 40:                    # a quoted word is a scare quote
                continue
            if not is_cited(text, (s, e)):
                uncited_quotations.append(
                    {"paragraph": i, "excerpt": _excerpt(text, s, e)})
    uncited_quotations = uncited_quotations[:MAX_QUOTE_FINDINGS]
    _step(2)

    # 3. Body ↔ reference list --------------------------------------------------
    body_text = "\n".join((paragraphs[i].text or "") for i in body_idx)
    ref_entries = [(paragraphs[i].text or "").strip() for i in sorted(refs)
                   if (paragraphs[i].text or "").strip()]
    cited = _citation_keys(body_text)
    listed = _reference_keys(ref_entries)
    all_listed_keys = {k for _e, keys in listed for k in keys}
    cited_not_listed = sorted(
        {f"{n.title()} ({y})" for n, y in cited
         if (n, y) not in all_listed_keys})[:MAX_QUOTE_FINDINGS]
    # An entry is uncited only when NONE of its keys were cited.
    listed_not_cited = sorted(
        {e for e, keys in listed if not (keys & cited)})[:MAX_QUOTE_FINDINGS]
    _step(3)

    # 4. Corpus, only if somebody bought one ------------------------------------
    corpus_matches: list[dict] = []
    corpus_checked = False
    corpus_error = None
    if provider is not None:
        try:
            raw = provider.check(body_text, language=language)
            corpus_matches = list(raw.get("matches") or [])
            corpus_checked = True
        except Exception:  # noqa: BLE001
            logger.exception("similarity_docx: provider failed")
            corpus_error = "provider_error"
        _step(4)

    flagged = ({d["paragraph_a"] for d in duplicates}
               | {d["paragraph_b"] for d in duplicates}
               | {q["paragraph"] for q in uncited_quotations})
    report = {
        "ok": True,
        # The load-bearing field. False means NOBODY LOOKED at the web or any
        # paper index — not that the document is clean.
        "corpus_checked": corpus_checked,
        "corpus_error": corpus_error,
        "provider": getattr(provider, "name", None) if provider is not None else None,
        "internal_duplication": duplicates,
        "uncited_quotations": uncited_quotations,
        "cited_not_in_references": cited_not_listed,
        "references_never_cited": listed_not_cited,
        "corpus_matches": corpus_matches,
        "counts": {
            "body_paragraphs": len(body_idx),
            "reference_entries": len(ref_entries),
            "flagged_paragraphs": len(flagged),
            "internal_duplication": len(duplicates),
            "uncited_quotations": len(uncited_quotations),
            "cited_not_in_references": len(cited_not_listed),
            "references_never_cited": len(listed_not_cited),
        },
    }

    try:
        out = _annotate(doc, flagged, report, language)
    except Exception:  # noqa: BLE001 — the findings are the product; the
        logger.exception("similarity_docx: annotation failed")  # marked-up copy is not
        return None, {**report, "annotation_failed": True}
    return out, report


# --- annotated copy ------------------------------------------------------------

_L = {
    "en": {"title": "Similarity & citation self-check",
           "not_turnitin": "This is a SELF-CHECK, not a Turnitin scan. No external "
                           "corpus was searched, so it cannot tell you whether this "
                           "text appears anywhere else. Highlighted paragraphs are "
                           "the ones listed below.",
           "corpus_on": "An external similarity provider ({p}) was also searched.",
           "dup": "Repeated inside this document",
           "quotes": "Quoted text with no citation nearby",
           "cited_missing": "Cited in the text, absent from the reference list",
           "listed_missing": "In the reference list, never cited",
           "none": "Nothing flagged.",
           "para": "paragraph"},
    "vi": {"title": "Tự kiểm tra trùng lặp & trích dẫn",
           "not_turnitin": "Đây là bản TỰ KIỂM TRA, không phải bản quét Turnitin. "
                           "Không có nguồn dữ liệu bên ngoài nào được tra, nên báo cáo "
                           "này KHÔNG cho biết đoạn văn có xuất hiện ở nơi khác hay "
                           "không. Các đoạn được tô sáng là các đoạn liệt kê bên dưới.",
           "corpus_on": "Đã tra thêm nguồn đối chiếu bên ngoài ({p}).",
           "dup": "Lặp lại bên trong tài liệu",
           "quotes": "Trích dẫn nguyên văn nhưng không có nguồn kèm theo",
           "cited_missing": "Có trích dẫn trong bài nhưng thiếu trong danh mục tham khảo",
           "listed_missing": "Có trong danh mục tham khảo nhưng không được trích dẫn",
           "none": "Không phát hiện vấn đề nào.",
           "para": "đoạn"},
}


def _heading(doc, text: str, size: int):
    """A heading that does not depend on the document defining one.

    add_heading() resolves "Heading 1" through the student's own styles.xml, and
    a thesis template that names its headings anything else raises KeyError —
    losing the entire annotated copy over a font choice. Bold text at a size is
    indistinguishable in the output and cannot fail.
    """
    from docx.shared import Pt  # noqa: PLC0415
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return para


def _annotate(doc, flagged: set[int], report: dict, language: str) -> bytes:
    """Highlight the flagged paragraphs and append the findings.

    Paragraph-level, not span-level: a match can start mid-run and end in
    another, and splitting runs to colour an exact span is how a rewriter
    corrupts formatting. The summary carries the exact excerpts, so nothing is
    lost by colouring a little more than was matched.
    """
    from docx.enum.text import WD_COLOR_INDEX  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    L = _L.get("vi" if str(language).lower().startswith("vi") else "en", _L["en"])
    paragraphs = doc.paragraphs
    for i in flagged:
        if 0 <= i < len(paragraphs):
            for run in paragraphs[i].runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_page_break()
    _heading(doc, L["title"], 16)
    note = doc.add_paragraph(L["not_turnitin"])
    for run in note.runs:
        run.font.size = Pt(10)
        run.font.italic = True
    if report.get("corpus_checked"):
        doc.add_paragraph(L["corpus_on"].format(p=report.get("provider") or "?"))

    sections = [
        (L["dup"], [f'{L["para"]} {d["paragraph_a"]} ↔ {d["paragraph_b"]} '
                    f'({d["tokens"]} tokens): {d["excerpt"]}'
                    for d in report["internal_duplication"]]),
        (L["quotes"], [f'{L["para"]} {q["paragraph"]}: {q["excerpt"]}'
                       for q in report["uncited_quotations"]]),
        (L["cited_missing"], report["cited_not_in_references"]),
        (L["listed_missing"], report["references_never_cited"]),
    ]
    empty = True
    for heading, items in sections:
        if not items:
            continue
        empty = False
        _heading(doc, f"{heading} ({len(items)})", 13)
        for item in items:
            # A literal bullet, not style="List Bullet". The student's document
            # supplies the styles, and a thesis template that does not define
            # that one raises KeyError — which on a real upload cost the whole
            # annotated copy. Same trap as the cover page's Author/Date styles.
            doc.add_paragraph(f"• {item}")
    if empty:
        doc.add_paragraph(L["none"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
