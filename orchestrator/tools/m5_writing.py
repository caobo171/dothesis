"""M5 — Writing & Export tools.

These tools are the bridge that lets the auto-mode orchestrator produce
the same final artifacts (docx + pdf) that today's `python -m engine` ships.
They delegate into engine/phases/compose, engine/phases/compile, and
engine/utils/* so we don't reimplement the layout/formatting logic.

Sub-project 1 note: where the engine's actual function signature doesn't match
the orchestrator's tool contract, the wrapper falls back to a minimal
implementation. Proper engine integration is a follow-on sub-project.
"""
from __future__ import annotations

import json
import logging
import os
import re
import sys
from pathlib import Path
from uuid import uuid4

from orchestrator.message_utils import text_of  # flatten Gemini 3.x list content

import boto3
from langchain_core.tools import tool

# Make engine package importable.
_ROOT = Path(__file__).resolve().parents[2]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))
# AND put engine/ itself on the path. The engine modules use repo-internal
# imports like `from utils.logging_config import get_logger` (engine/ as the
# root), so importing them as `engine.utils.*` from the api process fails at
# THAT line unless engine/ is also a sys.path root. This was the cause of M5
# exports landing as 24-byte placeholders (0 KB): _export_docx_via_engine's
# `from engine.utils.export_professional import …` raised ModuleNotFoundError
# on the internal `from utils.…`, the except branch wrote the placeholder.
_ENGINE = _ROOT / "engine"
if str(_ENGINE) not in sys.path:
    sys.path.insert(0, str(_ENGINE))

logger = logging.getLogger(__name__)

# Directory where M5 chapter prompt templates live.
_PROMPT_DIR = Path(__file__).resolve().parent.parent / "prompts" / "m5"


def _scratch_dir() -> Path:
    d = Path(os.getenv("ORCHESTRATOR_SCRATCH", "/tmp/orchestrator_scratch"))
    d.mkdir(parents=True, exist_ok=True)
    return d


def s3_from_env():
    """S3 client factory — mirrors the SP2 api/app/routers/uploads.py pattern.
    Indirection point so tests can monkeypatch easily.
    """
    return boto3.client(
        "s3",
        region_name=os.environ.get("AWS_REGION", "us-east-1"),
        aws_access_key_id=os.environ.get("AWS_ACCESS_KEY"),
        aws_secret_access_key=os.environ.get("AWS_SECRET_KEY"),
    )


_CONTENT_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":  "application/pdf",
}


def _upload_to_s3(local_path: str, project_id: str, kind: str, filename: str) -> tuple[str, int]:
    """Upload a local artifact to S3 under projects/{project_id}/exports/.

    Returns (s3_key, size_bytes). Deletes the local file after upload.
    Reading the file bytes before upload lets us capture size without a
    separate stat call, and avoids a second open() after write.
    """
    s3 = s3_from_env()
    # The project convention is S3_BUCKET (settings.py, job_runner, uploads,
    # engine all use it); job_runner sets S3_BUCKET when launching this auto
    # subprocess. Reading AWS_S3_BUCKET here silently broke M5 export in the real
    # flow. Prefer S3_BUCKET, keep AWS_S3_BUCKET as a back-compat fallback.
    bucket = os.environ.get("S3_BUCKET") or os.environ["AWS_S3_BUCKET"]
    s3_key = f"projects/{project_id}/exports/{filename}"
    content_type = _CONTENT_TYPES[kind]
    body = Path(local_path).read_bytes()
    size_bytes = len(body)
    s3.put_object(
        Bucket=bucket, Key=s3_key, Body=body,
        ContentType=content_type,
    )
    Path(local_path).unlink(missing_ok=True)
    return s3_key, size_bytes


_CITE_PATTERN = re.compile(r"\((?P<author>[A-Z][\w-]+(?: et al\.)?), (?P<year>\d{4})\)")


def _ref_author_label(ref: dict) -> str:
    """Derive the inline-citation author label from a reference record.

    M2 sources carry `authors` as a LIST (e.g. ["Nguyen", "Tran"] or ["NIH"]);
    earlier code read a singular `author` key that doesn't exist, so EVERY
    reference collapsed to "Anon" — which is exactly what the LLM then cited
    ("(Anon, 2011)"). This reads the real list: first author's surname, plus
    "et al." when there are co-authors. Falls back to a singular `author` field
    or "Anon" only when nothing usable is present.
    """
    authors = ref.get("authors")
    if isinstance(authors, list) and authors:
        first = str(authors[0]).strip()
        surname = first.split()[-1] if first else ""
        if surname:
            return f"{surname} et al." if len(authors) > 1 else surname
    single = str(ref.get("author") or "").strip()
    return single or "Anon"


def _ref_citation_key(ref: dict) -> tuple[str, str]:
    """(author_label, year_str) used as the canonical pool key for a reference,
    shared by the prompt formatter and the citation validators so they agree."""
    return (_ref_author_label(ref), str(ref.get("year", "")).strip())

# SP6.5: separate regex used by validate_citations_plain for the autosave PATCH
# endpoint. Broader than _CITE_PATTERN — accepts any author token and n.d. years
# so the inline autosave validator is tolerant of varied LLM citation styles.
_CITATION_REGEX = re.compile(r"\(([^)]+?),\s*(\d{4}|n\.d\.)\)")


def validate_citations_plain(prose: str, reference_pool: list[dict]) -> dict:
    """Extract (Author, Year) patterns; classify as used vs uncited based on pool.

    SP6.5: called directly by the autosave PATCH endpoint without @tool overhead.
    Uses _CITATION_REGEX (broader pattern, accepts n.d.) unlike the older
    _CITE_PATTERN. Returns {"citations_used": [...], "uncited_warnings": [...]}
    preserving first-occurrence order with no duplicates.
    """
    # Decision: strip whitespace from pool keys so "Smith " and "Smith" match,
    # and convert year to string to align with the regex group (always a string).
    pool_keys = {_ref_citation_key(r) for r in reference_pool}
    seen: dict[str, bool] = {}    # ordered dedupe via insertion-order dict
    citations_used: list[str] = []
    uncited: list[str] = []
    for match in _CITATION_REGEX.finditer(prose):
        author, year = match.group(1).strip(), match.group(2).strip()
        text = f"({author}, {year})"
        if text in seen:
            continue
        seen[text] = True
        if (author, year) in pool_keys:
            citations_used.append(text)
        else:
            uncited.append(text)
    return {"citations_used": citations_used, "uncited_warnings": uncited}


def validate_citations(prose: str, references: list[dict]) -> tuple[list[str], list[str]]:
    """Regex-scan prose for (Author, Year) patterns; partition into
    (cited_in_pool, uncited). Each returned list is de-duplicated and
    preserves first-occurrence order.

    Plain Python helper (not a @tool) — used by compose_chapter +
    rewrite_chapter post-validation.
    """
    # Build a pool of (author_label, year) tuples from references via the
    # shared key helper — reads the `authors` LIST, not a non-existent
    # `author` field (the bug that made every reference "Anon").
    pool = {_ref_citation_key(r) for r in references}
    cited: list[str] = []
    uncited: list[str] = []
    seen: set[tuple[str, str]] = set()
    for m in _CITE_PATTERN.finditer(prose):
        key = (m.group("author"), m.group("year"))
        if key in seen:
            continue
        seen.add(key)
        label = f"{m.group('author')}, {m.group('year')}"
        if key in pool:
            cited.append(label)
        else:
            uncited.append(label)
    return cited, uncited


# --- Wrappers that tests can monkeypatch ---------------------------------

def _compose_section_via_engine(section_name: str, context_store: dict) -> str:
    """Compose one section using engine/phases/compose helpers.

    If the engine's section composer isn't found or fails, return a fallback
    formed from the context_store (so the auto-mode pipeline keeps moving).
    """
    try:
        from engine.phases.context import DraftContext
        from engine.phases import compose as _compose
    except Exception as e:
        logger.warning("engine.phases.compose import failed: %s", e)
        return _fallback_section(section_name, context_store)

    try:
        ctx = DraftContext()
        ctx.topic = (context_store.get("m1_topic") or {}).get("research_title", "Untitled")
        ctx.language = (context_store.get("m1_topic") or {}).get("language", "en")
        ctx.scribe_output = (context_store.get("m2_literature") or {}).get("literature_review_doc", "")
        ctx.signal_output = "\n".join(
            g.get("description", "") for g in
            (context_store.get("m2_literature") or {}).get("research_gaps", [])
        )

        # Try documented private composers first; fall back to public run_compose_phase.
        composer = {
            "intro":       getattr(_compose, "_compose_intro", None),
            "lit_review":  getattr(_compose, "_compose_lit_review", None),
            "methodology": getattr(_compose, "_compose_methodology", None),
            "results":     getattr(_compose, "_compose_results", None),
            "conclusion":  getattr(_compose, "_compose_conclusion", None),
        }.get(section_name)
        if composer is None:
            return _fallback_section(section_name, context_store)
        return composer(ctx)
    except Exception as e:
        logger.warning("engine compose failed for %s: %s", section_name, e)
        return _fallback_section(section_name, context_store)


def _fallback_section(section_name: str, context_store: dict) -> str:
    """Minimal text we can emit when engine integration isn't wired up yet."""
    title = (context_store.get("m1_topic") or {}).get("research_title", "Untitled")
    return f"# {section_name.title()}\n\n[Auto-generated for '{title}']\n"


def _validate_via_engine(text: str) -> dict:
    """Run engine validation. Sub-project 1 fallback if signature differs."""
    try:
        from engine.phases.validate import quick_validate
        return quick_validate(text)
    except Exception:
        # Fallback: trivial sanity check
        issues = []
        if len(text) < 100:
            issues.append({"kind": "too_short", "message": "draft <100 chars"})
        return {"issues": issues, "score": 1.0 - 0.5 * len(issues)}


def _sections_to_markdown(sections: list[dict]) -> str:
    """Flatten [{title, prose}] into a single markdown document.

    The engine renderers (engine/utils/export_professional) consume a
    markdown *file*, not a Python list — so this is the missing adapter
    that lets the orchestrator's section list reach the real Pandoc-backed
    export. `prose` is assumed to already be markdown (chapter composers
    emit markdown); we only prepend an H1 per section title.
    """
    parts: list[str] = []
    for sec in sections:
        # Accept both section conventions: compose_all_sections emits
        # {title, prose}; the M5 agent's _build_sections_for_export historically
        # emitted {name, text}. Reading only title/prose silently dropped every
        # chapter's heading AND body for the agent path → a styles-only (blank)
        # docx/pdf. Normalize both here so neither producer can lose content.
        title = (sec.get("title") or "").strip()
        if not title and sec.get("name"):
            title = M5_CHAPTER_TITLES.get(
                sec["name"], str(sec["name"]).replace("_", " ").title())
        prose = (sec.get("prose") or sec.get("body")
                 or sec.get("text") or sec.get("content") or "").strip()
        # Defensive last line: strip internal placeholder/QA text so it can
        # never reach the rendered document, even on a forced export, then
        # normalize inline-bullet runs into proper markdown lists.
        prose = _sanitize_prose(_normalize_prose_markdown(
            _mermaid_to_prose(_split_run_on_hypotheses(_scrub_internal_markers(prose)))
        ))
        # Defense-in-depth against the duplicate chapter heading: the prompt
        # tells the LLM not to emit the chapter title, but Gemini sometimes
        # leads with its own "# Chương N: …" / "## <chapter title>" anyway, and
        # we prepend `# {title}` below — yielding the heading twice. Strip a
        # leading chapter-level heading (but never a numbered sub-section like
        # "## 4.1 …", which is real content).
        if title:
            prose = _strip_leading_chapter_heading(prose, title)
            parts.append(f"# {title}")
        if prose:
            parts.append(prose)
    md = "\n\n".join(parts) + "\n"
    # Normalize long dashes to a plain hyphen across the whole output (chapter
    # titles + prose): the LLM peppers prose with em-dashes and our headings use
    # them, which reads as machine-generated. Runs AFTER list-reflow (which must
    # distinguish en-dashes from ASCII hyphens), so it can't disturb that.
    md = md.replace("—", "-").replace("–", "-")
    return md


def _strip_leading_chapter_heading(prose: str, title: str | None = None) -> str:
    """Drop a leading chapter-title line the LLM emitted despite the prompt.

    We prepend `# {title}` to every section, so a chapter that opens with its
    own title renders the title twice (the "Chương N appears twice" bug). Remove
    the first non-blank line ONLY when it is a chapter title in one of these
    shapes, and never a numbered sub-section (`4.1`, `2.3.1`) which is content:
      (a) an ATX heading `#`..`######` (any level) that isn't a sub-section;
      (b) a NON-heading line the LLM wrote as plain/bold text, e.g.
          `CHƯƠNG 4: KẾT QUẢ NGHIÊN CỨU` or `**Chapter 4 - Results**` — this is
          what pandoc rendered as a stray `Normal` paragraph under the H1;
      (c) a line whose marker-stripped text equals the prepended `title`.
    """
    if not prose:
        return prose
    lines = prose.split("\n")
    i = 0
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i >= len(lines):
        return prose
    line = lines[i]
    rest = lambda: "\n".join(lines[i + 1:]).strip()

    # (a) ATX heading (widened to any level) that is not a numbered sub-section.
    m = re.match(r"^\s*#{1,6}\s+(.+?)\s*$", line)
    if m and not re.match(r"^\d+(\.\d+)+\b", m.group(1).strip()):
        return rest()

    # marker-stripped text: drop leading `#`s and surrounding `**` bold markers.
    bare = re.sub(r"^\s*#*\s*", "", line.strip())
    bare = re.sub(r"^\*\*\s*|\s*\*\*$", "", bare).strip()
    # (b) a title-shaped chapter line (bounded length so real sentences that
    #     merely start "Chương 4: trình bày…" — long — are never stripped).
    if re.match(r"(?i)^(CHƯƠNG|CHAPTER)\s+\d+\s*[:.\-–—]\s*.{0,80}$", bare) \
            and not re.match(r"^\d+(\.\d+)+\b", bare):
        return rest()
    # (c) exact match against the prepended title (any marker form).
    if title and bare and bare.casefold() == title.strip().casefold():
        return rest()
    return prose


def _scrub_internal_markers(prose: str) -> str:
    """Remove engine-internal placeholder/QA text that must never appear in a
    final thesis (composition stubs, the old uncited-citation warning block)."""
    if not prose:
        return prose
    # Drop the legacy "⚠️ … may be hallucinated …" blockquote if any survived.
    prose = re.sub(r"\n?>\s*⚠️[^\n]*\n?", "", prose)
    # Drop bracketed composition placeholders like "[Composition failed …]"
    # and "[Auto-generated for '…']".
    prose = re.sub(r"\[(?:Composition failed|Auto-generated for)[^\]]*\]", "", prose)
    # Defense at the last markdown boundary: DT placement tokens are consumed
    # by results_render.weave. If a chapter bypassed that path or the matching
    # M4 data does not exist yet, never print the internal token in the thesis.
    prose = re.sub(r"^[ \t]*\[\[DT:[a-z0-9_]+\]\][ \t]*$", "", prose,
                   flags=re.MULTILINE | re.IGNORECASE)
    return prose.strip()


def _split_run_on_hypotheses(prose: str) -> str:
    """Break a paragraph that lists H1/H2/H3… inline into ONE paragraph per H.

    The M5 discussion composer sometimes emits every hypothesis as a
    continuous run: "Cụ thể: - H1: KOL … (β=…). - H2: HATH … (β=…). - H3: …"
    That reads as a wall of prose because the separators aren't at line start
    (so `_normalize_prose_markdown` doesn't fire). Insert a real paragraph
    break before each `Hn:` marker so each hypothesis stands as its own block.

    Handles: `- H1:`, `– H1:`, `- **H1:**`, `**H1:**`, with the leading dash
    optional; the leading intro `Cụ thể: -` stays with its sentence.
    """
    if not prose or "H1" not in prose:
        return prose
    out: list[str] = []
    marker = re.compile(r"[ \t]+[-–—][ \t]+(?=(?:\*\*)?H\d{1,2}:)")
    for line in prose.split("\n"):
        # Decision: only split a genuinely inline RUN (H1 and H2 on the same
        # source line). A normal Markdown list item has one H marker, and the
        # generated model relationship has one inside parentheses. The old
        # global regex split both, leaving a dangling '-' that Pandoc interpreted
        # as malformed list/heading markup and produced giant page gaps.
        if len(re.findall(r"(?:\*\*)?H\d{1,2}:", line)) < 2:
            out.append(line)
            continue
        parts = marker.split(line)
        out.append(parts[0].rstrip())
        for part in parts[1:]:
            out.extend(["", part.strip()])
    return "\n".join(out)


def _render_mermaid_png(mmd_source: str, out_path: Path) -> bool:
    """Render mermaid source to a PNG via the local @mermaid-js/mermaid-cli.

    Puppeteer's own Chrome-for-Testing binary is used (the OS Chrome install
    fails to connect over the puppeteer WS socket on macOS). Returns True on a
    real PNG, False when the CLI is missing / times out — the caller falls back
    to a text list so the document still ships.
    """
    import os as _os
    import subprocess as _subprocess
    import shutil as _shutil

    engine_root = Path(__file__).resolve().parents[2] / "engine"
    mmdc_dir = engine_root / "tools" / "mermaid_cli"
    mmdc = mmdc_dir / "node_modules" / ".bin" / "mmdc"
    if not mmdc.exists():
        return False
    # Locate puppeteer's own Chrome. It gets installed under ~/.cache/puppeteer
    # (or the value of PUPPETEER_CACHE_DIR). Pick the highest-versioned dir.
    chrome_bin: Path | None = None
    cache_dir = Path(_os.environ.get("PUPPETEER_CACHE_DIR",
                                     Path.home() / ".cache" / "puppeteer"))
    chrome_root = cache_dir / "chrome"
    if chrome_root.exists():
        versions = sorted(chrome_root.iterdir(), key=lambda p: p.name, reverse=True)
        for v in versions:
            # macOS layout: <v>/chrome-mac-arm64/Google Chrome for Testing.app/…
            for candidate in v.rglob("Google Chrome for Testing"):
                if candidate.is_file():
                    chrome_bin = candidate
                    break
            if chrome_bin:
                break
            # Linux layout: <v>/chrome-linux64/chrome
            for candidate in v.rglob("chrome"):
                if candidate.is_file():
                    chrome_bin = candidate
                    break
            if chrome_bin:
                break

    mmd_path = out_path.with_suffix(".mmd")
    mmd_path.write_text(mmd_source, encoding="utf-8")
    puppeteer_cfg = mmdc_dir / "puppeteer.json"
    if not puppeteer_cfg.exists():
        puppeteer_cfg.write_text(
            '{ "args": ["--no-sandbox", "--disable-setuid-sandbox"] }',
            encoding="utf-8",
        )

    env = _os.environ.copy()
    if chrome_bin:
        env["PUPPETEER_EXECUTABLE_PATH"] = str(chrome_bin)
    cmd = [str(mmdc), "-p", str(puppeteer_cfg), "-i", str(mmd_path),
           "-o", str(out_path), "-w", "1100", "-b", "white"]
    try:
        result = _subprocess.run(cmd, env=env, capture_output=True, timeout=45)
    except (_subprocess.TimeoutExpired, OSError) as e:
        logger.warning("mermaid render failed: %s", e)
        return False
    finally:
        try:
            mmd_path.unlink(missing_ok=True)
        except Exception:
            pass
    if result.returncode != 0:
        logger.warning("mmdc exit=%s stderr=%s",
                       result.returncode, result.stderr[:200].decode(errors="ignore"))
        return False
    return out_path.exists() and out_path.stat().st_size > 200


def _mermaid_to_prose(prose: str) -> str:
    """DOCX/PDF can't render Mermaid syntax directly. Replace ```mermaid``` (or
    a bare `flowchart LR …`) with a rendered PNG image + a text-list caption
    (relationships) so the document has BOTH the visual and the source of truth.

    On render failure (mmdc missing / puppeteer timeout) the diagram falls back
    to just the text list — the doc still ships, just without the picture.
    """
    low = prose.lower()
    if "flowchart" not in low and "graph " not in low and "mermaid" not in low:
        return prose
    # Strip mermaid's own label quotes (`GO["Trust"]` → Trust) so the
    # reconstructed diagram doesn't double-wrap them ('Trust') in the figure.
    labels = {k: v.strip().strip('"').strip("'")
              for k, v in re.findall(r"([A-Za-z0-9_]+)\[([^\]]+)\]", prose)}
    edges: list[tuple[str, str, str]] = []
    for em in re.finditer(
        r"([A-Za-z0-9_]+)(?:\[[^\]]*\])?\s*-\.?-?->\s*(?:\|([^|]*)\|)?\s*([A-Za-z0-9_]+)(?:\[[^\]]*\])?",
        prose,
    ):
        s = (labels.get(em.group(1), em.group(1)) or "").strip()
        t = (labels.get(em.group(3), em.group(3)) or "").strip()
        edges.append((s, t, (em.group(2) or "").strip()))
    if not edges:
        return prose  # not a parseable diagram — don't risk mangling real prose

    # Reconstruct a clean mermaid source from what we parsed (drops any prose
    # accidentally interleaved with the diagram lines) — this is what we hand
    # to mmdc. Nodes use their long labels; edges keep the pipe-labeled arrow
    # so the hypothesis id ("H1: +") shows on the arrow.
    # Regex with THREE groups: src_id, edge_label (optional), tgt_id.
    edge_re = re.compile(
        r"([A-Za-z0-9_]+)(?:\[[^\]]*\])?\s*-\.?-?->\s*(?:\|([^|]*)\|)?\s*([A-Za-z0-9_]+)"
    )
    parsed_edges = [(m.group(1), (m.group(2) or "").strip(), m.group(3))
                    for m in edge_re.finditer(prose)]
    mmd_lines = ["flowchart LR"]
    seen_nodes: set[str] = set()
    for src_id, _e, tgt_id in parsed_edges:
        for nid in (src_id, tgt_id):
            if nid in seen_nodes:
                continue
            seen_nodes.add(nid)
            lbl = labels.get(nid, nid).replace('"', "'")
            mmd_lines.append(f'    {nid}["{lbl}"]')
    for s_id, e_lbl, t_id in parsed_edges:
        # Preserve the moderator's dashed arrow through the reconstruction so the
        # rendered PNG keeps it visually distinct from the IV→DV hypothesis arrows.
        arrow = "-.->" if _is_moderation_label(e_lbl) else "-->"
        if e_lbl:
            mmd_lines.append(f'    {s_id} {arrow}|{e_lbl}| {t_id}')
        else:
            mmd_lines.append(f'    {s_id} {arrow} {t_id}')
    mmd_source = "\n".join(mmd_lines)

    # Drop mermaid syntax lines from the surrounding prose.
    keep: list[str] = []
    for line in prose.split("\n"):
        st = line.strip().strip("`").strip()
        if re.match(r"^(flowchart|graph)\s+\w+", st, re.I):
            continue
        if st.lower() == "mermaid":
            continue
        if "-->" in st or "-.->" in st:
            continue
        if re.match(r"^[A-Za-z0-9_]+\[[^\]]*\]$", st):
            continue
        keep.append(line)

    # Try to render to PNG (in the same scratch dir the md file lives in, so
    # pandoc's relative-image resolution finds it via absolute path).
    img_line = ""
    try:
        png_path = _scratch_dir() / f"conceptmodel-{uuid4().hex[:8]}.png"
        if _render_mermaid_png(mmd_source, png_path):
            # Pandoc: `![caption](abs/path/img.png)`. Absolute path is fine.
            img_line = f'\n![Mô hình nghiên cứu]({png_path})\n'
    except Exception:
        logger.exception("mermaid render step failed")

    rel = ["", "**Mối quan hệ giả thuyết trong mô hình:**", ""]
    solid = [(s, t, e) for s, t, e in edges if not _is_moderation_label(e)]
    for s, t, e in solid:
        rel.append(f"- {s} → {t}" + (f" ({e})" if e else ""))
    # Moderator edges point at the predictors; collapse them into ONE sentence
    # per moderator, naming the moderated OUTCOME (the common target of the
    # solid IV→DV edges) instead of listing each predictor arrow separately.
    dv = solid[0][1] if solid else ""
    seen_mod: list[str] = []
    for s, t, e in edges:
        if _is_moderation_label(e) and s not in seen_mod:
            seen_mod.append(s)
            tail = f" và {dv}" if dv else ""
            rel.append(f"- {s} điều tiết mối quan hệ giữa các biến độc lập{tail}")
    return "\n".join(keep).rstrip() + img_line + "\n" + "\n".join(rel) + "\n"


def _mapping_or_empty(value: object) -> dict:
    """Return a mapping for structured state, including legacy JSON strings.

    Decision: older interactive M3 rows can contain a prose or JSON-string
    ``conceptual_model``. Export is a read boundary and must degrade gracefully
    instead of assuming every historical row already matches today's schema.
    """
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except (TypeError, ValueError, json.JSONDecodeError):
            return {}
        return parsed if isinstance(parsed, dict) else {}
    return {}


def _derive_scale_items(conceptual_model: dict | str | None,
                        instrument: dict | str | None = None) -> list[dict]:
    """Return the questionnaire as [{construct, items:[text,...]}, ...] robustly
    across BOTH M3 shapes we see in the wild:

      A) interactive widget → per-node Likert items on
         `conceptual_model.nodes[].questions`.
      B) headless/partner backfill → a flat `instrument.items`
         ([{id, text, construct}]) with the conceptual_model carrying only
         node metadata (id/label/definition) and NO `questions`.

    The old `_scale_items_from_conceptual_model` only handled (A), so headless
    reports silently shipped an EMPTY {scale_items} → Chapter 3 had no
    measurement table (reviewer feedback: "chưa có bảng hỏi"). Shape B is the
    fallback here. Construct ids are mapped to their human label when the
    conceptual_model declares one, so the table reads "Định hướng Mục tiêu"
    not "GO".
    """
    cm = _mapping_or_empty(conceptual_model)
    nodes = cm.get("nodes") or []
    label_by_id = {
        str(n.get("id")): (n.get("label") or n.get("id"))
        for n in nodes if isinstance(n, dict) and n.get("id")
    }

    # Shape A: items already live on the nodes.
    out: list[dict] = []
    for n in nodes:
        if not isinstance(n, dict):
            continue
        qs = [q for q in (n.get("questions") or []) if q]
        if qs:
            out.append({"construct": n.get("label") or n.get("id"), "items": qs})
    if out:
        return out

    # Shape B: group the flat instrument items by construct, preserving order.
    items = _mapping_or_empty(instrument).get("items") or []
    grouped: dict[str, list[str]] = {}
    order: list[str] = []
    for it in items:
        if not isinstance(it, dict):
            continue
        text = (it.get("text") or "").strip()
        if not text:
            continue
        cid = str(it.get("construct") or "")
        disp = label_by_id.get(cid, cid) or "—"
        if disp not in grouped:
            grouped[disp] = []
            order.append(disp)
        grouped[disp].append(text)
    return [{"construct": c, "items": grouped[c]} for c in order]


def _generate_scale_items(instrument: dict | str | None, language: str = "en",
                          conceptual_model: dict | str | None = None) -> list[dict]:
    """Generate real Likert items when the instrument is only a SPEC — i.e. it
    lists `constructs` + `items_per_construct` but no actual item texts (a shape
    headless often emits). Without this the questionnaire table is either empty
    or, worse, hallucinated inline by the chapter LLM. Generating the scale is
    legitimate instrument design (not fabricating data), and reuses the same
    tool the interactive flow uses so wording stays consistent.

    Fallback: when there is NO instrument at all (a common headless/partner
    outcome — the agent designed the model but never a survey), derive the
    construct list straight from the conceptual_model's node labels so a
    quantitative methodology always ships a real measurement scale for every
    latent construct instead of only describing one in prose (reviewer: "không
    hề có bảng hỏi").
    """
    inst = _mapping_or_empty(instrument)
    constructs = [str(c).strip() for c in (inst.get("constructs") or []) if str(c).strip()]
    if not constructs:
        # No instrument spec → measure every construct the model declares.
        nodes = _mapping_or_empty(conceptual_model).get("nodes") or []
        constructs = [str(n.get("label") or n.get("id")).strip()
                      for n in nodes if isinstance(n, dict) and (n.get("label") or n.get("id"))]
        constructs = [c for c in constructs if c]
    if not constructs:
        return []
    n = inst.get("items_per_construct")
    try:
        n = int(n)
    except (TypeError, ValueError):
        n = 4
    n = max(1, min(n, 8))
    try:  # lazy import — avoids an m5↔m3 import cycle at module load
        from orchestrator.tools.m3_design import suggest_scale_items_batch
        batch = suggest_scale_items_batch.invoke(
            {"constructs": constructs, "n": n, "language": language})
    except Exception:
        logger.warning("_generate_scale_items: suggest_scale_items_batch failed", exc_info=True)
        return []
    out: list[dict] = []
    for c in constructs:
        items = [str(it.get("text")).strip() for it in (batch.get(c) or [])
                 if isinstance(it, dict) and str(it.get("text") or "").strip()]
        if items:
            out.append({"construct": c, "items": items})
    return out


def _collect_construct_labels(conceptual_model: dict | str | None,
                              instrument: dict | str | None,
                              scale_items: list[dict] | None = None) -> list[str]:
    """All construct/variable labels that appear in the model + questionnaire,
    de-duplicated in first-seen order — the set we may need to localize."""
    cm = _mapping_or_empty(conceptual_model)
    labels: list[str] = []

    def add(x):
        if isinstance(x, dict):
            x = x.get("label") or x.get("name") or x.get("id")
        s = str(x or "").strip()
        if s:
            labels.append(s)

    for n in cm.get("nodes") or []:
        if isinstance(n, dict):
            add(n.get("label") or n.get("id"))
    for c in cm.get("constructs") or []:  # constructs+relationships shape
        add(c)
    add(cm.get("dependent_variable"))
    for v in cm.get("independent_variables") or []:
        add(v)
    add(cm.get("moderator"))
    for c in _mapping_or_empty(instrument).get("constructs") or []:
        add(c)
    for row in scale_items or []:
        add(row.get("construct"))

    seen, out = set(), []
    for lbl in labels:
        if lbl not in seen:
            seen.add(lbl)
            out.append(lbl)
    return out


def _localize_labels(labels: list[str], language: str) -> dict:
    """Translate construct/variable labels into the report language (one LLM
    call). Returns {original: translated}. Vietnamese-only for now, and skipped
    when nothing looks foreign (every label already has non-ASCII chars)."""
    if not labels or not str(language).lower().startswith("vi"):
        return {}
    if not any(l.isascii() for l in labels):
        return {}
    prompt = (
        "Translate each research construct / variable name below into natural "
        "academic Vietnamese (a concise noun phrase). Return ONLY a JSON object "
        "mapping each ORIGINAL string to its Vietnamese translation.\n"
        + json.dumps(labels, ensure_ascii=False)
    )
    try:
        raw = text_of(_get_llm().invoke(prompt)).strip()
        raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.M).strip()
        data = json.loads(raw)
        return {str(k): str(v).strip() for k, v in data.items()
                if str(k) in labels and str(v).strip()}
    except Exception:
        logger.warning("_localize_labels failed", exc_info=True)
        return {}


def _localize_title(title: str | None, language: str) -> str | None:
    """Translate an English research title into the report language for the cover.

    The M1 topic module sometimes stores an English `research_title` even when the
    project language is Vietnamese — it tags `language: vi` but writes the title
    (and questions/objectives) in English. The body then composes correctly in VN,
    but the cover would carry an English title, giving an English-cover /
    Vietnamese-body mismatch. This localizes the title at export time (same
    read-side spot as `_localize_labels`); it does NOT touch the stored M1 slice.

    No-ops when there's no title, the language isn't Vietnamese, or the title
    already contains non-ASCII (i.e. Vietnamese diacritics) — so a title the user
    genuinely wrote in Vietnamese is never round-tripped through the model."""
    if not title or not str(language).lower().startswith("vi"):
        return title
    if not str(title).isascii():
        return title
    prompt = (
        "Translate this thesis research title into natural academic Vietnamese. "
        "Keep brand/product names, platform names (e.g. MoMo, Shopee, TPBank), and "
        "model acronyms (e.g. TAM, UTAUT2, TPB, BNPL, SEM) unchanged. Return ONLY "
        "the translated title on a single line — no quotes, no explanation.\n\n"
        + str(title)
    )
    try:
        out = text_of(_get_llm().invoke(prompt)).strip()
        out = re.sub(r"^```(?:\w+)?|```$", "", out, flags=re.M).strip()
        out = (out.splitlines()[0] if out else "").strip().strip('"').strip()
        return out or title
    except Exception:
        logger.warning("_localize_title failed", exc_info=True)
        return title


def _apply_label_map_to_cm(conceptual_model: dict | str | None,
                           m: dict) -> dict | str | None:
    """Return a copy of the conceptual_model with every label the map covers
    replaced — handles both the nodes/edges and variable-decomposition shapes."""
    if not conceptual_model or not m:
        return conceptual_model
    normalized = _mapping_or_empty(conceptual_model)
    if not normalized:
        return conceptual_model
    cm = json.loads(json.dumps(normalized))  # deep copy
    for n in cm.get("nodes") or []:
        if isinstance(n, dict) and n.get("label") in m:
            n["label"] = m[n["label"]]
    if isinstance(cm.get("dependent_variable"), str) and cm["dependent_variable"] in m:
        cm["dependent_variable"] = m[cm["dependent_variable"]]
    if isinstance(cm.get("moderator"), str) and cm["moderator"] in m:
        cm["moderator"] = m[cm["moderator"]]
    if isinstance(cm.get("independent_variables"), list):
        cm["independent_variables"] = [m.get(v, v) if isinstance(v, str) else v
                                       for v in cm["independent_variables"]]
    return cm


def _variable_decomposition_to_graph(cm: dict) -> dict | None:
    """Coerce the {independent_variables, dependent_variable, moderator} M3
    shape (a variant the headless agent emits) into nodes/edges so the model
    still renders. Each independent variable gets a hypothesis edge to the
    dependent variable; a moderator (if any) also points at it."""
    dv = cm.get("dependent_variable")
    ivs = cm.get("independent_variables") or []
    mod = cm.get("moderator")

    def _lbl(x):
        return str(x.get("label") or x.get("name") or "").strip() if isinstance(x, dict) else str(x).strip()

    dv_lbl = _lbl(dv)
    iv_lbls = [_lbl(v) for v in ivs if _lbl(v)]
    if not dv_lbl or not iv_lbls:
        return None
    nodes = [{"id": "DV", "label": dv_lbl, "type": "dependent"}]
    edges = []
    for i, lbl in enumerate(iv_lbls, 1):
        nodes.append({"id": f"IV{i}", "label": lbl, "type": "independent"})
        edges.append({"source": f"IV{i}", "target": "DV", "hypothesis": f"H{i}"})
    mod_lbl = _lbl(mod)
    if mod_lbl:
        nodes.append({"id": "MOD", "label": mod_lbl, "type": "moderator"})
        # A moderator conditions each IV→DV path — so it must point at the
        # PREDICTORS it moderates, not sit as a plain arrow to the DV like an IV
        # (reviewer feedback). One dashed "Điều tiết" edge to every independent
        # variable.
        for i in range(1, len(iv_lbls) + 1):
            edges.append({"source": "MOD", "target": f"IV{i}", "effect": "moderates"})
    return {"nodes": nodes, "edges": edges}


def _constructs_relationships_to_graph(cm: dict) -> dict | None:
    """Coerce the {constructs:[{name,source}], relationships:[{from,to,hypothesis}]}
    M3 shape into nodes/edges. Relationships reference constructs BY NAME, so the
    construct name IS the node id and the from/to keys already match. A moderation
    relationship (effect/type says so) becomes a dashed moderator edge."""
    cons = cm.get("constructs") or []
    rels = cm.get("relationships") or cm.get("paths") or []
    if not rels:
        return None
    # Relationships reference constructs by NAME, which contain spaces — invalid
    # as mermaid node ids. Assign safe ids (N1, N2, …) and keep the name as label.
    id_of: dict[str, str] = {}
    nodes: list[dict] = []

    def _sid(name) -> str | None:
        name = str(name or "").strip()
        if not name:
            return None
        if name not in id_of:
            id_of[name] = f"N{len(id_of) + 1}"
            nodes.append({"id": id_of[name], "label": name})
        return id_of[name]

    for c in cons:
        _sid((c.get("name") or c.get("label") or c.get("id")) if isinstance(c, dict) else c)
    edges: list[dict] = []
    for r in rels:
        if not isinstance(r, dict):
            continue
        # a relationship may name a construct not in the constructs list — _sid adds it
        s = _sid(r.get("from") or r.get("source"))
        t = _sid(r.get("to") or r.get("target"))
        if not s or not t:
            continue
        edge: dict = {"from": s, "to": t}
        h = str(r.get("hypothesis") or r.get("label") or "").strip()
        if h:
            edge["label"] = h
        if (str(r.get("effect") or "").lower().startswith("moderat")
                or _is_moderation_label(r.get("type"))):
            edge["effect"] = "moderates"
        edges.append(edge)
    if not nodes or not edges:
        return None
    return {"nodes": nodes, "edges": edges}


def _coerce_cm(cm: dict | str | None) -> dict:
    """Normalize the alternative M3 conceptual_model shapes to nodes/edges so
    every figure renderer sees one grammar. No-op when nodes/edges already exist.
    Handles: variable-decomposition (independent/dependent/moderator) and
    constructs+relationships. Unknown shapes pass through unchanged."""
    # Preserve already drawable graphs and the older constructs/relationships
    # shape before normalization. The canonical normalizer intentionally drops
    # unknown legacy keys, which otherwise erased a drawable model here.
    if isinstance(cm, dict) and (cm.get("nodes") or cm.get("edges")):
        return cm
    if isinstance(cm, dict) and cm.get("constructs") and (
            cm.get("relationships") or cm.get("paths")):
        converted = _constructs_relationships_to_graph(cm)
        if converted:
            return converted
    # One canonical contract owns remaining legacy-shape recovery. In particular, prose
    # containing an explicit regression equation becomes a real graph instead
    # of crashing renderers or silently losing the required model figure.
    from agent.m3_contract import normalize_conceptual_model  # noqa: PLC0415
    cm, _ = normalize_conceptual_model(cm)
    if cm.get("nodes") or cm.get("edges"):
        return cm
    if cm.get("dependent_variable"):
        return _variable_decomposition_to_graph(cm) or cm
    return cm


def _is_moderation_label(label: str | None) -> bool:
    low = str(label or "").lower()
    return "điều tiết" in low or "moderat" in low


def _conceptual_model_to_mermaid(conceptual_model: dict | None,
                                 language: str = "vi") -> str | None:
    """Build a fenced ```mermaid``` flowchart from the STRUCTURED conceptual
    model (nodes/edges) in M3 state — so the research-model figure no longer
    depends on the LLM choosing to hand-write a diagram (headless qwen never
    did → reviewer feedback: "chưa có mô hình").

    Returns None if there aren't at least one node and one valid edge (a model
    with no drawn relationships isn't worth a figure). Node/edge grammar matches
    what `_mermaid_to_prose` parses, so the block renders to a PNG downstream.
    """
    cm = conceptual_model or {}
    # Headless sometimes emits a variable-decomposition shape instead of
    # nodes/edges (independent_variables / dependent_variable / moderator) —
    # coerce it so the diagram still renders instead of silently vanishing.
    cm = _coerce_cm(cm)
    nodes = cm.get("nodes") or []
    edges = cm.get("edges") or []
    node_type: dict[str, str] = {}
    lines = ["flowchart LR"]
    valid: set[str] = set()
    for n in nodes:
        if not isinstance(n, dict):
            continue
        nid = str(n.get("id") or "").strip()
        if not nid:
            continue
        label = str(n.get("label") or nid).replace('"', "'")
        lines.append(f'    {nid}["{label}"]')
        valid.add(nid)
        node_type[nid] = str(n.get("type") or "").lower()
    mod_word = "Điều tiết" if str(language).lower().startswith("vi") else "Moderates"
    edge_lines: list[str] = []
    for e in edges:
        if not isinstance(e, dict):
            continue
        # Two edge vocabularies coexist: headless/tools write from/to, the
        # interactive FlowChart widget writes source/target. Accept both so the
        # diagram renders regardless of which runtime authored the model.
        s = str(e.get("from") or e.get("source") or "").strip()
        t = str(e.get("to") or e.get("target") or "").strip()
        if s not in valid or t not in valid:
            continue
        # A moderator moderates the IV→DV paths — draw it as a DASHED arrow
        # labelled "Điều tiết", NOT a plain hypothesis arrow like an IV.
        is_mod = (str(e.get("effect") or "").lower().startswith("moderat")
                  or node_type.get(s) == "moderator")
        if is_mod:
            edge_lines.append(f'    {s} -.->|{mod_word}| {t}')
            continue
        lbl = str(e.get("label") or e.get("hypothesis") or "").strip()
        edge_lines.append(f'    {s} -->|{lbl}| {t}' if lbl else f'    {s} --> {t}')
    if not valid or not edge_lines:
        return None
    lines.extend(edge_lines)
    return "```mermaid\n" + "\n".join(lines) + "\n```"


def _render_model_figure(conceptual_model: dict | None, language: str = "vi") -> str | None:
    """Render the research-model figure (PNG) + a hypothesis list, built from the
    structured conceptual_model.

    A moderator is drawn per the standard conceptual-diagram convention (Hayes;
    Wikipedia "Moderation"): its arrow points at the IV→DV RELATIONSHIP, not at
    a variable box. Mermaid can't target an edge, so we insert a junction node
    on each moderated path and point the moderator at that junction — the arrow
    visibly meets the path. Returns None when there's no drawable model.

    We render the PNG here (not via `_mermaid_to_prose`) and return a plain
    image + text so the junctions never leak into the hypothesis list.
    """
    cm = conceptual_model or {}
    cm = _coerce_cm(cm)
    nodes = cm.get("nodes") or []
    edges = cm.get("edges") or []

    label: dict[str, str] = {}
    ntype: dict[str, str] = {}
    for n in nodes:
        if isinstance(n, dict) and str(n.get("id") or "").strip():
            nid = str(n["id"]).strip()
            label[nid] = str(n.get("label") or nid)
            ntype[nid] = str(n.get("type") or "").lower()

    def _end(e, a, b):
        return str(e.get(a) or e.get(b) or "").strip()

    solid: list[tuple[str, str, str]] = []   # (iv, dv, hypothesis)
    for e in edges:
        if not isinstance(e, dict):
            continue
        s, t = _end(e, "source", "from"), _end(e, "target", "to")
        if not s or not t or s not in label or t not in label:
            continue
        is_mod = (str(e.get("effect") or "").lower().startswith("moderat")
                  or ntype.get(s) == "moderator")
        if is_mod:
            continue  # moderator wiring is derived from node type below
        solid.append((s, t, str(e.get("hypothesis") or e.get("label") or "").strip()))
    if not solid:
        return None
    moderators = [nid for nid, ty in ntype.items() if ty == "moderator"]

    mod_word = "Điều tiết" if str(language).lower().startswith("vi") else "Moderates"
    lines = ["flowchart LR"]
    for nid, lbl in label.items():
        lines.append(f'    {nid}["{lbl.replace(chr(34), chr(39))}"]')
    if moderators:
        for i, (s, t, hyp) in enumerate(solid, 1):
            j = f"MJ{i}"
            lines.append(f'    {j}(( ))')
            lines.append(f'    {s} -->|{hyp}| {j}' if hyp else f'    {s} --> {j}')
            lines.append(f'    {j} --> {t}')
            for m in moderators:
                lines.append(f'    {m} -.->|{mod_word}| {j}')
    else:
        for s, t, hyp in solid:
            lines.append(f'    {s} -->|{hyp}| {t}' if hyp else f'    {s} --> {t}')

    img = ""
    try:
        png = _scratch_dir() / f"conceptmodel-{uuid4().hex[:8]}.png"
        if _render_mermaid_png("\n".join(lines), png):
            img = f'\n![Mô hình nghiên cứu]({png})\n'
    except Exception:
        logger.exception("model figure render failed")

    dv = solid[0][1]
    rel = ["", "**Mối quan hệ giả thuyết trong mô hình:**", ""]
    for s, t, hyp in solid:
        rel.append(f"- {label.get(s, s)} → {label.get(t, t)}" + (f" ({hyp})" if hyp else ""))
    for m in moderators:
        rel.append(f"- {label.get(m, m)} điều tiết mối quan hệ giữa các biến độc lập "
                   f"và {label.get(dv, dv)}")
    return img + "\n" + "\n".join(rel) + "\n"


def _svg_model_figure(conceptual_model: dict | None, language: str = "vi") -> str | None:
    """Render the research model as a clean hand-laid SVG → PNG (via cairosvg).

    Prettier than mermaid for the common star topology (N independent variables
    → one dependent variable, with an optional moderator whose dashed arrows fan
    onto each IV→DV path — the Hayes convention). Returns None for anything the
    star layout can't faithfully draw (mediators/chains, >1 outcome, >1
    moderator, or cairosvg unavailable) so the caller falls back to mermaid.
    """
    try:
        import cairosvg
    except Exception:
        return None
    import html as _html

    cm = conceptual_model or {}
    cm = _coerce_cm(cm)
    nodes, edges = cm.get("nodes") or [], cm.get("edges") or []
    label: dict[str, str] = {}
    ntype: dict[str, str] = {}
    for n in nodes:
        if isinstance(n, dict) and str(n.get("id") or "").strip():
            nid = str(n["id"]).strip()
            label[nid] = str(n.get("label") or nid)
            ntype[nid] = str(n.get("type") or "").lower()

    def _end(e, a, b):
        return str(e.get(a) or e.get(b) or "").strip()

    solid: list[tuple[str, str, str]] = []
    for e in edges:
        if not isinstance(e, dict):
            continue
        s, t = _end(e, "source", "from"), _end(e, "target", "to")
        if not s or not t or s not in label or t not in label:
            continue
        if str(e.get("effect") or "").lower().startswith("moderat") or ntype.get(s) == "moderator":
            continue
        solid.append((s, t, str(e.get("hypothesis") or e.get("label") or "").strip()))
    if not solid:
        return None
    targets = {t for _, t, _ in solid}
    sources = {s for s, _, _ in solid}
    if len(targets) != 1 or (sources & targets):
        return None  # not a clean star (mediator/chain/multi-outcome) → use mermaid
    dv = next(iter(targets))
    seen, iv_order = set(), []
    for s, _, _ in solid:
        if s not in seen:
            seen.add(s)
            iv_order.append(s)
    moderators = [nid for nid, ty in ntype.items() if ty == "moderator"]
    if len(moderators) > 1:
        return None

    def wrap(s, mx=24):
        words, lines, cur = str(s).split(), [], ""
        for w in words:
            if len(cur) + len(w) + 1 <= mx:
                cur = (cur + " " + w).strip()
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines[:3]

    def box(x, y, w, h, txt, fill, stroke, fs):
        lns = wrap(txt)
        ty0 = y + h/2 - (len(lns)-1)*fs*0.62
        t = "".join(
            f'<text x="{x+w/2:.0f}" y="{ty0+i*fs*1.25:.1f}" font-size="{fs}" text-anchor="middle" '
            f'font-family="DejaVu Sans, sans-serif" fill="#1a1a2e">{_html.escape(l)}</text>'
            for i, l in enumerate(lns))
        return (f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="6" fill="{fill}" '
                f'stroke="{stroke}" stroke-width="1.4"/>{t}')

    def bez(x1, y1, cx, x2, y2, tt):
        mt = 1 - tt
        return (mt**3*x1 + 3*mt*mt*tt*cx + 3*mt*tt*tt*cx + tt**3*x2,
                mt**3*y1 + 3*mt*mt*tt*y1 + 3*mt*tt*tt*y2 + tt**3*y2)

    BOX_W, BOX_H, GAP, IV_X, DV_W, W = 210, 56, 26, 34, 210, 930
    n = len(iv_order)
    stack_h = n*BOX_H + (n-1)*GAP
    has_mod = bool(moderators)
    H = int(stack_h + (120 if has_mod else 40) + 60)
    top, DV_X = 30, W - DV_W - 34
    dv_cy = top + stack_h/2
    s = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">',
         f'<rect width="{W}" height="{H}" fill="white"/>',
         '<defs><marker id="arr" markerWidth="9" markerHeight="9" refX="7" refY="3" orient="auto">'
         '<path d="M0,0 L7,3 L0,6 Z" fill="#333"/></marker>'
         '<marker id="arrm" markerWidth="9" markerHeight="9" refX="7" refY="3" orient="auto">'
         '<path d="M0,0 L7,3 L0,6 Z" fill="#8a6d3b"/></marker></defs>',
         box(DV_X, dv_cy-BOX_H/2, DV_W, BOX_H, label[dv], "#E7E9F7", "#4a4e8c", 14)]
    mids = []
    for i, iv in enumerate(iv_order):
        iy = top + i*(BOX_H+GAP)
        s.append(box(IV_X, iy, BOX_W, BOX_H, label[iv], "#EEF0FB", "#6B6FB0", 13))
        x1, y1, x2, y2 = IV_X+BOX_W, iy+BOX_H/2, DV_X, dv_cy
        cx = (x1+x2)/2
        s.append(f'<path d="M{x1},{y1} C{cx},{y1} {cx},{y2} {x2-2},{y2}" fill="none" '
                 f'stroke="#333" stroke-width="1.5" marker-end="url(#arr)"/>')
        hyp = solid[i][2] if i < len(solid) else ""
        if has_mod:
            mx, my = bez(x1, y1, cx, x2, y2, 0.5 + (i-(n-1)/2)*0.07)
            mids.append((mx, my))
            if hyp:
                s.append(f'<text x="{mx+8:.0f}" y="{my-7:.0f}" font-size="12" '
                         f'font-family="DejaVu Sans, sans-serif" fill="#333" font-weight="bold">{_html.escape(hyp)}</text>')
        elif hyp:
            mx, my = bez(x1, y1, cx, x2, y2, 0.5)
            s.append(f'<text x="{mx:.0f}" y="{my-7:.0f}" font-size="12" '
                     f'font-family="DejaVu Sans, sans-serif" fill="#333" font-weight="bold">{_html.escape(hyp)}</text>')
    if has_mod:
        mod_w = 190
        mod_x, mod_y, mcx = (W-mod_w)/2, H-66, W/2
        s.append(box(mod_x, mod_y, mod_w, 44, label[moderators[0]], "#FDF3E3", "#c79a3b", 13))
        for mx, my in mids:
            s.append(f'<path d="M{mcx},{mod_y} C{mcx+(mx-mcx)*0.35:.0f},{mod_y-40:.0f} '
                     f'{mx},{(my+mod_y)/2:.0f} {mx},{my+5:.0f}" fill="none" stroke="#8a6d3b" '
                     f'stroke-width="1.3" stroke-dasharray="5,4" marker-end="url(#arrm)"/>')
            s.append(f'<circle cx="{mx}" cy="{my}" r="3.4" fill="#8a6d3b"/>')
        mod_word = "Điều tiết" if str(language).lower().startswith("vi") else "Moderates"
        s.append(f'<text x="{mcx:.0f}" y="{mod_y-10:.0f}" font-size="12.5" text-anchor="middle" '
                 f'font-family="DejaVu Sans, sans-serif" fill="#8a6d3b" font-weight="bold" '
                 f'font-style="italic">{mod_word}</text>')
    s.append('</svg>')
    svg_str = "\n".join(s)

    try:
        png = _scratch_dir() / f"conceptmodel-{uuid4().hex[:8]}.png"
        cairosvg.svg2png(bytestring=svg_str.encode("utf-8"), write_to=str(png), scale=2.0)
    except Exception:
        logger.exception("cairosvg model figure render failed")
        return None

    rel = ["", "**Mối quan hệ giả thuyết trong mô hình:**", ""]
    for a, b, hyp in solid:
        rel.append(f"- {label[a]} → {label[b]}" + (f" ({hyp})" if hyp else ""))
    for m in moderators:
        rel.append(f"- {label.get(m, m)} điều tiết mối quan hệ giữa các biến độc lập và {label[dv]}")
    return f'\n![Mô hình nghiên cứu]({png})\n' + "\n" + "\n".join(rel) + "\n"


def _pillow_model_figure(conceptual_model: dict | str | None,
                         language: str = "vi") -> str | None:
    """Dependency-light PNG fallback for a direct-effects research model.

    Dev and slim deploy images may have neither CairoSVG nor Puppeteer's Chrome.
    Pillow is already required by the document stack, so a missing optional
    diagram engine must not silently remove the model from the thesis.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None
    cm = _coerce_cm(conceptual_model)
    nodes = [n for n in (cm.get("nodes") or []) if isinstance(n, dict) and n.get("id")]
    edges = [e for e in (cm.get("edges") or []) if isinstance(e, dict)]
    labels = {str(n["id"]): str(n.get("label") or n["id"]) for n in nodes}
    types = {str(n["id"]): str(n.get("type") or "").lower() for n in nodes}

    solid = []
    for edge in edges:
        source = str(edge.get("source") or edge.get("from") or "").strip()
        target = str(edge.get("target") or edge.get("to") or "").strip()
        if source not in labels or target not in labels:
            continue
        if str(edge.get("effect") or "").lower().startswith("moderat"):
            continue
        solid.append((source, target,
                      str(edge.get("hypothesis") or edge.get("label") or "").strip()))
    if not solid:
        return None

    # Decision: use a small layered-DAG layout instead of assuming every model
    # is a one-outcome star. Real thesis models commonly contain a mediator or
    # sequential outcome (A/B/C/D -> awareness -> intention); the old fallback
    # returned None for exactly those models whenever Chrome/Cairo was absent.
    node_ids = list(labels)
    depth = {node_id: 0 for node_id in node_ids}
    for _ in range(len(node_ids)):
        changed = False
        for source, target, _hypothesis in solid:
            candidate = depth[source] + 1
            if candidate > depth[target] and candidate < len(node_ids):
                depth[target] = candidate
                changed = True
        if not changed:
            break
    layers: dict[int, list[str]] = {}
    for node_id in node_ids:
        layers.setdefault(depth[node_id], []).append(node_id)
    ordered_depths = sorted(layers)

    box_w, box_h = 330, 104
    x_gap, y_gap, margin = 170, 64, 75
    width = max(1200, margin * 2 + len(ordered_depths) * box_w
                + max(0, len(ordered_depths) - 1) * x_gap)
    max_rows = max(len(layer) for layer in layers.values())
    height = max(420, margin * 2 + max_rows * box_h
                 + max(0, max_rows - 1) * y_gap)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_paths = (
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    )
    font_path = next((path for path in font_paths if Path(path).exists()), None)
    font = ImageFont.truetype(font_path, 28) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(font_path, 22) if font_path else ImageFont.load_default()

    positions: dict[str, tuple[int, int]] = {}
    for column, layer_depth in enumerate(ordered_depths):
        members = layers[layer_depth]
        layer_height = len(members) * box_h + max(0, len(members) - 1) * y_gap
        start_y = (height - layer_height) // 2
        x = margin + column * (box_w + x_gap)
        for row, node_id in enumerate(members):
            positions[node_id] = (x, start_y + row * (box_h + y_gap))

    def box(x: int, y: int, text: str, fill: str) -> None:
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=16,
                               fill=fill, outline="#315a9a", width=3)
        words, lines, current = text.split(), [], ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=font)[2] <= box_w - 34:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        lines = lines[:3]
        line_h = 32
        ty = y + (box_h - len(lines) * line_h) // 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            tx = x + (box_w - (bbox[2] - bbox[0])) // 2
            draw.text((tx, ty), line, fill="#17233a", font=font)
            ty += line_h

    # Edges go down first so box fills cover their endpoints cleanly.
    for source, target, hypothesis in solid:
        sx, sy = positions[source]
        tx, ty = positions[target]
        start = (sx + box_w, sy + box_h // 2)
        end = (tx, ty + box_h // 2)
        draw.line((start, end), fill="#315a9a", width=4)
        draw.polygon([(end[0], end[1]), (end[0] - 18, end[1] - 10),
                      (end[0] - 18, end[1] + 10)], fill="#315a9a")
        if hypothesis:
            mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
            short_hypothesis = hypothesis.split(":", 1)[0].strip() or hypothesis
            draw.text((mx - 24, my - 30), short_hypothesis,
                      fill="#203b67", font=small)

    for node_id, (x, y) in positions.items():
        fill = "#EAF1FF" if types.get(node_id) == "dependent" else "#F5F8FC"
        box(x, y, labels[node_id], fill)

    png = _scratch_dir() / f"conceptmodel-{uuid4().hex[:8]}.png"
    image.save(png, format="PNG", optimize=True)
    alt = "Mô hình nghiên cứu" if str(language).lower().startswith("vi") else "Research model"
    rel = ["", "**Mối quan hệ giả thuyết trong mô hình:**", ""]
    for source, target, hypothesis in solid:
        rel.append(f"- {labels[source]} → {labels[target]}" +
                   (f" ({hypothesis})" if hypothesis else ""))
    return f"\n![{alt}]({png})\n\n" + "\n".join(rel) + "\n"


def _ensure_model_diagram(prose: str, conceptual_model: dict | None,
                          language: str = "vi") -> str:
    """Guarantee the methodology chapter carries the research-model figure.

    If the LLM already drew a diagram (prose contains a flowchart/mermaid
    block) we leave it alone. Otherwise we render one deterministically: a clean
    SVG (cairosvg) for the common star topology, falling back to the mermaid
    builder for shapes the SVG layout can't draw (or if cairosvg is missing).
    """
    low = prose.lower()
    if ("```mermaid" in low or "flowchart" in low
            or re.search(r"\bgraph\s+\w", low)
            or "![mô hình nghiên cứu]" in low
            or "![research model]" in low
            or "conceptmodel-" in low):
        return prose
    # A prior optional renderer could emit the caption + relationship list but
    # no image when Chrome was unavailable. Remove that trailing orphan before
    # inserting the deterministic Pillow image; otherwise exports contain two
    # Figure 3.1 captions and two relationship lists, with the first appearing
    # to point at a missing figure.
    captions = ("**Hình 3.1: Mô hình nghiên cứu đề xuất**",
                "**Figure 3.1: Proposed research model**")
    for existing_caption in captions:
        marker_at = prose.rfind(existing_caption)
        if marker_at < 0:
            continue
        suffix = prose[marker_at:]
        if "![" not in suffix and (
                "Mối quan hệ giả thuyết trong mô hình" in suffix
                or "Hypothesized relationships" in suffix):
            prose = prose[:marker_at].rstrip()
            break
    fig = _svg_model_figure(conceptual_model, language) or \
        _pillow_model_figure(conceptual_model, language) or \
        _render_model_figure(conceptual_model, language)
    if not fig:
        return prose
    caption = ("**Hình 3.1: Mô hình nghiên cứu đề xuất**"
               if str(language).lower().startswith("vi")
               else "**Figure 3.1: Proposed research model**")
    return prose.rstrip() + "\n\n" + caption + "\n" + fig


def _ensure_export_model_diagrams(sections: list[dict], context_store: dict,
                                  language: str) -> list[dict]:
    """Export-time safety net for stored Methodology chapters.

    Chapters composed before the M3 legacy backfill can contain valid prose but
    no figure. Reusing that prose must not bypass the model-image requirement.
    """
    if not isinstance(sections, list) or not isinstance(context_store, dict):
        return sections
    m3 = context_store.get("m3_design")
    cm = m3.get("conceptual_model") if isinstance(m3, dict) else None
    out = []
    for section in sections:
        if not isinstance(section, dict):
            out.append(section)
            continue
        name = (section.get("chapter_name") or "").lower()
        title = str(section.get("title") or section.get("name") or "").lower()
        is_methodology = name == "methodology" or "methodolog" in title or "phương pháp" in title
        key = "prose" if section.get("prose") is not None else "content"
        prose = section.get(key)
        if not is_methodology or not isinstance(prose, str):
            out.append(section)
            continue
        updated = dict(section)
        updated[key] = _ensure_model_diagram(prose, cm, language)
        out.append(updated)
    return out


def _normalize_prose_markdown(prose: str) -> str:
    """Fix the most common LLM markdown mistake: bullet items mushed onto one
    line ("Goals: * a * b * c"), which pandoc renders as literal asterisks in a
    paragraph instead of a real list. Splits such runs into a proper
    newline-separated markdown list with a blank line before it.

    Conservative: only triggers when a line has 2+ inline markers
    (a near-certain inlined list), so it won't disturb `*emphasis*`/`**bold**`
    (no surrounding spaces) or a single stray asterisk. The marker set covers
    `*`, `-`, the en-dash `–`, the em-dash `—`, and the bullet `•` — LLMs emit
    any of these for an inlined list (e.g. hypotheses "… – H1: … – H2: …"), and
    only the hyphen was caught before, leaving en-dash lists run together.
    """
    if not prose:
        return prose

    out_lines: list[str] = []
    for line in prose.split("\n"):
        # A real list line already starts with a marker — leave it alone.
        if re.match(r"^\s*([*\-+–—•]|\d+\.)\s+", line):
            out_lines.append(line)
            continue
        # `*` / `•` are unambiguous inline-list markers. `-`/`–`/`—` double as
        # PARENTHETICAL dashes ("text — aside — more"), so only treat them as a
        # list when the items are bold-led labels — otherwise a dash-parenthetical
        # sentence gets wrongly chopped into bullets.
        star = re.findall(r"\s[*•]\s+\S", line)
        dash = re.findall(r"\s[-–—]\s+\S", line)
        parts = None
        require_bold = False
        if len(star) >= 2:
            parts = re.split(r"\s+[*•]\s+", line)
        elif len(dash) >= 2:
            parts = re.split(r"\s+[-–—]\s+", line)
            require_bold = True
        if parts is not None:
            head = parts[0].strip()
            items = [p.strip() for p in parts[1:] if p.strip()]
            if len(items) >= 2 and (not require_bold or all(it.startswith("**") for it in items)):
                if head:
                    out_lines.append(head)
                    out_lines.append("")
                out_lines.extend(f"- {it}" for it in items)
                continue
        out_lines.append(line)
    return "\n".join(out_lines)


# YAML block that makes pandoc citeproc emit EVERY bibliography entry, even ones
# not matched by an inline `[@key]`. Uses the literal-block form (the documented
# way — `-M nocite=@*` on the CLI is a plain MetaString citeproc won't parse).
_NOCITE_FRONTMATTER = "---\nnocite: |\n  @*\n---"


def _write_markdown_tmp(sections: list[dict], frontmatter: str | None = None) -> Path:
    md = _sections_to_markdown(sections)
    if frontmatter:
        md = frontmatter.rstrip("\n") + "\n\n" + md
    md_path = _scratch_dir() / f"draft-{uuid4().hex[:8]}.md"
    md_path.write_text(md, encoding="utf-8")
    return md_path


def _compile_pdf_via_engine(sections: list[dict], output_path: str, **kw) -> str:
    """Render sections to a PDF via engine/utils/export_professional.export_pdf.

    The engine fn takes (md_file, output_pdf) and has its own multi-engine
    fallback (libreoffice → pandoc → weasyprint); if none are available it
    returns False and we drop a minimal placeholder so the pipeline still
    produces a downloadable artifact instead of raising.
    """
    try:
        from engine.utils.export_professional import export_pdf
        md_path = _write_markdown_tmp(sections)
        ok = export_pdf(md_path, Path(output_path), **kw)
        if ok and Path(output_path).exists():
            return output_path
        logger.warning("engine export_pdf returned falsy — writing placeholder")
    except Exception as e:
        logger.warning("engine export_pdf failed: %s — writing placeholder", e)
    Path(output_path).write_bytes(b"%PDF-1.4\n%% placeholder - engine renderer unavailable\n")
    return output_path


def _export_docx_via_engine(sections: list[dict], output_path: str,
                            frontmatter: str | None = None,
                            populate_toc: bool = False, **kw) -> str:
    """Render to .docx via engine/utils/export_professional.export_docx.

    Same (md_file, output) contract + internal Pandoc fallback as the PDF
    path. NOTE: the earlier import target (docx_post_processor) had no
    export_docx — that's why every prior export produced a placeholder.

    `frontmatter` is prepended to the markdown (used for the `nocite: @*` block).
    `populate_toc` fills the pandoc-inserted (otherwise empty) Word TOC field
    with static entries so the DOCX opens with a visible table of contents.
    """
    try:
        from engine.utils.export_professional import export_docx as _real
        md_path = _write_markdown_tmp(sections, frontmatter=frontmatter)
        ok = _real(md_path, Path(output_path), **kw)
        if ok and Path(output_path).exists():
            if populate_toc:
                try:
                    _populate_docx_toc(output_path)
                except Exception as e:  # noqa: BLE001 — TOC is best-effort
                    logger.warning("populate TOC failed: %s", e)
            return output_path
        logger.warning("engine export_docx returned falsy — writing placeholder")
    except Exception as e:
        logger.warning("engine export_docx failed: %s — writing placeholder", e)
    Path(output_path).write_bytes(b"PK\x03\x04 placeholder docx\n")
    return output_path


class CitationCompiler:
    """Wrapper exposing a `(style)` constructor + `.compile(items) -> str`.

    Engine's CitationCompiler takes a CitationDatabase, not a style string, so
    sub-project 1 uses a minimal formatter here. Proper engine integration ships later.
    """
    def __init__(self, style: str):
        self.style = style

    def compile(self, items: list[dict]) -> str:
        if not items:
            return ""
        lines = []
        for it in items:
            author = it.get("author") or "Anon"
            year = it.get("year") or "n.d."
            title = it.get("title") or ""
            lines.append(f"{author} ({year}). {title}.".strip())
        return "\n".join(lines)


# --- M5 LLM helpers -------------------------------------------------------

def _format_references_for_prompt(refs: list[dict]) -> str:
    """One reference per line, with the EXACT inline-citation form to use.

    Shows the LLM the precise `(Author, Year)` string it must reproduce so it
    cites real sources instead of inventing "(Anon, 2011)". The author label
    comes from `_ref_author_label` (reads the `authors` list), so it matches
    what `validate_citations` will accept.
    """
    if not refs:
        return "(no references available — write WITHOUT inline citations)"
    lines = [
        "Cite ONLY from this list, using the exact (Author, Year) form shown. "
        "Do NOT invent citations. If no source fits a claim, state it without a citation."
    ]
    for i, r in enumerate(refs, 1):
        label = _ref_author_label(r)
        year = r.get("year", "n.d.")
        title = r.get("title", "")
        venue = r.get("venue", "")
        cite = f"({label}, {year})"
        lines.append(f"[{i}] {cite} — {title}{(' · ' + venue) if venue else ''}".strip())
    return "\n".join(lines)


def _safe_format_kwargs(context_slice: dict) -> dict:
    """Convert dict/list values to JSON strings so str.format() doesn't break."""
    out = {}
    for k, v in context_slice.items():
        if isinstance(v, (dict, list)):
            out[k] = json.dumps(v, ensure_ascii=False, default=str)
        elif v is None:
            out[k] = ""
        else:
            out[k] = str(v)
    return out


# Matches ONLY `{valid_python_identifier}` — leaves `{N+1}`, `{"json": 1}`, and
# any other stray braces in a prompt template untouched.
_PLACEHOLDER_RE = re.compile(r"\{([a-zA-Z_][a-zA-Z0-9_]*)\}")


def _fill_template(template: str, kwargs: dict) -> str:
    """Substitute `{key}` placeholders without str.format()'s brittleness.

    str.format() parses EVERY brace in the template, so a literal `{N+1}` or a
    JSON example like `{"x": 1}` in a prompt file raises KeyError/ValueError and
    aborts the whole chapter (the "[Composition failed — please retry]" bug —
    only results.md had a stray `4.{N+1}`). This fills known identifier
    placeholders and leaves any other brace content exactly as written.
    """
    def _repl(m: "re.Match") -> str:
        key = m.group(1)
        return str(kwargs[key]) if key in kwargs else m.group(0)

    return _PLACEHOLDER_RE.sub(_repl, template)


# A "citation-shaped" parenthetical: `(<author text with a letter>, <YYYY>)`.
# Broader than _CITE_PATTERN so it catches every form a Vietnamese-language LLM
# emits — `(Hair & Ringle, 2019)`, `(Nguyễn và cộng sự, 2021)`, `(Đặng, 2019)`,
# `(Smith, 2020; Jones, 2021)` — while NOT matching non-citations that lack a
# real year: `(xem Bảng 4.1)`, `(β = 0.302, p = 0.000)`, `(N = 188)`.
# Exclude ';' so a multi-citation like "(TUNÇ, 2022; Wu, 2025)" is left intact
# (not greedily matched and stripped when only the LAST pair fails to validate).
_CITE_SHAPED = re.compile(r"\(\s*([^()';]*?[^\W\d_][^()';]*?),\s*((?:19|20)\d{2}[a-z]?)\s*\)")


def _cite_surname_year(author_text: str, year: str) -> tuple[str, str]:
    """First alphabetic surname token (Unicode-aware) + 4-digit year, for
    tolerant matching against the reference pool. Handles `&` / `và cộng sự`
    joined authors and diacritic names (Đặng, Nguyễn)."""
    toks = re.findall(r"[^\W\d_]+", author_text, flags=re.UNICODE)
    surname = toks[0].casefold() if toks else ""
    y = re.match(r"(\d{4})", year)
    return surname, (y.group(1) if y else year)


def _pool_surname_years(references: list[dict]) -> set[tuple[str, str]]:
    """Pool keyed by (first-author surname, 4-digit year) for tolerant lookup."""
    out: set[tuple[str, str]] = set()
    for r in references:
        toks = re.findall(r"[^\W\d_]+", _ref_author_label(r), flags=re.UNICODE)
        surname = toks[0].casefold() if toks else ""
        year = str(r.get("year", "")).strip()
        ym = re.match(r"(\d{4})", year)
        if surname and ym:
            out.add((surname, ym.group(1)))
    return out


def _strip_uncited_citations(prose: str, references: list[dict]) -> str:
    """Remove inline citations not backed by the reference pool, so a
    hallucinated "(Anon, 2011)" never reaches the rendered document.

    When the pool is EMPTY (e.g. Crossref returned nothing), every
    citation-shaped parenthetical is stripped — there is nothing to back any of
    them. When the pool is non-empty, a citation is kept only when its
    (surname, year) matches a pool entry (tolerant of `&`/`và cộng sự` joins and
    diacritics). We delete the parenthetical plus one leading space; real,
    pool-backed citations are untouched.
    """
    pool = _pool_surname_years(references)

    def _repl(m: "re.Match") -> str:
        surname, year = _cite_surname_year(m.group(1), m.group(2))
        return m.group(0) if (pool and (surname, year) in pool) else ""

    # Consume an optional leading space so "research (Anon, 2011)." → "research."
    cleaned = re.sub(r"\s?" + _CITE_SHAPED.pattern, _repl, prose)
    # Collapse any double spaces / space-before-punctuation the removal created.
    cleaned = re.sub(r"  +", " ", cleaned)
    cleaned = re.sub(r"\s+([.,;:])", r"\1", cleaned)
    return cleaned


# Appended to every chapter prompt. The composed prose is rendered to DOCX/PDF
# via pandoc, which needs well-formed markdown — the LLM otherwise mushes list
# items onto one line and emits broken tables.
# Appended to every chapter prompt so the academic-style + table-interpretation
# rules apply uniformly across all 5 chapters from one place, rather than being
# duplicated (and drifting) across each prompts/m5/<chapter>.md file.
_MARKDOWN_FORMAT_RULES = """

---
OUTPUT FORMATTING (strict — the text is rendered to a Word document via Markdown):
- ALL section headings and numbered subsection titles MUST be written in the
  SAME language as the body prose. If you write the body in Vietnamese, the
  headings are Vietnamese too (e.g. "1.1 Bối cảnh và động lực nghiên cứu", NOT
  "1.1 Background and motivation"). Translate any English section name from the
  instructions into that language — never copy it verbatim.
- Write academic prose in full paragraphs. Separate paragraphs with a blank line.
- For any list, put EACH item on its OWN line starting with "- ", and leave a
  blank line BEFORE the list. NEVER put multiple bullet items on one line.
- Prefer prose over tables for NARRATIVE content. Do NOT wrap ordinary
  discussion in tables.
- BUT statistical / measurement results ARE tabular data and MUST be presented
  as Markdown tables, not buried in prose: reliability & validity (Cronbach's
  alpha, composite reliability, AVE, item loadings), correlation/HTMT matrices,
  descriptives, and path/regression coefficients. Put each row on its own line
  with correct "| col | col |" pipes and a "| --- | --- |" separator row, then
  interpret the table in prose AFTER it.
- ONLY build a table when the ACTUAL numbers are present in the provided
  analysis. NEVER output a table containing placeholder cells ("…", "...",
  "TBD", "N/A", blanks) or invented numbers. If the specific values for a table
  (e.g. the full HTMT matrix) are NOT in the data, describe that finding in
  prose only — do not emit a table shell of placeholders.
- Use "## Heading" on its own line for sub-sections.
- Do not output the chapter title as an H1 — it is added automatically.
- Cite sources ONLY as plain "(Author, Year)" in the body. NEVER use the
  frontend pill markup "{{cite: … | … | url}}" — it renders as raw braces in the
  Word/PDF document. The bibliography is generated automatically; do not paste
  titles or URLs inline.
- Never write meta-commentary about the writing process, missing inputs, or
  assumptions you made to fill gaps (e.g. "the author assumed", "due to limited
  information", "giả định", "thiếu thông tin"). Write as a finished scholarly
  document — if a fact is not in the inputs, omit it, do not narrate its absence.

ACADEMIC STYLE (applies to every chapter):
- Impersonal voice only — never "I", "we", "the author thinks". Attribute claims
  to evidence: "the results indicate...", "the analysis shows...".
- No unsupported certainties — replace "certainly / will definitely / proves"
  with calibrated phrasing ("the results suggest", "the findings indicate").
- No absolutes about the sample — "all students" becomes "most respondents in
  the sample".
- No emotive or colloquial words ("sadly", "amazing", "a lot of") — stay neutral
  and precise.
- Split any sentence over ~40 words into shorter, single-idea sentences.
- Vary vocabulary — don't repeat the same verb (rotate study/examine/analyse/
  assess; affect/influence/shape; show/indicate/demonstrate).
- Expand each abbreviation on first use, then use the short form.

CITATIONS (strict — invalid citations are removed automatically):
- Cite ONLY sources from the provided reference list. NEVER invent an author or
  a source that is not on that list.
- Always cite in PARENTHETICAL form at the END of the sentence/clause:
  "… mối quan hệ này (Author, Year)." Use the EXACT (Author, Year) string shown
  in the list.
- NEVER write a citation as the sentence subject / narrative form
  ("Author (Year) đã chỉ ra…", "Author (Year) showed…"). Rephrase so the
  citation sits in parentheses at the end.
- If no provided source supports a claim, state the claim without a citation —
  do not attribute it to an invented source.

GOLDEN RULE FOR TABLES: every statistical table MUST be followed by an
interpreting paragraph. Pattern: [overview sentence] -> [caption] -> [table] ->
[source] -> [detailed interpretation of the key values] -> [closing comment].
Never leave a table standing alone with no prose.

TABLE CAPTION & SOURCE (match a standard thesis exactly):
- Immediately BEFORE each statistical table, put a bold caption on its OWN line.
  Vietnamese: `**Bảng <chương>.<số>: <tên bảng>**` (e.g. `**Bảng 4.1: Thống kê
  mô tả mẫu**`). English: `**Table <số>: <name>**`.
- Immediately AFTER each table, put an italic source line on its OWN line.
  Vietnamese: `*Nguồn: Kết quả phân tích từ SmartPLS/SPSS, tác giả tổng hợp.*`
  English: `*Source: Author's analysis.*`.
- Number tables sequentially per chapter (Bảng 4.1, 4.2, 4.3 …).
"""


def _get_llm():
    """LLM factory for M5 tools. Monkeypatchable in tests.

    Delegates to the engine-wide factory so ORCHESTRATOR_LLM_ROUTE routes every
    tool at once; temperature 0.4 is this tool's original per-site setting.
    """
    from orchestrator.llm import get_orchestrator_llm  # lazy: keep import-light
    return get_orchestrator_llm(temperature=0.4)


# --- Public tools --------------------------------------------------------

@tool
def compose_section(section_name: str, context_store: dict) -> str:
    """Compose one section of the thesis from the project's context_store.

    `section_name` is one of: intro, lit_review, methodology, results,
    conclusion (M5_CHAPTER_ORDER).
    """
    return _compose_section_via_engine(section_name, context_store)


# --- Per-module "M5-mini" composer (for single/multi-module exports) ---------
# A real LLM write-up of ONE module's slice as academic prose — used by the
# Export-to-Word actions (scope=M1..M4). Distinct from `_compose_section_via_
# engine`, which targets the auto-mode full-thesis pipeline and falls back to
# placeholder stubs when the engine's private composers aren't available.
_MODULE_COMPOSE_SKIP = {"confirmed_at", "needs_review", "module_status", "focus"}
_MODULE_COMPOSE_GUIDE = {
    "M1": "Write the **Introduction & Research Focus** section: 2-3 short paragraphs "
          "establishing the background and problem, then state the research title and "
          "present the research questions in prose (you may list the RQs).",
    "M2": "Write the **Literature Review** section: synthesize the provided sources into "
          "flowing paragraphs (what is known, debates, and the research gaps). Use inline "
          "citations in (Author, Year) form. CITE ONLY the sources provided below — never "
          "invent a source. Do NOT include a reference list (it is appended automatically).",
    "M3": "Write the **Research Design & Methodology** section in prose: the conceptual "
          "model and variables, the hypotheses, the methodology (paradigm, design, "
          "sampling, planned analysis), and a short description of the instrument. Cite "
          "(Author, Year) where a construct/scale comes from a source.",
    "M4": "Write the **Data Analysis** section: describe the analysis plan and report only "
          "results that are present in the data below. Never invent statistics.",
}


def compose_module_prose(module: str, slice_: dict, title: str = "Untitled thesis") -> str:
    """Compose ONE module's slice as a standalone academic write-up (markdown body,
    heading stripped — the caller supplies the section title). Real LLM compose;
    returns "" on unknown module or failure so the caller can flag needs_data."""
    guide = _MODULE_COMPOSE_GUIDE.get((module or "").upper())
    if not guide:
        return ""
    payload = {
        k: v for k, v in (slice_ or {}).items()
        if k not in _MODULE_COMPOSE_SKIP and not str(k).startswith("_")
    }
    prompt = (
        "You are writing one section of a Master's thesis for a student to submit "
        "to their professor. Formal academic register, concise (about 300-600 "
        "words). Markdown only: '## ' for the section heading, '### ' for "
        "sub-headings, '- ' for bullets; no tables.\n\n"
        f"## Thesis: {title}\n## Section: {module}\n\n"
        f"Instructions: {guide}\n\n"
        f"Module data (JSON):\n{json.dumps(payload, ensure_ascii=False, default=str)[:9000]}\n\n"
        "Write the section now (markdown):"
    )
    try:
        msg = _get_llm().invoke(prompt)
        c = getattr(msg, "content", "")
        text = (
            "".join(b.get("text", "") if isinstance(b, dict) else str(b) for b in c)
            if isinstance(c, list) else str(c)
        )
        return _strip_leading_chapter_heading(text.strip())
    except Exception:
        logger.exception("compose_module_prose failed for %s", module)
        return ""


@tool
def validate_draft(text: str) -> dict:
    """Run engine's validation pipeline on a draft section. Returns issues + score."""
    return _validate_via_engine(text)


@tool
def compile_pdf(sections: list[dict], project_id: str) -> dict:
    """SP6: render sections to PDF, upload to S3, return {s3_key, size_bytes}.

    Returning size_bytes avoids a separate S3 HeadObject call in the agent and
    lets ExportArtifact carry real file sizes instead of hardcoded 0.
    """
    if not project_id:
        raise ValueError("compile_pdf requires project_id")
    filename = f"thesis-{uuid4().hex[:8]}.pdf"
    local_path = _scratch_dir() / filename
    _compile_pdf_via_engine(sections, str(local_path))
    s3_key, size_bytes = _upload_to_s3(str(local_path), project_id, "pdf", filename)
    return {"s3_key": s3_key, "size_bytes": size_bytes}


@tool
def export_docx(sections: list[dict], project_id: str) -> dict:
    """SP6: render sections to DOCX, upload to S3, return {s3_key, size_bytes}.

    Returning size_bytes avoids a separate S3 HeadObject call in the agent and
    lets ExportArtifact carry real file sizes instead of hardcoded 0.
    """
    if not project_id:
        raise ValueError("export_docx requires project_id")
    filename = f"thesis-{uuid4().hex[:8]}.docx"
    local_path = _scratch_dir() / filename
    _export_docx_via_engine(sections, str(local_path), populate_toc=True)
    s3_key, size_bytes = _upload_to_s3(str(local_path), project_id, "docx", filename)
    return {"s3_key": s3_key, "size_bytes": size_bytes}


# Canonical 5-chapter order + display titles, shared by every caller that
# turns an m5_writing slice into exporter sections (the auto-export hook in
# api/app/agent_state.py, the /m5/export route, and the agent's export tool).
# One source of truth so the paths can't drift.
#
# FIVE, not six. A Vietnamese quantitative thesis ends at "Chương 5 — Kết luận
# và Kiến nghị" and writes the discussion of findings INSIDE it (5.1 summary,
# 5.2 discussion, 5.3 contributions, ...). The order used to declare a separate
# `discussion` chapter and compose_export patched it back to five at export
# time — but only on 2 of the 5 export paths, so auto-mode and the editor
# shipped a Chapter 6 no supervisor asked for. Removing the chapter from the
# canonical order is what makes five hold everywhere by construction.
M5_CHAPTER_ORDER = ["intro", "lit_review", "methodology", "results", "conclusion"]
M5_CHAPTER_TITLES = {
    "intro":       "Chapter 1 — Introduction",
    "lit_review":  "Chapter 2 — Literature Review",
    "methodology": "Chapter 3 — Methodology",
    "results":     "Chapter 4 — Results",
    "conclusion":  "Chapter 5 — Conclusions and Recommendations",
}
# Vietnamese chapter titles — the document language must be consistent, so the
# headings match the (Vietnamese) chapter prose instead of staying English.
# "Kiến nghị" (Recommendations), not "Hàm ý" (Implications): implications sit at
# subsection level (5.3), not in the chapter title.
M5_CHAPTER_TITLES_VI = {
    "intro":       "Chương 1 — Giới thiệu",
    "lit_review":  "Chương 2 — Tổng quan tài liệu",
    "methodology": "Chương 3 — Phương pháp nghiên cứu",
    "results":     "Chương 4 — Kết quả",
    "conclusion":  "Chương 5 — Kết luận và Kiến nghị",
}
# Chapter names that no longer exist canonically, mapped to the one that
# replaced them. Projects composed before the five-chapter collapse hold their
# final chapter under `discussion`; reads alias it forward so a student
# mid-thesis does not lose a written chapter. Writes only ever use canonical
# names, so this never grows a second direction.
LEGACY_CHAPTER_ALIASES = {"discussion": "conclusion"}
# Chapter TITLES that no longer exist, mapped to the chapter key they were
# written under at the time. The agent path stores only `title` (no
# `chapter_name`), so for a pre-branch project these headings are the ONLY
# handle on its prose — unmapped means the section is dropped whole.
# Values are the ORIGINAL keys, not the canonical ones, so `merge_chapter_prose`
# can still tell a legacy `discussion` block from a `conclusion` block and order
# them correctly; `canonical_chapter` resolves them afterwards.
_LEGACY_CHAPTER_TITLES = {
    # The six-chapter era: two separate closing chapters.
    "chapter 5 - discussion":  "discussion",
    "chương 5 - thảo luận":    "discussion",
    "chapter 6 - conclusion":  "conclusion",
    "chương 6 - kết luận":     "conclusion",
    # The old export-time merge (compose_export.wants_merged_conclusion) wrote
    # the discussion prose out under THESE titles, so a legacy export that was
    # re-imported carries them.
    "chapter 5 - conclusion":  "discussion",
    "chương 5 - kết luận":     "discussion",
}
_REFERENCES_TITLE = {"vi": "Tài liệu tham khảo", "en": "References"}


def _chapter_titles(language: str) -> dict:
    """Chapter-title map matching the prose language (vi → Vietnamese)."""
    return M5_CHAPTER_TITLES_VI if str(language).lower().startswith("vi") else M5_CHAPTER_TITLES


def canonical_chapter(name: str | None) -> str | None:
    """Canonical chapter key for `name`, resolving retired aliases.

    Returns None for anything that is not a chapter, so callers can use it as
    the single "is this a chapter, and which one" test instead of each
    re-implementing the alias rule.
    """
    if not name:
        return None
    key = str(name).strip()
    key = LEGACY_CHAPTER_ALIASES.get(key, key)
    return key if key in M5_CHAPTER_ORDER else None


def _title_key(title: str | None) -> str:
    """Normalized lookup key for a chapter title.

    Dash and spacing variants are noise, not identity: our own titles use an
    em-dash but a re-imported or hand-edited heading may carry a hyphen, an
    en-dash or doubled spaces (`_sections_to_markdown` itself rewrites em-dashes
    to hyphens on the way out, so a round-tripped export comes back hyphenated).
    Normalizing here is what lets ONE title map serve every producer.
    """
    t = str(title or "").strip().lower()
    t = t.replace("—", "-").replace("–", "-")
    return re.sub(r"\s+", " ", t)


def chapter_title_lookup() -> dict[str, str]:
    """Normalized chapter title -> stored chapter key, current AND retired.

    Indexes BOTH language maps separately (a dict merge would collapse them —
    they share the canonical-name keys, so the second map's titles would win and
    the first language's titles would be lost), then folds in the retired titles
    so a pre-branch project's headings still resolve.
    """
    out: dict[str, str] = {}
    for mapping in (M5_CHAPTER_TITLES, M5_CHAPTER_TITLES_VI):
        for name, title in mapping.items():
            out[_title_key(title)] = name
    # Retired titles are added last but can never shadow a current one: no
    # canonical title is also a legacy title (the Chapter 5 heading changed).
    for title, name in _LEGACY_CHAPTER_TITLES.items():
        out.setdefault(_title_key(title), name)
    return out


# Blank line = a markdown paragraph break, so the exporter renders the two
# blocks as continuous prose rather than one run-on paragraph.
_PROSE_JOIN = "\n\n"


def merge_chapter_prose(items) -> dict[str, str]:
    """Fold ``(stored_chapter_name, prose)`` pairs onto canonical chapter keys.

    THE rule for what happens when a legacy project holds prose under a retired
    key that now shares a canonical home with a live one — implemented once,
    here, because it was previously hand-rolled in three places (this module
    twice, orchestrator/artifacts.py once) and all three had it backwards.

    We CONCATENATE, we never pick a winner. In the six-chapter era `discussion`
    and `conclusion` were two DISTINCT written chapters: the discussion ran
    1200-2000 words (5.1 summary → 5.6 future research) and carried the
    `[[DT:limitations]]` disclosure, while the conclusion was 500-800 words of
    restatement and closing remarks. Dropping either deletes real written work
    from a student's in-flight thesis — and rescuing in-flight projects is the
    whole reason the aliases exist. So Chapter 5 is the legacy prose FIRST (it
    holds the 5.1→5.6 flow) then the canonical prose, under one canonical title.

    Ordering is by legacy-first, then first-seen — never by input order, because
    neither `chapters` (a dict) nor `final_sections` (a list) guarantees one.
    Returns {canonical_name: prose}; blank prose and non-chapters are skipped.
    """
    buckets: dict[str, list[tuple[int, int, str]]] = {}
    for i, (stored, prose) in enumerate(items or []):
        name = canonical_chapter(stored)
        if name is None:
            continue
        text = (prose or "").strip()
        if not text:
            continue
        # Rank 0 = written under a retired key, so it leads.
        rank = 0 if str(stored or "").strip() in LEGACY_CHAPTER_ALIASES else 1
        buckets.setdefault(name, []).append((rank, i, text))
    out: dict[str, str] = {}
    for name, blocks in buckets.items():
        blocks.sort(key=lambda b: (b[0], b[1]))
        kept: list[str] = []
        for _rank, _i, text in blocks:
            # The two old prompts overlapped heavily and some projects hold the
            # same text under both keys; printing it twice is not "losing
            # nothing", it is a visible defect. Exact duplicates only — near
            # duplicates are real edits and stay.
            if text not in kept:
                kept.append(text)
        out[name] = _PROSE_JOIN.join(kept)
    return out


# Which canonical chapters each module OWNS. This is the pivot from "M5 writes
# the whole thesis" to "every module composes its own chapter as it completes":
# M1–M4 map 1:1 to Chapters 1–4, and M5 owns the closing chapter. Single source
# of truth for per-module composition and the module→chapter mapping the
# export/UI share. Keep consistent with M5_CHAPTER_ORDER — every chapter must be
# owned by exactly one module.
MODULE_CHAPTERS = {
    "M1": ["intro"],
    "M2": ["lit_review"],
    "M3": ["methodology"],
    "M4": ["results"],
    "M5": ["conclusion"],
}


def chapters_for_module(module: str) -> list[str]:
    """Canonical chapter keys a module owns; [] for an unknown module."""
    return list(MODULE_CHAPTERS.get(str(module).upper(), []))


def module_for_chapter(chapter: str) -> str | None:
    """The module that owns a canonical chapter, or None if unowned."""
    for mod, names in MODULE_CHAPTERS.items():
        if chapter in names:
            return mod
    return None


def _references_title(language: str) -> str:
    return _REFERENCES_TITLE["vi"] if str(language).lower().startswith("vi") else _REFERENCES_TITLE["en"]


def chapters_from_final_sections(final_sections: list[dict]) -> dict:
    """Map the conversational `final_sections` list onto the editor's canonical
    chapter dict: ``{intro: {name, prose}, lit_review: {...}, …}``.

    The editor (OutlineRail) only knows the five canonical chapter names, but the
    conversational / export path stores M5 as a flat `final_sections` list. We
    resolve each section's stored name from its explicit `chapter_name`
    (compose path) first, then fall back to a title reverse-lookup across the EN,
    VI and RETIRED title maps (agent path, which carries only `title`). Sections
    that map to no canonical chapter — e.g. References — are dropped: they aren't
    editable chapters. Returns {} when nothing maps, so callers can fall through.

    Two closing chapters from a pre-branch project are CONCATENATED into the one
    canonical Chapter 5 by `merge_chapter_prose` — see its docstring for why
    nothing may be discarded.
    """
    title_to_name = chapter_title_lookup()
    pairs: list[tuple[str, str]] = []
    sources: dict[str, str] = {}
    for sec in final_sections or []:
        if not isinstance(sec, dict):
            continue
        stored = sec.get("chapter_name")
        if canonical_chapter(stored) is None:
            stored = title_to_name.get(_title_key(sec.get("title") or sec.get("name")))
        name = canonical_chapter(stored)
        if name is None:
            continue
        prose = (sec.get("prose") or sec.get("body") or sec.get("content") or "").strip()
        if not prose:
            continue
        pairs.append((stored, prose))
        # `source` rides along: it marks the student's own imported prose, and
        # dropping it here would strip the mark on the first compose — the same
        # way dropping `chapter_name` used to strip the canonical name. A
        # protection that dissolves the first time it is read is not one.
        if sec.get("source") and name not in sources:
            sources[name] = sec["source"]
    out: dict = {}
    for name, prose in merge_chapter_prose(pairs).items():
        out[name] = {"name": name, "prose": prose}
        if name in sources:
            out[name]["source"] = sources[name]
    return out


def sections_from_m5_slice(m5_slice: dict) -> list[dict]:
    """Build exporter sections [{title, prose}] from an m5_writing slice.

    Tolerates both shapes the writers produce:
    - `chapters: {intro: {prose: "…"}, …}` — the canonical auto-mode shape.
    - `final_sections: [{title, body|prose}, …]` — the conversational agent
      shape (the M5 skill's DocumentSection list).
    Returns [] when neither carries usable prose, so callers can short-circuit
    instead of exporting an empty document.
    """
    chapters = (m5_slice or {}).get("chapters") or {}
    if chapters:
        # One rule, one home: a legacy `discussion` entry is concatenated ahead
        # of a real `conclusion` under the single canonical Chapter 5 rather
        # than either one being dropped (see merge_chapter_prose).
        def _prose_of(ch):
            if isinstance(ch, dict):
                return ch.get("prose") or ch.get("body") or ""
            return str(ch or "")

        merged = merge_chapter_prose(
            (stored, _prose_of(ch)) for stored, ch in chapters.items())
        out = [{"chapter_name": name, "title": M5_CHAPTER_TITLES[name],
                "prose": merged[name]}
               for name in M5_CHAPTER_ORDER if merged.get(name)]
        if out:
            return out
    final_sections = (m5_slice or {}).get("final_sections") or []
    # Resolve every section's canonical identity — from `chapter_name` when the
    # producer set one, else from the (now retired-title-aware) reverse lookup —
    # so a legacy slice cannot export a sixth chapter. Passing titles through
    # verbatim is what let a literal "Chapter 6 — Conclusion" heading reach the
    # document. Non-chapter sections (References) have no canonical identity and
    # keep their own title.
    title_to_name = chapter_title_lookup()
    chapter_pairs: list[tuple[str, str]] = []
    lineage: dict[str, object] = {}
    vi_titled: set[str] = set()
    slots: list[dict | str] = []   # a rendered non-chapter dict, or a chapter key
    for sec in final_sections:
        if not isinstance(sec, dict):
            continue
        title = (sec.get("title") or sec.get("name") or "").strip()
        prose = (sec.get("prose") or sec.get("body") or sec.get("content") or "").strip()
        if not prose:
            continue
        stored = sec.get("chapter_name")
        if canonical_chapter(stored) is None:
            stored = title_to_name.get(_title_key(title))
        name = canonical_chapter(stored)
        if name is None:
            extra = {"title": title or "Section", "prose": prose}
            # A non-canonical `chapter_name` is not a chapter, but it is still
            # the producer's own label — pass it through rather than drop it.
            if sec.get("chapter_name"):
                extra["chapter_name"] = sec["chapter_name"]
            slots.append(extra)
            continue
        chapter_pairs.append((stored, prose))
        # Re-titling must not anglicize a Vietnamese thesis. `language` is not
        # threaded into this function (deliberately — separate ticket), so infer
        # it from the heading the project actually wrote: a "Chương N …" title
        # keeps the Vietnamese canonical title, everything else stays English.
        if re.match(r"(?i)^\s*chương\b", title):
            vi_titled.add(name)
        # Lineage keeps the student's own imported prose traceable through
        # rendering; keep the first one seen for the chapter.
        if sec.get("lineage") and name not in lineage:
            lineage[name] = sec["lineage"]
        # Each chapter occupies ONE slot, at its first occurrence, so the caller's
        # section order survives (the two legacy closing blocks collapse into the
        # discussion's slot) instead of being reshuffled into canonical order.
        if name not in slots:
            slots.append(name)
    merged = merge_chapter_prose(chapter_pairs)
    out = []
    for slot in slots:
        if isinstance(slot, dict):
            out.append(slot)
            continue
        if not merged.get(slot):
            continue
        titles = M5_CHAPTER_TITLES_VI if slot in vi_titled else M5_CHAPTER_TITLES
        rendered = {"chapter_name": slot, "title": titles[slot],
                    "prose": merged[slot]}
        if slot in lineage:
            rendered["lineage"] = lineage[slot]
        out.append(rendered)
    return out


# Markers that signal a chapter could NOT be properly composed. These must
# never reach the exported document (the "[Composition failed]" / "[Auto-
# generated for …]" weirdness the user saw). When detected we treat the
# chapter as missing and ask the user to fill the gap instead of shipping it.
_STUB_MARKERS = ("[Composition failed", "[Auto-generated for", "[Composition failed — please retry]")


def _is_stub_prose(prose: str) -> bool:
    """True when prose is a placeholder/failure stub, not real chapter content."""
    if not prose or not prose.strip():
        return True
    if any(marker in prose for marker in _STUB_MARKERS):
        return True
    # A real chapter is at least a few sentences; anything tiny is a stub.
    return len(prose.strip()) < 120


# Minimal data each chapter needs to be writable. Used to decide whether to
# ask the user to fill gaps BEFORE spending ~1 min composing — and to tell
# them exactly what's missing in plain language.
def assess_export_readiness(context_store: dict, chapters: list[str] | None = None) -> list[str]:
    """Return human-readable missing-data items (empty = ready).

    When `chapters` is given, only report items whose owning chapter is in the
    requested set, so a subset compose (partner "analysis_report" = intro/
    results/conclusion) isn't blocked by data a skipped chapter would have used.
    Title/RQs and literature are needed by every academic chapter (owner ANY);
    methodology + results are chapter-specific. This is the ONE gate for full
    AND subset composes — it replaces partner's second _has_sufficient_m4_data
    store-gate.
    """
    m1 = context_store.get("m1_topic") or {}
    m2 = context_store.get("m2_literature") or {}
    m3 = context_store.get("m3_design") or {}
    m4 = context_store.get("m4_analysis") or {}

    ANY = None  # relevant to every chapter
    checks = [
        (ANY, not str(m1.get("research_title") or "").strip(), "M1 — research title"),
        (ANY, not (m1.get("research_questions") or []), "M1 — research questions"),
        (ANY, not m2_references(m2),
         "M2 — literature sources (no references to cite)"),
        ("methodology", not (m3.get("methodology") or m3.get("conceptual_model")),
         "M3 — methodology / conceptual model"),
        ("results", not (m4.get("analysis_results") or m4.get("qual_themes") or m4.get("qual_codes")),
         "M4 — analysis results (the Results chapter has no data)"),
    ]
    req = set(chapters) if chapters is not None else None
    missing: list[str] = []
    for owner, is_missing, label in checks:
        if not is_missing:
            continue
        if req is None or owner is ANY or owner in req:
            missing.append(label)
    return missing


def m2_references(m2_slice: dict | None) -> list[dict]:
    """The M2 reference pool, from whichever of its two keys is filled.

    M2 stores one concept under two names and different readers picked
    different ones, so a project could be simultaneously "M2 done" and "no
    references to cite":

      * `dod_literature` (orchestrator/artifacts.py) counts `citation_list`;
      * this module's export gate and the References builder read
        `literature_sources`;
      * `M2Output` — the schema the backfill's LLM reconstruction fills — has
        only `citation_list`, so an inferred M2 can never populate the other;
      * `literature_sources` is written by the grounded scout, which runs on a
        120s budget (orchestrator/backfill.py) and returns [] when it loses.

    Measured on a real project: citation_list held 6 sources, literature_sources
    did not exist, M2 read `done`, and the export refused for want of
    references. Preferring the richer key and falling back to the other makes
    every reader agree, including on state already stored.
    """
    m2 = m2_slice or {}
    return (m2.get("literature_sources") or m2.get("citation_list") or [])


def _references_section_body(references: list[dict]) -> str:
    """Build the References list body (no heading) as markdown.

    Each entry ends with a markdown link to the source's DOI or URL — pandoc
    renders `[text](url)` as a real clickable hyperlink in the DOCX/PDF, which
    is what makes the bibliography clickable. Sorted by author then year.
    """
    if not references:
        return ""

    def _key(r: dict) -> tuple:
        return (_ref_author_label(r).lower(), str(r.get("year", "")))

    entries: list[str] = []
    for r in sorted(references, key=_key):
        label = _ref_author_label(r)
        year = r.get("year", "n.d.")
        title = (r.get("title") or "").strip()
        venue = (r.get("venue") or "").strip()
        doi = (r.get("doi") or "").strip()
        url = (r.get("url") or "").strip()
        link = f"https://doi.org/{doi}" if doi else url

        # Build the title clause only when there IS a title. The one-line form
        # rendered "Ohanian (1990). ." for a titleless entry, and titleless is
        # the normal shape of a `citation_list` record (author + year + what it
        # was used for) — which m2_references now legitimately falls back to.
        entry = f"{label} ({year})."
        if title:
            entry += f" {title.rstrip('.')}."
        if venue:
            entry += f" *{venue}*."
        if link:
            entry += f" [{link}]({link})"
        entries.append(entry)
    # Blank line between entries so pandoc treats each as its own paragraph.
    return "\n\n".join(entries)


def _match_language(prose: str, chapter_name: str, language: str) -> str:
    """Return a preserved chapter in `language`, translating only if it differs.

    Translating is the ONLY transformation a preserved chapter may undergo. The
    student asked for their thesis in English, not for a new Chapter 4 — so the
    tables, the numbers and the argument all have to come through unchanged and
    only the prose changes language. Recomposing instead is what dropped the
    imported tables in the first place.

    Fail-open in both directions: an undetectable language or a failed
    translation returns the original text. A chapter in the wrong language is a
    problem the student can see and ask about; a chapter silently replaced by an
    error string is not.
    """
    from orchestrator.tools.humanize import detect_language  # noqa: PLC0415

    target = (language or "").strip().lower()[:2]
    source = detect_language(prose)
    if not target or not source or source[:2] == target:
        return prose
    try:
        # translate_markdown, not the editor's translate_selection. Two reasons,
        # and the second one shipped: a whole chapter needs batching and
        # structure checks that the inline tool does not do, and
        # translate_selection is a @tool — a StructuredTool object, which is not
        # callable. Calling it like a function raised TypeError on every single
        # preserved chapter, the except below swallowed it, and the student got
        # their Vietnamese Chapter 4 back inside an otherwise English thesis.
        # The unit test missed it by monkeypatching translate_selection with a
        # plain function, i.e. by asserting the interface the bug assumed.
        from orchestrator.tools.m5_inline import translate_markdown  # noqa: PLC0415
        out = translate_markdown(
            chapter_name=chapter_name, target_lang=target, markdown=prose)
        return out if (out or "").strip() else prose
    except Exception:
        logger.exception("compose_all_sections: translating preserved %s failed", chapter_name)
        return prose


def compose_all_sections(context_store: dict,
                         chapters: list[str] | None = None) -> list[dict]:
    """Compose all 5 chapters from a nested context_store → [{title, prose}].

    `context_store` is the nested module shape ({m1_topic, m2_literature,
    m3_design, m4_analysis}). Each chapter is written by `compose_chapter`
    (real LLM composition against the orchestrator/prompts/m5/<name>.md
    templates), grounded in the project state. On a per-chapter LLM failure
    we drop in a minimal fallback so one bad chapter can't abort the whole
    export — but the happy path is full prose, not stubs.

    Used by the export tool to generate a draft on demand when the user asks
    for the file but nothing was written yet.
    """
    m1 = context_store.get("m1_topic") or {}
    m2 = context_store.get("m2_literature") or {}
    m3 = context_store.get("m3_design") or {}
    m4 = context_store.get("m4_analysis") or {}

    methodology = m3.get("methodology") if isinstance(m3.get("methodology"), dict) else {}
    paradigm = (methodology or {}).get("paradigm", "") or ""
    references = m2_references(m2)
    language = m1.get("language") or "vi"
    citation_style = "apa7"

    # Merge every module's keys into one flat slice for the prompt templates.
    # compose_chapter JSON-encodes nested values and fills missing keys with
    # "", so an over-broad merge is safe; we just make sure the canonical
    # `results` key points at M4's analysis output.
    context_slice: dict = {**m1, **m2, **m3, **m4}
    context_slice.setdefault("results", m4.get("analysis_results"))

    titles = _chapter_titles(language)
    # Compose only the chapters a partner ordered (interactive leaves the scope
    # unset → the whole thesis) — skips the fabricated Results/Discussion an
    # analysis-only order never bought.
    from agent.run_context import scoped_chapters  # noqa: PLC0415
    # A chat request may need complete Chapters 1–3 without fabricating or
    # composing Results/Discussion. Explicit chapters win; run-context scope
    # remains the default for auto/partner runs.
    names = [n for n in (chapters or scoped_chapters(M5_CHAPTER_ORDER))
             if n in M5_CHAPTER_ORDER]

    # Chapters the student already wrote (an imported thesis lands its results
    # and conclusion here verbatim). Composing over them is pure loss: their
    # Chapter 4 carries the EFA, KMO, correlation and regression tables, and a
    # rewrite from the summarised analysis_results reproduces none of it. Reuse
    # the real chapter and spend the LLM only on what is genuinely missing.
    preserved = chapters_from_final_sections(
        (context_store.get("m5_writing") or {}).get("final_sections") or [])

    def _one(name):
        kept = preserved.get(name)
        if kept and kept.get("prose", "").strip():
            return name, _match_language(kept["prose"], name, language)
        try:
            draft = compose_chapter.invoke({
                "chapter_name": name,
                "paradigm": paradigm,
                "context_slice": context_slice,
                "references": references,
                "citation_style": citation_style,
                "language": language,
            })
            prose = (draft or {}).get("prose") or ""
        except Exception:
            logger.exception("compose_all_sections: compose_chapter failed for %s", name)
            prose = ""
        if not prose.strip():
            prose = _fallback_section(name, context_store)
        return name, prose

    # Chapters are independent LLM calls — this per-chapter loop is the ~6-8 min
    # bottleneck of a partner report. Compose concurrently (capped for the Ofox
    # gateway) and reassemble in canonical order.
    import concurrent.futures as _cf  # noqa: PLC0415
    proses: dict[str, str] = {}
    with _cf.ThreadPoolExecutor(max_workers=max(1, min(len(names), 5))) as ex:
        for name, prose in ex.map(_one, names):
            proses[name] = prose
    # `chapter_name` travels with each section, not just its title.
    #
    # Without it the canonical name is destroyed the first time these sections
    # are persisted back into final_sections, and chapters_from_final_sections
    # then has only a title reverse-lookup to work with — which matches its own
    # EN/VI title maps and nothing else. An imported chapter titled "CHƯƠNG 4:
    # KẾT QUẢ NGHIÊN CỨU" matched nothing, so `preserved` came back empty on the
    # next compose: the composer wrote a fresh Chapter 4 AND a fresh Chapter 5
    # in English while the student's untranslated originals sat alongside them
    # as unnamed sections. Preservation that self-destructs on first compose is
    # not preservation.
    #
    # `source` travels with it for the same reason: a preserved chapter that has
    # only been translated is still the student's work, and the shrink guard at
    # the commit edge can only protect what is still marked.
    out: list[dict] = [{"chapter_name": name, "title": titles[name], "prose": proses[name],
                        **({"source": (preserved.get(name) or {}).get("source")}
                           if (preserved.get(name) or {}).get("source") else {})}
                       for name in names]

    # Append a References section built from the M2 sources, with clickable
    # DOI/URL links. Without this the document has inline "(Author, Year)"
    # citations but no bibliography to back them.
    refs_body = _references_section_body(references)
    if refs_body:
        out.append({"title": _references_title(language), "prose": refs_body})
    return out


def compose_module_chapters(context_store: dict, module: str) -> dict:
    """Compose the chapter(s) a single module owns → {name: {name, prose, ...}}.

    The per-module writing step behind "every module contributes to the docx as
    it completes." Reuses compose_all_sections (which already composes a chapter
    subset via its `chapters=` arg, grounded in the flattened M1–M4 slice), then
    reshapes the result to merge straight into context_store.m5_writing.chapters.

    Returns {} for a module that owns no chapters, or on any failure — fail-open
    so composing a chapter can never block a module from completing.
    """
    names = chapters_for_module(module)
    if not names:
        return {}
    try:
        # Explicit chapter subset → compose_all_sections composes exactly these
        # (it still appends a References section, which we filter out below).
        sections = compose_all_sections(context_store, chapters=names)
    except Exception:
        logger.exception("compose_module_chapters: compose failed for %s", module)
        return {}
    out: dict = {}
    for sec in sections:
        name = sec.get("chapter_name")
        if name not in names:
            continue  # drops the appended References section
        prose = (sec.get("prose") or "").strip()
        if not prose:
            continue
        entry = {"name": name, "prose": prose}
        if sec.get("source"):
            entry["source"] = sec["source"]
        out[name] = entry
    return out


def _assign_citation_keys(references: list[dict]) -> tuple[list[dict], dict]:
    """Build (CSL-JSON items, {(author_label, year): citekey}) from references.

    Citekeys are surname+year (deduped with a/b/c suffixes), the form pandoc
    citeproc uses. The label→key map lets us rewrite the LLM's "(Author, Year)"
    inline citations into pandoc `[@key]` syntax that becomes clickable.
    """
    csl_items: list[dict] = []
    ly_to_key: dict[tuple, str] = {}
    used: set[str] = set()
    for r in references:
        label = _ref_author_label(r)
        surname = re.sub(r"[^a-z0-9]", "", label.replace(" et al.", "").lower()) or "ref"
        year = str(r.get("year", "")).strip()
        base = f"{surname}{year or 'nd'}"
        key, suffix = base, ord("a")
        while key in used:
            key = f"{base}{chr(suffix)}"
            suffix += 1
        used.add(key)
        ly_to_key.setdefault((label, year), key)

        authors = r.get("authors") or []
        csl_authors = [{"family": str(a)} for a in authors] if authors else [{"family": label}]
        item: dict = {
            "id": key,
            "type": "article-journal",
            "title": r.get("title") or "Untitled",
            "author": csl_authors,
        }
        if year.isdigit():
            item["issued"] = {"date-parts": [[int(year)]]}
        if r.get("venue"):
            item["container-title"] = r["venue"]
        if r.get("doi"):
            item["DOI"] = r["doi"]
        if r.get("url"):
            item["URL"] = r["url"]
        csl_items.append(item)
    return csl_items, ly_to_key


def _convert_inline_citations(prose: str, ly_to_key: dict) -> str:
    """Rewrite "(Author, Year; Author2, Year2)" → pandoc "[@key1; @key2]".

    Accepts both "Author, Year" and "Author Year" (the LLM frequently drops the
    comma — "(Hilman 2024; Nocker 2024)" — and without this the citation stayed
    plain text and never became a clickable link in the DOCX/PDF).

    Only rewrites a parenthetical when EVERY part maps to a known reference;
    otherwise the parenthetical is left exactly as written (so "(see Figure 1)"
    or an unknown citation is never mangled).
    """
    if not prose or not ly_to_key:
        return prose

    def _repl(m: "re.Match") -> str:
        parts = [p.strip() for p in m.group(1).split(";")]
        keys: list[str] = []
        for p in parts:
            # Author/year separator may be a comma+space OR just a space; an
            # optional trailing letter handles disambiguated years like "2024a".
            cm = re.match(r"^(.*?)[,\s]\s*(\d{4})[a-z]?$", p)
            if not cm:
                return m.group(0)
            key = ly_to_key.get((cm.group(1).strip(), cm.group(2)))
            if not key:
                return m.group(0)
            keys.append(f"@{key}")
        return "[" + "; ".join(keys) + "]"

    return re.sub(r"\(([^()]*\d{4}[^()]*)\)", _repl, prose)


def _populate_docx_toc(docx_path: str) -> None:
    """Replace the DOCX's auto-updating Word TOC field with STATIC entries.

    Pandoc inserts a `TOC \\o "1-3" \\h \\z \\u` field code that Word (and only
    Word) fills in on "Update Field". LibreOffice's headless PDF converter
    doesn't run field codes → the PDF shows only the "Table of Contents" title
    with nothing under it. So walk the doc, collect every Heading 1/2/3, and
    write those in place of the empty field placeholder. Best-effort — silent
    return on failure leaves the (still-working, TOC-less) PDF path intact.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except Exception:
        return
    try:
        doc = Document(docx_path)
    except Exception:
        logger.warning("populate-toc: could not open docx", exc_info=True)
        return

    # Collect Heading 1..3 entries in document order.
    from docx.oxml import OxmlElement
    # Collect Heading 1..3 entries in document order, dropping a bookmark on each
    # heading so the static TOC entries below can hyperlink STRAIGHT to it — a
    # clickable TOC in the LibreOffice-rendered PDF. Without per-heading anchors
    # every TOC line resolved to the same spot (or nothing) when clicked.
    entries: list[tuple[int, str, str]] = []
    _levels = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}
    _bid = 8000
    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        level = _levels.get(style)
        text = p.text.strip()
        if not level or not text:
            continue
        anchor = f"_Toc_dt_{len(entries)}"
        p_el = p._p
        bstart = OxmlElement("w:bookmarkStart")
        bstart.set(qn("w:id"), str(_bid))
        bstart.set(qn("w:name"), anchor)
        bend = OxmlElement("w:bookmarkEnd")
        bend.set(qn("w:id"), str(_bid))
        # bookmarkStart must sit AFTER pPr (pPr must be the paragraph's 1st child).
        pPr = p_el.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(bstart)
        else:
            p_el.insert(0, bstart)
        p_el.append(bend)
        _bid += 1
        entries.append((level, text, anchor))
    if not entries:
        return

    # Pandoc wraps the whole TOC in a `w:sdt` (structured document tag) block:
    #   <w:sdt><w:sdtPr><w:docPartObj><w:docPartGallery val="Table of Contents"/>
    # …</w:sdtPr><w:sdtContent>…w:p with TOC field…</w:sdtContent></w:sdt>
    # Find that sdt so we can gut its content and refill with static entries
    # (while keeping the "Table of Contents" title paragraph LibreOffice already
    # renders correctly).
    body = doc.element.body
    from docx.oxml import OxmlElement

    def _find_toc_sdt():
        for sdt in body.iter(qn("w:sdt")):
            for gallery in sdt.iter(qn("w:docPartGallery")):
                if gallery.get(qn("w:val")) == "Table of Contents":
                    return sdt
        return None

    toc_sdt = _find_toc_sdt()
    if toc_sdt is not None:
        # Inside sdtContent, keep the FIRST paragraph (the "Table of Contents"
        # title, styled TOCHeading) and drop everything else (the field-code
        # placeholders + any pandoc hint paragraphs).
        content = toc_sdt.find(qn("w:sdtContent"))
        if content is None:
            return
        children = list(content)
        # First para = title. Remove all children EXCEPT that first paragraph.
        title_el = children[0] if children else None
        for child in children[1:]:
            content.remove(child)
    else:
        # Fallback: no sdt wrapper — look for a bare TOC field paragraph.
        target = None
        for p_el in body.iterfind(qn("w:p")):
            instr = p_el.find(f".//{qn('w:instrText')}")
            if instr is not None and instr.text and "TOC " in instr.text:
                target = p_el
                break
        if target is None:
            return
        target.getparent().remove(target)
        content = None
        title_el = None

    def _mk_para(level: int, text: str, anchor: str):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        # A tiny left indent per level so H2/H3 read hierarchical.
        if level > 1:
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(360 * (level - 1)))  # 360 dxa ≈ 0.25"
            pPr.append(ind)
        p.append(pPr)
        # Wrap the entry in an INTERNAL hyperlink to the heading's bookmark so
        # clicking it jumps to that section in the PDF.
        hyper = OxmlElement("w:hyperlink"); hyper.set(qn("w:anchor"), anchor)
        r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        if level == 1:
            b = OxmlElement("w:b"); rPr.append(b)
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
        r.append(t); hyper.append(r); p.append(hyper)
        return p

    if content is not None:
        # Append fresh entries inside w:sdtContent (below the title paragraph
        # that survived the pruning above).
        for level, text, anchor in entries:
            content.append(_mk_para(level, text, anchor))
    else:
        # Insert at the same body index the deleted TOC field paragraph held.
        # (This branch runs only when there's no w:sdt wrapper.)
        for level, text, anchor in entries:
            body.append(_mk_para(level, text, anchor))

    try:
        doc.save(docx_path)
    except Exception:
        logger.warning("populate-toc: could not save docx", exc_info=True)


def _docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert a DOCX to PDF via LibreOffice headless so the PDF inherits the
    citeproc-rendered clickable citations from the DOCX (rather than
    re-rendering from markdown, which LibreOffice can't citeproc).

    Populates the pandoc-inserted TOC field with static entries first —
    LibreOffice's field engine doesn't run Word's TOC codes, so without this
    the PDF's TOC is empty.
    """
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False
    # Fill the TOC in place so both the DOCX and the resulting PDF have entries.
    _populate_docx_toc(docx_path)
    outdir = str(Path(pdf_path).parent)
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
        capture_output=True, timeout=180, check=True,
    )
    produced = Path(outdir) / (Path(docx_path).stem + ".pdf")
    if produced.exists() and str(produced) != pdf_path:
        produced.replace(pdf_path)
    return Path(pdf_path).exists()


def _artifact_dict(kind: str, pid: str, s3_key: str, size_bytes: int) -> dict:
    return {
        "kind": kind,
        "s3_key": s3_key,
        "size_bytes": size_bytes,
        "download_url": f"/api/v1/projects/{pid}/exports/{s3_key.split('/')[-1]}",
        "uri": "",
    }


def run_export(sections: list[dict], project_id: str,
               references: list[dict] | None = None,
               language: str = "en", title: str | None = None,
               context_store: dict | None = None) -> list[dict]:
    """Render docx + pdf, upload to S3, return ContextPanel-ready artifacts.

    The single export entrypoint shared by the auto-export hook, the
    /m5/export route, and the agent's export_docx tool.

    When `references` is provided, renders via the citeproc path so inline
    "(Author, Year)" citations become clickable links to the auto-generated
    References section. `language` keeps the generated References heading
    consistent with the (possibly Vietnamese) document. Falls back to the plain
    render on any citeproc failure so export never breaks.

    When `context_store` (the nested store) is provided, `ensure_rendered` weaves
    any missing verified-state tables/cleaning/limitations blocks into the
    chapters before rendering — the export-time safety net for sections that
    reached export without passing through compose_chapter. `None` → byte-
    identical to the prior behavior. Fail-open.
    """
    pid = str(project_id)
    if context_store is not None:
        try:
            from orchestrator.tools.results_render import ensure_rendered  # noqa: PLC0415
            sections = ensure_rendered(sections, context_store, language)
            sections = _ensure_export_model_diagrams(sections, context_store, language)
        except Exception:
            logger.debug("run_export: export safety nets skipped", exc_info=True)
    # Take the title off the store when the caller did not pass one — which is
    # ALL SEVEN of them. `title` has been an optional parameter that nothing has
    # ever supplied, so the cover page has been built from None since it was
    # written. M1 has had the real title the whole time.
    if not title and isinstance(context_store, dict):
        title = ((context_store.get("m1_topic") or {}).get("research_title") or None)
    # Localize an English title onto a Vietnamese cover (M1 sometimes stores the
    # title in English despite language=vi). Read-side only; leaves the M1 slice
    # untouched. Done here so both the citeproc and plain paths get the same title.
    title = _localize_title(title, language)
    if references:
        try:
            return _run_export_citeproc(sections, pid, references, language, title,
                                        cover_fields(context_store, language))
        except Exception:
            logger.exception("citeproc export failed — falling back to plain render")

    docx_res = export_docx.invoke({"sections": sections, "project_id": pid})
    pdf_res = compile_pdf.invoke({"sections": sections, "project_id": pid})
    return [
        _artifact_dict("docx", pid, docx_res["s3_key"], docx_res["size_bytes"]),
        _artifact_dict("pdf", pid, pdf_res["s3_key"], pdf_res["size_bytes"]),
    ]


# All language variants of the References heading — used to strip a manually
# built bibliography before citeproc generates its own.
_ALL_REFERENCE_TITLES = {t.lower() for t in _REFERENCES_TITLE.values()}


# Standard Word hyperlink blue. Applied inline so links look like hyperlinks
# even if the reference template's "Hyperlink" character style is missing/plain.
_HYPERLINK_COLOR = "0563C1"


def _style_link_runs(docx_path: str) -> None:
    """Give every <w:hyperlink> run a blue color + single underline.

    pandoc's `link-citations` wraps in-text "(Author, Year)" citations in real
    hyperlinks to the bibliography, but whether they LOOK like links depends on
    the reference doc's Hyperlink style. We set the color/underline inline so the
    look is guaranteed (citations, the bibliography's DOI/URL links, the TOC).
    No-ops silently if python-docx isn't available or the file isn't a real docx.
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:  # noqa: BLE001
        return
    try:
        doc = Document(docx_path)
    except Exception:  # noqa: BLE001 — placeholder / non-docx; skip
        return

    changed = False
    for hl in doc.element.body.iter(qn("w:hyperlink")):
        for run in hl.findall(qn("w:r")):
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run.insert(0, rpr)
            if rpr.find(qn("w:color")) is None:
                c = OxmlElement("w:color")
                c.set(qn("w:val"), _HYPERLINK_COLOR)
                rpr.append(c)
            if rpr.find(qn("w:u")) is None:
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rpr.append(u)
            changed = True
    if changed:
        doc.save(docx_path)


# Cover-page fields, in the YAML names engine.utils.export_professional already
# reads (it builds PDFGenerationOptions from the markdown metadata when no
# options object is passed, and docx_post_processor turns those into the
# institution block above the title and the degree/supervisor block below the
# date). The whole cover machinery was already there; nothing ever fed it.
#
# `degree` and `advisor` are the engine's names for course/instructor. Order
# matters only for readability of the emitted YAML.
_COVER_FIELDS = ("author", "institution", "faculty", "department", "degree",
                 "project_type", "advisor", "second_examiner", "student_id",
                 "location")

# What a thesis is, per M1's research_type, when the student hasn't said. Used
# only to fill `project_type` — the line a cover page reads as "A Master's
# Thesis". Never guesses the degree, the university or the supervisor: those are
# facts about the student, and a plausible-looking wrong university on a cover
# page is worse than a missing line.
_PROJECT_TYPE = {
    "en": {"quantitative": "A Master's Thesis", "qualitative": "A Master's Thesis",
           "mixed": "A Master's Thesis"},
    "vi": {"quantitative": "Luận văn Thạc sĩ", "qualitative": "Luận văn Thạc sĩ",
           "mixed": "Luận văn Thạc sĩ"},
}


def cover_fields(context_store: dict | None, language: str = "en") -> dict:
    """Title-page facts for this project: whatever M1's `cover` block holds,
    plus the few we can derive from M1 itself.

    Explicit values always win — this only fills a blank.
    """
    m1 = (context_store or {}).get("m1_topic") or {}
    cover = m1.get("cover") if isinstance(m1.get("cover"), dict) else {}
    out = {k: v for k, v in cover.items() if k in _COVER_FIELDS and str(v or "").strip()}
    lang = "vi" if str(language).lower().startswith("vi") else "en"
    if not out.get("department") and str(m1.get("field") or "").strip():
        out["department"] = str(m1["field"]).strip()
    if not out.get("project_type"):
        rt = str(m1.get("research_type") or "").strip().lower()
        default = _PROJECT_TYPE[lang].get(rt)
        if default:
            out["project_type"] = default
    return out


def _title_block_frontmatter_lines(title: str | None, language: str,
                                   cover: dict | None = None) -> list[str]:
    """YAML lines for a Pandoc title block → a real cover page. Emits nothing
    without a title (no title block is better than an empty one). Always pairs
    the title with a date, because docx_post_processor._find_title_block needs
    both a Title and a Date paragraph to build the cover."""
    def esc(s):
        # `str(None)` is "None" — five characters, truthy, and it went straight
        # onto the cover page of every export as the thesis title. The guard
        # below was written to catch exactly this and could never fire.
        return "" if s is None else str(s).replace("\\", "\\\\").replace('"', '\\"').strip()

    t = esc(title)
    if not t:
        return []
    from datetime import datetime  # server-side; real wall clock is fine here
    year = datetime.now().year
    date = f"Năm {year}" if str(language).lower().startswith("vi") else str(year)
    # `lang` is pandoc's own key AND how the cover post-processor knows which
    # language its fixed words go in — without it a Vietnamese cover carried
    # "submitted in partial fulfillment of the requirements for the degree of".
    lang = "vi" if str(language).lower().startswith("vi") else "en"
    lines = [f'title: "{t}"', f'date: "{date}"', f'lang: "{lang}"']
    for key in _COVER_FIELDS:
        v = esc((cover or {}).get(key))
        if v:
            lines.append(f'{key}: "{v}"')
    return lines


def _run_export_citeproc(sections: list[dict], pid: str, references: list[dict],
                         language: str = "en", title: str | None = None,
                         cover: dict | None = None) -> list[dict]:
    """Render DOCX with pandoc citeproc (clickable citations + auto bibliography),
    then convert that DOCX to PDF via LibreOffice so both formats match."""
    csl_items, ly_to_key = _assign_citation_keys(references)

    # Drop any manually-built References section (any language) — citeproc
    # generates its own — and rewrite inline citations to [@key] in each chapter.
    body: list[dict] = []
    for s in sections:
        if (s.get("title") or "").strip().lower() in _ALL_REFERENCE_TITLES:
            continue
        body.append({
            "title": s.get("title", ""),
            "prose": _convert_inline_citations(s.get("prose", ""), ly_to_key),
        })
    # Trailing heading (in the doc's language) tells pandoc where to place the
    # generated bibliography.
    body.append({"title": _references_title(language), "prose": ""})

    bib_path = _scratch_dir() / f"refs-{uuid4().hex[:8]}.json"
    bib_path.write_text(json.dumps(csl_items, ensure_ascii=False), encoding="utf-8")

    docx_name = f"thesis-{uuid4().hex[:8]}.docx"
    docx_local = _scratch_dir() / docx_name
    # nocite:@* → citeproc prints the FULL bibliography (never an empty heading,
    # regardless of which inline [@key]s matched). populate_toc → filled TOC.
    # toc-title makes the TOC heading match the document language.
    toc_title = "Mục lục" if str(language).lower().startswith("vi") else "Contents"
    _fm = _title_block_frontmatter_lines(title, language, cover) + \
        ["nocite: |", "  @*", f'toc-title: "{toc_title}"']
    frontmatter = "---\n" + "\n".join(_fm) + "\n---"
    _export_docx_via_engine(body, str(docx_local), bibliography=bib_path,
                            frontmatter=frontmatter, populate_toc=True)
    # Force the hyperlink look (blue + underline) on every link run, so in-text
    # citation links read as hyperlinks regardless of the reference template's
    # Hyperlink style. Done before PDF conversion so the PDF inherits it.
    _style_link_runs(str(docx_local))

    pdf_name = f"thesis-{uuid4().hex[:8]}.pdf"
    pdf_local = _scratch_dir() / pdf_name
    pdf_ok = _docx_to_pdf(str(docx_local), str(pdf_local))

    docx_key, docx_size = _upload_to_s3(str(docx_local), pid, "docx", docx_name)
    out = [_artifact_dict("docx", pid, docx_key, docx_size)]
    if pdf_ok:
        pdf_key, pdf_size = _upload_to_s3(str(pdf_local), pid, "pdf", pdf_name)
        out.append(_artifact_dict("pdf", pid, pdf_key, pdf_size))
    else:
        # No LibreOffice → weasyprint can't resolve `[@key]`. Render the ORIGINAL
        # sections (plain "(Author, Year)" text + the deterministic References
        # section) so the PDF never ships raw `[@key]` markers or an empty
        # bibliography — instead of the citeproc-rewritten `body`.
        pdf_res = compile_pdf.invoke({"sections": sections, "project_id": pid})
        out.append(_artifact_dict("pdf", pid, pdf_res["s3_key"], pdf_res["size_bytes"]))
    bib_path.unlink(missing_ok=True)
    return out


@tool
def format_citations(items: list[dict], style: str = "apa7") -> str:
    """Format a citation list using the requested style."""
    return CitationCompiler(style).compile(items)


@tool
def compile_bibliography(references: list[dict], citation_style: str) -> str:
    """Format M2 references as a bibliography section using the existing
    CitationCompiler. Returns the formatted block as a multi-line string,
    or '(No references)' on empty input.
    """
    if not references:
        return "(No references)"
    return CitationCompiler(citation_style).compile(references)


# --- Shared prose sanitation -------------------------------------------------
# Moved here from api/app/partner_report_service so EVERY compose path (auto-mode
# graph, chat agent, partner) ships normalized prose from ONE place instead of
# each surface re-cleaning (or, for auto-mode/agent, not cleaning at all). The
# regexes + helpers are copied verbatim from the partner clone; sanitize_prose is
# the public merged pass applied inside compose_chapter below.
_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+(.+?)\s*$")
_HYP_RE = re.compile(r"^(H\s*\d+)\s*[:.\)]\s*(.+)$", re.IGNORECASE)
_BOLD_HYP_RE = re.compile(r"^\s*\*\*\s*(H\s*\d+)\s*[:.\)]\s*(.+?)\s*\*\*\s*$", re.IGNORECASE)
# Inline list markers the composer sometimes emits on a single line.
#  * a spaced single ``*`` is unambiguous (``**bold**``/``*italic*`` never have a
#    space directly inside the asterisks) -> always a bullet.
#  * a spaced ASCII hyphen ``-`` is ambiguous with a parenthetical dash and MUST
#    only match U+002D (en-dash ``–`` / em-dash ``—`` in "Fornell–Larcker" etc.
#    are different codepoints and never match) -> only treated as a bullet when
#    every item is a bold-led label, so real dashes in prose are left alone.
_STAR_BULLET_RE = re.compile(r"\s\*\s+")
_DASH_BULLET_RE = re.compile(r"\s-\s+")
# Escape literal pipes that live INSIDE a parenthesis group in a table row
# (the SmartPLS header ``T Statistics (|O/STDEV|)``). Scoping to parentheses
# avoids the earlier bug where a tightly-formatted row like ``|QD|0.5|0.6|-|``
# had its *delimiter* pipes escaped and collapsed into a single cell.
_PAREN_GROUP_RE = re.compile(r"\(([^()]*)\)")

# A markdown table row that is a placeholder shell: a cell whose whole content is
# an ellipsis / dots (…, ...) — the LLM emits these when it lacks the real
# numbers. Such tables are dropped (see _drop_placeholder_tables).
_PLACEHOLDER_CELL_RE = re.compile(r"(?:^|\|)\s*(?:…|\.{2,})\s*(?=\||$)")


def _reflow_inline_bullets(line: str) -> list[str]:
    """Turn inline ``* a * b`` / ``- a - b`` list markers into real one-per-line
    list items.

    The composer sometimes emits a list on a single line (``Intro: * item one
    * item two`` or ``... . - **Label:** text - **Label2:** text``). Markdown
    only makes a list when each marker starts its own line, so inline markers
    render literally and the items collapse into one paragraph. We split them
    into a proper bullet list (blank line before/after so it parses in both
    pandoc and weasyprint).
    """
    stripped = line.lstrip()
    # leave tables, real headings, blockquotes and already-list lines alone
    if not stripped or stripped[0] in "|>#" or stripped[:2] in ("- ", "* ", "+ "):
        return [line]

    require_bold = False
    if len(_STAR_BULLET_RE.split(line)) >= 2:
        parts = _STAR_BULLET_RE.split(line)
    else:
        parts = _DASH_BULLET_RE.split(line)
        if len(parts) < 2:
            return [line]
        # ASCII-hyphen lists only when items are bold-led (avoids splitting a
        # parenthetical "abc - def" or a name range "Hà Nội - Hải Phòng").
        require_bold = True

    head = parts[0].rstrip()
    items = [p.strip() for p in parts[1:] if p.strip()]
    if not items:
        return [line]
    if require_bold and not all(it.startswith("**") for it in items):
        return [line]

    out: list[str] = []
    if head:
        out.append(head)
        out.append("")
    out.extend(f"- {it}" for it in items)
    out.append("")
    return out


# A frontend/chat "source pill": `{{cite: label | title | url}}` (see
# agent.runtime SYSTEM_PROMPT §"inline source pills"). It's a CHAT convention the
# browser renders as a clickable pill — it must NEVER reach the exported thesis,
# where pandoc prints the raw braces verbatim ("{{cite: Davis 1989 | … | https…}}"
# in the PDF). The composer occasionally borrows the format; we convert every pill
# to a plain APA "(Author, Year)" inline cite (the bibliography is emitted
# separately from the M2 pool via citeproc nocite:@*).
_CITE_PILL_RE = re.compile(r"\{\{\s*cite:\s*(?P<body>[^{}]*?)\s*\}\}", re.IGNORECASE)

# A markdown list item: "- ", "* ", "+ " bullets or "1." / "1)" ordered markers.
_LIST_ITEM_RE = re.compile(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)")


def _pill_label_to_citation(label: str) -> str:
    """'Davis 1989' / 'Nguyen et al. 2021' / 'Roslan, 2023' -> '(Author, Year)'.
    No trailing 4-digit year -> '(label)'; empty -> '' (drop the pill entirely)."""
    label = label.strip().strip(".,; ")
    if not label:
        return ""
    m = re.search(r"^(?P<author>.*?)[,\s]+(?P<year>(?:19|20)\d{2}[a-z]?)$", label)
    if m:
        return f"({m.group('author').strip().rstrip(',')}, {m.group('year')})"
    return f"({label})"


def _convert_cite_pills(prose: str) -> str:
    """Replace `{{cite: label | title | url}}` pills with `(Author, Year)`.

    Keeps only the label (first `|`-separated field) and drops the title/url —
    those live in the reference list, not inline. Idempotent and pool-agnostic:
    runs on every export path via _sanitize_prose so a stray pill can never render
    as raw markup again."""
    return _CITE_PILL_RE.sub(lambda m: _pill_label_to_citation(m.group("body").split("|", 1)[0]),
                             prose)


def _sanitize_prose(prose: str) -> str:
    """Normalize LLM markdown quirks before export.

    The composer sometimes emits each hypothesis (``H1: ...`` full sentence) as a
    Markdown *heading*. Word then (a) renders it as an oversized bold line and
    (b) pulls it into the Table of Contents with a page number — both wrong. It
    also sometimes wraps the whole hypothesis sentence in ``**bold**``.

    We demote those to a normal body paragraph with only the ``Hn:`` label bold,
    and demote any heading whose text is a full sentence (a real section title is
    short and doesn't end in a period) — never a legitimate heading.
    """
    # Convert `{{cite: …}}` frontend pills to plain "(Author, Year)" BEFORE any
    # line normalization so no raw pill markup can survive into the export.
    prose = _convert_cite_pills(prose)
    out: list[str] = []
    # Expand any inline "* a * b" lists into one-item-per-line first, then run
    # the per-line normalizations over the expanded lines.
    expanded: list[str] = []
    for raw in prose.split("\n"):
        expanded.extend(_reflow_inline_bullets(raw))
    for ln in expanded:
        # 0) table rows: escape a literal ``|`` that lives INSIDE a parenthesis
        #    group (the SmartPLS header ``T Statistics (|O/STDEV|)``) so it isn't
        #    read as an extra column delimiter. Scoping to parentheses avoids
        #    touching real delimiter pipes in a tight row like ``|QD|0.5|-|``.
        if ln.lstrip().startswith("|") and set(ln.strip()) - set("|-: "):
            ln = _PAREN_GROUP_RE.sub(lambda m: "(" + m.group(1).replace("|", "\\|") + ")", ln)
        # 1) whole-sentence bold hypothesis:  **H1: ....**  ->  **H1:** ....
        bm = _BOLD_HYP_RE.match(ln)
        if bm:
            out.append(f"**{bm.group(1).strip().upper().replace(' ', '')}:** {bm.group(2).strip()}")
            continue
        # 2) heading form
        hm = _HEADING_RE.match(ln)
        if hm:
            text = hm.group(1).strip()
            hyp = _HYP_RE.match(text)
            if hyp:  # "### H1: full sentence"  ->  "**H1:** full sentence"
                out.append(f"**{hyp.group(1).strip().upper().replace(' ', '')}:** {hyp.group(2).strip()}")
                continue
            # a heading that is actually a full sentence -> plain bold paragraph
            if len(text) > 60 and text.rstrip().endswith((".", ")", ":")):
                out.append(f"**{text}**")
                continue
        # 3) blank line BEFORE a pipe table. Pandoc only parses a pipe table when
        #    a blank line precedes it; the composer often glues the table header
        #    directly under its "**Bảng 4.x: …**" caption, so pandoc reads the
        #    whole block as one paragraph and ships literal "| a | b |" text
        #    instead of a rendered table. Insert the missing separator.
        if (ln.lstrip().startswith("|") and out and out[-1].strip()
                and not out[-1].lstrip().startswith("|")):
            out.append("")
        # 4) blank line BEFORE a bullet/numbered list glued under a paragraph.
        #    Markdown only parses a list when a blank line precedes it; the
        #    composer routinely writes the lead-in ("...sau:") directly above
        #    "- item", so pandoc reads the whole block as ONE paragraph and the
        #    "- "/"1." markers render literally (the RQ1–RQ4 / "Về mặt …" runs
        #    that collapsed into a wall of text). Insert the missing separator —
        #    only when the previous line is prose, never between list items.
        elif (_LIST_ITEM_RE.match(ln) and out and out[-1].strip()
              and not _LIST_ITEM_RE.match(out[-1])):
            out.append("")
        out.append(ln)
    return _drop_placeholder_tables("\n".join(out))


def _drop_placeholder_tables(prose: str) -> str:
    """Remove a Markdown table whose cells are placeholder dots/ellipsis.

    When the uploaded analysis lacks the real numbers for a table (e.g. only a
    prose "all HTMT < 0.85" with no matrix), the composer sometimes emits the
    table SHELL filled with "…"/"..." cells. A table of placeholders is worse
    than none, so drop it — along with its bold "Bảng x.y" caption and italic
    "Nguồn:" source line — and keep the surrounding interpretive prose.
    """
    lines = prose.split("\n")
    out: list[str] = []
    i, n = 0, len(lines)
    while i < n:
        if lines[i].lstrip().startswith("|"):
            j = i
            block: list[str] = []
            while j < n and lines[j].lstrip().startswith("|"):
                block.append(lines[j])
                j += 1
            if any(_PLACEHOLDER_CELL_RE.search(b) for b in block):
                # Drop a preceding "**Bảng …**" caption (and blank lines).
                while out and not out[-1].strip():
                    out.pop()
                if out and re.match(r"^\s*\*\*\s*(Bảng|Table)\b", out[-1]):
                    out.pop()
                    while out and not out[-1].strip():
                        out.pop()
                # Skip a following blank + italic "*Nguồn…*" source line.
                k = j
                while k < n and not lines[k].strip():
                    k += 1
                if k < n and re.match(r"^\s*\*\s*(Nguồn|Source)\b", lines[k]):
                    k += 1
                i = k
                if out and out[-1].strip():
                    out.append("")
                continue
            out.extend(block)
            i = j
        else:
            out.append(lines[i])
            i += 1
    return "\n".join(out)


# Public: the single prose-normalization pass shared by every compose path.
# Moved here from partner_report_service so auto-mode + agent + partner all get
# the same cleanup (inline-bullet reflow, hypothesis-heading demotion, dropping
# placeholder "…" tables).
def sanitize_prose(prose: str) -> str:
    return _sanitize_prose(prose)


def _weave_verified_blocks(chapter_name: str, prose: str, context_slice: dict,
                           language: str = "en") -> str:
    """Splice renderer blocks into a chapter's prose (vision §3.6). Pure/fail-open;
    imports results_render lazily so m5_writing's heavy deps never load it."""
    if chapter_name not in ("results", "methodology", "conclusion"):
        return prose
    from orchestrator.tools.results_render import (  # noqa: PLC0415
        render_cleaning_section, render_limitations, render_results_tables, weave)
    cs = context_slice if isinstance(context_slice, dict) else {}
    ar = cs.get("results") or cs.get("analysis_results")
    blocks = []
    if chapter_name == "results":
        # host_prose so the rendered captions continue the chapter's own table
        # numbering instead of restarting at 4.1 inside a chapter that already
        # has one.
        blocks = render_results_tables(ar, language, host_prose=prose)
    elif chapter_name == "methodology":
        b = render_cleaning_section(ar, language)
        blocks = [b] if b else []
    else:  # conclusion → limitations
        nested = {"m3_design": {"sample_plan": cs.get("sample_plan")},
                  "m4_analysis": {"analysis_results": ar}}
        b = render_limitations(nested, language=language)
        blocks = [b] if b else []
    if not blocks:
        return prose
    return weave(prose, blocks, drop_llm_tables=(chapter_name == "results"))


@tool
def compose_chapter(
    chapter_name: str, paradigm: str, context_slice: dict,
    references: list[dict], citation_style: str, language: str,
) -> dict:
    """Compose one chapter via LLM, returns ChapterDraft dict.

    Loads orchestrator/prompts/m5/<chapter_name>.md as the prompt template;
    fills it with the context_slice + references; calls the LLM; runs
    validate_citations on the result; returns
    {name, prose, citations_used, uncited_warnings}.
    """
    prompt_template = (_PROMPT_DIR / f"{chapter_name}.md").read_text(encoding="utf-8")
    refs_block = _format_references_for_prompt(references)
    safe_kwargs = _safe_format_kwargs(context_slice)
    safe_kwargs.setdefault("paradigm", paradigm)
    safe_kwargs.setdefault("language", language)
    safe_kwargs.setdefault("citation_style", citation_style)
    safe_kwargs["references_list"] = refs_block
    # str.format may KeyError on placeholders not in safe_kwargs — pre-extract
    # all expected keys and fall back to empty string for missing context.
    expected_keys = (
        "research_title", "field", "paradigm", "research_type",
        "objectives", "research_questions", "target_population", "scope",
        "literature_review_doc", "research_gaps",
        "design", "tool", "conceptual_model", "scale_items",
        "themes", "interview_guide", "purposive_criteria",
        "sampling_strategy", "target_sample_size", "mixed_design_type",
        "data_type_detected", "results", "qual_codes", "qual_themes",
        "custom_analyses",
        "language", "citation_style", "references_list",
    )
    for k in expected_keys:
        safe_kwargs.setdefault(k, "")

    # Questionnaire recovery: if scale_items arrived empty (headless backfill
    # stores items on the flat `instrument` key, not on conceptual_model nodes)
    # rebuild them so Chapter 3 has a real measurement table instead of a blank.
    if chapter_name == "methodology" and not safe_kwargs.get("scale_items"):
        safe_kwargs["scale_items"] = _derive_scale_items(
            context_slice.get("conceptual_model"), context_slice.get("instrument"))
        # Instrument is only a spec (constructs + count, no item texts): generate
        # real items so Chapter 3 ships a grounded scale table. Methodology-only
        # to avoid an extra LLM call on chapters that never render {scale_items}.
        if not safe_kwargs["scale_items"]:
            safe_kwargs["scale_items"] = _generate_scale_items(
                context_slice.get("instrument"), language,
                conceptual_model=context_slice.get("conceptual_model"))

    # Localize construct/variable NAMES so the model figure, the scale table,
    # and the prose all speak the report language — M3 state often carries
    # English labels even for a Vietnamese report. Translate once and reuse the
    # same map everywhere so a construct reads identically in all three places.
    localized_cm = context_slice.get("conceptual_model")
    if chapter_name == "methodology":
        _lmap = _localize_labels(
            _collect_construct_labels(localized_cm, context_slice.get("instrument"),
                                      safe_kwargs.get("scale_items")),
            language)
        if _lmap:
            for _row in (safe_kwargs.get("scale_items") or []):
                _row["construct"] = _lmap.get(_row.get("construct"), _row.get("construct"))
            localized_cm = _apply_label_map_to_cm(localized_cm, _lmap)
            safe_kwargs["conceptual_model"] = localized_cm

    try:
        prompt = _fill_template(prompt_template, safe_kwargs) + _MARKDOWN_FORMAT_RULES
        prose = text_of(_get_llm().invoke(prompt)).strip()
    except Exception as e:
        logger.warning("compose_chapter LLM call failed for %s: %s", chapter_name, e)
        prose = f"# {chapter_name.title()}\n\n[Composition failed — please retry]"

    # Strip any citation the LLM invented that isn't in the reference pool, so
    # hallucinated "(Anon, 2011)" never reaches the rendered document. The
    # warning stays as RETURNED metadata (uncited_warnings) for QA/logging — it
    # must NOT be appended into the prose, which gets rendered verbatim.
    cited_in_pool, uncited = validate_citations(prose, references)
    # ALWAYS run the pool-based stripper (not only when the narrow validator
    # flagged something): it removes every parenthetical citation not backed by
    # the reference pool — the authoritative cleaner for hallucinated cites.
    prose = _strip_uncited_citations(prose, references)
    # Sanitize here so EVERY caller (auto-mode graph, chat agent, partner) ships
    # normalized prose from one place instead of each surface re-cleaning.
    prose = sanitize_prose(prose)
    # Methodology must SHOW the research model. If the LLM didn't draw one,
    # inject a diagram built from the structured conceptual_model so the figure
    # ships regardless of the model's diagramming habits.
    if chapter_name == "methodology":
        prose = _ensure_model_diagram(
            prose, localized_cm, safe_kwargs.get("language", "vi"))
    # Renderer over verified state (vision §3.6): splice the Chapter 4 tables /
    # cleaning paragraph / limitations bullets — rendered VERBATIM from the
    # persisted analysis_results — into the LLM's prose at its [[DT:kind]] tokens
    # (or appended if it omitted them). The numbers ship from the renderer, not
    # the model. Fail-open: any renderer hiccup → compose exactly as before.
    try:
        prose = _weave_verified_blocks(chapter_name, prose, context_slice,
                                       safe_kwargs.get("language", "en"))
    except Exception:
        logger.debug("compose_chapter: verified-block weave skipped for %s", chapter_name,
                     exc_info=True)
    return {
        "name": chapter_name,
        "prose": prose,
        "citations_used": cited_in_pool,
        "uncited_warnings": uncited,
    }


@tool
def rewrite_chapter(
    chapter_name: str, current_prose: str, instruction: str,
    context_slice: dict, references: list[dict], language: str,
) -> dict:
    """Rewrite one chapter per user instruction. Returns new ChapterDraft dict.

    Used when the user says e.g. "rewrite the intro to be less formal".
    Same prompt template as compose_chapter + the instruction + current prose
    as anchor. On LLM error, returns the original prose unchanged.
    """
    # Decision: load the same chapter prompt template as compose_chapter so the
    # rewrite stays grounded in the chapter's structural requirements, then
    # append the user's instruction and existing prose as a rewrite anchor.
    prompt_template = (_PROMPT_DIR / f"{chapter_name}.md").read_text(encoding="utf-8")
    refs_block = _format_references_for_prompt(references)
    safe_kwargs = _safe_format_kwargs(context_slice)
    safe_kwargs.setdefault("paradigm", context_slice.get("paradigm", "quantitative"))
    safe_kwargs.setdefault("language", language)
    safe_kwargs.setdefault("citation_style", "apa7")
    safe_kwargs["references_list"] = refs_block
    expected_keys = (
        "research_title", "field", "paradigm", "research_type",
        "objectives", "research_questions", "target_population", "scope",
        "literature_review_doc", "research_gaps",
        "design", "tool", "conceptual_model", "scale_items",
        "themes", "interview_guide", "purposive_criteria",
        "sampling_strategy", "target_sample_size", "mixed_design_type",
        "data_type_detected", "results", "qual_codes", "qual_themes",
        "custom_analyses",
        "language", "citation_style", "references_list",
    )
    for k in expected_keys:
        safe_kwargs.setdefault(k, "")

    try:
        base_prompt = _fill_template(prompt_template, safe_kwargs)
        rewrite_prompt = (
            f"{base_prompt}\n\n"
            f"## User rewrite instruction\n{instruction}\n\n"
            f"## Current chapter prose (rewrite based on the instruction; preserve good content):\n"
            f"{current_prose}\n\n"
            f"Output ONLY the rewritten chapter prose."
        )
        prose = text_of(_get_llm().invoke(rewrite_prompt)).strip()
    except Exception as e:
        # Decision: on any LLM failure, return the original prose unchanged so
        # the user never loses work they've already reviewed or produced.
        logger.warning("rewrite_chapter LLM call failed for %s: %s", chapter_name, e)
        prose = current_prose  # unchanged on failure

    # Strip any citation the LLM invented that isn't in the reference pool, so
    # hallucinated "(Anon, 2011)" never reaches the rendered document. The
    # warning stays as RETURNED metadata (uncited_warnings) for QA/logging — it
    # must NOT be appended into the prose, which gets rendered verbatim.
    cited_in_pool, uncited = validate_citations(prose, references)
    # ALWAYS run the pool-based stripper (not only when the narrow validator
    # flagged something): it removes every parenthetical citation not backed by
    # the reference pool — the authoritative cleaner for hallucinated cites.
    prose = _strip_uncited_citations(prose, references)
    return {
        "name": chapter_name,
        "prose": prose,
        "citations_used": cited_in_pool,
        "uncited_warnings": uncited,
    }
