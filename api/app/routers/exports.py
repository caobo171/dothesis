"""SP6: download endpoint for M5 export artifacts.

Mounted under /api/v1 by app/main.py only when ORCHESTRATOR_ENABLED=true.
Resolves the s3_key from the project's M5Output.export_artifacts and
302-redirects the browser to a fresh 5-minute signed URL.
"""
from __future__ import annotations

import io
import json
import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import RedirectResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from ..agent_state import DbProjectStateStore
from ..db import db_session
from ..deps import current_user, stream_user_factory
from ..models import ContextStore, Export, Project, User
from ..routers.uploads import s3_from_env

router = APIRouter(tags=["exports"])

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Modules that can be exported on their own as a teacher-report Word doc. M5 is
# excluded — it already has the full-thesis docx/pdf export.
_MODULE_EXPORT = {
    "M1": ("m1_topic", "Topic Discovery"),
    "M2": ("m2_literature", "Literature Review"),
    "M3": ("m3_design", "Research Design"),
    "M4": ("m4_analysis", "Data Analysis"),
}


def _humanize(key: str) -> str:
    return key.replace("_", " ").strip().title()


def _slug(name: str | None, max_len: int = 60) -> str:
    # Transliterate diacritics to ASCII first so a Vietnamese title like
    # "Các yếu tố ảnh hưởng…" slugs to "cac-yeu-to-anh-huong…" instead of
    # losing every accented letter to a hyphen. NFKD drops the combining marks;
    # đ/Đ don't decompose, so map them explicitly.
    import unicodedata
    raw = (name or "thesis").replace("đ", "d").replace("Đ", "D")
    ascii_ = unicodedata.normalize("NFKD", raw).encode("ascii", "ignore").decode()
    s = re.sub(r"[^a-zA-Z0-9]+", "-", ascii_.lower()).strip("-")
    return (s[:max_len].strip("-") or "thesis")


def _stringify(v: Any) -> str:
    if isinstance(v, (list, tuple)):
        return "; ".join(_stringify(x) for x in v)
    if isinstance(v, dict):
        return ", ".join(f"{_humanize(k)}: {_stringify(val)}" for k, val in v.items())
    return str(v)


_SKIP_KEYS = {"confirmed_at", "needs_review", "module_status", "focus"}
_LIKERT_TYPES = {"scale", "likert"}


def _authors_str(src: dict) -> str:
    a = src.get("authors") or src.get("author")
    if isinstance(a, list):
        return ", ".join(str(x) for x in a)
    return str(a) if a else "Anon"


def _citation_line(src: dict) -> str:
    """A clean reference line: Authors (Year). Title. Venue."""
    yr = src.get("year") or "n.d."
    parts = [f"{_authors_str(src)} ({yr})."]
    if src.get("title"):
        parts.append(f"{str(src['title']).rstrip('.')}.")
    if src.get("venue") or src.get("journal"):
        parts.append(f"{src.get('venue') or src.get('journal')}.")
    return " ".join(parts)


def _render_obj(doc, obj: Any) -> None:
    """Render a free-shaped value as readable prose/bullets (for model/methodology)."""
    if isinstance(obj, dict):
        for k, v in obj.items():
            if k in _SKIP_KEYS or str(k).startswith("_"):
                continue
            p = doc.add_paragraph()
            p.add_run(f"{_humanize(k)}: ").bold = True
            p.add_run(_stringify(v))
    elif isinstance(obj, list):
        for it in obj:
            doc.add_paragraph(_stringify(it), style="List Bullet")
    elif obj not in (None, ""):
        for line in str(obj).split("\n"):
            doc.add_paragraph(line)


def _render_m1(doc, s: dict) -> None:
    if s.get("research_title"):
        doc.add_heading("Research Title", level=1)
        doc.add_paragraph(str(s["research_title"]))
    rqs = s.get("research_questions") or []
    if rqs:
        doc.add_heading("Research Questions", level=1)
        for i, q in enumerate(rqs):
            p = doc.add_paragraph(style="List Number")
            p.add_run(("Main — " if i == 0 else "") + str(q))


def _render_m2(doc, s: dict) -> None:
    sources = s.get("literature_sources") or []
    if sources:
        doc.add_heading("Reviewed Literature", level=1)
        doc.add_paragraph(f"{len(sources)} verified academic sources.")
        for src in sources:
            p = doc.add_paragraph(style="List Number")
            p.add_run(_citation_line(src))
            if src.get("doi"):
                p.add_run(f"  https://doi.org/{src['doi']}").italic = True
    gaps = s.get("research_gaps") or []
    if gaps:
        doc.add_heading("Identified Research Gaps", level=1)
        for g in gaps:
            if isinstance(g, dict):
                doc.add_heading(str(g.get("title", "Gap")), level=2)
                if g.get("description"):
                    doc.add_paragraph(str(g["description"]))
            else:
                doc.add_paragraph(str(g), style="List Bullet")


def _render_m3(doc, s: dict) -> None:
    if s.get("conceptual_model"):
        doc.add_heading("Conceptual Model", level=1)
        _render_obj(doc, s["conceptual_model"])
    hyps = s.get("hypotheses") or []
    if hyps:
        doc.add_heading("Hypotheses", level=1)
        for h in hyps:
            doc.add_paragraph(str(h), style="List Bullet")
    if s.get("methodology"):
        doc.add_heading("Methodology", level=1)
        _render_obj(doc, s["methodology"])
    inst = s.get("instrument")
    if inst:
        title = inst.get("title") if isinstance(inst, dict) else None
        doc.add_heading(f"Instrument — {title}" if title else "Instrument", level=1)
        questions = inst.get("questions") if isinstance(inst, dict) else None
        if questions:
            for q in questions:
                p = doc.add_paragraph(style="List Number")
                p.add_run(str(q.get("text", ""))).bold = True
                if q.get("required"):
                    p.add_run("  (required)").italic = True
                qtype = str(q.get("type", "")).lower()
                opts = q.get("options") or []
                if qtype in _LIKERT_TYPES:
                    lo, hi = q.get("scale_min", 1), q.get("scale_max", 5)
                    lab = ""
                    if q.get("scale_min_label") or q.get("scale_max_label"):
                        lab = f" ({q.get('scale_min_label','')} → {q.get('scale_max_label','')})"
                    doc.add_paragraph(f"{lo}–{hi} Likert scale{lab}").italic = True
                else:
                    for o in opts:
                        doc.add_paragraph(f"☐ {o}")
        else:
            _render_obj(doc, inst)


def _render_m4(doc, s: dict) -> None:
    if s.get("analysis_outline"):
        doc.add_heading("Analysis Plan", level=1)
        _render_obj(doc, s["analysis_outline"])
    if s.get("analysis_results"):
        doc.add_heading("Results", level=1)
        _render_obj(doc, s["analysis_results"])


_MODULE_RENDERERS = {"M1": _render_m1, "M2": _render_m2, "M3": _render_m3, "M4": _render_m4}


# ── "M5-mini" prose composition ──────────────────────────────────────────────
# Per-module instructions: write the section as flowing academic prose (like the
# M5 chapter composer, but a short single-module write-up), NOT a field dump.
_COMPOSE_GUIDE = {
    "M1": "Write the **Introduction & Research Focus** section: 2-3 short paragraphs "
          "establishing the background and problem, then state the research title and "
          "present the research questions in prose (you may list the RQs).",
    "M2": "Write the **Literature Review** section: synthesize the provided sources into "
          "flowing paragraphs (what is known, debates, and the research gaps). Use inline "
          "citations in (Author, Year) form. CITE ONLY the sources provided below — never "
          "invent a source. Do NOT include a reference list (it is appended automatically).",
    "M3": "Write the **Research Design & Methodology** section in prose: the conceptual "
          "model and variables, the hypotheses, the methodology (paradigm, design, "
          "sampling, planned analysis), and a short description of the instrument. Cite "
          "(Author, Year) where a construct/scale comes from a source.",
    "M4": "Write the **Data Analysis** section: describe the analysis plan and report only "
          "results that are present in the data below. Never invent statistics.",
}


def _llm_text(prompt: str) -> str:
    """Single LLM call → plain text (flattens Gemini 3.x list content)."""
    from orchestrator.llm import get_orchestrator_llm
    # Route through the shared factory so the export composer follows the
    # configured provider (native/Ofox) instead of a hardcoded Google client —
    # a hardcoded ChatGoogleGenerativeAI + a provider/model id (e.g. Ofox's
    # bailian/qwen-plus) would crash. Found by the live E2E tier.
    llm = get_orchestrator_llm(temperature=0.3, timeout=70)
    msg = llm.invoke(prompt)
    c = getattr(msg, "content", "")
    if isinstance(c, list):
        return "".join(b.get("text", "") if isinstance(b, dict) else str(b) for b in c)
    return str(c)


def _compose_module_prose(module: str, label: str, title: str, slice_: dict) -> str:
    guide = _COMPOSE_GUIDE.get(module)
    if not guide:
        return ""
    payload = {k: v for k, v in slice_.items() if k not in _SKIP_KEYS and not str(k).startswith("_")}
    prompt = (
        "You are writing one section of a Master's thesis for a student to submit to "
        "their professor. Formal academic register, concise (about 300-600 words). "
        "Markdown only: use '## ' for the section heading and '### ' for sub-headings, "
        "'- ' for bullets; no tables.\n\n"
        f"## Thesis: {title}\n## Section to write: {module} — {label}\n\n"
        f"Instructions: {guide}\n\n"
        f"Module data (JSON):\n{json.dumps(payload, ensure_ascii=False, default=str)[:9000]}\n\n"
        "Write the section now (markdown):"
    )
    try:
        return _llm_text(prompt).strip()
    except Exception:  # noqa: BLE001 — fall back to structured render
        return ""


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_BULLET_RE = re.compile(r"^\s*[-*]\s+")
_NUM_RE = re.compile(r"^\s*\d+\.\s+")


def _add_md_runs(p, text: str) -> None:
    for i, seg in enumerate(_BOLD_RE.split(text)):
        if seg:
            run = p.add_run(seg)
            if i % 2 == 1:
                run.bold = True


def _md_to_docx(doc, md: str) -> None:
    """Render lightweight markdown (headings, bullets, numbered, **bold**) into the doc."""
    for raw in md.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif _BULLET_RE.match(line):
            _add_md_runs(doc.add_paragraph(style="List Bullet"), _BULLET_RE.sub("", line))
        elif _NUM_RE.match(line):
            _add_md_runs(doc.add_paragraph(style="List Number"), _NUM_RE.sub("", line))
        else:
            _add_md_runs(doc.add_paragraph(), line)


def _build_module_docx(project_name: str, module: str, label: str, slice_: dict) -> io.BytesIO:
    """Render one module as a teacher-ready academic write-up — an "M5-mini": the LLM
    composes the section as prose, then we append a real reference list (M2) so
    citations stay grounded. Falls back to a structured render if composition fails."""
    from docx import Document  # local import: keeps cold-start light

    doc = Document()
    title = str(slice_.get("research_title") or project_name or "Untitled thesis")
    doc.add_heading(title, level=0)
    doc.add_paragraph().add_run(f"{module} — {label}").bold = True
    doc.add_paragraph().add_run("Module report · generated by DoThesis").italic = True

    if not slice_:
        doc.add_paragraph("No content has been produced for this module yet.")
        buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf

    prose = _compose_module_prose(module, label, title, slice_)
    if prose:
        _md_to_docx(doc, prose)
        # Grounding: append the REAL reference list for M2 (never LLM-generated).
        if module == "M2" and slice_.get("literature_sources"):
            doc.add_heading("References", level=1)
            for src in slice_["literature_sources"]:
                p = doc.add_paragraph(style="List Number")
                p.add_run(_citation_line(src))
                if src.get("doi"):
                    p.add_run(f"  https://doi.org/{src['doi']}").italic = True
    else:
        renderer = _MODULE_RENDERERS.get(module)
        if renderer:
            renderer(doc, slice_)
        else:
            for key, value in slice_.items():
                if key in _SKIP_KEYS or str(key).startswith("_"):
                    continue
                doc.add_heading(_humanize(key), level=1)
                _render_obj(doc, value)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class ModuleExportBody(BaseModel):
    module: str
    # Declared so the JSON body validates; the value is consumed by current_user.
    access_token: str | None = None


@router.post("/projects/{project_id}/export/module")
def export_module_docx(
    project_id: uuid.UUID,
    body: ModuleExportBody,
    user: User = Depends(current_user),
    db: Session = Depends(db_session),
):
    """Stream a single module (M1–M4) as a Word doc for a teacher report.

    Built on the fly from the project's context_store slice and returned
    directly (no S3) so it works regardless of object-storage config.
    """
    proj = _owned_project(db, user, project_id)
    module = body.module.upper()
    if module not in _MODULE_EXPORT:
        raise HTTPException(400, detail={"error": {"code": "bad_module",
                                                   "message": f"exportable: {list(_MODULE_EXPORT)}"}})
    column, label = _MODULE_EXPORT[module]
    store = DbProjectStateStore(db.bind, project_id, Path(tempfile.gettempdir()))
    slice_ = (store.load_full_context_store().get(column)) or {}
    buf = _build_module_docx(proj.name or "Untitled thesis", module, label, slice_)
    filename = f"{_slug(proj.name)}-{module.lower()}-{column}.docx"
    return StreamingResponse(
        buf, media_type=_DOCX_MIME,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _owned_project(db: Session, user: User, project_id: uuid.UUID) -> Project:
    # Raise 404 (not 403) to avoid leaking project existence to non-owners.
    p = db.get(Project, project_id)
    if not p or p.user_id != user.id:
        raise HTTPException(
            status_code=404,
            detail={"error": {"code": "not_found"}},
        )
    return p


@router.get("/projects/{project_id}/exports/{filename}")
def download_export(
    project_id: uuid.UUID, filename: str,
    # GET-only (browser <a download>). Auth via a short-lived ?st= token scoped
    # to exactly this artifact, keeping the long-lived JWT out of the URL/logs.
    user: User = Depends(stream_user_factory(
        lambda project_id, filename: f"project-export:{project_id}/{filename}")),
    db: Session = Depends(db_session),
):
    """302-redirect to a fresh 5-minute signed URL for the requested artifact."""
    proj = _owned_project(db, user, project_id)
    expected_key = f"projects/{project_id}/exports/{filename}"
    # Only redirect if this artifact is a recorded export for the project —
    # prevents guessing other S3 keys. Authoritative source is now the exports
    # table (was m5_writing.export_artifacts).
    row = (
        db.query(Export)
        .filter(Export.project_id == project_id, Export.s3_key == expected_key)
        .first()
    )
    if row is None:
        raise HTTPException(
            404, detail={"error": {"code": "artifact_not_found"}},
        )
    # Meaningful download name from the project title + scope, instead of the
    # opaque storage name ("thesis-233f5248.docx"). e.g.
    # "gen-z-tiktok-cosmetics-M2-M3.docx" or "<title>.docx" for a full export.
    ext = filename.rsplit(".", 1)[-1] if "." in filename else (row.kind or "docx")
    scope_suffix = "" if (row.scope or "full") == "full" else "-" + (row.scope or "").replace(",", "-")
    download_name = f"{_slug(proj.name)}{scope_suffix}.{ext}"
    s3 = s3_from_env()
    signed_url = s3.generate_presigned_url(
        "get_object",
        # Project convention is S3_BUCKET; AWS_S3_BUCKET kept as a fallback.
        # ResponseContentDisposition overrides the browser "Save As" name on the
        # signed GET without renaming the stored object.
        Params={"Bucket": os.environ.get("S3_BUCKET") or os.environ["AWS_S3_BUCKET"],
                "Key": expected_key,
                "ResponseContentDisposition": f'attachment; filename="{download_name}"'},
        ExpiresIn=300,
    )
    return RedirectResponse(url=signed_url, status_code=302)


@router.post("/projects/{project_id}/exports/list")
def list_exports(
    project_id: uuid.UUID,
    user: User = Depends(current_user),
    db: Session = Depends(db_session),
):
    """The LATEST export document per (scope, kind), newest first — powers the
    module-agnostic Exports panel. Every re-export used to add another row, so a
    project that ran "Full thesis" 4× showed 4 identical PDF rows; users only
    ever want the most recent one. We keep one per (kind, scope) pair. Older
    documents stay in the DB (still downloadable via a direct download URL) —
    they're just hidden from the panel."""
    _owned_project(db, user, project_id)
    rows = (
        db.query(Export)
        .filter(Export.project_id == project_id)
        .order_by(Export.created_at.desc())
        .all()
    )
    seen: set[tuple[str, str]] = set()
    latest: list[Export] = []
    for r in rows:
        key = ((r.kind or "").lower(), r.scope or "full")
        if key in seen:
            continue
        seen.add(key)
        latest.append(r)
    return [
        {
            "id": str(r.id),
            "scope": r.scope,
            "kind": r.kind,
            "filename": r.filename,
            "size_bytes": r.size_bytes,
            "created_at": r.created_at.isoformat() if r.created_at else None,
            "download_url": f"/api/v1/projects/{project_id}/exports/{r.filename}",
        }
        for r in latest
    ]
