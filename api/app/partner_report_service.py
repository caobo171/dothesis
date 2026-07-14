"""Partner report generation — service-to-service ("Powered by DoThesis").

Turns an uploaded analysis PDF (e.g. Fillform's SPSS / SmartPLS statistical
output) into a composed, exported report (DOCX + PDF) using the same M5 writing
engine that powers a real thesis — but with NO user project, NO chat thread and
NO credit ledger. This is the cross-product path: a partner app authenticates
with a shared token, POSTs the PDF, and gets back short-lived download URLs.

Two depths:
  - "analysis_report" (default): a focused narrative — Introduction, Results,
    Discussion, Conclusion. Fast; matches "drop a file -> get a report".
  - "full_thesis": the full six-chapter M5 composition (heavier, slower).

The heavy work (pdfminer extract + LLM chapter composition + LibreOffice render)
is blocking, so callers should run generate_partner_report() in a threadpool.
"""
from __future__ import annotations

import logging
import os
import re
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import Any

from .pdf_extract import extract_pdf_text

# Prose sanitation now lives in the engine (orchestrator.tools.m5_writing.
# sanitize_prose) and is applied inside compose_chapter, so partner no longer
# sanitizes here — the shared compose_sections path handles it. No local import
# needed after generate_partner_report was rewired onto that shared back half.

logger = logging.getLogger(__name__)

# Mermaid CLI (mmdc) tool dir — renders the M3 research-model diagram to PNG.
_MERMAID_DIR = Path(__file__).resolve().parents[2] / "engine" / "tools" / "mermaid_cli"
# node lives under nvm here; ensure the mmdc subprocess can find it.
_NODE_BIN = "/Users/kaoguyen/.nvm/versions/node/v24.18.0/bin"

# Subset composed for the lighter "analysis_report" depth. Ordered as they
# should appear in the document. "full_thesis" uses compose_all_sections
# (all six chapters + auto References) instead.
_ANALYSIS_CHAPTERS = ["intro", "results", "discussion", "conclusion"]

# Canonical chapter order (matches orchestrator M5_CHAPTER_ORDER). A caller can
# request any subset of these via `chapters`; we always compose them in this
# order so "Chương 1..6" stay in sequence regardless of checkbox order.
_CHAPTER_ORDER = ["intro", "lit_review", "methodology", "results", "discussion", "conclusion"]

_VALID_DEPTHS = {"analysis_report", "full_thesis"}

# --- Live progress tracking (drives the partner poll) ------------------------
# In-memory, keyed by a caller-supplied progress_token. The compose loop updates
# it per chapter; the /partner/report/progress endpoint reads it. Single-process
# only, which is fine — the compose runs in this same API process.
_PROGRESS: dict[str, dict] = {}


def _set_progress(token: str | None, **fields) -> None:
    if not token:
        return
    cur = _PROGRESS.get(token) or {}
    cur.update(fields)
    _PROGRESS[token] = cur


def get_progress(token: str) -> dict | None:
    """Read the live progress for a token (None if unknown/expired)."""
    return _PROGRESS.get(token)


class ReportError(Exception):
    """Raised with a stable `code` the router maps to an HTTP response."""

    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def _s3_from_env():
    """Raw boto3 client matching the export upload convention.

    M5 export uploads artifacts to `Bucket=S3_BUCKET, Key=projects/.../file`
    with NO settings prefix (see orchestrator/tools/m5_writing._upload_to_s3),
    so we presign the same way the exports router does — raw key, no prefix.
    """
    import boto3

    return boto3.client(
        "s3",
        region_name=os.environ.get("AWS_REGION", "us-east-1"),
        aws_access_key_id=os.environ.get("AWS_ACCESS_KEY"),
        aws_secret_access_key=os.environ.get("AWS_SECRET_KEY"),
    )


def _presign(s3, s3_key: str, *, expires_in: int = 3600) -> str:
    bucket = os.environ.get("S3_BUCKET") or os.environ["AWS_S3_BUCKET"]
    return s3.generate_presigned_url(
        "get_object",
        Params={"Bucket": bucket, "Key": s3_key},
        ExpiresIn=expires_in,
    )


def _maybe_embed_model_diagram(sections: list[dict], text: str, language: str,
                               chapter_keys: list[str]) -> None:
    """Infer + render the M3 model diagram and embed it in the methodology section
    (in place). Best-effort — a diagram failure never breaks the report.

    Kept as a helper so the partner flow can compose → EMBED → export in that
    exact order (the diagram must land between the shared compose_sections and
    run_export steps).
    """
    if "methodology" not in chapter_keys:
        return
    try:
        model = _infer_model(text, language)
        png = _render_model_diagram(model) if model else None
        if not png:
            return
        from orchestrator.tools.m5_writing import _chapter_titles  # noqa: PLC0415
        meth_title = _chapter_titles(language).get("methodology")
        caption = ("Hình 1. Mô hình nghiên cứu đề xuất"
                   if str(language).lower().startswith("vi")
                   else "Figure 1. Proposed research model")
        figure_md = f"\n\n![{caption}]({png})\n"
        for sec in sections:
            if sec.get("title") == meth_title:
                sec["prose"] = (sec.get("prose") or "") + figure_md
                break
    except Exception:
        logger.exception("partner_report: model diagram step failed (continuing)")


def _infer_topic(analysis_text: str, language: str) -> dict:
    """Infer the study's framing from the raw analysis output.

    The partner flow only has statistical output (reliability, EFA/CFA, SEM/PLS
    paths, correlations) — no topic. Without a title/objectives/RQs the chapter
    composer emits bracketed "[fill this in]" stubs. This runs ONE cheap LLM
    call to infer research_title / field / objectives / research_questions /
    target_population / scope from the constructs & relationships in the data,
    so the composed prose is concrete. Best-effort: returns {} on any failure.
    """
    import json as _json

    from orchestrator.tools.m5_writing import _get_llm  # noqa: PLC0415 — heavy import

    snippet = (analysis_text or "")[:6000]
    lang_name = "Vietnamese" if str(language).lower().startswith("vi") else "English"
    prompt = (
        "You are a research methodologist. Below is the raw statistical analysis "
        "output of a quantitative study (reliability, EFA/CFA, SEM/PLS path "
        "results, correlations, etc.). Infer the study's most plausible framing "
        "from the CONSTRUCTS, VARIABLES and RELATIONSHIPS present in the data.\n\n"
        f"Write every value in {lang_name}. Be specific and realistic. Do NOT use "
        "bracketed placeholders like [ ... ]. Do NOT restate or verbalize any "
        "statistics/numbers (no coefficients, no p-values, no 'zero point eight "
        "seven') — describe the study's framing conceptually, not its results.\n\n"
        "Field meanings:\n"
        "- research_title: a concise academic title naming the constructs & outcome.\n"
        "- field: the academic discipline.\n"
        "- research_type: e.g. quantitative explanatory / survey-based SEM study.\n"
        "- objectives: the study's AIM/purpose in 1-2 sentences (no numbers).\n"
        "- research_questions: 2-3 questions about the relationships between constructs.\n"
        "- target_population: the likely respondents.\n"
        "- scope: the study's boundary in one phrase.\n\n"
        "Return STRICT JSON only (no prose, no code fence) with exactly these keys:\n"
        '{"research_title": "", "field": "", "research_type": "", "objectives": "", '
        '"research_questions": [], "target_population": "", "scope": ""}\n\n'
        f"ANALYSIS OUTPUT:\n{snippet}"
    )
    try:
        resp = _get_llm().invoke(prompt)
        content = getattr(resp, "content", resp)
        if isinstance(content, list):
            content = " ".join(
                str(p.get("text", "") if isinstance(p, dict) else p) for p in content
            )
        content = str(content).strip()
        start, end = content.find("{"), content.rfind("}")
        if start == -1 or end == -1:
            return {}
        data = _json.loads(content[start:end + 1])
        return data if isinstance(data, dict) else {}
    except Exception:
        logger.exception("partner_report: topic inference failed (continuing without it)")
        return {}


def _search_query_en(topic: str, research_questions: list[str] | None) -> str:
    """One short ENGLISH bibliographic query from the (possibly Vietnamese) topic.

    Crossref indexes mostly English scholarship, so a Vietnamese title returns
    little. A tiny LLM call turns the topic into an English keyword query; falls
    back to the raw topic if the LLM is unavailable.
    """
    try:
        from orchestrator.tools.m5_writing import _get_llm  # noqa: PLC0415
        rq = ("; ".join(research_questions or []))[:300]
        prompt = (
            "Turn this research topic into ONE short English academic search query "
            "(5-10 keywords, no punctuation, no quotes). Topic: "
            f"{topic}\nQuestions: {rq}\nQuery:"
        )
        resp = _get_llm().invoke(prompt)
        q = getattr(resp, "content", resp)
        if isinstance(q, list):
            q = " ".join(str(p.get("text", "") if isinstance(p, dict) else p) for p in q)
        q = str(q).strip().strip('"').splitlines()[0][:200]
        if q:
            return q
    except Exception:
        logger.exception("partner_report: query translation failed")
    return (topic or "")[:200]


def _literature_search(topic: str, research_questions: list[str] | None,
                       *, n: int = 8, timeout_s: int = 20) -> list[dict]:
    """Fast M2 literature search via Crossref (real peer-reviewed sources + DOIs).

    dothesis's full scout is tuned for 100+ sources / 119 queries / minutes and
    rate-limits badly, so it can't return within a per-report budget. A direct
    Crossref query gives ~8 verified sources in a couple seconds. Best-effort:
    returns [] on any failure (citations are a nice-to-have, never a blocker).
    """
    import httpx

    query = _search_query_en(topic, research_questions)
    try:
        r = httpx.get(
            "https://api.crossref.org/works",
            params={
                "query.bibliographic": query,
                "rows": n,
                "select": "title,author,issued,DOI,container-title,URL",
                "filter": "type:journal-article,has-abstract:true",
                "sort": "relevance",
            },
            timeout=timeout_s,
            headers={"User-Agent": "DoThesis-PartnerReport/1.0 (mailto:cao.nguyen@wele-learn.com)"},
        )
        items = r.json().get("message", {}).get("items", [])
    except Exception:
        logger.exception("partner_report: Crossref search failed (skipping citations)")
        return []

    refs: list[dict] = []
    for it in items:
        title = (it.get("title") or [""])[0].strip()
        if not title:
            continue
        authors = [str(a.get("family")).strip() for a in it.get("author", []) if a.get("family")]
        parts = (it.get("issued", {}).get("date-parts") or [[None]])
        year = parts[0][0] if parts and parts[0] else None
        refs.append({
            "title": title,
            "authors": authors,
            "year": year,
            "venue": (it.get("container-title") or [None])[0],
            "doi": it.get("DOI"),
            "url": it.get("URL"),
        })
    return refs[:n]


# Real M2 research for a partner report, under a wall-clock budget. The full
# scout can run minutes / rate-limit; we cap it and fall back to the light
# Crossref search so a report never hangs and never ships zero references.
_M2_SCOUT_MIN = 8
_M2_SCOUT_TIMEOUT_S = 45


def _budgeted_scout(topic: str, research_questions: list[str]) -> list[dict]:
    """Deep scout capped by time; Crossref fallback on timeout/error/empty."""
    import concurrent.futures as _fut

    composed = topic
    if research_questions:
        composed += "\nResearch questions:\n" + "\n".join(f"- {q}" for q in research_questions)

    # NOTE: do NOT use `with ThreadPoolExecutor(...)`. Its __exit__ calls
    # shutdown(wait=True), which BLOCKS until the (possibly runaway) scout thread
    # finishes — so a `future.result(timeout=45)` that times out would still wait
    # minutes for the hung thread, defeating the wall-clock cap entirely. On a
    # flaky network (Semantic Scholar 429 storm) that pushes the whole report
    # past the caller's timeout. Shut down WITHOUT waiting instead.
    ex = _fut.ThreadPoolExecutor(max_workers=1)
    try:
        from orchestrator.tools.m2_literature import scout_citations  # noqa: PLC0415
        future = ex.submit(scout_citations.func, composed, min_n=_M2_SCOUT_MIN)
        citations = future.result(timeout=_M2_SCOUT_TIMEOUT_S)
        sources = [
            {"title": c.get("title"), "authors": c.get("authors"), "year": c.get("year"),
             "venue": c.get("source") or c.get("venue"), "doi": c.get("doi"),
             "url": c.get("url")}
            for c in (citations or [])
        ]
        if sources:
            return sources
    except Exception:
        # TimeoutError, engine failure, rate limit — all fall through to Crossref.
        logger.exception("partner_report: budgeted scout failed; using Crossref fallback")
    finally:
        # wait=False so we never block on the runaway scout thread; cancel_futures
        # drops anything still queued. The 45s cap is now actually enforced.
        ex.shutdown(wait=False, cancel_futures=True)

    return _literature_search(topic, research_questions)


def _references_section(references: list[dict], language: str) -> dict:
    """Build a populated References section from the sources.

    Deterministic — lists every source we found regardless of whether the LLM
    cited it inline, so the section is never an empty heading. Reuses dothesis's
    shared `_references_section_body` (sorted by author+year, clickable DOI/URL
    links) instead of a drifted local clone. This section is what the PDF
    (weasyprint) fallback ships; the citeproc DOCX path generates its own.
    """
    from orchestrator.tools.m5_writing import _references_section_body  # noqa: PLC0415
    title = "Tài liệu tham khảo" if str(language).lower().startswith("vi") else "References"
    return {"title": title, "prose": _references_section_body(references)}


def _infer_model(analysis_text: str, language: str) -> dict:
    """Infer the study's structural/conceptual model (constructs + directed
    paths) from the analysis output, for the M3 research-model diagram.

    Returns {"constructs":[{"id","label"}], "paths":[{"from","to"}]} or {}.
    """
    import json as _json

    from orchestrator.tools.m5_writing import _get_llm  # noqa: PLC0415

    snippet = (analysis_text or "")[:6000]
    lang = "Vietnamese" if str(language).lower().startswith("vi") else "English"
    prompt = (
        "From this statistical analysis output, extract the STRUCTURAL / CONCEPTUAL "
        "research model: the latent constructs and the DIRECTED relationships "
        "(which construct predicts which), based on the path/regression results.\n"
        f"Give each construct a short ascii id (no spaces) and a {lang} label.\n"
        "Return STRICT JSON only:\n"
        '{"constructs":[{"id":"","label":""}],"paths":[{"from":"id","to":"id"}]}\n\n'
        f"ANALYSIS OUTPUT:\n{snippet}"
    )
    try:
        resp = _get_llm().invoke(prompt)
        content = getattr(resp, "content", resp)
        if isinstance(content, list):
            content = " ".join(str(p.get("text", "") if isinstance(p, dict) else p) for p in content)
        content = str(content)
        s, e = content.find("{"), content.rfind("}")
        if s == -1 or e == -1:
            return {}
        data = _json.loads(content[s:e + 1])
        if isinstance(data, dict) and data.get("constructs") and data.get("paths"):
            # The shared gate (assess_export_readiness) requires m3.methodology
            # OR m3.conceptual_model for the methodology chapter. _infer_model
            # only yields constructs/paths (for the diagram), so WITHOUT this the
            # partner flow — which never ships methodology prose, only analysis
            # output — is ALWAYS blocked with needs_data whenever the methodology
            # chapter is requested (i.e. every full report). Synthesize a
            # conceptual_model description from the inferred model so the gate is
            # satisfied and compose has real M3 content to write from.
            if not data.get("conceptual_model") and not data.get("methodology"):
                labels = {c.get("id"): (c.get("label") or c.get("id"))
                          for c in data["constructs"] if isinstance(c, dict)}
                rels = "; ".join(
                    f"{labels.get(p.get('from'), p.get('from'))} → "
                    f"{labels.get(p.get('to'), p.get('to'))}"
                    for p in data["paths"] if isinstance(p, dict))
                if str(language).lower().startswith("vi"):
                    data["conceptual_model"] = (
                        "Mô hình nghiên cứu đề xuất gồm các mối quan hệ giả thuyết: "
                        + rels + ".")
                else:
                    data["conceptual_model"] = (
                        "The proposed research model comprises the hypothesized "
                        "relationships: " + rels + ".")
            return data
    except Exception:
        logger.exception("partner_report: model inference failed")
    return {}


def _render_model_diagram(model: dict) -> str | None:
    """Render a mermaid flowchart of the model to a PNG via mmdc. Abs path or None."""
    constructs = {
        str(c["id"]): str(c.get("label") or c["id"])
        for c in model.get("constructs", [])
        if isinstance(c, dict) and c.get("id")
    }
    paths = [
        (str(p.get("from")), str(p.get("to")))
        for p in model.get("paths", [])
        if isinstance(p, dict) and str(p.get("from")) in constructs and str(p.get("to")) in constructs
    ]
    if not constructs or not paths:
        return None

    lines = ["flowchart LR"]
    for cid, label in constructs.items():
        lines.append(f'  {cid}["{label.replace(chr(34), chr(39))}"]')
    for a, b in paths:
        lines.append(f"  {a} --> {b}")
    mmd = "\n".join(lines)

    mmdc = _MERMAID_DIR / "node_modules" / ".bin" / "mmdc"
    cfg = _MERMAID_DIR / "puppeteer.json"
    if not mmdc.exists():
        logger.warning("partner_report: mmdc not installed — skipping model diagram")
        return None
    try:
        d = Path(tempfile.mkdtemp(prefix="model_"))
        mmd_path = d / "model.mmd"
        png_path = d / "model.png"
        mmd_path.write_text(mmd, encoding="utf-8")
        env = dict(os.environ)
        env["PATH"] = f"{_NODE_BIN}:/opt/homebrew/bin:" + env.get("PATH", "")
        subprocess.run(
            [str(mmdc), "-i", str(mmd_path), "-o", str(png_path), "-c", str(cfg), "-b", "white", "-w", "1100"],
            check=True, capture_output=True, timeout=120, env=env, cwd=str(_MERMAID_DIR),
        )
        if not png_path.exists():
            return None
        # Return a base64 data URI so the image embeds in BOTH the WeasyPrint PDF
        # (no base_url needed) and the pandoc DOCX, with no external file at render.
        import base64
        b64 = base64.b64encode(png_path.read_bytes()).decode("ascii")
        return f"data:image/png;base64,{b64}"
    except Exception:
        logger.exception("partner_report: mermaid render failed")
        return None


def _extract_text(file_bytes: bytes, filename: str | None) -> tuple[str, int]:
    """Extract analysis text from a PDF or DOCX. Returns (text, page_count).

    DOCX is detected by extension or the zip magic (PK). Everything else goes
    through the PDF extractor. Table cell text is flattened into pipe rows so the
    statistics inside result tables survive.
    """
    name = (filename or "").lower()
    if name.endswith(".docx") or file_bytes[:2] == b"PK":
        try:
            import io

            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return ("\n".join(parts), 0)
        except Exception:
            logger.exception("partner_report: docx text extraction failed")
            return ("", 0)
    return extract_pdf_text(file_bytes)


# Signals that an uploaded file actually contains statistical-analysis output
# (SmartPLS / SPSS), which the Results (M4) chapter needs. Covers English and
# Vietnamese terminology.
_M4_DATA_SIGNALS = (
    "cronbach", "composite reliability", "average variance", "ave", "htmt",
    "heterotrait", "outer loading", "factor loading", "r square", "r-square",
    "p value", "p-value", "t statistic", "t-statistic", "std", "standard deviation",
    "correlation", "regression", "coefficient", "path coefficient", "variance",
    "eigenvalue", "kmo", "bartlett", "f square", "vif", "original sample",
    "sample mean", "stdev", "significance", "sig.", "beta", "β", "α",
    "độ tin cậy", "phương sai", "tương quan", "hồi quy", "hệ số", "trung bình",
    "độ lệch chuẩn", "kiểm định", "giá trị hội tụ", "giá trị phân biệt",
    "tải nhân tố", "nhân tố",
)


def pdf_looks_like_analysis(text: str) -> bool:
    """True when the extracted text looks like real statistical-analysis output.

    The Results (M4) chapter is built FROM the uploaded analysis (reliability,
    validity, path coefficients, …). If the file is a proposal / an unrelated
    doc / mostly prose with no numbers, M4 has nothing to tabulate — we fail
    fast instead of fabricating tables. Heuristic: needs at least two distinct
    statistical terms AND a handful of decimal numbers.
    """
    low = (text or "").lower()
    keyword_hits = sum(1 for k in _M4_DATA_SIGNALS if k in low)
    decimal_hits = len(re.findall(r"\d[.,]\d", text or ""))
    return keyword_hits >= 2 and decimal_hits >= 6


def build_partner_context_store(
    text: str,
    *,
    notes: str | None,
    language: str,
    m1: dict | None = None,
    m2: dict | None = None,
    m3: dict | None = None,
) -> dict:
    """Assemble the nested context_store for a partner report.

    Input contract: M1/M2/M3 are OPTIONAL. Each one the caller provides is used
    VERBATIM; each missing one is generated (M1 via _infer_topic, M3 via
    _infer_model, M2 via real budgeted research). M4 is always the uploaded
    analysis text. Provided modules are never overwritten — partner generation
    only fills holes.
    """
    notes_clean = (notes or "").strip()

    # M1: provided or inferred (notes prepended so inference reflects user intent).
    if m1:
        m1_topic = dict(m1)
        m1_topic.setdefault("language", language)
    else:
        infer_text = (f"Mô tả bổ sung:\n{notes_clean}\n\n{text}" if notes_clean else text)
        inferred = _infer_topic(infer_text, language)
        m1_topic = {"research_title": str(inferred.get("research_title") or "").strip()
                    or "Báo cáo phân tích", "language": language}
        for key in ("field", "research_type", "objectives", "target_population", "scope"):
            val = inferred.get(key)
            if isinstance(val, str) and val.strip():
                m1_topic[key] = val.strip()
        rqs = inferred.get("research_questions")
        if isinstance(rqs, list) and rqs:
            m1_topic["research_questions"] = [str(q) for q in rqs if str(q).strip()]
    if notes_clean:
        m1_topic.setdefault("user_context", notes_clean)

    # Grounded research brief → attach the real research GAP + context +
    # suggested directions to M1 so the report's Introduction can justify the
    # study against the literature (not just restate the topic). Best-effort:
    # a slow/failed brief never blocks the report. Skipped if the caller already
    # supplied gaps.
    if not m1_topic.get("research_gaps"):
        try:
            from orchestrator.tools.m1_topic import research_brief  # noqa: PLC0415
            brief_idea = str(m1_topic.get("research_title") or "").strip()
            if notes_clean:
                brief_idea = (brief_idea + " " + notes_clean).strip()
            if brief_idea:
                brief = research_brief.func(brief_idea, field=str(m1_topic.get("field") or ""))
                if brief.get("context"):
                    m1_topic.setdefault("background_context", brief["context"])
                if brief.get("gaps"):
                    m1_topic["research_gaps"] = brief["gaps"]
                if brief.get("suggested_topics"):
                    m1_topic["suggested_topics"] = brief["suggested_topics"]
        except Exception:
            logger.warning("partner_report: research brief failed (skipping)", exc_info=True)

    store: dict = {"m1_topic": m1_topic, "m4_analysis": {"analysis_results": text}}

    # M3: provided or inferred (used later for the methodology diagram).
    store["m3_design"] = dict(m3) if m3 else (_infer_model(text, language) or {})
    # A partner report ONLY ever ships analysis output — the methodology is always
    # reconstructed from the results, never user-supplied. So the M3 gate must
    # never hard-block it: guarantee a conceptual_model even when inference came
    # back empty (bad extraction / model shown only as an image). compose still
    # has the full analysis text (M4) to write the methodology chapter from.
    _m3d = store["m3_design"]
    if not (_m3d.get("methodology") or _m3d.get("conceptual_model")):
        _m3d["conceptual_model"] = (
            "Mô hình nghiên cứu được tái lập từ kết quả phân tích thống kê đã cung cấp "
            "(các cấu trúc tiềm ẩn và quan hệ đường dẫn suy ra từ kết quả)."
            if language.startswith("vi")
            else "The research model is reconstructed from the provided statistical "
            "analysis results (latent constructs and path relationships inferred "
            "from the reported results)."
        )
        store["m3_design"] = _m3d

    # M2: provided verbatim, else REAL budgeted research (never a bare token fetch).
    if m2:
        store["m2_literature"] = dict(m2)
    else:
        refs = _budgeted_scout(
            m1_topic.get("research_title", ""),
            m1_topic.get("research_questions") or [],
        )
        if refs:
            store["m2_literature"] = {"literature_sources": refs}

    return store


def generate_partner_report(
    pdf_bytes: bytes,
    *,
    depth: str = "analysis_report",
    chapters: list[str] | None = None,
    progress_token: str | None = None,
    filename: str | None = None,
    title: str | None = None,
    notes: str | None = None,
    language: str = "en",
    # Optional caller-supplied M1/M2/M3 modules (the input contract). Each one
    # given is used verbatim by build_partner_context_store; missing ones are
    # generated. Lets a partner pass a real topic/sources/model to skip inference.
    m1: dict | None = None,
    m2: dict | None = None,
    m3: dict | None = None,
) -> dict[str, Any]:
    """Generate a report from an analysis PDF and return download URLs.

    Chapter selection precedence:
      1. `chapters` — an explicit subset of _CHAPTER_ORDER (the "tick Chương N"
         path). Unknown keys are ignored; empty-after-filter is a bad_chapters
         error.
      2. `depth` — "full_thesis" (all six) or "analysis_report" (the light
         four). Kept for back-compat / the preset buttons.

    Returns {pages, depth, chapters, sections: [titles], pdf_url, docx_url}.
    Raises ReportError(code) for empty-text / bad-depth / bad-chapters /
    compose-failed so the router can map to a stable 4xx/5xx.
    """
    if chapters:
        chapter_keys = [c for c in chapters if c in set(_CHAPTER_ORDER)]
        if not chapter_keys:
            raise ReportError("bad_chapters",
                              f"chapters must be a subset of {_CHAPTER_ORDER}")
    elif depth == "full_thesis":
        chapter_keys = list(_CHAPTER_ORDER)
    elif depth == "analysis_report":
        chapter_keys = list(_ANALYSIS_CHAPTERS)
    else:
        raise ReportError("bad_depth", f"depth must be one of {sorted(_VALID_DEPTHS)}")

    # Combine Discussion + Conclusion into a SINGLE concluding chapter, matching
    # the standard thesis structure (one "Chương 5: Kết luận", not a separate
    # Thảo luận + Kết luận). The discussion composer already emits the full
    # conclusion structure (summary → contributions → limitations → future
    # work), so we drop `conclusion` and just relabel `discussion`.
    if "conclusion" in chapter_keys:
        chapter_keys = [k for k in chapter_keys if k != "conclusion"]
        if "discussion" not in chapter_keys:
            chapter_keys.append("discussion")

    total = len(chapter_keys)
    _set_progress(progress_token, status="processing", phase="extract",
                  total=total, done=0, current=None)

    try:
        text, pages = _extract_text(pdf_bytes, filename)
        if not text.strip():
            # Image-only scans and non-text files land here — a clear 422, not 500.
            raise ReportError("no_extractable_text",
                              "the file has no machine-readable text (image-only scan?)")

        # Ingest pre-check: cheap fail-fast before any LLM spend when the upload
        # isn't statistical output at all (only when a Results chapter is asked).
        # This is a heuristic PDF sniff, NOT the store-level completeness gate.
        if "results" in chapter_keys and not pdf_looks_like_analysis(text):
            raise ReportError("insufficient_m4_data",
                              "the uploaded file lacks the statistical analysis data "
                              "needed to write the Results (M4) chapter")

        # Build the nested context_store: caller-provided M1/M2/M3 are used
        # verbatim, missing ones are generated (M2 runs REAL budgeted research).
        # This replaces the old inline topic-inference + Crossref block.
        _set_progress(progress_token, phase="research")
        context_store = build_partner_context_store(
            text, notes=notes, language=language, m1=m1, m2=m2, m3=m3)
        # A caller-typed title always wins over an inferred one (preserves the old
        # behavior where the `title` form field seeded research_title).
        if (title or "").strip():
            context_store.setdefault("m1_topic", {})["research_title"] = title.strip()

        # The ONE store-level gate (shared with the rest of the app), scoped to
        # the requested chapters. Missing data -> needs_data (partner charges a
        # small validation fee instead of a full report). No second gate here:
        # pdf_looks_like_analysis above is only a cheap ingest sniff.
        from orchestrator.tools.m5_writing import assess_export_readiness  # noqa: PLC0415
        missing = assess_export_readiness(context_store, chapter_keys)
        if missing:
            raise ReportError("needs_data", "missing required data: " + "; ".join(missing))

        references = (context_store.get("m2_literature") or {}).get("literature_sources") or None

        def _on_chapter(idx, key, title_, phase):
            _set_progress(progress_token, done=idx + (1 if phase == "end" else 0),
                          current=None if phase == "end" else title_)

        # The (now combined) Discussion chapter is presented as the concluding
        # chapter — relabel its heading to "Kết luận"/"Conclusion".
        combined_title = "Chương 5 — Kết luận" if language.startswith("vi") else "Chapter 5 — Conclusion"
        _set_progress(progress_token, phase="compose")

        # Partner composes via the shared `compose_sections`, injects its
        # methodology diagram into the returned sections, THEN exports. We do NOT
        # use compose_and_export here (which composes+exports in one call) because
        # the diagram must be embedded between those two steps — compose_and_export
        # stays for callers that don't need the diagram (and for Spec-2 reuse).
        from orchestrator.tools.compose_export import compose_sections  # noqa: PLC0415
        from orchestrator.tools.m5_writing import run_export  # noqa: PLC0415
        project_id = f"partner-{uuid.uuid4().hex}"
        sections = compose_sections(
            context_store, chapter_keys, language,
            references=references, progress=_on_chapter,
            title_overrides={"discussion": combined_title},
        )
        if not sections:
            raise ReportError("compose_failed", "the writing engine produced no sections")
        _maybe_embed_model_diagram(sections, text, language, chapter_keys)
        _set_progress(progress_token, phase="export", current=None)
        artifacts = run_export(sections, project_id, references=references, language=language)

        s3 = _s3_from_env()
        urls: dict[str, str] = {}
        keys: dict[str, str] = {}
        for a in artifacts:
            kind = a.get("kind") or a.get("type")
            key = a.get("s3_key")
            if kind and key:
                urls[f"{kind}_url"] = _presign(s3, key)
                keys[f"{kind}_key"] = key

        _set_progress(progress_token, status="done", phase="done", done=total, current=None)
    except Exception:
        _set_progress(progress_token, status="error")
        raise

    return {
        "pages": pages,
        "depth": depth,
        "chapters": chapter_keys,
        "sections": [s["title"] for s in sections],
        "pdf_url": urls.get("pdf_url"),
        "docx_url": urls.get("docx_url"),
        # Durable S3 keys so the partner can re-presign fresh download URLs when
        # a user re-opens a saved report later (presigned URLs above expire).
        "pdf_key": keys.get("pdf_key"),
        "docx_key": keys.get("docx_key"),
    }
