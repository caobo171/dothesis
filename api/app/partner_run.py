"""Partner report = a headless client of the deep agent (convergence spec §3).

The old partner_report_service was a THIRD generation engine: inline prompts,
a private compose loop, zero tools/skills/state. This module replaces it with
the same brain every surface runs: create a system-owned project row, seed the
store through commit_slice (the ONLY write path), run the deep agent headless
in a Job subprocess, compose the requested report shape through the SHARED
compose/export path, and presign from the shared exports rows.

What partner gains the day this lands: all ~20 tools, all 8 skills, threshold
checks, questionnaire audit, rubric review, preflight — everything it lacked.
"""
from __future__ import annotations

import logging
import os
import re  # used by the moved pdf_looks_like_analysis
import secrets

from sqlalchemy import select
from sqlalchemy.orm import Session

from agent.state import SLICE_OWNERSHIP
# Module-level (not lazy) so tests can monkeypatch partner_run.compose_sections /
# partner_run.run_export; m5_writing defers its own heavy LLM deps internally.
from orchestrator.tools.compose_export import compose_sections
from orchestrator.tools.m5_writing import M5_CHAPTER_ORDER, run_export

from .models import User
from .pdf_extract import extract_pdf_text

logger = logging.getLogger(__name__)

# Subset composed for the lighter "analysis_report" depth (single copy now —
# the old _CHAPTER_ORDER clone is gone; the canonical order is M5_CHAPTER_ORDER).
ANALYSIS_CHAPTERS = ["intro", "results", "discussion", "conclusion"]

PARTNER_USER_EMAIL = "partner-system@dothesis.internal"


class ReportError(Exception):
    """Raised with a stable `code` the router maps to an HTTP response."""

    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def mint_partner_token() -> str:
    """Mint the progress-polling token SERVER-SIDE.

    Task 9 made `jobs.partner_token` UNIQUE, but the token was still whatever
    the CALLER put in the `progress_token` form field — and partner auth is one
    globally shared secret, so there is no partner identity to scope tokens by.
    Two callers holding that same secret and picking the same string (a fixed
    "job-1", a colliding client-side uuid, a retry replaying an old value) would
    now collide at INSERT and 500. Minting here removes the caller from the
    equation entirely: uniqueness is guaranteed by the generator, not hoped for.

    It doubles as the capability to READ that run's progress, so it must also be
    unguessable — token_urlsafe(32) is 256 bits of CSPRNG, not a uuid4 (which
    leaks version/variant bits and is routinely treated as non-secret).
    """
    return secrets.token_urlsafe(32)


def resolve_chapters(depth: str, chapters: list[str] | None) -> list[str]:
    """Chapter selection precedence unchanged from the old service: explicit
    `chapters` subset wins over `depth`; unknown keys are ignored; an
    empty-after-filter list and an unknown depth are clean 422 codes."""
    if chapters:
        keys = [c for c in chapters if c in set(M5_CHAPTER_ORDER)]
        if not keys:
            raise ReportError("bad_chapters",
                              f"chapters must be a subset of {M5_CHAPTER_ORDER}")
        return keys
    if depth == "full_thesis":
        return list(M5_CHAPTER_ORDER)
    if depth == "analysis_report":
        return list(ANALYSIS_CHAPTERS)
    raise ReportError("bad_depth",
                      "depth must be one of ['analysis_report', 'full_thesis']")


def required_modules_for(chapters: list[str]) -> frozenset[str]:
    """The modules a headless run must finish to serve THIS chapter request.

    Fed to RunProfile.required_modules so a partner asking for 4 chapters isn't
    forced to drive a full M1-M5 run — a seeded project's M2 is usually empty
    (payloads rarely carry literature), and demanding a literature review before
    an analysis_report meant the partner got a hard error because work they never
    requested stalled.

    The chapter -> owning-module map is REUSED from the agent's advisor-feedback
    router (`agent.tools.state_tools`) rather than restated here: two maps of the
    same fact drift, and the disagreement would be invisible until a run ended in
    the wrong place. M1 owns no chapter, so it is never in the required set — its
    intake framing is a partner-supplied INPUT (seed_partner_store writes it), not
    work the run has to produce, and the M1 data the chapters actually need is
    already gated by m5_writing.assess_export_readiness at compose time.
    """
    from agent.tools.state_tools import _chapter_to_module  # noqa: PLC0415
    return frozenset(_chapter_to_module(c) for c in chapters)


def ensure_partner_user(db: Session) -> User:
    """Find-or-create the system user that owns partner projects.

    Partner runs have no end-user relationship (the partner owns billing), but
    Project.user_id is NOT NULL and the whole app authorizes through ownership
    — one well-known system row keeps every query intact, and its permanent
    0-credit balance makes job_runner._charge_auto_run a guaranteed no-op
    (charge = min(cost, credit or 0) = 0).

    password_hash is a bcrypt-shaped impossibility, not a hash of anything: no
    password can verify against it, so the row cannot be logged into even though
    it is a normal `users` row to every other query.
    """
    u = db.scalar(select(User).where(User.email == PARTNER_USER_EMAIL))
    if u is None:
        u = User(email=PARTNER_USER_EMAIL, username="partner-system",
                 password_hash="!disabled", email_verified=True, credit=0)
        db.add(u)
        db.flush()
    return u


def seed_partner_store(
    store,
    *,
    analysis_text: str,
    m1: dict | None = None,
    m2: dict | None = None,
    m3: dict | None = None,
    title: str | None = None,
    notes: str | None = None,
    language: str = "en",
) -> None:
    """Seed the project store from the partner payload — through commit_slice
    only, filtered to OWNED keys (an unowned key would either raise or, worse,
    silently never reach prod rows). Caller-provided modules are used verbatim;
    missing ones stay empty for the agent to reconstruct (backfill tool).

    Seed order M1 -> M4 matters: commit_slice flags STARTED downstream modules
    needs_review, and seeding forward means every commit's downstream is still
    locked — no spurious review flags on a brand-new project."""
    m1 = dict(m1 or {})
    if (title or "").strip():
        # A caller-typed title always wins over anything inferred later.
        m1["research_title"] = title.strip()
    m1.setdefault("language", language)
    if (notes or "").strip():
        m1.setdefault("user_context", notes.strip())
    writes1 = {k: v for k, v in m1.items() if k in set(SLICE_OWNERSHIP["M1"])}
    if writes1:
        store.commit_slice("M1", writes1, "partner payload: topic framing seed")
    if m2:
        writes2 = {k: v for k, v in m2.items() if k in set(SLICE_OWNERSHIP["M2"])}
        if writes2:
            store.commit_slice("M2", writes2, "partner payload: literature seed")
    if m3:
        writes3 = {k: v for k, v in m3.items() if k in set(SLICE_OWNERSHIP["M3"])}
        if writes3:
            store.commit_slice("M3", writes3, "partner payload: design seed")
    store.commit_slice("M4", {"analysis_results": analysis_text},
                       "partner payload: uploaded analysis output")


def run_partner_export(store, project_id, params: dict) -> dict:
    """Compose the REQUESTED report shape from the finished run's store and
    export through the shared renderer.

    The run itself already produced a full-thesis export via the M5 done-hook
    (DbProjectStateStore._auto_export_m5) — that stays in the project's Exports
    as a bonus. This renders the partner's requested subset/merged shape:
    chapters/depth are presentation choices, not a pipeline fork (spec §3)."""
    chapters = resolve_chapters(params.get("depth") or "analysis_report",
                                params.get("chapters"))
    language = params.get("language") or "en"
    full_cs = store.load_full_context_store()
    references = (full_cs.get("m2_literature") or {}).get("literature_sources") or None
    # merge_conclusion=True carries BOTH of the old service's merge sites (the
    # chapter-key drop and the separate retitle) as one Task 8 argument.
    sections = compose_sections(full_cs, chapters, language,
                                references=references, merge_conclusion=True)
    if not sections:
        raise ReportError("compose_failed", "the writing engine produced no sections")
    artifacts = run_export(sections, str(project_id),
                           references=references, language=language)
    store.persist_export_artifacts(artifacts, scope="partner")
    return {
        "sections": [s["title"] for s in sections],
        "chapters": chapters,
        "artifact_keys": {a.get("kind"): a.get("s3_key")
                          for a in artifacts if a.get("s3_key")},
    }


def _presign(s3, s3_key: str, *, expires_in: int = 3600) -> str:
    """Same raw-key presign convention as the exports router — M5 export
    uploads with Bucket=S3_BUCKET and NO settings prefix (see
    orchestrator/tools/m5_writing._upload_to_s3)."""
    bucket = os.environ.get("S3_BUCKET") or os.environ["AWS_S3_BUCKET"]
    return s3.generate_presigned_url(
        "get_object", Params={"Bucket": bucket, "Key": s3_key},
        ExpiresIn=expires_in,
    )


# --- Ingest helpers (moved verbatim from partner_report_service) --------------


def _extract_text(file_bytes: bytes, filename: str | None) -> tuple[str, int]:
    """Extract analysis text from a PDF or DOCX. Returns (text, page_count).

    DOCX is detected by extension or the zip magic (PK). Everything else goes
    through the PDF extractor. Table cell text is flattened into pipe rows so the
    statistics inside result tables survive.
    """
    name = (filename or "").lower()
    if name.endswith(".docx") or file_bytes[:2] == b"PK":
        try:
            import io

            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return ("\n".join(parts), 0)
        except Exception:
            logger.exception("partner_report: docx text extraction failed")
            return ("", 0)
    return extract_pdf_text(file_bytes)


# Signals that an uploaded file actually contains statistical-analysis output
# (SmartPLS / SPSS), which the Results (M4) chapter needs. Covers English and
# Vietnamese terminology.
_M4_DATA_SIGNALS = (
    "cronbach", "composite reliability", "average variance", "ave", "htmt",
    "heterotrait", "outer loading", "factor loading", "r square", "r-square",
    "p value", "p-value", "t statistic", "t-statistic", "std", "standard deviation",
    "correlation", "regression", "coefficient", "path coefficient", "variance",
    "eigenvalue", "kmo", "bartlett", "f square", "vif", "original sample",
    "sample mean", "stdev", "significance", "sig.", "beta", "β", "α",
    "độ tin cậy", "phương sai", "tương quan", "hồi quy", "hệ số", "trung bình",
    "độ lệch chuẩn", "kiểm định", "giá trị hội tụ", "giá trị phân biệt",
    "tải nhân tố", "nhân tố",
)


def pdf_looks_like_analysis(text: str) -> bool:
    """True when the extracted text looks like real statistical-analysis output.

    The Results (M4) chapter is built FROM the uploaded analysis (reliability,
    validity, path coefficients, …). If the file is a proposal / an unrelated
    doc / mostly prose with no numbers, M4 has nothing to tabulate — we fail
    fast instead of fabricating tables. Heuristic: needs at least two distinct
    statistical terms AND a handful of decimal numbers.
    """
    low = (text or "").lower()
    keyword_hits = sum(1 for k in _M4_DATA_SIGNALS if k in low)
    decimal_hits = len(re.findall(r"\d[.,]\d", text or ""))
    return keyword_hits >= 2 and decimal_hits >= 6
