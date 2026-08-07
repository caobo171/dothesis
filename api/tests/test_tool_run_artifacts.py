"""A tool run keeps its input, its output, and its progress.

Spec: docs/superpowers/specs/2026-08-05-tool-run-artifacts-design.md

The history could say "-89 credit" and nothing else: `tool_runs` stored counts
only, and /document/humanize streamed the .docx back without keeping it. A
student who closed the tab had lost the document they paid for.

Two lines are load-bearing here and are asserted rather than assumed:
  - the ROW outlives the FILES. Purging nulls the URIs; the billing record stays.
  - reads are owner-or-super-admin, writes are owner-only — the same asymmetry
    every other route has held since 40cec09.
"""
from __future__ import annotations

import io
import uuid
from datetime import datetime, timedelta, timezone

import pytest

from app.db import get_session_factory
from app.models import ToolRun
from tests.conftest import make_user


def _docx_bytes() -> bytes:
    """A minimal .docx with one body paragraph long enough to be eligible."""
    from docx import Document
    d = Document()
    d.add_paragraph("Kết quả phân tích cho thấy mô hình đề xuất phù hợp với dữ "
                    "liệu thu thập được từ mẫu khảo sát của nghiên cứu này.")
    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue()


def _run(db, user, **overrides) -> ToolRun:
    """A finished humanize-docx run with both files stored."""
    fields = {
        "user_id": user.id, "surface": "web", "tool": "humanize-docx", "ok": True,
        "input_s3_uri": "s3://b/users/x/tool-runs/y/input/thesis.docx",
        "output_s3_uri": "s3://b/users/x/tool-runs/y/output/thesis.docx",
        "input_filename": "thesis.docx",
        "files_expire_at": datetime.now(timezone.utc) + timedelta(days=30),
        "metrics": {"rewritten": 80, "skipped": 52},
    }
    fields.update(overrides)
    row = ToolRun(**fields)
    db.add(row); db.commit(); db.refresh(row)
    return row


# --- schema ---------------------------------------------------------------

def test_tool_run_carries_its_artifacts_and_progress():
    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        r = _run(s, u)
        assert r.metrics["rewritten"] == 80
        assert r.input_filename == "thesis.docx"
        assert r.parent_run_id is None
        # Defaults: a row written the old way is a finished run, not a stuck one.
        assert r.status == "done"
        assert r.progress_done == 0 and r.progress_total == 0


def test_a_run_can_point_at_the_run_it_was_rerun_from():
    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        first = _run(s, u)
        again = _run(s, u, parent_run_id=first.id)
        assert again.parent_run_id == first.id


# --- storage --------------------------------------------------------------

def test_store_run_files_puts_both_halves_under_one_run_directory(monkeypatch):
    from unittest.mock import MagicMock
    from app import tool_artifacts as A

    fake = MagicMock()
    monkeypatch.setattr(A, "s3_from_env", lambda: fake)
    monkeypatch.setenv("S3_BUCKET", "b")
    uid = uuid.uuid4()

    r = A.store_run_files(user_id=uid, filename="thesis.docx",
                          input_bytes=b"IN", output_bytes=b"OUT")

    assert r.input_uri.startswith(f"s3://b/users/{uid}/tool-runs/")
    assert r.input_uri.endswith("/input/thesis.docx")
    assert r.output_uri.endswith("/output/thesis.docx")
    # Both halves of one run belong together — same directory, so a purge that
    # walks a prefix cannot take one and leave the other.
    assert r.input_uri.rsplit("/input/", 1)[0] == r.output_uri.rsplit("/output/", 1)[0]
    assert fake.put_object.call_count == 2
    assert r.expires_at > datetime.now(timezone.utc) + timedelta(days=29)


def test_store_run_files_never_raises_when_s3_is_down(monkeypatch):
    """record_tool_run's contract, extended here: a bookkeeping failure must not
    cost the student the document they have already been charged for."""
    from app import tool_artifacts as A

    def boom():
        raise RuntimeError("no s3")

    monkeypatch.setattr(A, "s3_from_env", boom)
    r = A.store_run_files(user_id=uuid.uuid4(), filename="t.docx",
                          input_bytes=b"IN", output_bytes=b"OUT")
    assert r.input_uri is None and r.output_uri is None and r.expires_at is None


def test_a_failed_run_stores_its_input_anyway(monkeypatch):
    """The input is what makes a failure reproducible — and re-runnable without
    asking the student to find the file again."""
    from unittest.mock import MagicMock
    from app import tool_artifacts as A

    fake = MagicMock()
    monkeypatch.setattr(A, "s3_from_env", lambda: fake)
    monkeypatch.setenv("S3_BUCKET", "b")
    r = A.store_run_files(user_id=uuid.uuid4(), filename="t.docx",
                          input_bytes=b"IN", output_bytes=None)
    assert r.input_uri is not None
    assert r.output_uri is None
    assert fake.put_object.call_count == 1


# --- purge ----------------------------------------------------------------

def test_purge_deletes_the_objects_and_keeps_the_row(monkeypatch):
    """The row is a billing record. Only the files age out."""
    from unittest.mock import MagicMock
    from app import tool_artifacts as A

    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        stale = _run(s, u, files_expire_at=datetime.now(timezone.utc) - timedelta(days=1))
        fresh = _run(s, u)
        fake = MagicMock()
        n = A.purge_expired(s, s3=fake, now=datetime.now(timezone.utc))

        assert n == 1
        assert fake.delete_object.call_count == 2      # input + output
        s.refresh(stale); s.refresh(fresh)
        assert stale.id is not None                     # row survives
        assert stale.input_s3_uri is None and stale.output_s3_uri is None
        assert fresh.input_s3_uri is not None           # untouched


def test_purge_is_idempotent(monkeypatch):
    from unittest.mock import MagicMock
    from app import tool_artifacts as A

    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        _run(s, u, files_expire_at=datetime.now(timezone.utc) - timedelta(days=1))
        fake = MagicMock()
        assert A.purge_expired(s, s3=fake, now=datetime.now(timezone.utc)) == 1
        assert A.purge_expired(s, s3=fake, now=datetime.now(timezone.utc)) == 0


# --- recording ------------------------------------------------------------

def test_the_run_row_points_at_the_stored_files_and_what_it_did():
    from app.tool_artifacts import RunFiles
    from app.tool_billing import record_tool_run

    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        files = RunFiles(input_uri="s3://b/in.docx", output_uri="s3://b/out.docx",
                         expires_at=datetime.now(timezone.utc) + timedelta(days=30))
        res = record_tool_run(s, u, tool="humanize-docx", ok=True, files=files,
                              input_filename="thesis.docx",
                              metrics={"rewritten": 3, "skipped": 1})
        row = s.get(ToolRun, res.run_id)
        assert row.input_s3_uri == "s3://b/in.docx"
        assert row.output_s3_uri == "s3://b/out.docx"
        assert row.input_filename == "thesis.docx"
        assert row.metrics == {"rewritten": 3, "skipped": 1}
        assert row.status == "done"


# --- progress -------------------------------------------------------------

def test_begin_then_finish_updates_one_row_rather_than_writing_two():
    """The row is created before the work so progress can be polled — but the
    run must still be ONE row, or the history shows every document twice."""
    from app.tool_billing import begin_tool_run, record_tool_run

    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        run_id = begin_tool_run(s, u, tool="humanize-docx", surface="web", total=70)
        row = s.get(ToolRun, run_id)
        assert row.status == "running"
        assert row.progress_total == 70 and row.progress_done == 0

        record_tool_run(s, u, tool="humanize-docx", ok=True, run_id=run_id)
        s.expire_all()
        assert s.query(ToolRun).filter_by(user_id=u.id).count() == 1
        assert s.get(ToolRun, run_id).status == "done"


def test_a_failed_run_is_marked_failed_not_left_running():
    from app.tool_billing import begin_tool_run, record_tool_run

    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        run_id = begin_tool_run(s, u, tool="humanize-docx", surface="web", total=4)
        record_tool_run(s, u, tool="humanize-docx", ok=False, error="rewrite_failed",
                        run_id=run_id)
        s.expire_all()
        assert s.get(ToolRun, run_id).status == "failed"


def test_progress_is_visible_while_the_run_is_still_open():
    """bump_progress opens its OWN session on purpose: the request's session is
    inside a long transaction, so a write on it would not be visible to the
    poll until the whole document finished — which is the entire point."""
    from app.tool_billing import begin_tool_run, bump_progress

    Session = get_session_factory()
    with Session() as s:
        u = make_user(s); s.commit()
        run_id = begin_tool_run(s, u, tool="humanize-docx", surface="web", total=70)

    bump_progress(run_id, done=12)

    with Session() as s2:
        row = s2.get(ToolRun, run_id)
        assert row.progress_done == 12
        assert row.status == "running"


def test_bump_progress_never_raises_on_a_missing_run():
    from app.tool_billing import bump_progress
    bump_progress(10**12, done=1)      # no such row — must be a no-op


# --- access gate ----------------------------------------------------------
# admin_config._SEED — the allowlist is by email.
ADMIN_EMAIL = "caotest171@gmail.com"


def test_readable_run_admits_the_owner_and_an_admin_but_not_a_stranger():
    from fastapi import HTTPException
    from app.auth_admin import readable_run, owned_run

    Session = get_session_factory()
    with Session() as s:
        student = make_user(s)
        admin = make_user(s, email=ADMIN_EMAIL)
        stranger = make_user(s)
        s.commit()
        run = _run(s, student)

        assert readable_run(s, student, run.id).id == run.id
        # Debugging a bad run is a READ — same gate as opening the thread.
        assert readable_run(s, admin, run.id).id == run.id
        with pytest.raises(HTTPException) as e:
            readable_run(s, stranger, run.id)
        assert e.value.status_code == 404
        # 404 for a row that does not exist either — never an existence oracle.
        with pytest.raises(HTTPException) as e2:
            readable_run(s, student, 10**12)
        assert e2.value.status_code == 404


def test_rerunning_someone_elses_document_is_not_a_read():
    """The write half of the asymmetry: an admin may look at a student's run,
    never spend their credits re-running it."""
    from fastapi import HTTPException
    from app.auth_admin import owned_run

    Session = get_session_factory()
    with Session() as s:
        student = make_user(s)
        admin = make_user(s, email=ADMIN_EMAIL)
        s.commit()
        run = _run(s, student)
        assert owned_run(s, student, run.id).id == run.id
        with pytest.raises(HTTPException) as e:
            owned_run(s, admin, run.id)
        assert e.value.status_code == 404


# --- download -------------------------------------------------------------

@pytest.fixture
def dl_client(monkeypatch):
    from fastapi.testclient import TestClient
    from app.main import create_app
    monkeypatch.setenv("ORCHESTRATOR_ENABLED", "true")
    return TestClient(create_app(), follow_redirects=False)


def _login(client, email: str | None = None):
    from app.security import create_session
    Session = get_session_factory()
    with Session() as db:
        u = make_user(db, **({"email": email} if email else {}))
        db.commit(); db.refresh(u)
        token = create_session(db, u)
        db.expunge(u)
    client.headers["Authorization"] = f"Bearer {token}"
    return u, token


def _st(client, token: str, scope: str) -> str:
    return client.post("/api/v1/auth/stream-token",
                       json={"access_token": token, "scope": scope}).json()["stream_token"]


def test_owner_downloads_the_output(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    fake = MagicMock()
    fake.generate_presigned_url.return_value = "https://s3.example/x?sig=y"
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: fake)

    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    st = _st(dl_client, token, f"tool-run-file:{run.id}/output")
    r = dl_client.get(f"/api/v1/tools/runs/{run.id}/file/output?st={st}")
    assert r.status_code == 302, r.text


def test_a_stranger_cannot_download_it(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    owner, _ = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, owner)
    _, token2 = _login(dl_client)           # a different user
    st = _st(dl_client, token2, f"tool-run-file:{run.id}/output")
    assert dl_client.get(f"/api/v1/tools/runs/{run.id}/file/output?st={st}").status_code == 404


def test_a_purged_run_answers_410_not_404(dl_client, monkeypatch):
    """"It aged out" is a different fact from "no such run", and the student is
    entitled to the difference."""
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u, input_s3_uri=None, output_s3_uri=None)
    st = _st(dl_client, token, f"tool-run-file:{run.id}/output")
    r = dl_client.get(f"/api/v1/tools/runs/{run.id}/file/output?st={st}")
    assert r.status_code == 410
    assert "30" in r.text          # the retention window is named


def test_a_token_for_the_input_does_not_open_the_output(dl_client, monkeypatch):
    """The scope names the half, so a leaked URL opens one file, not both."""
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    st = _st(dl_client, token, f"tool-run-file:{run.id}/input")
    assert dl_client.get(f"/api/v1/tools/runs/{run.id}/file/output?st={st}").status_code == 401


# --- progress endpoint ----------------------------------------------------

def test_progress_endpoint_reports_where_the_run_is(dl_client):
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u, status="running", progress_done=12, progress_total=70)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/progress",
                       json={"access_token": token})
    assert r.status_code == 200
    b = r.json()
    assert b["status"] == "running" and b["done"] == 12 and b["total"] == 70


def test_a_stranger_gets_no_progress(dl_client):
    owner, _ = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, owner, status="running")
    _, token2 = _login(dl_client)
    assert dl_client.post(f"/api/v1/tools/runs/{run.id}/progress",
                          json={"access_token": token2}).status_code == 404


# --- re-run ---------------------------------------------------------------

def test_rerun_starts_a_new_run_that_remembers_its_parent(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    fake = MagicMock()
    fake.get_object.return_value = {"Body": io.BytesIO(_docx_bytes())}
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: fake)
    monkeypatch.setattr("app.tool_artifacts.s3_from_env", lambda: fake)
    monkeypatch.setattr(
        "orchestrator.tools.humanize_docx.humanize_docx",
        lambda body, **kw: (body, {"ok": True, "rewritten": 1, "skipped": 0,
                                   "usage": [], "failures": []}))

    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/rerun",
                       json={"access_token": token})
    assert r.status_code == 200, r.text

    with Session() as s:
        fresh = (s.query(ToolRun).filter(ToolRun.parent_run_id == run.id).one())
        assert fresh.tool == "humanize-docx"
        assert fresh.status == "done"


def test_an_admin_cannot_rerun_someone_elses_document(dl_client):
    owner, _ = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, owner)
    _, admin_token = _login(dl_client, email=ADMIN_EMAIL)
    assert dl_client.post(f"/api/v1/tools/runs/{run.id}/rerun",
                          json={"access_token": admin_token}).status_code == 404


def test_rerunning_a_purged_run_answers_410(dl_client):
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u, input_s3_uri=None)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/rerun",
                       json={"access_token": token})
    assert r.status_code == 410


def test_a_non_document_tool_cannot_be_rerun(dl_client):
    """Only the tools that take a file in and give a file back."""
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u, tool="writing-rhythm")
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/rerun",
                       json={"access_token": token})
    assert r.status_code == 422


def test_the_failure_breakdown_survives_into_the_run_metrics(dl_client, monkeypatch):
    """Regression, from a real run: humanize_docx computed the exact reason each
    batch kept its original text and the router persisted only
    {"rewritten", "skipped"} — so a dissertation that came back with 4,990 words
    untouched had a row that could not answer "why". The per-kind counts (and
    the coverage ratio) must land in ToolRun.metrics: counts, never prose, the
    line models.py::ToolRun draws."""
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.tool_artifacts.s3_from_env", lambda: MagicMock())
    monkeypatch.setenv("S3_BUCKET", "b")
    monkeypatch.setattr(
        "orchestrator.tools.humanize_docx.humanize_docx",
        lambda body, **kw: (body, {
            "ok": True, "rewritten": 5, "skipped": 3, "coverage": 0.625,
            "failures": [
                {"paragraphs": 2, "error": "frozen_violation"},
                {"paragraphs": 1, "error": "llm_failed", "retried": True},
            ],
            "usage": []}))

    u, token = _login(dl_client)
    r = dl_client.post(
        "/api/v1/tools/document/humanize",
        files={"file": ("t.docx", _docx_bytes(),
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document")})
    assert r.status_code == 200, r.text
    # The headless surfaces read headers, not the DB — the ratio rides there too.
    assert r.headers["X-Rewrite-Coverage"] == "0.625"

    Session = get_session_factory()
    with Session() as s:
        row = (s.query(ToolRun)
               .filter(ToolRun.user_id == u.id, ToolRun.tool == "humanize-docx")
               .one())
        assert row.metrics["rewritten"] == 5
        assert row.metrics["skipped"] == 3
        assert row.metrics["coverage"] == 0.625
        # Aggregated by KIND in PARAGRAPHS — 2 lost to a frozen violation, 1 to
        # a provider failure — so "why did the run skip text" is answerable
        # from the row alone.
        assert row.metrics["failures"] == {"frozen_violation": 2, "llm_failed": 1}


def test_an_already_human_document_answers_200_not_422(dl_client, monkeypatch):
    """Owner-measured flaw in the first coverage floor: a re-uploaded, already-
    humanized document skips everything as flatter_than_original — the guard
    declining to make good prose worse — and was answered with a 422 plus a
    token charge. The headless surfaces (auto-mode, partner API) branch on
    status codes, so this must be a 200 carrying its own distinct signal."""
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.tool_artifacts.s3_from_env", lambda: MagicMock())
    monkeypatch.setenv("S3_BUCKET", "b")
    monkeypatch.setattr(
        "orchestrator.tools.humanize_docx.humanize_docx",
        lambda body, **kw: (body, {
            "ok": True, "rewritten": 0, "skipped": 5, "declined": 5,
            "already_human": True, "coverage": 0.0,
            "failures": [{"paragraphs": 5, "error": "flatter_than_original"}],
            "usage": []}))

    u, token = _login(dl_client)
    r = dl_client.post(
        "/api/v1/tools/document/humanize",
        files={"file": ("t.docx", _docx_bytes(),
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document")})
    assert r.status_code == 200, r.text
    assert r.headers["X-Already-Human"] == "1"
    assert r.headers["X-Paragraphs-Declined"] == "5"

    Session = get_session_factory()
    with Session() as s:
        row = (s.query(ToolRun)
               .filter(ToolRun.user_id == u.id, ToolRun.tool == "humanize-docx")
               .one())
        assert row.ok is True
        assert row.metrics["declined"] == 5
        # The decline is still counted in the breakdown — visible, not damning.
        assert row.metrics["failures"] == {"flatter_than_original": 5}


# --- delete now -----------------------------------------------------------

def test_the_owner_can_delete_the_files_before_they_expire(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    fake = MagicMock()
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: fake)

    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/files/delete",
                       json={"access_token": token})
    assert r.status_code == 200 and r.json()["deleted"] == 2
    assert fake.delete_object.call_count == 2
    with Session() as s:
        row = s.get(ToolRun, run.id)
        assert row is not None                      # the billing record stays
        assert row.input_s3_uri is None and row.output_s3_uri is None


def test_deleting_twice_is_not_an_error(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u, input_s3_uri=None, output_s3_uri=None)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/files/delete",
                       json={"access_token": token})
    assert r.status_code == 200 and r.json()["deleted"] == 0


def test_an_admin_cannot_delete_someone_elses_files(dl_client, monkeypatch):
    """Reading a student's run is allowed; destroying it is not."""
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    owner, _ = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, owner)
    _, admin_token = _login(dl_client, email=ADMIN_EMAIL)
    assert dl_client.post(f"/api/v1/tools/runs/{run.id}/files/delete",
                          json={"access_token": admin_token}).status_code == 404
    with Session() as s:
        assert s.get(ToolRun, run.id).input_s3_uri is not None


def test_a_failed_s3_delete_leaves_the_row_pointing_at_the_file(dl_client, monkeypatch):
    """Clearing the URI on a failed delete would orphan the object with nothing
    left to retry it — the nightly purge finds it only while the row points."""
    from unittest.mock import MagicMock
    fake = MagicMock()
    fake.delete_object.side_effect = RuntimeError("s3 down")
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: fake)
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/files/delete",
                       json={"access_token": token})
    assert r.status_code == 502
    with Session() as s:
        assert s.get(ToolRun, run.id).input_s3_uri is not None


# --- cite progress --------------------------------------------------------

def test_cite_docx_reports_progress_across_both_phases():
    """Phase A is a CrossRef lookup per citation and phase B a model call per
    batch. A bar that sits at 0% through all of phase A reads as a stuck tool."""
    from orchestrator.tools.cite_docx import cite_docx

    seen: list[tuple[int, int]] = []
    body = _docx_bytes()
    cite_docx(body, add_missing=False,
              resolve_fn=lambda c, hint: None,
              on_progress=lambda done, total: seen.append((done, total)))
    # Always at least the opening call, so the denominator is known up front.
    assert seen and seen[0][0] == 0
    assert all(d <= tot for d, tot in seen)


def test_rerun_round_trips_a_real_document(dl_client, monkeypatch):
    """The version above stubs humanize_docx wholesale, so it proves the row
    lineage and nothing about the document. This one replaces only the MODEL —
    the real walk opens the stored .docx, rewrites its paragraphs and rebuilds
    the file — so the S3 read → parse → rewrite → rebuild → stream path is
    actually exercised end to end.
    """
    from unittest.mock import MagicMock
    from docx import Document

    original = _docx_bytes()
    fake = MagicMock()
    fake.get_object.return_value = {"Body": io.BytesIO(original)}
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: fake)
    monkeypatch.setattr("app.tool_artifacts.s3_from_env", lambda: fake)
    monkeypatch.setenv("S3_BUCKET", "b")
    # Only the model is faked. humanize_docx itself runs for real.
    monkeypatch.setattr(
        "orchestrator.tools.humanize.humanize_prose",
        lambda text, **kw: {"ok": True, "text": "ĐÃ VIẾT LẠI HOÀN TOÀN.",
                            "usage": [{"model": "m", "prompt_tokens": 5,
                                       "completion_tokens": 5}]})

    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)

    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/rerun",
                       json={"access_token": token})
    assert r.status_code == 200, r.text
    assert r.headers["content-disposition"].endswith('-humanized.docx"')

    # What came back is a real .docx, and it holds the rewrite.
    out = Document(io.BytesIO(r.content))
    assert any("ĐÃ VIẾT LẠI" in p.text for p in out.paragraphs)

    with Session() as s:
        fresh = s.query(ToolRun).filter(ToolRun.parent_run_id == run.id).one()
        assert fresh.status == "done" and fresh.ok
        # The walk reported its own denominator, and the run was billed for the
        # tokens the (faked) model reported.
        assert fresh.progress_total >= 1
        assert fresh.prompt_tokens == 5
        # Its own files were stored, so the re-run is itself re-runnable.
        assert fresh.input_s3_uri and fresh.output_s3_uri


# --- the diff view --------------------------------------------------------

def _two_docx(before: list[str], after: list[str]) -> tuple[bytes, bytes]:
    from docx import Document

    def mk(paras):
        d = Document()
        for p in paras:
            d.add_paragraph(p)
        buf = io.BytesIO(); d.save(buf)
        return buf.getvalue()
    return mk(before), mk(after)


def _stub_s3_pair(monkeypatch, before: bytes, after: bytes):
    from unittest.mock import MagicMock
    fake = MagicMock()
    fake.get_object.side_effect = [
        {"Body": io.BytesIO(before)}, {"Body": io.BytesIO(after)}]
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: fake)
    return fake


def test_the_diff_shows_which_words_moved(dl_client, monkeypatch):
    before, after = _two_docx(
        ["Giữ nguyên đoạn này.", "Kết quả cho thấy rằng mô hình phù hợp."],
        ["Giữ nguyên đoạn này.", "Kết quả cho thấy mô hình phù hợp."])
    _stub_s3_pair(monkeypatch, before, after)
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)

    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/diff",
                       json={"access_token": token})
    assert r.status_code == 200, r.text
    b = r.json()
    assert b["aligned"] and b["total"] == 2 and b["changed"] == 1
    # Only the changed paragraph is carried, and it names the removed word.
    assert [i["index"] for i in b["items"]] == [1]
    dels = [s["text"] for s in b["items"][0]["segments"] if s["op"] == "del"]
    assert any("rằng" in d for d in dels)


def test_an_admin_can_see_the_diff_but_a_stranger_cannot(dl_client, monkeypatch):
    before, after = _two_docx(["Một đoạn."], ["Một đoạn khác."])
    owner, _ = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, owner)

    _stub_s3_pair(monkeypatch, before, after)
    _, admin_token = _login(dl_client, email=ADMIN_EMAIL)
    assert dl_client.post(f"/api/v1/tools/runs/{run.id}/diff",
                          json={"access_token": admin_token}).status_code == 200

    _, other = _login(dl_client)
    assert dl_client.post(f"/api/v1/tools/runs/{run.id}/diff",
                          json={"access_token": other}).status_code == 404


def test_the_diff_of_a_purged_run_is_410(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u, output_s3_uri=None)
    r = dl_client.post(f"/api/v1/tools/runs/{run.id}/diff",
                       json={"access_token": token})
    assert r.status_code == 410


# --- exporting the diff ---------------------------------------------------

def test_the_diff_exports_as_self_contained_html(dl_client, monkeypatch):
    before, after = _two_docx(
        ["Giữ nguyên.", "Kết quả cho thấy rằng mô hình phù hợp."],
        ["Giữ nguyên.", "Kết quả cho thấy mô hình phù hợp."])
    _stub_s3_pair(monkeypatch, before, after)
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    st = _st(dl_client, token, f"tool-run-diff:{run.id}")

    r = dl_client.get(f"/api/v1/tools/runs/{run.id}/diff.html?st={st}")
    assert r.status_code == 200, r.text
    assert r.headers["content-disposition"].endswith('-diff.html"')
    body = r.text
    assert "<del" in body and "rằng" in body
    # Self-contained: nothing to fetch, so it opens anywhere.
    assert "http://" not in body and "<script" not in body
    # The export carries the WHOLE document, not only what changed.
    assert "Giữ nguyên." in body


def test_the_html_export_escapes_document_text(dl_client, monkeypatch):
    """A thesis can contain < and &; neither may become markup."""
    before, after = _two_docx(["Điều kiện a < b và c & d."], ["Điều kiện a < b, c & d."])
    _stub_s3_pair(monkeypatch, before, after)
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    st = _st(dl_client, token, f"tool-run-diff:{run.id}")
    body = dl_client.get(f"/api/v1/tools/runs/{run.id}/diff.html?st={st}").text
    assert "&lt;" in body and "&amp;" in body


def test_a_token_for_the_diff_does_not_open_the_files(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    st = _st(dl_client, token, f"tool-run-diff:{run.id}")
    assert dl_client.get(
        f"/api/v1/tools/runs/{run.id}/file/output?st={st}").status_code == 401


def test_an_unknown_export_format_is_404(dl_client, monkeypatch):
    from unittest.mock import MagicMock
    monkeypatch.setattr("app.routers.tools.s3_from_env", lambda: MagicMock())
    u, token = _login(dl_client)
    Session = get_session_factory()
    with Session() as s:
        run = _run(s, u)
    st = _st(dl_client, token, f"tool-run-diff:{run.id}")
    assert dl_client.get(
        f"/api/v1/tools/runs/{run.id}/diff.docx?st={st}").status_code == 404
