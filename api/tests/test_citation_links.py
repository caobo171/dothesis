"""Inline-citation → pandoc [@key] rewrite (drives clickable links in DOCX/PDF).

Regression: the LLM frequently writes "(Hilman 2024; Nocker 2024)" without the
comma the original regex required, so citations stayed plain text and never
became clickable links in the export.
"""
from orchestrator.tools.m5_writing import (
    _assign_citation_keys,
    _convert_inline_citations,
)

_REFS = [
    {"authors": ["Hilman"], "year": 2024, "title": "A"},
    {"authors": ["Nocker"], "year": 2024, "title": "B"},
]


def _ly():
    _csl, ly_to_key = _assign_citation_keys(_REFS)
    return ly_to_key


def test_no_comma_multi_citation_is_linkified():
    ly = _ly()
    out = _convert_inline_citations("... quan tâm (Hilman 2024; Nocker 2024).", ly)
    assert "[@hilman2024; @nocker2024]" in out
    assert "(Hilman 2024" not in out


def test_comma_form_still_works():
    ly = _ly()
    assert "[@hilman2024]" in _convert_inline_citations("text (Hilman, 2024).", ly)


def test_non_citation_parenthetical_untouched():
    ly = _ly()
    assert _convert_inline_citations("see (Figure 1) here", ly) == "see (Figure 1) here"


def test_unknown_reference_left_plain():
    ly = _ly()
    # Smith isn't in the pool — leave it exactly as written, don't mangle.
    assert _convert_inline_citations("x (Smith 2099) y", ly) == "x (Smith 2099) y"


def test_style_link_runs_adds_blue_underline(tmp_path):
    """Hyperlink runs (in-text citation links) get blue + single underline so
    they visibly look like hyperlinks in the DOCX."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from orchestrator.tools.m5_writing import _style_link_runs, _HYPERLINK_COLOR

    path = tmp_path / "doc.docx"
    doc = Document()
    p = doc.add_paragraph()
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), "ref-hilman2024")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "(Hilman 2024)"
    run.append(t)
    hl.append(run)
    p._p.append(hl)
    doc.save(str(path))

    _style_link_runs(str(path))

    doc2 = Document(str(path))
    hl2 = next(doc2.element.body.iter(qn("w:hyperlink")))
    run2 = hl2.find(qn("w:r"))
    rpr = run2.find(qn("w:rPr"))
    assert rpr is not None
    assert rpr.find(qn("w:color")).get(qn("w:val")) == _HYPERLINK_COLOR
    assert rpr.find(qn("w:u")).get(qn("w:val")) == "single"
