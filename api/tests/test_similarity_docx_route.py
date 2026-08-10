"""/tools/document/similarity — the .docx self-check, billed like humanize.

The invariant these routes exist to protect: a student must never be able to
read "we found nothing" out of "we did not look". Without a configured
provider the run still happens (the offline half is real work), but
X-Corpus-Checked is false and the charge is the offline price only.
"""
import io
import uuid

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.db import get_session_factory
from app.main import create_app
from app.models import User
from app.security import create_session

_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_DUP = ("Nghiên cứu sử dụng phương pháp chọn mẫu phi xác suất có chủ đích kết hợp "
        "thuận tiện, phát bảng câu hỏi trực tuyến đến người tiêu dùng đang sinh sống "
        "và làm việc tại Thành phố Hồ Chí Minh trong khoảng thời gian từ tháng ba "
        "đến tháng năm năm 2025.")


@pytest.fixture
def client(monkeypatch):
    # The test-support router (used below to give the user a balance) only
    # mounts under this flag — create_app() calls reset_settings(), so setting
    # it here is what makes /test/set-credit exist.
    monkeypatch.setenv("DOTHESIS_TEST_SUPPORT", "1")
    return TestClient(create_app())


@pytest.fixture
def auth(client):
    """A logged-in user WITH credit.

    Balance matters here: record_tool_run caps the charge at what the student
    actually has, so a broke user is charged 0 no matter what the tool costs —
    which would make every billing assertion below pass for the wrong reason.
    """
    sf = get_session_factory()
    with sf() as db:
        u = User(email=f"u{uuid.uuid4().hex[:6]}@x", username=f"u{uuid.uuid4().hex[:6]}",
                 password_hash="x", email_verified=True)
        db.add(u); db.commit(); db.refresh(u)
        token = create_session(db, u)
    client.headers["Authorization"] = f"Bearer {token}"
    r = client.post("/api/v1/test/set-credit", json={"access_token": token, "credit": 500})
    assert r.status_code == 200, r.text
    return token


def _thesis() -> bytes:
    d = Document()
    d.add_heading("CHƯƠNG 3", level=1)
    d.add_paragraph(_DUP)
    d.add_heading("CHƯƠNG 4", level=1)
    d.add_paragraph(_DUP)
    d.add_paragraph("Kết quả cho thấy cả ba giả thuyết đều được chấp nhận (Nguyen, 2021).")
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def _upload(name="thesis.docx"):
    return {"file": (name, io.BytesIO(_thesis()), _DOCX)}


def test_the_scan_is_free_and_quotes_the_run(client, auth):
    r = client.post("/api/v1/tools/document/similarity-scan", files=_upload())
    assert r.status_code == 200, r.text
    out = r.json()
    assert out["ok"] and out["body_paragraphs"] >= 3
    assert out["check_cost"] > 0            # the run is quoted…
    assert out["corpus_available"] is False  # …and honest about what it can do


def test_the_run_streams_a_docx_and_says_nobody_searched_a_corpus(client, auth):
    r = client.post("/api/v1/tools/document/similarity", files=_upload())
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == _DOCX
    # THE header. false ≠ clean.
    assert r.headers["X-Corpus-Checked"] == "false"
    assert int(r.headers["X-Duplication"]) == 1
    assert int(r.headers["X-Flagged"]) >= 2
    doc = Document(io.BytesIO(r.content))
    assert "không phải bản quét Turnitin" in "\n".join(p.text for p in doc.paragraphs)


def test_the_offline_run_is_not_billed_at_the_vendor_price(client, auth):
    from app.tool_billing import tool_cost

    r = client.post("/api/v1/tools/document/similarity", files=_upload())
    assert int(r.headers["X-Credits-Charged"]) == tool_cost("similarity-docx")
    assert tool_cost("similarity-docx") < tool_cost("similarity-docx-corpus")


def test_a_pdf_is_refused_rather_than_silently_downgraded(client, auth):
    r = client.post("/api/v1/tools/document/similarity",
                    files={"file": ("thesis.pdf", io.BytesIO(b"%PDF-1.4"), "application/pdf")})
    assert r.status_code == 415


def test_an_unreadable_docx_fails_the_scan_honestly(client, auth):
    r = client.post("/api/v1/tools/document/similarity-scan",
                    files={"file": ("x.docx", io.BytesIO(b"not a docx"), _DOCX)})
    assert r.status_code == 200
    assert r.json()["ok"] is False and r.json()["error"] == "unreadable"


def test_a_configured_provider_is_charged_for_and_reported(client, auth, monkeypatch):
    class _P:
        name = "fake"

        def check(self, text, *, language="vi"):
            return {"score": 0.1, "matches": [{"source": "s", "overlap": 0.1}],
                    "provider": "fake"}

    import orchestrator.tools.plagiarism as P
    monkeypatch.setattr(P, "get_provider", lambda: _P())

    from app.tool_billing import tool_cost
    r = client.post("/api/v1/tools/document/similarity", files=_upload())
    assert r.status_code == 200, r.text
    assert r.headers["X-Corpus-Checked"] == "true"
    assert int(r.headers["X-Credits-Charged"]) == tool_cost("similarity-docx-corpus")


def test_a_broken_provider_does_not_bill_the_corpus_price_or_claim_it_ran(
        client, auth, monkeypatch):
    class _Boom:
        name = "fake"

        def check(self, text, *, language="vi"):
            raise RuntimeError("vendor down")

    import orchestrator.tools.plagiarism as P
    monkeypatch.setattr(P, "get_provider", lambda: _Boom())

    from app.tool_billing import tool_cost
    r = client.post("/api/v1/tools/document/similarity", files=_upload())
    assert r.status_code == 200, r.text
    assert r.headers["X-Corpus-Checked"] == "false"
    assert int(r.headers["X-Credits-Charged"]) == tool_cost("similarity-docx")
