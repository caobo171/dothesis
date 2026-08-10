"""Stateless helper tools — routers/tools.py.

The assertions that matter here are about HONESTY, not plumbing. Both endpoints
report on things a student will act on: whether their writing looks machine-
generated, and whether a reference is real. Overstating either is worse than
returning nothing — one invites false confidence, the other invites deleting a
genuine source.
"""
import pytest
from fastapi.testclient import TestClient

from app.db import get_session_factory
from app.main import create_app
from tests.conftest import make_user

EVEN = ("This is a sentence of words here. This is a sentence of words here. "
        "This is a sentence of words here. This is a sentence of words here.")
VARIED = ("Short. " + "The committee, having reviewed the sampling frame in "
          "considerable detail across three separate sessions, concluded that "
          "the stratification was defensible. It was not. " * 2)


@pytest.fixture
def user():
    Session = get_session_factory()
    with Session() as s:
        u = make_user(s, email="w@e.com")
        s.commit(); s.refresh(u); s.expunge(u)
        return u


def _as(user):
    app = create_app()
    from app.deps import current_user
    app.dependency_overrides[current_user] = lambda: user
    return TestClient(app)


# --- writing rhythm ---------------------------------------------------------

def test_uniform_sentences_score_high(user):
    r = _as(user).post("/api/v1/tools/writing-rhythm",
                       json={"access_token": "x", "text": EVEN})
    assert r.status_code == 200
    b = r.json()
    assert b["ok"] and b["score"] > 0.6
    assert b["verdict"] == "very_even"


def test_varied_sentences_score_lower(user):
    even = _as(user).post("/api/v1/tools/writing-rhythm",
                          json={"access_token": "x", "text": EVEN}).json()
    varied = _as(user).post("/api/v1/tools/writing-rhythm",
                            json={"access_token": "x", "text": VARIED}).json()
    assert varied["score"] < even["score"]


def test_too_short_is_declined_not_guessed(user):
    """Three sentences is the floor. Scoring one sentence would be a number with
    nothing behind it, which is worse than saying no."""
    r = _as(user).post("/api/v1/tools/writing-rhythm",
                       json={"access_token": "x", "text": "One sentence."})
    assert r.json()["ok"] is False
    assert r.json()["verdict"] == "too_short"


def test_the_response_never_claims_to_be_a_detector(user):
    """detector.py calls this a WEAK signal that must not be read as a verdict.
    If wording here ever implies Turnitin/GPTZero agreement, that is a product
    honesty regression, not a copy tweak."""
    b = _as(user).post("/api/v1/tools/writing-rhythm",
                       json={"access_token": "x", "text": EVEN}).json()
    blob = " ".join(str(v) for v in b.values()).lower()
    for banned in ("turnitin", "gptzero", "detector will", "will be flagged",
                   "ai-generated", "plagiar"):
        assert banned not in blob, f"response implies detection: {banned!r}"
    assert "burstiness" in b["basis"]


# --- citation verification --------------------------------------------------

def test_a_real_doi_is_confirmed(user, monkeypatch):
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_doi", lambda doi: {
        "DOI": doi, "title": ["Attention Is All You Need"],
        "author": [{"family": "Vaswani", "given": "Ashish"}],
        "issued": {"date-parts": [[2017]]},
        "container-title": ["NeurIPS"], "URL": "https://doi.org/x"})
    r = _as(user).post("/api/v1/tools/verify-citation",
                       json={"access_token": "x",
                             "reference": "Vaswani et al. 10.5555/3295222.3295349"})
    b = r.json()
    assert b["found"] is True
    assert b["matched_by"] == "doi"
    assert b["year"] == 2017
    # An exact DOI hit carries no fuzzy-match caveat.
    assert b["warning"] is None


def test_a_text_match_is_flagged_as_fuzzy(user, monkeypatch):
    """CrossRef returns its best guess for ANY query, so a title hit is evidence,
    not proof. Flattening that into found=true would be the dangerous shortcut."""
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_text", lambda q: {
        "DOI": "10.1/x", "title": ["Something Vaguely Similar"],
        "author": [], "issued": {"date-parts": [[2020]]}})
    b = _as(user).post("/api/v1/tools/verify-citation",
                       json={"access_token": "x",
                             "reference": "Some paper about widgets, 2020"}).json()
    assert b["found"] is True
    assert b["matched_by"] == "search"
    assert "Fuzzy" in b["warning"]


def test_a_missing_doi_does_not_call_it_fabricated(user, monkeypatch):
    """A typo looks identical to an invention here. The wording must not push a
    student into deleting a real source."""
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_doi", lambda doi: None)
    b = _as(user).post("/api/v1/tools/verify-citation",
                       json={"access_token": "x", "reference": "10.9999/nope"}).json()
    assert b["found"] is False
    assert "typo" in b["detail"].lower()


def test_no_crossref_match_explains_the_limits_of_crossref(user, monkeypatch):
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_text", lambda q: None)
    b = _as(user).post("/api/v1/tools/verify-citation",
                       json={"access_token": "x", "reference": "A book, 1998"}).json()
    assert b["found"] is False
    assert "not proof" in b["detail"].lower()


def test_a_network_failure_is_never_reported_as_not_found(user, monkeypatch):
    """The worst possible error: telling a student a real citation is fake
    because CrossRef was briefly unreachable."""
    import app.routers.tools as t

    def _boom(*a, **k):
        raise RuntimeError("connection reset")
    monkeypatch.setattr(t, "_crossref_by_text", _boom)
    b = _as(user).post("/api/v1/tools/verify-citation",
                       json={"access_token": "x", "reference": "Something, 2020"}).json()
    assert b["ok"] is False
    assert b["found"] is False
    assert "no conclusion" in b["detail"].lower()


# --- citation verification, whole list --------------------------------------
# This is the flow a student actually has: attach the finished thesis, ask which
# of its sources are real. The risks are the mirror image of the single check —
# checking things that are NOT references (noise the student learns to ignore)
# and silently checking only some of them (a clean bill of health for a list
# that was never fully read).

REF_DOC = """CHƯƠNG 4: KẾT QUẢ NGHIÊN CỨU

Chương 4 trình bày kết quả xử lý dữ liệu khảo sát năm 2023 từ những du khách
đã trực tiếp trải nghiệm sản phẩm City Tour Phố Cổ.

TÀI LIỆU THAM KHẢO

[1] Nguyen, T. (2021). Service quality in heritage tourism. Tourism Review.
[2] Vaswani, A. (2017). Attention is all you need. NeurIPS.
[3] Hair, J. F. (2019). Multivariate data analysis. Cengage.
"""


def test_a_reference_list_is_split_into_its_entries():
    from app.routers.tools import extract_references

    refs = extract_references(REF_DOC)
    assert len(refs) == 3
    # The numbering is stripped — "[1] " is formatting, not part of the citation.
    assert refs[0].startswith("Nguyen, T. (2021)")


def test_body_prose_before_the_heading_is_not_checked():
    """The chapter text above carries a year ("dữ liệu khảo sát năm 2023"), so
    without the heading cut it would be looked up as if it were a citation."""
    from app.routers.tools import extract_references

    refs = extract_references(REF_DOC)
    assert not any("City Tour" in r for r in refs)


def test_a_reference_wrapped_across_lines_is_rejoined():
    """PDF extraction breaks one reference over several lines. Checking each
    fragment separately would report three failures for one real source."""
    from app.routers.tools import extract_references

    refs = extract_references(
        "References\n\nNguyen, T., & Tran, H.\nService quality in heritage\n"
        "tourism contexts. Tourism Review, 2021.\n")
    assert len(refs) == 1
    assert "Tourism Review, 2021" in refs[0]


def test_lines_without_a_year_or_doi_are_left_alone():
    from app.routers.tools import extract_references

    assert extract_references("References\n\nSee the appendix.\nAcknowledgements\n") == []


def test_every_entry_gets_the_same_verdict_the_single_check_would_give(user, monkeypatch):
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_text", lambda q: {
        "DOI": "10.1/x", "title": ["Something Vaguely Similar"],
        "author": [], "issued": {"date-parts": [[2020]]}})

    b = _as(user).post("/api/v1/tools/verify-citations",
                       json={"access_token": "x", "text": REF_DOC}).json()
    assert b["ok"] is True
    assert b["detected"] == 3
    assert b["checked"] == 3
    assert b["truncated"] is False
    # Fuzzy stays fuzzy in bulk. A batch result that quietly upgraded these to
    # "confirmed" would be the whole danger of this endpoint.
    assert all(i["matched_by"] == "search" for i in b["items"])
    assert all("Fuzzy" in i["warning"] for i in b["items"])
    # Each verdict is attributable to the line it came from.
    assert b["items"][0]["reference"].startswith("Nguyen, T. (2021)")


def test_a_long_list_reports_that_it_was_truncated(user, monkeypatch):
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_text", lambda q: None)

    doc = "References\n\n" + "".join(
        f"Author{i}, A. (20{i:02d}). A paper about widgets number {i}. Journal.\n"
        for i in range(60))
    b = _as(user).post("/api/v1/tools/verify-citations",
                       json={"access_token": "x", "text": doc}).json()
    assert b["detected"] == 60
    assert b["checked"] == 50
    assert b["truncated"] is True


def test_a_full_length_thesis_is_accepted(user, monkeypatch):
    """The mode exists to take a whole thesis, and a 100-page Vietnamese one
    extracts to ~200k characters. The first cap was 40k, which 422'd the exact
    input the attach button produces."""
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_text", lambda q: None)

    filler = "Chương 4 trình bày kết quả xử lý dữ liệu khảo sát.\n" * 4000
    doc = filler + "\nTÀI LIỆU THAM KHẢO\n\nNguyen, T. (2021). Heritage tourism. Review.\n"
    assert len(doc) > 40000
    r = _as(user).post("/api/v1/tools/verify-citations",
                       json={"access_token": "x", "text": doc})
    assert r.status_code == 200, r.text
    assert r.json()["detected"] == 1


def test_a_document_with_no_reference_list_says_so(user):
    b = _as(user).post("/api/v1/tools/verify-citations",
                       json={"access_token": "x",
                             "text": "Chương 4 trình bày kết quả xử lý dữ liệu."}).json()
    assert b["ok"] is True
    assert b["detected"] == 0
    assert b["items"] == []
    assert "No references found" in b["detail"]


# --- transport ---------------------------------------------------------------

@pytest.mark.parametrize("path,body", [
    ("/api/v1/tools/writing-rhythm", {"text": EVEN}),
    ("/api/v1/tools/verify-citation", {"reference": "10.1/x"}),
])
def test_header_only_auth_works(user, path, body, monkeypatch):
    """The MCP path: the token is an Authorization header and there is NO
    access_token in the body. An AuthedBody schema here makes Pydantic 422 the
    request before current_user ever reads the header — which is exactly how
    both of these shipped broken and were caught only against production.

    Every test above passes access_token in the body, so none of them covered
    the transport the connector actually uses.
    """
    import app.routers.tools as t
    monkeypatch.setattr(t, "_crossref_by_doi", lambda doi: None)

    client = _as(user)
    r = client.post(path, json=body, headers={"Authorization": "Bearer irrelevant"})
    assert r.status_code == 200, r.text


# --- similarity / plagiarism ------------------------------------------------
# The one behaviour worth pinning: an unconfigured or broken provider must never
# be reported as "nothing matched". A student reads that as a clean bill of
# health and submits on it.

def test_unconfigured_provider_is_not_a_clean_result(user, monkeypatch):
    monkeypatch.delenv("PLAGIARISM_PROVIDER", raising=False)
    r = _as(user).post("/api/v1/tools/plagiarism-check",
                       json={"access_token": "x", "text": "Some thesis prose."})
    assert r.status_code == 200
    b = r.json()
    assert b["ok"] is False
    assert b["error"] == "provider_not_configured"
    # Not 0.0 — a zero score is indistinguishable from "we checked, you're clean".
    assert b["score"] is None
    assert b["matches"] == []
    assert "NOT checked" in b["detail"]


def test_provider_failure_is_not_a_clean_result(user, monkeypatch):
    class Boom:
        name = "boom"

        def check(self, text, *, language="vi"):
            raise RuntimeError("vendor down")

    monkeypatch.setattr("orchestrator.tools.plagiarism.get_provider", lambda: Boom())
    r = _as(user).post("/api/v1/tools/plagiarism-check",
                       json={"access_token": "x", "text": "Some thesis prose."})
    b = r.json()
    assert b["ok"] is False
    assert b["error"] == "provider_error"
    assert b["score"] is None


def test_configured_provider_result_is_passed_through(user, monkeypatch):
    class Fake:
        name = "fake"

        def check(self, text, *, language="vi"):
            return {"score": 0.23, "provider": "fake", "matches": [
                {"source": "Nguyen 2021", "url": None, "overlap": 0.23,
                 "excerpt": "…"},
            ]}

    monkeypatch.setattr("orchestrator.tools.plagiarism.get_provider", lambda: Fake())
    b = _as(user).post("/api/v1/tools/plagiarism-check",
                       json={"access_token": "x", "text": "Some thesis prose."}).json()
    assert b["ok"] is True
    assert b["score"] == 0.23
    assert b["provider"] == "fake"
    assert b["matches"][0]["source"] == "Nguyen 2021"


# --- writing anchor ---------------------------------------------------------
# Saving used to be a SIDE EFFECT of a successful humanize, so the sample the
# feature refuses to run without could only be stored by first paying for a
# rewrite — and a rewrite that failed verification discarded it.

def test_anchor_saves_and_reads_back(user):
    c = _as(user)
    sample = " ".join(["từ"] * 160)
    saved = c.post("/api/v1/tools/writing-anchor/save",
                   json={"access_token": "x", "anchor": sample}).json()
    assert saved["ok"] is True and saved["words"] == 160

    read = c.post("/api/v1/tools/writing-anchor", json={"access_token": "x"}).json()
    assert read["has_anchor"] is True
    assert read["words"] == 160


def test_anchor_rejects_a_sample_too_short_to_carry_rhythm(user):
    c = _as(user)
    r = c.post("/api/v1/tools/writing-anchor/save",
               json={"access_token": "x", "anchor": "quá ngắn"}).json()
    assert r["ok"] is False
    assert r["error"] == "too_short"
    # And nothing was stored, so a later humanize still asks for a real sample.
    assert c.post("/api/v1/tools/writing-anchor",
                  json={"access_token": "x"}).json()["has_anchor"] is False


def test_no_anchor_reads_as_absent_not_as_an_error(user):
    r = _as(user).post("/api/v1/tools/writing-anchor", json={"access_token": "x"}).json()
    assert r["ok"] is True
    assert r["has_anchor"] is False


# --- text extraction --------------------------------------------------------

def test_extract_reads_a_plain_text_file(user):
    r = _as(user).post(
        "/api/v1/tools/extract-text",
        files={"file": ("draft.txt", "Kết quả cho thấy tác động tích cực.".encode(), "text/plain")},
    )
    assert r.status_code == 200
    b = r.json()
    assert b["ok"] is True
    assert "Kết quả" in b["text"]
    assert b["filename"] == "draft.txt"


def test_extract_rejects_an_unsupported_type(user):
    r = _as(user).post(
        "/api/v1/tools/extract-text",
        files={"file": ("photo.png", b"\x89PNG\r\n", "image/png")},
    )
    assert r.status_code == 415


def test_empty_extraction_names_the_scan_problem(user):
    """A scanned PDF is images. "No text found" sends the student hunting for a
    bug in our parser; naming the cause tells them what to actually do."""
    r = _as(user).post(
        "/api/v1/tools/extract-text",
        files={"file": ("blank.txt", b"   \n  ", "text/plain")},
    ).json()
    assert r["ok"] is False
    assert r["error"] == "no_text"
    assert "scan" in (r["detail"] or "")


# --- document humanize ------------------------------------------------------
# The point of this path is that the student gets their DOCUMENT back. The text
# extractor flattens headings to plain lines and moves every table to the end,
# which is fine for feeding the agent and useless as the input to a rewrite
# someone will hand to a supervisor.

def _sample_docx() -> bytes:
    import io as _io
    from docx import Document
    d = Document()
    d.add_heading("4.3.3. Thang đo Chuyên môn của KOLs (EX)", level=3)
    d.add_paragraph(
        "Thang đo EX gồm 5 biến quan sát đo lường mức độ người tiêu dùng cảm nhận "
        "KOLs có kiến thức, kỹ năng và sự am hiểu về sản phẩm trong lĩnh vực họ "
        "giới thiệu tới khách hàng của mình.")
    d.add_paragraph("Bảng 4.7: Chi tiết độ tin cậy")   # caption — too short
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "EXP_1"
    t.rows[0].cells[1].text = "0.694"
    buf = _io.BytesIO(); d.save(buf); return buf.getvalue()


def test_scan_counts_what_would_be_touched(user):
    r = _as(user).post("/api/v1/tools/document/scan",
                       files={"file": ("t.docx", _sample_docx(),
                                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
    assert r.status_code == 200
    b = r.json()
    assert b["ok"] is True
    assert b["body_paragraphs"] == 1      # only the real prose paragraph
    assert b["headings"] == 1             # heading skipped
    assert b["short_or_captions"] == 1    # "Bảng 4.7:" caption skipped
    assert b["tables"] == 1               # table skipped — it's data
    assert b["passages"] == 1


def test_scan_rejects_a_pdf(user):
    r = _as(user).post("/api/v1/tools/document/scan",
                       files={"file": ("t.pdf", b"%PDF-1.4", "application/pdf")})
    assert r.status_code == 415


def test_rewrite_replaces_prose_and_leaves_structure_alone(monkeypatch):
    """Headings, captions and tables must come out byte-identical."""
    from docx import Document
    import io as _io
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake_humanize(text, **kw):
        return {"ok": True, "text": "ĐÃ VIẾT LẠI.", "usage": [{"model": "m",
                "prompt_tokens": 1, "completion_tokens": 1}]}

    out, report = humanize_docx(_sample_docx(), humanize_fn=fake_humanize)
    assert report["ok"] and report["rewritten"] == 1

    d = Document(_io.BytesIO(out))
    texts = [p.text for p in d.paragraphs]
    assert "4.3.3. Thang đo Chuyên môn của KOLs (EX)" in texts   # heading intact
    assert "Bảng 4.7: Chi tiết độ tin cậy" in texts              # caption intact
    assert "ĐÃ VIẾT LẠI." in texts                               # prose rewritten
    assert d.tables[0].rows[0].cells[1].text == "0.694"          # table untouched
    # And the heading is still a heading, not a flattened line.
    heading = next(p for p in d.paragraphs if p.text.startswith("4.3.3."))
    assert heading.style.name.lower().startswith("heading")


def _para_docx(n: int, sep_headings: bool = False) -> bytes:
    """`n` eligible body paragraphs, optionally forced into separate batches.

    With `sep_headings` a heading sits between every pair, which breaks the
    index run — _batches never spans a gap — so each paragraph becomes its own
    single-paragraph batch. That is the shape the no-second-pass and coverage
    tests need, and building it here beats five copies of the same loop."""
    import io as _io
    from docx import Document
    d = Document()
    for i in range(n):
        if sep_headings and i:
            d.add_heading(f"Mục {i}", level=2)
        d.add_paragraph(f"Đoạn văn số {i} có đủ số từ để được coi là một đoạn "
                        f"văn thân bài thực sự trong tài liệu này nhé.")
    buf = _io.BytesIO(); d.save(buf)
    return buf.getvalue()


def test_a_batch_that_changes_paragraph_count_keeps_the_originals():
    """The batch is reassembled positionally, so a model that merges two
    paragraphs would shift every later one into the wrong slot. Refuse instead —
    at BOTH levels: the whole-batch pass and the per-paragraph retry run the
    same count check, so a model that pads every reply with an extra paragraph
    can never misalign the document, first try or second."""
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    # Always one MORE paragraph than it was sent: the batch fails 3-vs-2 and
    # each single-paragraph retry fails 2-vs-1.
    out, report = humanize_docx(
        _para_docx(2),
        humanize_fn=lambda t, **k: {"ok": True, "text": t + "\n\nTHÊM MỘT ĐOẠN."})
    assert report["rewritten"] == 0
    assert report["skipped"] == 2
    assert all(f["error"] == "paragraph_count_changed" for f in report["failures"])
    assert "Đoạn văn số 0" in [p.text for p in Document(_io.BytesIO(out)).paragraphs][0]


def test_a_failed_batch_is_retried_paragraph_by_paragraph():
    """Regression, from a real run: a 10,950-word dissertation came back with
    ~70% of its eligible prose byte-identical, because a batch was all-or-
    nothing — one bad reply for a 4-16 paragraph batch silently kept a page of
    original text. A failed multi-paragraph batch must be retried one paragraph
    at a time, keeping whichever succeed."""
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake(text, **kw):
        if "\n\n" in text:  # the whole-batch pass: merge, tripping the count check
            return {"ok": True, "text": "Gộp hết lại làm một."}
        return {"ok": True, "text": f"VIẾT LẠI: {text}"}

    out, report = humanize_docx(_para_docx(3), humanize_fn=fake)
    assert report["rewritten"] == 3 and report["skipped"] == 0
    assert report["failures"] == []       # nothing was LOST, so nothing to report
    assert report["ok"]
    body = [p.text for p in Document(_io.BytesIO(out)).paragraphs if p.text.strip()]
    assert all(t.startswith("VIẾT LẠI:") for t in body)


def test_a_paragraph_that_fails_its_own_retry_keeps_only_itself_original():
    """The fallback is per paragraph, so one stubborn paragraph costs exactly
    itself — its neighbours in the same failed batch still get their rewrite."""
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake(text, **kw):
        if "\n\n" in text:
            return {"ok": False, "error": "flatter_than_original"}
        if "số 1" in text:
            return {"ok": False, "error": "frozen_violation"}
        return {"ok": True, "text": f"VIẾT LẠI: {text}"}

    out, report = humanize_docx(_para_docx(3), humanize_fn=fake)
    assert report["rewritten"] == 2 and report["skipped"] == 1
    assert report["failures"] == [
        {"paragraphs": 1, "error": "frozen_violation", "retried": True}]
    body = [p.text for p in Document(_io.BytesIO(out)).paragraphs if p.text.strip()]
    assert body[0].startswith("VIẾT LẠI:") and body[2].startswith("VIẾT LẠI:")
    assert body[1].startswith("Đoạn văn số 1")  # the one that failed, unchanged


def test_a_single_paragraph_batch_gets_no_second_pass():
    """Retrying one paragraph by itself would repeat the identical call — the
    fallback exists to shrink the passage, and one paragraph cannot shrink."""
    from orchestrator.tools.humanize_docx import humanize_docx

    calls = []

    def fake(text, **kw):
        calls.append(text)
        return {"ok": False, "error": "llm_failed"}

    out, report = humanize_docx(_para_docx(1), humanize_fn=fake)
    assert len(calls) == 1
    assert report["skipped"] == 1
    assert report["failures"] == [{"paragraphs": 1, "error": "llm_failed"}]


def test_a_configuration_failure_is_not_retried_per_paragraph():
    """no_anchor is a property of the deployment, not of the batch shape —
    every retry would learn the same fact once per paragraph, for money."""
    from orchestrator.tools.humanize_docx import humanize_docx

    calls = []

    def fake(text, **kw):
        calls.append(text)
        return {"ok": False, "error": "no_anchor"}

    out, report = humanize_docx(_para_docx(2), humanize_fn=fake)
    assert len(calls) == 1
    assert report["skipped"] == 2
    assert report["failures"] == [{"paragraphs": 2, "error": "no_anchor"}]


def test_a_run_that_leaves_most_prose_untouched_is_not_ok():
    """Regression, from the same real run: `ok = rewritten > 0` let a document
    with 70% of its prose byte-identical report success, charge the student and
    fail Turnitin. Below half coverage the run must SAY it failed — the partial
    output still comes back (and is stored), but not under an "ok"."""
    import pytest as _pytest
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake(text, **kw):
        if "số 0" in text:
            return {"ok": True, "text": f"VIẾT LẠI: {text}"}
        return {"ok": False, "error": "llm_failed"}

    # Headings force 3 single-paragraph batches: 1 rewritten, 2 skipped.
    out, report = humanize_docx(_para_docx(3, sep_headings=True), humanize_fn=fake)
    assert report["rewritten"] == 1 and report["skipped"] == 2
    assert report["coverage"] == _pytest.approx(1 / 3, abs=0.01)
    assert report["ok"] is False
    assert report["error"] == "mostly_skipped"
    assert out is not None  # the partial document survives for the run history


def test_a_below_floor_run_hands_back_the_original_not_the_partial_rewrite():
    """Measured on a real submission: a partial rewrite is WORSE than no rewrite.

    Turnitin classifies overlapping ~5-10 sentence segments, so a segment
    straddling the join between a rewritten paragraph and an untouched one is
    scored as a unit. On the run this came from, paragraphs that were never
    touched — byte-identical before and after — went from 0.0% flagged to 8.2%
    purely because of what ended up next to them. Whole document 23% -> 30%.

    So below the floor the honest output is the document the student gave us.
    Shipping the partial rewrite hands back something measurably worse than
    what they uploaded.
    """
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake(text, **kw):
        if "số 0" in text:
            return {"ok": True, "text": f"VIẾT LẠI: {text}"}
        return {"ok": False, "error": "llm_failed"}

    source = _para_docx(3, sep_headings=True)
    out, report = humanize_docx(source, humanize_fn=fake)
    assert report["ok"] is False
    assert report["error"] == "mostly_skipped"
    assert report["reverted"] is True
    assert out == source          # byte-identical to what was uploaded
    assert b"VIET LAI" not in out and "VIẾT LẠI".encode() not in out


def test_a_healthy_run_still_returns_the_rewritten_document():
    """Guard on the above: reverting must fire ONLY below the floor."""
    from orchestrator.tools.humanize_docx import humanize_docx

    source = _para_docx(3, sep_headings=True)
    out, report = humanize_docx(
        source, humanize_fn=lambda text, **kw: {"ok": True, "text": f"VIẾT LẠI: {text}"})
    assert report["ok"] is True
    assert report.get("reverted") is False
    assert out != source


def test_a_document_that_is_already_human_is_not_a_failed_run():
    """Regression, owner-measured on a re-uploaded already-humanized
    dissertation: every rewrite came back flatter_than_original, because the
    burstiness guard is a RELATIVE test — a rewrite ships only if it is at
    least as varied as what it was handed, so the better the input prose, the
    more rewrites get refused. That is the guard WORKING, not the run breaking.
    The run must report success, say so distinctly (already_human), and return
    the unchanged bytes — not 422 a student for having good prose."""
    from orchestrator.tools.humanize_docx import humanize_docx

    out, report = humanize_docx(
        _para_docx(3),
        humanize_fn=lambda t, **k: {"ok": False, "error": "flatter_than_original"})
    assert report["rewritten"] == 0 and report["skipped"] == 3
    assert report["declined"] == 3
    assert report["ok"] is True
    assert report["already_human"] is True
    assert "error" not in report
    # Still COUNTED — the per-kind breakdown stays complete; the kind just
    # does not condemn the run.
    assert sum(f["paragraphs"] for f in report["failures"]
               if f["error"] == "flatter_than_original") == 3
    assert out is not None


def test_declines_do_not_drag_a_healthy_run_under_the_floor():
    """One rewritten + three already-good paragraphs was 25% coverage under the
    old floor — an HTTP 422 for a run in which nothing went wrong."""
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake(text, **kw):
        if "số 0" in text:
            return {"ok": True, "text": f"VIẾT LẠI: {text}"}
        return {"ok": False, "error": "flatter_than_original"}

    out, report = humanize_docx(_para_docx(4, sep_headings=True), humanize_fn=fake)
    assert report["rewritten"] == 1 and report["declined"] == 3
    assert report["ok"] is True
    assert report["already_human"] is False   # something WAS changed
    assert "error" not in report


def test_the_reference_list_is_never_rewritten():
    """Regression from a real re-run: _batches packed the bibliography into two
    16-paragraph batches, the whole-batch rewrites failed as a unit, and the
    references survived BY ACCIDENT. The per-paragraph fallback then dutifully
    rewrote every entry one by one — re-voicing bibliographic DATA, where a
    changed word corrupts a title, an edition, a page range. Everything after
    the references heading is ineligible like a table; the appendix that
    FOLLOWS the references is body prose again and must stay eligible."""
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx, scan_docx

    d = Document()
    # The table of contents carries the same words — last match must win.
    d.add_paragraph("References")
    body = ("The hotel sector depends on frontline behaviour, which is why "
            "leadership style matters for daily operating results in hotels.")
    d.add_paragraph(body)
    d.add_heading("REFERENCES", level=1)
    refs = [
        "Bass, B.M. (1985) Leadership and Performance Beyond Expectations. "
        "New York, NY: Free Press.",
        "Braun, S., Peus, C., Weisweiler, S. and Frey, D. (2013) "
        "'Transformational leadership, job satisfaction, and team performance', "
        "The Leadership Quarterly, 24(1), pp. 270-283.",
    ]
    for rtext in refs:
        d.add_paragraph(rtext)
    d.add_heading("APPENDIX", level=1)
    appendix = ("The table reports the loading of every indicator on its "
                "assigned construct together with bootstrap significance "
                "levels for each estimate.")
    d.add_paragraph(appendix)
    buf = _io.BytesIO(); d.save(buf)

    # The scan quotes the price off the same eligibility rule — a student must
    # not be quoted for reference entries that will never be touched.
    scan = scan_docx(buf.getvalue())
    assert scan["body_paragraphs"] == 2      # the body + the appendix, no refs
    # `words` is what partners charge on, so it has to obey the SAME rule: the
    # body plus the appendix, and not one token of the bibliography.
    assert scan["words"] == len(body.split()) + len(appendix.split())

    def fake(text, **kw):
        return {"ok": True, "text": "\n\n".join(
            "ĐÃ SỬA " + s for s in text.split("\n\n"))}

    out, report = humanize_docx(buf.getvalue(), humanize_fn=fake)
    texts = [p.text for p in Document(_io.BytesIO(out)).paragraphs]
    for rtext in refs:
        assert rtext in texts, "reference entries must come back byte-identical"
    assert "ĐÃ SỬA " + body in texts
    assert "ĐÃ SỬA " + appendix in texts
    assert report["rewritten"] == 2


def test_a_plain_text_vietnamese_references_heading_still_ends_the_prose():
    """TÀI LIỆU THAM KHẢO typed as an ordinary paragraph — no Heading style —
    must still start the frozen region: students paste, they do not style."""
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    body = ("Kết quả phân tích cho thấy phong cách lãnh đạo chuyển đổi có ảnh "
            "hưởng tích cực đến sự hài lòng của nhân viên khách sạn.")
    ref = ("Nguyễn Văn A (2019) Ảnh hưởng của phong cách lãnh đạo đến sự hài "
           "lòng của nhân viên khách sạn tại Hà Nội. Hà Nội: NXB Kinh tế.")
    d = Document()
    d.add_paragraph(body)
    d.add_paragraph("TÀI LIỆU THAM KHẢO")
    d.add_paragraph(ref)
    buf = _io.BytesIO(); d.save(buf)

    out, report = humanize_docx(
        buf.getvalue(),
        humanize_fn=lambda t, **k: {"ok": True, "text": "ĐÃ SỬA " + t})
    texts = [p.text for p in Document(_io.BytesIO(out)).paragraphs]
    assert ref in texts, "the entry after the plain-text heading must survive"
    assert "ĐÃ SỬA " + body in texts
    assert report["rewritten"] == 1


def test_markdown_emphasis_is_stripped_markers_only():
    """The model emits *Title* / **bold** despite the prompt, and the document
    writer copies characters verbatim — a real re-run shipped literal asterisks
    around every book title. Markers go; words stay; and the legitimate uses of
    the same characters (footnote stars, significance marks, multiplication,
    snake_case indicator names) must survive untouched."""
    from orchestrator.tools.humanize_docx import _strip_markdown_emphasis as S

    assert S("*Improving Organizational Effectiveness*. Thousand Oaks") == \
        "Improving Organizational Effectiveness. Thousand Oaks"
    assert S("**Kết quả** cho thấy *tác động* và _một ý_ nữa") == \
        "Kết quả cho thấy tác động và một ý nữa"
    assert S("__đậm__ và bình thường") == "đậm và bình thường"
    # Legitimate content, character-for-character:
    assert S("p < 0.05*") == "p < 0.05*"
    assert S("3 * 4 và 3*4") == "3 * 4 và 3*4"
    assert S("TL_1, JS_2 và EX_3 là các biến quan sát") == \
        "TL_1, JS_2 và EX_3 là các biến quan sát"
    assert S("* a footnote marker opening the line") == \
        "* a footnote marker opening the line"


def test_markdown_emphasis_never_reaches_the_document():
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    def fake(text, **kw):
        return {"ok": True, "text": "Kết quả **rất rõ** cho thấy *tác động* "
                                    "của TL_1 với p < 0.05* trong mô hình này."}

    out, report = humanize_docx(_para_docx(1), humanize_fn=fake)
    body = [p.text for p in Document(_io.BytesIO(out)).paragraphs][0]
    assert body == ("Kết quả rất rõ cho thấy tác động của TL_1 với p < 0.05* "
                    "trong mô hình này.")


def test_an_unknown_skip_kind_still_counts_as_a_failure():
    """A gate added upstream must announce itself in the partition before it
    can be excused — anything unrecognised must not silently become "fine"."""
    from orchestrator.tools.humanize_docx import humanize_docx

    out, report = humanize_docx(
        _para_docx(2),
        humanize_fn=lambda t, **k: {"ok": False, "error": "mystery_gate"})
    assert report["ok"] is False
    assert report["already_human"] is False
    assert report["declined"] == 0


def test_an_english_document_is_not_translated(monkeypatch):
    """Regression, from a real run: an English dissertation went through
    /tools/document/humanize and came back in Vietnamese.

    The route defaulted `language` to "vi" and the rewrite prompt reads that as
    "Rewrite the user's text in Vietnamese", so the model translated 69
    paragraphs. Numbers and citations survived, so every existing gate passed.
    The walk must hand the paragraph's OWN language down, and a translated
    rewrite must be refused even if it does not.
    """
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx
    from orchestrator.tools.humanize import detect_language

    english = ("Prior research has confirmed that leadership matters in the "
               "hospitality sector, but three limitations restrict what it can "
               "tell a hotel manager about daily practice.")
    d = Document()
    d.add_paragraph(english)
    buf = _io.BytesIO(); d.save(buf)

    seen = {}

    def fake_humanize(text, **kw):
        # What the real pass resolves the language to, without the model call.
        seen["language"] = detect_language(text) or kw.get("language") or "vi"
        return {"ok": True, "text": text.replace("Prior research", "Earlier work")}

    out, report = humanize_docx(buf.getvalue(), humanize_fn=fake_humanize)
    assert seen["language"] == "en"
    assert report["ok"]
    body = [p.text for p in Document(_io.BytesIO(out)).paragraphs][0]
    assert detect_language(body) == "en"


def test_the_university_declaration_is_never_rewritten():
    """Regression, from a real run: the humanizer re-voiced Bolton Business
    School's plagiarism declaration — the form the student SIGNS, including the
    sentence swearing no AI tool was used.

    It is a template the university supplies verbatim, not the student's prose,
    so changing its wording makes the signed statement stop matching the form.
    Turnitin flagged the paragraph anyway, so the edit bought nothing.
    """
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    declaration = (
        "I confirm that I have read the University policy on plagiarism and the "
        "work presented here is my own. I acknowledge that the University uses "
        "plagiarism detection software. No Generative AI tools were used in the "
        "development or completion of this assessment (Category A - GAI not "
        "permitted).")
    body = ("The hotel industry operates on a service logic in which the product "
            "and the person delivering it cannot be separated, which is why "
            "operating results depend so directly on frontline behaviour.")
    d = Document()
    d.add_paragraph(declaration)
    d.add_paragraph(body)
    buf = _io.BytesIO(); d.save(buf)

    def fake_humanize(text, **kw):
        return {"ok": True,
                "text": text.replace("I confirm", "I hereby confirm")
                            .replace("The hotel industry", "Hotel work")}

    out, report = humanize_docx(buf.getvalue(), humanize_fn=fake_humanize)
    paras = [p.text for p in Document(_io.BytesIO(out)).paragraphs]
    assert paras[0] == declaration, "the signed declaration must survive verbatim"
    assert paras[1] != body, "real body prose must still be rewritten"
    assert report["rewritten"] == 1


def test_body_prose_that_merely_echoes_a_form_phrase_is_still_rewritten():
    """The boilerplate markers must not swallow ordinary sentences.

    From the same dissertation: the ethics paragraph says an approval form "was
    submitted to and approved by the supervising tutor". That is body prose that
    happens to contain a cover-page phrase. A marker generic enough to match it
    anywhere in a paragraph freezes real writing the student paid to have
    rewritten — so the generic ones only count when they OPEN the paragraph.
    """
    import io as _io
    from docx import Document
    from orchestrator.tools.humanize_docx import humanize_docx

    ethics = ("Data were stored securely and used solely for academic purposes, "
              "and an Ethics Approval Form was submitted to and approved by the "
              "supervising tutor before data collection began.")
    d = Document()
    d.add_paragraph(ethics)
    buf = _io.BytesIO(); d.save(buf)

    out, report = humanize_docx(
        buf.getvalue(),
        humanize_fn=lambda t, **k: {"ok": True, "text": t.replace("Data were", "The data was")})
    assert report["rewritten"] == 1, "body prose must not be frozen as boilerplate"
    assert [p.text for p in Document(_io.BytesIO(out)).paragraphs][0] != ethics
