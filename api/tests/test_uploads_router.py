"""Tests for /api/v1/projects/{pid}/uploads + /api/v1/uploads/{id}."""
import io
import uuid
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from fastapi.testclient import TestClient

from app.db import get_session_factory
from app.main import create_app
from app.models import PaperUpload, Project, User
from app.security import create_session

FIXTURE = Path(__file__).parent / "fixtures" / "sample.pdf"


@pytest.fixture
def client(monkeypatch):
    monkeypatch.setenv("ORCHESTRATOR_ENABLED", "true")
    return TestClient(create_app())


def _login(client, email: str | None = None) -> uuid.UUID:
    # `email` is optional so a test can log in AS a specific identity — the
    # super-admin allowlist is keyed by email (app/admin_config._SEED), so the
    # admin-read tests below can't use the random address.
    sf = get_session_factory()
    with sf() as db:
        u = User(email=email or f"u{uuid.uuid4().hex[:6]}@x",
                 username=f"u{uuid.uuid4().hex[:6]}",
                 password_hash="x", email_verified=True)
        db.add(u); db.commit()
        client.headers["Authorization"] = f"Bearer {create_session(db, u)}"
        return u.id


def _project(client) -> uuid.UUID:
    return uuid.UUID(client.post("/api/v1/projects", json={"name": "T"}).json()["id"])


def test_upload_pdf_returns_id_and_extracted_text(client, monkeypatch):
    fake_s3 = MagicMock()
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: fake_s3)

    _login(client)
    pid = _project(client)

    with FIXTURE.open("rb") as f:
        r = client.post(
            f"/api/v1/projects/{pid}/uploads",
            files={"file": ("sample.pdf", f, "application/pdf")},
        )
    assert r.status_code == 200, r.text
    body = r.json()
    assert "upload_id" in body
    assert body["filename"] == "sample.pdf"
    assert body["size_bytes"] > 0
    assert body["page_count"] == 1

    assert fake_s3.put_object.call_count == 2

    sf = get_session_factory()
    with sf() as db:
        row = db.query(PaperUpload).filter_by(project_id=pid).one()
        assert row.text_extracted_at is not None
        assert row.page_count == 1


def test_upload_rejects_oversized_file(client, monkeypatch):
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: MagicMock())
    monkeypatch.setenv("M2_UPLOAD_MAX_BYTES", "100")

    _login(client)
    pid = _project(client)

    payload = b"x" * 200
    r = client.post(
        f"/api/v1/projects/{pid}/uploads",
        files={"file": ("big.pdf", io.BytesIO(payload), "application/pdf")},
    )
    assert r.status_code == 413


def test_upload_rejects_disallowed_mime_type(client, monkeypatch):
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: MagicMock())

    _login(client)
    pid = _project(client)

    r = client.post(
        f"/api/v1/projects/{pid}/uploads",
        files={"file": ("data.bin", io.BytesIO(b"x"), "application/octet-stream")},
    )
    assert r.status_code == 415


def test_list_uploads_returns_project_scoped(client, monkeypatch):
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: MagicMock())
    _login(client)
    pid = _project(client)
    with FIXTURE.open("rb") as f:
        client.post(f"/api/v1/projects/{pid}/uploads",
                    files={"file": ("a.pdf", f, "application/pdf")})

    r = client.post(f"/api/v1/projects/{pid}/uploads/list")
    assert r.status_code == 200
    items = r.json()
    assert len(items) == 1
    assert items[0]["filename"] == "a.pdf"


def test_delete_upload_removes_row(client, monkeypatch):
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: MagicMock())
    _login(client)
    pid = _project(client)
    with FIXTURE.open("rb") as f:
        upload_id = client.post(
            f"/api/v1/projects/{pid}/uploads",
            files={"file": ("a.pdf", f, "application/pdf")},
        ).json()["upload_id"]

    r = client.delete(f"/api/v1/uploads/{upload_id}")
    assert r.status_code == 204

    sf = get_session_factory()
    with sf() as db:
        assert db.query(PaperUpload).filter_by(id=uuid.UUID(upload_id)).count() == 0


def test_get_upload_text_returns_extracted_body(client, monkeypatch):
    fake_s3 = MagicMock()
    fake_s3.get_object.return_value = {"Body": io.BytesIO(b"extracted text body")}
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: fake_s3)
    _login(client)
    pid = _project(client)
    with FIXTURE.open("rb") as f:
        upload_id = client.post(
            f"/api/v1/projects/{pid}/uploads",
            files={"file": ("a.pdf", f, "application/pdf")},
        ).json()["upload_id"]

    r = client.post(f"/api/v1/uploads/{upload_id}/text")
    assert r.status_code == 200
    assert "extracted" in r.text


def test_raw_upload_supports_unicode_filename(client, monkeypatch):
    """A Vietnamese filename must not crash Starlette's Latin-1 headers."""
    fake_s3 = MagicMock()
    fake_s3.get_object.return_value = {"Body": io.BytesIO(b"docx bytes")}
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: fake_s3)
    _login(client)
    pid = _project(client)
    filename = "BẢNG HỎI.docx"  # decomposed marks reproduce the production crash
    upload_id = client.post(
        f"/api/v1/projects/{pid}/uploads",
        files={"file": (filename, io.BytesIO(b"not-a-real-docx"),
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    ).json()["upload_id"]

    st = _stream_token(client, f"project-upload:{upload_id}")
    r = client.get(f"/api/v1/uploads/{upload_id}/raw?st={st}")

    assert r.status_code == 200, r.text
    assert "filename*=UTF-8''" in r.headers["content-disposition"]
    assert r.content == b"docx bytes"


# --- owner-or-admin reads --------------------------------------------------
# An email on the super-admin seed allowlist (app/admin_config.py).
ADMIN_EMAIL = "caotest171@gmail.com"


def _student_upload(client, monkeypatch) -> str:
    """A student's project with one uploaded PDF. Returns the upload id.

    Leaves `client` authenticated as the student; the admin tests re-login.
    """
    fake_s3 = MagicMock()
    fake_s3.generate_presigned_url.return_value = "https://s3.example/x?sig=y"
    fake_s3.get_object.return_value = {"Body": io.BytesIO(b"extracted text body")}
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: fake_s3)
    _login(client)
    pid = _project(client)
    with FIXTURE.open("rb") as f:
        return client.post(f"/api/v1/projects/{pid}/uploads",
                           files={"file": ("a.pdf", f, "application/pdf")}).json()["upload_id"]


def _stream_token(client, scope: str) -> str:
    token = client.headers["Authorization"].split(" ", 1)[1]
    return client.post("/api/v1/auth/stream-token",
                       json={"access_token": token, "scope": scope}).json()["stream_token"]


def test_download_upload_allowed_for_super_admin(client, monkeypatch):
    """A super admin can download a file from another user's project.

    Regression: uploads/list was widened to owner-or-admin (readable_project)
    so the chat context panel renders for an admin debugging a student's run —
    but this download route kept the owner-only check, so the panel listed the
    file and the download button answered `project not found`. Downloading is a
    read, so it takes the read gate. Writes below stay owner-only.
    """
    upload_id = _student_upload(client, monkeypatch)

    _login(client, email=ADMIN_EMAIL)
    st = _stream_token(client, f"project-upload:{upload_id}")
    r = client.get(f"/api/v1/uploads/{upload_id}/download?st={st}",
                   follow_redirects=False)
    assert r.status_code == 302, r.text


def test_download_upload_presigns_unicode_filename_with_ascii_header(client, monkeypatch):
    fake_s3 = MagicMock()
    fake_s3.generate_presigned_url.return_value = "https://s3.example/signed"
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: fake_s3)
    _login(client)
    pid = _project(client)
    filename = "BẢNG HỎI.docx"
    upload_id = client.post(
        f"/api/v1/projects/{pid}/uploads",
        files={"file": (filename, io.BytesIO(b"not-a-real-docx"),
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    ).json()["upload_id"]

    st = _stream_token(client, f"project-upload:{upload_id}")
    response = client.get(f"/api/v1/uploads/{upload_id}/download?st={st}",
                          follow_redirects=False)
    assert response.status_code == 302
    disposition = fake_s3.generate_presigned_url.call_args.kwargs["Params"][
        "ResponseContentDisposition"]
    assert disposition.isascii()
    assert "filename*=UTF-8''" in disposition


def test_get_upload_text_allowed_for_super_admin(client, monkeypatch):
    """Same gate for the text sibling — the panel's preview must not 404 either."""
    upload_id = _student_upload(client, monkeypatch)

    _login(client, email=ADMIN_EMAIL)
    r = client.post(f"/api/v1/uploads/{upload_id}/text")
    assert r.status_code == 200, r.text
    assert "extracted" in r.text


def test_download_upload_404_for_unrelated_user(client, monkeypatch):
    """The gate widened for admins only — a normal stranger still gets 404."""
    upload_id = _student_upload(client, monkeypatch)

    _login(client)  # random, non-admin email
    st = _stream_token(client, f"project-upload:{upload_id}")
    r = client.get(f"/api/v1/uploads/{upload_id}/download?st={st}",
                   follow_redirects=False)
    assert r.status_code == 404


def test_admin_cannot_delete_someone_elses_upload(client, monkeypatch):
    """The write half of the asymmetry: reading a student's file is allowed,
    destroying it is not."""
    upload_id = _student_upload(client, monkeypatch)

    _login(client, email=ADMIN_EMAIL)
    assert client.delete(f"/api/v1/uploads/{upload_id}").status_code == 404

    sf = get_session_factory()
    with sf() as db:
        assert db.query(PaperUpload).filter_by(id=uuid.UUID(upload_id)).count() == 1


def test_upload_with_no_extractable_text_leaves_text_extracted_at_null(client, monkeypatch):
    monkeypatch.setattr("app.routers.uploads.s3_from_env", lambda: MagicMock())
    monkeypatch.setattr("app.routers.uploads.extract_pdf_text", lambda b: ("", 0))

    _login(client)
    pid = _project(client)
    with FIXTURE.open("rb") as f:
        r = client.post(
            f"/api/v1/projects/{pid}/uploads",
            files={"file": ("a.pdf", f, "application/pdf")},
        )
    assert r.status_code == 200

    sf = get_session_factory()
    with sf() as db:
        row = db.query(PaperUpload).filter_by(project_id=pid).one()
        assert row.text_extracted_at is None
        assert row.text_extract_uri is None


def _docx_with_table_between_chapters():
    """chapter 4 heading, its result table, then the chapter 5 heading —
    the shape of a finished quantitative thesis."""
    import io
    from docx import Document
    d = Document()
    d.add_paragraph("CHƯƠNG 4: KẾT QUẢ NGHIÊN CỨU")
    # Long enough on BOTH sides to clear split_final_chapter's 400-char gates —
    # a shorter fixture makes the split decline and the test pass for the wrong
    # reason.
    d.add_paragraph("Kết quả phân tích độ tin cậy được trình bày dưới đây. " * 12)
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Thang đo"; t.cell(0, 1).text = "Cronbach's Alpha"
    t.cell(1, 0).text = "ATT"; t.cell(1, 1).text = "0.8431"
    d.add_paragraph("CHƯƠNG 5: KẾT LUẬN VÀ KHUYẾN NGHỊ")
    d.add_paragraph("Nghiên cứu đóng góp vào lý thuyết hiện có. " * 12)
    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue()


def test_docx_text_keeps_tables_in_document_order():
    """A result table must stay where it was written, not be moved to the end.

    Extraction used to emit every paragraph and THEN every table, so a thesis's
    result tables all landed in one block after the final chapter. The import's
    chapter split then cut on the last chapter heading and sent every table to
    the M5 side, leaving M4 — the analysis module — with no numbers; when the
    writer regenerated chapter 5, the student's tables were overwritten and
    vanished from the export.
    """
    from app.routers.uploads import _extract_docx_text
    text, _ = _extract_docx_text(_docx_with_table_between_chapters())

    assert "0.8431" in text                       # the table survived at all
    # and it sits BETWEEN the two chapter headings, where it was written.
    assert text.index("CHƯƠNG 4") < text.index("0.8431") < text.index("CHƯƠNG 5")


def test_the_chapter_split_leaves_result_tables_with_the_analysis():
    """The property that actually matters to the student: after the import
    splits chapter 5 off, the numbers are still on the analysis side."""
    from app.routers.uploads import _extract_docx_text
    from orchestrator.chapter_split import split_final_chapter

    text, _ = _extract_docx_text(_docx_with_table_between_chapters())
    split = split_final_chapter(text)
    assert split is not None
    head, tail = split
    assert "0.8431" in head                       # M4 keeps its results
    assert "0.8431" not in tail
