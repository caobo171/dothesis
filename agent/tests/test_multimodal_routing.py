"""Capability routing (headless convergence spec §2). The routing-table test is
the test that would have caught defect 1: env-sniffing detect_provider ignored
DOTHESIS_MODEL_ROUTE and emitted Gemini media blocks into OpenAI-compat
endpoints, breaking EVERY attachment on route=ofox."""
import pytest

from agent.model_factory import ModelSpec
from agent import multimodal
from agent.multimodal import Attachment, build_user_message, detect_provider


@pytest.mark.parametrize("route,model,anthropic_key,expected", [
    ("native", "gemini-3.5-flash", None, "gemini"),
    ("native", "claude-sonnet-4-6", "sk-x", "anthropic"),
    ("ofox", "qwen/qwen-plus", None, "openai"),      # defect 1's exact case
    ("ofox", "google/gemini-2.5-flash", None, "openai"),
    ("openrouter", "meta-llama/llama-3", None, "openai"),
])
def test_provider_derives_from_spec(monkeypatch, route, model, anthropic_key, expected):
    if anthropic_key:
        monkeypatch.setenv("ANTHROPIC_API_KEY", anthropic_key)
    else:
        monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    assert detect_provider(ModelSpec(route=route, model=model)) == expected


def _png(name="chart.png"):
    return Attachment(filename=name, bytes=b"\x89PNG fakebytes", mime_type="image/png")


def _pdf(name="doc.pdf"):
    return Attachment(filename=name, bytes=b"%PDF-1.4 fake", mime_type="application/pdf")


def test_text_only_brain_transcribes_image(monkeypatch):
    # qwen-plus can't see: the image must arrive as TEXT (vision sidecar
    # transcription), never as a media/image block the endpoint rejects.
    monkeypatch.setattr(multimodal, "_transcribe_via_vision",
                        lambda att: f"[transcribed {att.filename}]")
    msg = build_user_message("look", [_png()], "openai", supports_vision=False)
    assert isinstance(msg.content, str)
    assert "[transcribed chart.png]" in msg.content


def test_text_only_brain_extracts_pdf_text(monkeypatch):
    monkeypatch.setattr("agent.pdf_extract.extract_pdf_text",
                        lambda b: ("Cronbach alpha table " * 20, 3))
    msg = build_user_message("read", [_pdf()], "openai", supports_vision=False)
    assert isinstance(msg.content, str)
    assert "Cronbach alpha table" in msg.content


def test_scanned_pdf_falls_back_to_vision(monkeypatch):
    # Risk 3: extract_pdf_text has no OCR — a scan yields near-empty text.
    # Proceeding with a hollow message is the silent failure; the fallback
    # transcribes via the vision sidecar instead.
    monkeypatch.setattr("agent.pdf_extract.extract_pdf_text", lambda b: ("", 0))
    monkeypatch.setattr(multimodal, "_transcribe_via_vision",
                        lambda att: "[vision transcription of scan]")
    msg = build_user_message("read", [_pdf("scan.pdf")], "openai", supports_vision=False)
    assert "[vision transcription of scan]" in msg.content


def test_vision_capable_openai_keeps_image_blocks():
    msg = build_user_message("look", [_png()], "openai", supports_vision=True)
    kinds = [b.get("type") for b in msg.content]
    assert "image_url" in kinds


def test_openai_non_image_never_raises(monkeypatch):
    # The NotImplementedError landmine (multimodal.py:200-209) is gone: a CSV
    # on the openai provider becomes text, whatever the vision capability.
    csv = Attachment(filename="data.csv", bytes=b"a,b\n1,2", mime_type="text/csv")
    msg = build_user_message("data", [csv], "openai", supports_vision=True)
    flat = msg.content if isinstance(msg.content, str) else str(msg.content)
    assert "a,b" in flat


def test_anthropic_pdf_document_block():
    msg = build_user_message("read", [_pdf()], "anthropic")
    kinds = [b.get("type") for b in msg.content]
    assert "document" in kinds  # no NotImplementedError anymore


def test_gemini_path_unchanged():
    msg = build_user_message("look", [_png()], "gemini")
    kinds = [b.get("type") for b in msg.content]
    assert "media" in kinds


_DOCX_MIME = ("application/vnd.openxmlformats-officedocument"
              ".wordprocessingml.document")


def _thesis_docx():
    """A results chapter: heading, prose, and the table holding the numbers."""
    import io
    from docx import Document
    d = Document()
    d.add_paragraph("CHƯƠNG 4: KẾT QUẢ NGHIÊN CỨU")
    d.add_paragraph("Kết quả phân tích độ tin cậy được trình bày dưới đây.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Thang đo"; t.cell(0, 1).text = "Cronbach's Alpha"
    t.cell(1, 0).text = "ATT"; t.cell(1, 1).text = "0.8431"
    d.add_paragraph("Kết luận chương 4.")
    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue()


def test_a_docx_attachment_is_read_as_text_not_zip_bytes():
    """A .docx used to fall through to `bytes.decode(errors="replace")`.

    A .docx IS a zip, so the brain was handed the mojibake of a compressed
    archive and told the student it had received "compressed Word package data
    rather than readable thesis text". Word is the format a thesis actually
    arrives in — it was the attachment type that mattered most and the only
    document type with no branch of its own.
    """
    from agent.multimodal import Attachment, _textualize
    label, body = _textualize(Attachment(
        filename="thesis.docx", bytes=_thesis_docx(), mime_type=_DOCX_MIME))

    assert label == "Word text"
    assert "CHƯƠNG 4" in body
    assert "0.8431" in body                    # the result table came through
    assert "PK" not in body[:20]               # not the raw zip header


def test_a_docx_attachment_keeps_its_tables_in_place():
    """Same document-order property the upload path needs: the table stays
    between the prose around it, not moved to the end."""
    from agent.multimodal import Attachment, _textualize
    _label, body = _textualize(Attachment(
        filename="thesis.docx", bytes=_thesis_docx(), mime_type=_DOCX_MIME))
    assert body.index("CHƯƠNG 4") < body.index("0.8431") < body.index("Kết luận chương 4")


def test_a_docx_recognised_by_extension_alone():
    """Browsers and partner clients do not always send the Office mime type."""
    from agent.multimodal import Attachment, _textualize
    label, body = _textualize(Attachment(
        filename="thesis.docx", bytes=_thesis_docx(),
        mime_type="application/octet-stream"))
    assert label == "Word text"
    assert "0.8431" in body


def test_an_unreadable_docx_degrades_instead_of_dying():
    """Never raise on one odd file — the turn is worth more than the attachment."""
    from agent.multimodal import Attachment, _textualize
    label, _body = _textualize(Attachment(
        filename="broken.docx", bytes=b"not really a docx", mime_type=_DOCX_MIME))
    assert label == "file content"             # fell through, no exception
