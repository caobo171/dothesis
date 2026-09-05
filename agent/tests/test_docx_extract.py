"""What the .docx walk must see.

A quantitative thesis keeps its findings in tables, and in real uploads those
tables are very often pasted screenshots of SmartPLS/SPSS output rather than
Word tables. So "did we read the file" means "did we read the pictures", and
every gap here loses numbers silently: the extraction still looks fine, it is
just missing the results.
"""
from __future__ import annotations

import io
import os

import pytest

from agent.docx_extract import _MAX_IMAGES, extract_docx_text


def _png(width: int, height: int, *, noisy: bool = False) -> bytes:
    """A real PNG python-docx will accept. `noisy` makes it incompressible, so
    the byte size crosses the small-image floor."""
    from PIL import Image

    if noisy:
        img = Image.frombytes("RGB", (width, height), os.urandom(width * height * 3))
    else:
        img = Image.new("RGB", (width, height), (255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _docx(build) -> bytes:
    from docx import Document

    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def vision(monkeypatch):
    """Stub the vision pass and record what it was asked to read."""
    import agent.multimodal as mm

    seen: list[str] = []

    def fake(att, prompt=None):
        seen.append(att.filename)
        return f"TRANSCRIBED {len(seen)}"

    monkeypatch.setattr(mm, "_transcribe_via_vision", fake)
    return seen


def test_an_image_pasted_inside_a_table_cell_is_transcribed(vision):
    """Centring a screenshot by dropping it into a 1x1 table is an ordinary Word
    habit. The walk read `cell.text` and nothing else, so that image, and every
    number in it, was invisible."""
    def build(doc):
        doc.add_paragraph("4.2. Kiểm định độ tin cậy")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).paragraphs[0].add_run().add_picture(
            io.BytesIO(_png(400, 200, noisy=True)))

    out = extract_docx_text(_docx(build))

    assert "4.2. Kiểm định độ tin cậy" in out
    assert "TRANSCRIBED 1" in out, "an image inside a table cell was never read"
    assert len(vision) == 1


def test_a_cell_image_is_read_once_even_when_the_cell_is_merged(vision):
    """`row.cells` repeats a merged cell, and a repeat here means paying for the
    same vision call twice and printing the same table twice."""
    def build(doc):
        table = doc.add_table(rows=1, cols=2)
        merged = table.cell(0, 0).merge(table.cell(0, 1))
        merged.paragraphs[0].add_run().add_picture(
            io.BytesIO(_png(400, 200, noisy=True)))

    out = extract_docx_text(_docx(build))

    assert out.count("TRANSCRIBED") == 1
    assert len(vision) == 1


def test_images_past_the_cap_are_reported_rather_than_dropped(vision):
    """The cap bounds spend on one upload, which is right. Dropping the excess
    without a word is not: a results chapter with 30 screenshots came back
    looking complete and missing its last five tables."""
    over = 2

    def build(doc):
        for _ in range(_MAX_IMAGES + over):
            doc.add_paragraph().add_run().add_picture(
                io.BytesIO(_png(120, 120, noisy=True)))

    out = extract_docx_text(_docx(build))

    assert out.count("TRANSCRIBED") == _MAX_IMAGES
    assert len(vision) == _MAX_IMAGES          # the cap still bounds the spend
    assert f"{over} more image" in out, "the dropped images are not reported"
    assert f"{_MAX_IMAGES}-image limit" in out


def test_a_small_but_readable_screenshot_is_not_mistaken_for_a_logo(vision):
    """The byte floor exists to skip logos and bullets. A tightly cropped table
    of four rows compresses well below it while still being 600px of readable
    numbers, and skipping it loses a whole construct's reliability figures."""
    def build(doc):
        doc.add_paragraph().add_run().add_picture(io.BytesIO(_png(600, 180)))

    small = _png(600, 180)
    assert len(small) < 3000, "fixture must sit under the byte floor to be a test"

    out = extract_docx_text(_docx(build))

    assert "TRANSCRIBED 1" in out
    assert len(vision) == 1


def test_a_genuine_icon_is_still_skipped(vision):
    """The other half of the same rule: small in bytes AND small on the page is
    a logo, and it must not cost a model call."""
    def build(doc):
        doc.add_paragraph("Trường Đại học")
        doc.add_paragraph().add_run().add_picture(io.BytesIO(_png(40, 40)))

    out = extract_docx_text(_docx(build))

    assert "Trường Đại học" in out
    assert "TRANSCRIBED" not in out
    assert vision == []


def test_document_order_survives(vision):
    """Tables and images must come out WHERE THEY WERE WRITTEN. Emitting all the
    prose then all the tables is what broke the import's chapter split."""
    def build(doc):
        doc.add_paragraph("before")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Construct"
        table.cell(0, 1).text = "AVE"
        doc.add_paragraph().add_run().add_picture(
            io.BytesIO(_png(300, 300, noisy=True)))
        doc.add_paragraph("after")

    lines = [ln for ln in extract_docx_text(_docx(build)).splitlines() if ln.strip()]

    assert lines[0] == "before"
    assert lines[1] == "Construct | AVE"
    assert "TRANSCRIBED 1" in lines[3]
    assert lines[-1] == "after"


def test_the_vision_pass_can_be_turned_off(vision):
    """Callers that only need prose must not pay the latency."""
    def build(doc):
        doc.add_paragraph("prose")
        doc.add_paragraph().add_run().add_picture(
            io.BytesIO(_png(400, 200, noisy=True)))

    out = extract_docx_text(_docx(build), transcribe_images=False)

    assert out.strip() == "prose"
    assert vision == []
