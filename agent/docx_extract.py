"""Read a .docx into plain text, IN DOCUMENT ORDER.

The agent twin of pdf_extract: one place that knows how to turn a Word file
into something a text-only brain can read.

Two properties matter, and both were learned the hard way:

  - Tables must come out AT ALL. A quantitative thesis keeps its results in
    them — Cronbach's alpha, EFA loadings, KMO, path coefficients. Text
    without the tables is a results chapter with no results.
  - Tables must come out WHERE THEY WERE WRITTEN. python-docx exposes
    `doc.paragraphs` and `doc.tables` as two flat lists with no interleaving,
    so the obvious implementation emits every paragraph and then every table,
    which moves a thesis's tables into one block at the very end. Downstream
    that broke the import's chapter split (every table landed on the final
    chapter's side, leaving the analysis module with none) and left the writer
    unable to tell which table belonged to which section.

  - Third, learned the same way: a results table is very often not a table at
    all. Students paste the SmartPLS/SPSS output straight in as a SCREENSHOT,
    so the numbers live in an image and this walk saw nothing. One real upload
    carried eleven result screenshots and extracted as a page of prose holding
    exactly one number; the writer then produced a Results chapter that cited
    tables it had never been given. Images are transcribed with the vision
    model and emitted in place, for the same document-order reason as tables.

  - Fourth: "we transcribe images" was not the same claim as "we read every
    image", and the gap between them lost numbers quietly. An image is skipped
    when it sits inside a table cell (a screenshot centred in a 1x1 table is an
    ordinary Word habit and `cell.text` cannot see it), when it is small on
    disk but large on the page (a cropped four-row table compresses under the
    byte floor), or when it falls past the per-document cap. Only the cap is a
    real decision, and it now says so in the output instead of dropping the
    tail in silence.

So walk the body XML instead of the convenience lists.
"""
from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

# Namespaces needed to find an image reference inside a paragraph.
_NS_A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
_NS_R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"

# Bound the vision spend on one document: a deck of screenshots must not turn a
# single upload into an unbounded run of model calls. Small images are skipped
# outright — logos, bullets and signature scribbles carry no statistics.
_MAX_IMAGES = 25
_MIN_IMAGE_BYTES = 3000
# ...but bytes alone got this wrong in the direction that loses data. A tightly
# cropped four-row reliability table is mostly white, compresses to well under
# 2 KB, and is still 600px of readable numbers. So an image is only a logo when
# it is small BOTH on disk and on the page.
_MIN_IMAGE_PX = 200

_IMAGE_PROMPT = (
    "This image comes from a thesis (often SmartPLS/SPSS output). "
    "If it shows a TABLE, transcribe it verbatim as a Markdown table — keep every "
    "row label, column header and number EXACTLY as shown. "
    "If it is a path/model diagram, describe the constructs and the arrows between "
    "them, including any numbers on the paths. "
    "If it carries no data, reply with the single word NONE. "
    "Never invent, round or guess a number; mark unreadable cells [unreadable]."
)


def _image_rids(element) -> list[str]:
    """Relationship ids of every image embedded anywhere under one element.

    Takes any block element, not just a paragraph: a screenshot centred by
    dropping it into a 1x1 table is an ordinary Word habit, and a table cell is
    where this has to look to find it.
    """
    return [
        rid
        for blip in element.iter(f"{_NS_A}blip")
        if (rid := blip.get(f"{_NS_R}embed"))
    ]


def _longest_side_px(part) -> int:
    """Longest side of an image part in pixels, or 0 when it will not say."""
    try:
        image = part.image
        return max(int(image.px_width or 0), int(image.px_height or 0))
    except Exception:
        return 0


def _transcribe(doc, rid: str, index: int) -> str | None:
    """Vision-transcribe one embedded image, or None to skip it.

    Never raises: a missing key, an unreachable model or an odd part degrades to
    the text-only extraction this module did before.
    """
    try:
        part = doc.part.related_parts[rid]
        data = part.blob
    except Exception:
        return None
    if not data:
        return None
    if len(data) < _MIN_IMAGE_BYTES and _longest_side_px(part) < _MIN_IMAGE_PX:
        return None
    try:
        from agent.multimodal import Attachment, _transcribe_via_vision  # noqa: PLC0415 — heavy/lazy

        name = str(getattr(part, "partname", f"image{index}")).rsplit("/", 1)[-1]
        att = Attachment(
            filename=name,
            bytes=data,
            mime_type=getattr(part, "content_type", None) or "image/png",
        )
        text = (_transcribe_via_vision(att, prompt=_IMAGE_PROMPT) or "").strip()
    except Exception:
        logger.exception("docx image transcription failed (rid=%s)", rid)
        return None
    if not text or text.upper().startswith("NONE"):
        return None
    return f"[Hình {index}]\n{text}"


def extract_docx_text(data: bytes, *, transcribe_images: bool = True) -> str:
    """Paragraphs, table rows and embedded images as text, in document order.

    Table rows are flattened to `a | b | c` so the numbers inside them survive
    as text, and pasted result screenshots are transcribed where they appear.
    Best-effort: returns "" rather than raising, because the callers are an
    upload path and a chat turn, and neither should die on one odd file.

    `transcribe_images=False` skips the vision pass for callers that only need
    the prose and cannot afford the latency.
    """
    try:
        from docx import Document  # noqa: PLC0415 — heavy, and only needed here
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        images_done = 0
        over_cap = 0
        # One rid is one image, however many places reference it. `row.cells`
        # repeats a merged cell, so without this a merged screenshot is
        # transcribed twice: two vision calls, and the same table printed twice
        # into text a model then reads as two findings.
        seen_rids: set[str] = set()

        def take_images(element) -> None:
            nonlocal images_done, over_cap
            if not transcribe_images:
                return
            for rid in _image_rids(element):
                if rid in seen_rids:
                    continue
                seen_rids.add(rid)
                if images_done >= _MAX_IMAGES:
                    over_cap += 1
                    continue
                block = _transcribe(doc, rid, images_done + 1)
                if block:
                    images_done += 1
                    parts.append(block)

        for child in doc.element.body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                text = Paragraph(child, doc).text
                if text and text.strip():
                    parts.append(text)
                take_images(child)
            elif tag == "tbl":
                for row in Table(child, doc).rows:
                    cells = [c.text.strip() for c in row.cells
                             if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
                    # `cell.text` cannot see a pasted screenshot, and a results
                    # table inside a bordered cell is exactly where one lives.
                    for cell in row.cells:
                        take_images(cell._tc)
        # Say what was left out. The cap is right — it bounds the spend on one
        # upload — but dropping the excess silently is not: a results chapter
        # with thirty screenshots came back looking complete and missing its
        # last five tables, and nothing downstream could tell.
        if over_cap:
            parts.append(
                f"[docx: {over_cap} more image(s) not read: over the "
                f"{_MAX_IMAGES}-image limit for one document]")
            logger.warning("docx extraction: %d image(s) over the %d-image cap",
                           over_cap, _MAX_IMAGES)
        if images_done:
            logger.info("docx extraction: transcribed %d embedded image(s)", images_done)
        return "\n".join(parts)
    except Exception:
        logger.exception("docx text extraction failed")
        return ""
