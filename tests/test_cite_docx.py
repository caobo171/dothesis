"""Citing a .docx — orchestrator/tools/citations.py + cite_docx.py.

The assertions that matter are about FABRICATION. This feature writes references
into a student's thesis, which is the exact artifact a fabricated citation
destroys, so the tests below are mostly about what it must REFUSE to write: no
reference that CrossRef did not return, no citation the model would not confirm,
and a visible marker instead of a guess.
"""
import re

import pytest

from orchestrator.tools import citations as C
from orchestrator.tools.cite_docx import _insert_citation, cite_docx, scan_cite_docx

docx = pytest.importorskip("docx", reason="python-docx not installed")


# --- parsing ----------------------------------------------------------------

def test_vietnamese_and_english_intext_forms_are_both_found():
    """A Vietnamese thesis mixes both freely in the same chapter."""
    text = ("Bốn điều kiện cần đạt gồm hệ số KMO nằm trong khoảng 0.5 đến 1 "
            "(Hair và cộng sự, 2010). Fornell & Larcker (1981) đề xuất ngưỡng AVE. "
            "Kết quả tương tự đã được báo cáo (Nunnally, 1978; Peterson, 1994).")
    got = {(c.authors, c.year) for c in C.parse_intext_citations(text)}
    assert ("Hair và cộng sự", "2010") in got
    assert ("Fornell & Larcker", "1981") in got
    assert ("Nunnally", "1978") in got
    assert ("Peterson", "1994") in got


def test_numbers_in_brackets_are_not_read_as_sources():
    """"(0.5 đến 1)" and "(n = 2010)" are everywhere in a results chapter."""
    text = "Hệ số nằm trong khoảng (0.5 đến 1) với cỡ mẫu (n = 2010) hợp lệ."
    assert C.parse_intext_citations(text) == []


def test_the_same_source_cited_twice_is_one_reference():
    text = "(Hair và cộng sự, 2010) ... (Hair và cộng sự, 2010) ... Hair và cộng sự (2010)"
    assert len(C.dedupe(C.parse_intext_citations(text))) == 1


def test_the_documents_own_convention_is_detected():
    assert C.uses_vietnamese_convention("(Hair và cộng sự, 2010)") is True
    assert C.uses_vietnamese_convention("(Hair et al., 2010)") is False


# --- resolution -------------------------------------------------------------

CROSSREF_HIT = {
    "DOI": "10.1/x", "title": ["Multivariate Data Analysis"],
    "author": [{"family": "Hair", "given": "Joseph F"},
               {"family": "Black", "given": "William C"}],
    "issued": {"date-parts": [[2010]]}, "container-title": ["Pearson"],
}


def test_a_year_mismatch_is_not_accepted_as_the_source(monkeypatch):
    """CrossRef answers every query with something. A record from a different
    year is not the work the student cited, and taking it would put a wrong
    reference in the list under a right-looking name."""
    monkeypatch.setattr(C, "search", lambda q, rows=3: [
        {**CROSSREF_HIT, "issued": {"date-parts": [[1998]]}}])
    assert C.resolve(C.InText("Hair và cộng sự", "2010", "(Hair và cộng sự, 2010)")) is None


def test_a_matching_year_resolves(monkeypatch):
    monkeypatch.setattr(C, "search", lambda q, rows=3: [CROSSREF_HIT])
    assert C.resolve(C.InText("Hair và cộng sự", "2010", "x")) is CROSSREF_HIT


def test_a_reference_line_is_apa_formatted():
    out = C.format_reference(CROSSREF_HIT)
    assert out.startswith("Hair, J. F., & Black, W. C. (2010).")
    assert "https://doi.org/10.1/x" in out


def test_intext_form_follows_the_documents_convention():
    assert C.intext_form(CROSSREF_HIT, vietnamese=True) == "(Hair và Black, 2010)"
    assert C.intext_form(CROSSREF_HIT, vietnamese=False) == "(Hair & Black, 2010)"


# --- insertion --------------------------------------------------------------

def test_a_citation_lands_before_the_full_stop():
    out = _insert_citation("Alpha đo độ tin cậy. Câu sau.", "Alpha đo độ tin cậy.",
                           "(Hair, 2010)")
    assert out == "Alpha đo độ tin cậy (Hair, 2010). Câu sau."


def test_a_sentence_that_cannot_be_located_is_skipped():
    """The model echoes sentences back; a paraphrased echo must not be pasted in
    somewhere approximate — that silently rewrites a student's prose."""
    assert _insert_citation("Nguyên văn của sinh viên.", "Câu model tự viết lại.",
                            "(X, 2010)") is None


# --- the document walk ------------------------------------------------------

def _doc(*paragraphs):
    from docx import Document
    d = Document()
    for text, style in paragraphs:
        d.add_paragraph(text, style=style) if style else d.add_paragraph(text)
    import io
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _read(blob):
    import io
    from docx import Document
    return [p.text for p in Document(io.BytesIO(blob)).paragraphs]


BODY = ("Bốn điều kiện cần đạt gồm hệ số KMO nằm trong khoảng 0.5 đến 1, kiểm định "
        "Bartlett có Sig nhỏ hơn 0.05 và tổng phương sai trích đạt từ 50 phần trăm "
        "trở lên (Hair và cộng sự, 2010).")


def test_scan_counts_without_touching_anything():
    out = scan_cite_docx(_doc((BODY, None)))
    assert out["ok"] is True
    assert out["intext_citations"] == 1
    assert out["distinct_sources"] == 1
    assert out["has_reference_section"] is False


def test_phase_a_builds_the_reference_list_from_crossref(monkeypatch):
    blob, report = cite_docx(
        _doc((BODY, None)), add_missing=False,
        resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    assert report["resolved"] == 1
    assert report["usage"] == []          # phase A calls no model at all
    lines = _read(blob)
    assert "TÀI LIỆU THAM KHẢO" in lines
    assert any(ln.startswith("Hair, J. F., & Black, W. C. (2010).") for ln in lines)


def test_an_unresolvable_citation_is_marked_not_invented():
    """The dangerous shortcut would be composing a reference out of the in-text
    mention. It looks perfect and cites nothing that exists."""
    blob, report = cite_docx(
        _doc((BODY, None)), add_missing=False, resolve_fn=lambda cit, line=None: None)
    assert report["unresolved"] == 1
    lines = _read(blob)
    assert any("chưa đối chiếu được" in ln for ln in lines)
    assert not any("https://doi.org" in ln for ln in lines)


def test_an_existing_reference_section_is_replaced_not_duplicated():
    blob, _ = cite_docx(
        _doc((BODY, None),
             ("TÀI LIỆU THAM KHẢO", None),
             ("Hair, J. (2010). Bản cũ sai chính tả. Nxb.", None)),
        add_missing=False, resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    lines = [ln for ln in _read(blob) if ln.strip()]
    assert lines.count("TÀI LIỆU THAM KHẢO") == 1
    assert not any("Bản cũ sai chính tả" in ln for ln in lines)


class FakeMessage:
    """What langchain hands back — `text_of` reads `.content`, so a bare string
    here would pass a test the real provider fails."""

    def __init__(self, content: str):
        self.content = content


class FakeLLM:
    """Drives the two model passes without a provider."""

    model = "fake"

    def __init__(self, claim_reply: str, verify_reply: str):
        self.claim_reply, self.verify_reply = claim_reply, verify_reply

    def invoke(self, prompt: str):
        return FakeMessage(self.verify_reply if "CLAIM:" in prompt else self.claim_reply)


UNCITED = ("Cronbach Alpha chỉ cho biết các biến trong cùng một nhóm có nhất quán "
           "với nhau hay không, chứ chưa xác nhận ba nhóm biến độc lập có thực sự "
           "tách bạch thành ba khái niệm riêng biệt hay không.")


def test_a_confirmed_source_is_inserted_and_added_to_the_list():
    llm = FakeLLM(
        claim_reply='{"claims":[{"sentence":"%s","query":"cronbach alpha"}]}' % UNCITED,
        verify_reply='{"index": 0}')
    blob, report = cite_docx(
        _doc((UNCITED, None)), add_missing=True, llm=llm,
        search_fn=lambda q, rows=3: [CROSSREF_HIT],
        resolve_fn=lambda cit, line=None: None)
    assert report["added"] == 1
    assert report["marked"] == 0
    lines = _read(blob)
    assert any("(Hair và Black, 2010)" in ln for ln in lines)
    assert any(ln.startswith("Hair, J. F., & Black, W. C. (2010).") for ln in lines)


def test_an_unconfirmed_claim_is_marked_and_nothing_is_cited():
    """The whole safety property: the model declining to confirm must produce a
    visible marker, never a citation and never a reference entry."""
    llm = FakeLLM(
        claim_reply='{"claims":[{"sentence":"%s","query":"cronbach alpha"}]}' % UNCITED,
        verify_reply='{"index": null}')
    blob, report = cite_docx(
        _doc((UNCITED, None)), add_missing=True, llm=llm,
        search_fn=lambda q, rows=3: [CROSSREF_HIT],
        resolve_fn=lambda cit, line=None: None)
    assert report["added"] == 0
    assert report["marked"] == 1
    lines = _read(blob)
    assert any("[cần nguồn]" in ln for ln in lines)
    assert not any("Hair, J. F." in ln for ln in lines)


def test_no_crossref_candidate_means_no_citation():
    """Nothing reaches the document that CrossRef did not return first."""
    llm = FakeLLM(
        claim_reply='{"claims":[{"sentence":"%s","query":"cronbach alpha"}]}' % UNCITED,
        verify_reply='{"index": 0}')
    blob, report = cite_docx(
        _doc((UNCITED, None)), add_missing=True, llm=llm,
        search_fn=lambda q, rows=3: [], resolve_fn=lambda cit, line=None: None)
    assert report["added"] == 0
    assert report["marked"] == 1


def test_headings_and_body_text_survive_the_walk():
    blob, _ = cite_docx(
        _doc(("4.4. Phân tích nhân tố khám phá EFA", "Heading 2"), (BODY, None)),
        add_missing=False, resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    lines = _read(blob)
    assert "4.4. Phân tích nhân tố khám phá EFA" in lines
    assert any(ln.startswith("Bốn điều kiện cần đạt") for ln in lines)


# --- the junk CrossRef returns ----------------------------------------------
#
# Everything below reproduces a reference line that actually reached a student's
# document. CrossRef indexes supplemental files, figure components and
# peer-review reports, and its bibliographic search ranks them alongside papers.

COMPONENT = {
    "type": "component", "DOI": "10.7717/peerj.15542/supp-6",
    "title": ["Supplemental Information 6: Reliability and validity tests "
              "(<i>N</i>\n = 2,175)"],
    "author": [], "issued": {"date-parts": [[None]]},
}

AUTHORLESS = {
    "type": "book-chapter", "DOI": "10.4135/9781529693706",
    "title": ["Interpreting Correlation Matrix &amp; Unrotated Factor Solution"],
    "author": [], "issued": {"date-parts": [[2017]]},
}


def test_an_authorless_record_is_not_a_reference():
    """It produced "(2017). Interpreting Correlation Matrix..." — a line with no
    author, for a source the student never cited."""
    assert C.is_citable(AUTHORLESS) is False
    assert C.is_citable(COMPONENT) is False
    assert C.is_citable(CROSSREF_HIT) is True


def test_a_record_by_someone_else_entirely_is_refused(monkeypatch):
    """The gate that was missing. The student wrote a NAME; a 2017 record whose
    authors share none of it is not the work they cited, however well it ranks."""
    other = {**CROSSREF_HIT, "issued": {"date-parts": [[2010]]},
             "author": [{"family": "Nguyen", "given": "T"}]}
    monkeypatch.setattr(C, "search", lambda q, rows=3: [other])
    assert C.resolve(C.InText("Hair và cộng sự", "2010", "x")) is None


def test_jats_markup_never_reaches_the_document():
    """CrossRef titles arrive as JATS. Written raw, the student's reference list
    shows "&amp;" and "<i>N</i>" where a title should be."""
    assert C.clean("Interpreting Correlation Matrix &amp; Unrotated Factor") == \
        "Interpreting Correlation Matrix & Unrotated Factor"
    assert C.clean("Reliability tests (<i>N</i>\n = 2,175)") == \
        "Reliability tests ( N = 2,175)"


def test_an_escaped_less_than_survives_cleaning():
    """"p &lt; 0.05" is a real title fragment, not markup to strip."""
    assert C.clean("Effects at p &lt; 0.05") == "Effects at p < 0.05"


def test_the_reference_line_is_cleaned(monkeypatch):
    hit = {**CROSSREF_HIT, "title": ["Data Analysis &amp; <i>Method</i>"]}
    assert "Data Analysis & Method." in C.format_reference(hit)


# --- "Theo" is not an author ------------------------------------------------

def test_a_vietnamese_lead_in_word_is_not_part_of_the_author():
    """"Theo Hair và cộng sự (2019)" means "According to Hair et al. (2019)".
    Left in, it went to CrossRef as part of the name, never resolved, and the
    reference list printed "Theo Hair và cộng sự" as if that were a person."""
    got = C.parse_intext_citations("Theo Hair và cộng sự (2019) thì mô hình đạt yêu cầu.")
    assert [(c.authors, c.year) for c in got] == [("Hair và cộng sự", "2019")]


def test_the_typo_for_cong_su_is_still_read_as_et_al():
    """"công sự" for "cộng sự" is what students and PDF exports actually write."""
    cit = C.parse_intext_citations("Theo Hair và công sự (2019) thì đạt.")[0]
    assert cit.key == C.parse_intext_citations("Hair và cộng sự (2019) thì đạt.")[0].key


def test_a_genitive_lead_in_is_dropped_too():
    got = C.parse_intext_citations("Mô hình của Fornell & Larcker (1981) được dùng.")
    assert got[0].authors == "Fornell & Larcker"


# --- the student's own reference list ---------------------------------------
#
# An in-text citation carries a surname and a year, and that is a poor query:
# CrossRef's bibliographic search answers "Fornell & Larcker 1981" with five
# TABLES from other people's papers that mention the Fornell-Larcker criterion.
# The student's own reference entry carries the title, and finds the real paper.

APA_LINE = ("Fornell, C., & Larcker, D. F. (1981). Evaluating structural equation "
            "models with unobservable variables and measurement error. Journal of "
            "Marketing Research, 18(1), 39-50.")


def test_the_students_own_entry_is_matched_to_their_citation():
    cit = C.InText("Fornell & Larcker", "1981", "x")
    assert C.match_reference_line(cit, [APA_LINE]) == APA_LINE


def test_an_entry_for_a_different_year_is_not_matched():
    cit = C.InText("Fornell & Larcker", "1994", "x")
    assert C.match_reference_line(cit, [APA_LINE]) is None


def test_the_entry_that_opens_with_the_surname_wins():
    """A line that merely mentions Hair in the middle is somebody else's paper."""
    mine = "Hair, J. F. (2019). Multivariate data analysis. Cengage."
    theirs = "Sarstedt, M. (2019). Revisiting Hair et al.'s multivariate data analysis."
    cit = C.InText("Hair và cộng sự", "2019", "x")
    assert C.match_reference_line(cit, [theirs, mine]) == mine


def test_the_reference_line_is_used_as_the_query(monkeypatch):
    seen = []

    def fake_search(text, rows=3):
        seen.append(text)
        return []

    monkeypatch.setattr(C, "search", fake_search)
    C.resolve(C.InText("Fornell & Larcker", "1981", "x"), APA_LINE)
    assert seen[0] == APA_LINE  # the full line first, author+year only as fallback


def test_a_doi_in_the_students_line_is_looked_up_exactly(monkeypatch):
    """The one path that is not a search. The student handed us the identifier."""
    hit = {**CROSSREF_HIT, "DOI": "10.1177/002224378101800104"}
    monkeypatch.setattr(C, "by_doi", lambda doi: hit)
    monkeypatch.setattr(C, "search", lambda text, rows=3: [])
    got = C.resolve(C.InText("Fornell & Larcker", "1981", "x"),
                    APA_LINE + " https://doi.org/10.1177/002224378101800104")
    assert got is hit


def test_an_unresolved_citation_keeps_the_students_own_entry():
    """Replacing a real entry with "Tác giả (2019). [chưa đối chiếu được]" throws
    away a title, a journal and page numbers the student already had."""
    blob, report = cite_docx(
        _doc((BODY, None),
             ("TÀI LIỆU THAM KHẢO", None),
             ("Hair, J. F. (2010). Multivariate data analysis. Pearson.", None)),
        add_missing=False, resolve_fn=lambda cit, line=None: None)
    assert report["unresolved"] == 1
    kept = [ln for ln in _read(blob) if "Multivariate data analysis" in ln]
    assert kept and "chưa đối chiếu được" in kept[0]


def test_a_name_and_year_only_match_is_labelled_weak(monkeypatch):
    """A citation the student never listed can only be matched on surname and
    year, and "Nunnally 1978" is both "Psychometric theory" and "1K Delay Line
    Digitizer". Presenting that as a checked reference is the guess this whole
    feature exists to refuse."""
    monkeypatch.setattr(C, "search", lambda text, rows=3: [CROSSREF_HIT])
    _, how = C.resolve_verbose(C.InText("Hair và cộng sự", "2010", "x"))
    assert how == "author-year"

    blob, _ = cite_docx(_doc((BODY, None)), add_missing=False,
                        resolve_fn=lambda cit, line=None: (CROSSREF_HIT, "author-year"))
    assert any("khớp theo tên và năm" in ln for ln in _read(blob))


def test_a_match_against_the_students_line_carries_no_warning():
    blob, _ = cite_docx(_doc((BODY, None)), add_missing=False,
                        resolve_fn=lambda cit, line=None: (CROSSREF_HIT, "line"))
    assert not any("khớp theo tên và năm" in ln for ln in _read(blob))


def test_an_entry_nothing_cites_is_kept_and_flagged():
    """Deleting a line the student typed is not this tool's call — and an
    uncited entry is the other half of what a supervisor checks."""
    blob, _ = cite_docx(
        _doc((BODY, None),
             ("TÀI LIỆU THAM KHẢO", None),
             ("Nguyễn Văn A. (1999). Một cuốn không ai trích. Nxb Trẻ.", None)),
        add_missing=False, resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    lines = _read(blob)
    orphan = [ln for ln in lines if "Một cuốn không ai trích" in ln]
    assert orphan and "không thấy trích dẫn trong bài" in orphan[0]


# --- clickable citations ----------------------------------------------------

def _links(blob):
    """(anchor, linked text) for every internal hyperlink in the body."""
    import io
    import re
    from docx import Document
    doc = Document(io.BytesIO(blob))
    out = []
    for p in doc.paragraphs:
        xml = p._element.xml
        for m in re.finditer(r'<w:hyperlink[^>]*w:anchor="([^"]+)"[^>]*>(.*?)</w:hyperlink>',
                             xml, re.S):
            out.append((m.group(1), "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>",
                                                       m.group(2), re.S))))
    return out


def _bookmarks(blob):
    import io
    import re
    from docx import Document
    doc = Document(io.BytesIO(blob))
    return set(re.findall(r'w:name="(dtref\d+)"', doc.element.body.xml))


def test_a_citation_links_to_its_own_reference_entry():
    blob, report = cite_docx(_doc((BODY, None)), add_missing=False,
                             resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    links = _links(blob)
    assert report["linked"] == 1
    assert links[0][1] == "(Hair và cộng sự, 2010)"
    # The anchor exists, and it is on the reference paragraph.
    assert links[0][0] in _bookmarks(blob)


def test_two_sources_in_one_bracket_link_to_different_entries():
    text = ("Kết quả tương tự đã được nhiều nghiên cứu trước đây báo cáo trong "
            "cùng bối cảnh (Nunnally, 1978; Peterson, 1994).")

    def resolve(cit, line=None):
        return {**CROSSREF_HIT, "author": [{"family": cit.authors, "given": "A"}],
                "issued": {"date-parts": [[int(cit.year)]]},
                "DOI": f"10.1/{cit.year}"}

    blob, report = cite_docx(_doc((text, None)), add_missing=False, resolve_fn=resolve)
    links = _links(blob)
    assert report["linked"] == 2
    assert {t for _, t in links} == {"Nunnally, 1978", "Peterson, 1994"}
    assert len({a for a, _ in links}) == 2  # two entries, not one


def test_an_unresolved_citation_still_links_to_its_marked_entry():
    """Clicking it should land on "[chưa đối chiếu được]" — that is the student
    seeing exactly which of their sources did not check out."""
    blob, report = cite_docx(_doc((BODY, None)), add_missing=False,
                             resolve_fn=lambda cit, line=None: None)
    assert report["unresolved"] == 1
    assert report["linked"] == 1
    assert any("chưa đối chiếu được" in ln for ln in _read(blob))


def test_the_body_text_is_unchanged_by_linking():
    """A link must be formatting, never an edit. If the visible text moves, the
    tool has rewritten a sentence the student did not ask it to touch."""
    blob, _ = cite_docx(_doc((BODY, None)), add_missing=False,
                        resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    assert _read(blob)[0] == BODY


def test_existing_bookmarks_are_not_overwritten():
    """A thesis arrives full of them — Word writes one per table-of-contents
    entry — and reusing an id silently breaks that TOC."""
    import io

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = Document()
    p = d.add_paragraph(BODY)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "42")
    start.set(qn("w:name"), "_Toc12345")
    p._element.insert(0, start)
    buf = io.BytesIO()
    d.save(buf)

    blob, _ = cite_docx(buf.getvalue(), add_missing=False,
                        resolve_fn=lambda cit, line=None: CROSSREF_HIT)
    ids = re.findall(r'w:bookmarkStart w:id="(\d+)"',
                     Document(io.BytesIO(blob)).element.body.xml)
    assert len(ids) == len(set(ids))
    assert "42" in ids
