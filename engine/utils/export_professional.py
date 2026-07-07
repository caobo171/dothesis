#!/usr/bin/env python3
"""
ABOUTME: Production-grade academic document export utility with strategy pattern
ABOUTME: Supports multiple PDF engines (LibreOffice, Pandoc, WeasyPrint) with auto-fallback
"""

import sys
import argparse
from pathlib import Path
from typing import Optional, Literal

# Use centralized logging system
from utils.logging_config import get_logger
logger = get_logger(__name__)

# Add project root to path for subprocess/different CWD execution
project_root = Path(__file__).resolve().parent.parent
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

from utils.pdf_engines import (
    PDFGenerationOptions,
    PDFEngineFactory,
    get_available_engines,
    get_recommended_engine
)


def extract_metadata_from_yaml(md_file: Path) -> dict:
    """
    Extract metadata from YAML frontmatter in markdown file.

    Normalizes German/Spanish/French field names to English.

    Args:
        md_file: Path to markdown file with YAML frontmatter

    Returns:
        dict: Normalized metadata (title, author, date, institution, department, degree)
    """
    import yaml

    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract YAML frontmatter
        if not content.strip().startswith('---'):
            return {}

        parts = content.split('---', 2)
        if len(parts) < 3:
            return {}

        yaml_content = parts[1]
        metadata = yaml.safe_load(yaml_content) or {}

        # Normalize localized field names to English
        field_map = {
            # German
            'titel': 'title',
            'untertitel': 'subtitle',
            'autor': 'author',
            'datum': 'date',
            # Spanish
            'título': 'title',
            'subtítulo': 'subtitle',
            'autor': 'author',
            'fecha': 'date',
            # French
            'titre': 'title',
            'sous-titre': 'subtitle',
            'auteur': 'author',
            'date': 'date',
            # English (pass-through)
            'title': 'title',
            'subtitle': 'subtitle',
            'author': 'author',
            'date': 'date',
            'institution': 'institution',
            'department': 'department',
            'degree': 'degree',
            'advisor': 'advisor',
            'student_id': 'student_id',
            'project_type': 'project_type',
            'system_credit': 'system_credit',
            'generated_by': 'system_credit',  # Map generated_by to system_credit
        }

        normalized = {}
        for key, value in metadata.items():
            normalized_key = field_map.get(key.lower(), key.lower())
            normalized[normalized_key] = value

        return normalized

    except Exception as e:
        logger.warning(f"Could not extract YAML metadata: {e}")
        return {}


def export_pdf(
    md_file: Path,
    output_pdf: Path,
    engine: Literal['auto', 'libreoffice', 'pandoc', 'weasyprint'] = 'auto',
    options: Optional[PDFGenerationOptions] = None
) -> bool:
    """
    Export markdown to PDF using specified engine.

    Uses strategy pattern with automatic fallback to other engines on failure.
    This follows SOLID principles and provides production-grade error handling.

    Args:
        md_file: Path to input markdown file
        output_pdf: Path for output PDF file
        engine: PDF engine to use ('auto' for automatic selection)
        options: PDF generation options (uses academic defaults if None)

    Returns:
        bool: True if PDF generation succeeded, False otherwise

    Examples:
        >>> export_pdf(Path('draft.md'), Path('draft.pdf'))
        >>> export_pdf(Path('draft.md'), Path('draft.pdf'), engine='libreoffice')
        >>> options = PDFGenerationOptions(line_spacing=1.5)
        >>> export_pdf(Path('draft.md'), Path('draft.pdf'), options=options)
    """
    md_file = Path(md_file)
    output_pdf = Path(output_pdf)

    # Extract metadata from YAML frontmatter
    metadata = extract_metadata_from_yaml(md_file)

    # Create options with metadata if not provided
    if options is None:
        options = PDFGenerationOptions(
            title=metadata.get('title'),
            subtitle=metadata.get('subtitle'),
            author=metadata.get('author'),
            date=metadata.get('date'),
            institution=metadata.get('institution'),
            department=metadata.get('department'),
            course=metadata.get('degree'),  # Map degree to course field
            instructor=metadata.get('advisor'),  # Map advisor to instructor field
            student_id=metadata.get('student_id'),
            project_type=metadata.get('project_type'),
            system_credit=metadata.get('system_credit')
        )

    logger.info("="*70)
    logger.info(f"Generating PDF: {output_pdf.name}")
    logger.info(f"Input: {md_file}")
    logger.info(f"Engine: {engine}")
    logger.info("="*70)

    # Use factory to generate with automatic fallback
    result = PDFEngineFactory.generate_with_fallback(
        md_file=md_file,
        output_pdf=output_pdf,
        options=options,
        preferred_engine=engine if engine != 'auto' else None
    )

    # Display result
    if result.success:
        logger.info(f"PDF created successfully: {result.output_path}")
        logger.info(f"Engine used: {result.engine_name}")

        if result.warnings:
            logger.warning("Warnings during PDF generation:")
            for warning in result.warnings:
                logger.warning(f"  - {warning}")

        return True
    else:
        logger.error("PDF generation failed")
        logger.error(f"Error: {result.error_message}")
        return False


def _normalize_yaml_for_pandoc(md_content: str) -> str:
    """
    Normalize YAML frontmatter field names to English for Pandoc compatibility.

    Pandoc only recognizes English field names (title, subtitle, author, date)
    for creating styled paragraphs in DOCX/PDF output. This function translates
    localized field names back to English while preserving the field values.

    Supported translations:
    - German: titel→title, untertitel→subtitle, autor→author, datum→date
    - Spanish: título→title, subtítulo→subtitle, autor→author, fecha→date
    - French: titre→title, sous-titre→subtitle, auteur→author, date→date

    Args:
        md_content: Markdown content with YAML frontmatter

    Returns:
        Markdown content with normalized YAML field names
    """
    import re

    # Translation map: localized → English
    field_translations = {
        # German
        'titel:': 'title:',
        'untertitel:': 'subtitle:',
        'autor:': 'author:',
        'datum:': 'date:',
        'wortzahl:': 'word_count:',
        'seitenzahl:': 'page_count:',
        'sprache:': 'language:',
        'thema:': 'topic:',
        'schlagwörter:': 'keywords:',
        'qualitäts_bewertung:': 'quality_score:',
        'system_ersteller:': 'system_creator:',
        'zitate_verifiziert:': 'citations_verified:',
        'visuelle_elemente:': 'visual_elements:',
        'generierungs_methode:': 'generation_method:',
        'beschreibung_showcase:': 'showcase_description:',
        'system_fähigkeiten:': 'system_capabilities:',
        'aufruf_zur_aktion:': 'call_to_action:',
        'lizenz:': 'license:',

        # Spanish
        'título:': 'title:',
        'subtítulo:': 'subtitle:',
        'fecha:': 'date:',
        'recuento_de_palabras:': 'word_count:',
        'idioma:': 'language:',

        # French
        'titre:': 'title:',
        'sous-titre:': 'subtitle:',
        'auteur:': 'author:',
        'nombre_de_mots:': 'word_count:',
        'langue:': 'language:',
    }

    # Only process if YAML frontmatter exists
    if not md_content.strip().startswith('---'):
        return md_content

    # Extract YAML frontmatter
    parts = md_content.split('---', 2)
    if len(parts) < 3:
        return md_content

    yaml_content = parts[1]
    rest_content = parts[2]

    # Translate field names (case-insensitive)
    for localized, english in field_translations.items():
        yaml_content = re.sub(
            f'^{re.escape(localized)}',
            english,
            yaml_content,
            flags=re.MULTILINE | re.IGNORECASE
        )

    # Reconstruct markdown
    return f'---{yaml_content}---{rest_content}'


def _is_table_separator(line: str) -> bool:
    """A markdown table separator looks like `|---|---|` (with optional :---, ---:, :--:)."""
    s = line.strip()
    if not s.startswith("|") or not s.endswith("|"):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    if not cells:
        return False
    for cell in cells:
        if not cell:
            return False
        # Allow leading/trailing colons for alignment
        if not all(ch in "-:" for ch in cell):
            return False
        if "-" not in cell:
            return False
    return True


def _split_table_row(line: str) -> list[str]:
    """Split `| a | b | c |` into ['a', 'b', 'c'], stripping whitespace."""
    inner = line.strip().strip("|")
    return [cell.strip() for cell in inner.split("|")]


def _add_docx_table(doc, rows: list[list[str]]) -> None:
    """Insert a python-docx Table from a list of rows (header first)."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell_text = row[c_idx] if c_idx < len(row) else ""
            # Clear the default paragraph then add a properly-formatted one.
            cell.text = ""
            para = cell.paragraphs[0]
            _emit_inline_runs(para, cell_text, base_bold=(r_idx == 0))


# Inline markdown patterns. Order matters — longer / more specific patterns first.
# Each tuple is (pattern, kind) where kind tells _emit_inline_runs how to style the captured text.
import re as _re

_INLINE_PATTERNS = [
    # Bold-italic ***x*** or ___x___
    (_re.compile(r"\*\*\*(.+?)\*\*\*"), "bold_italic"),
    (_re.compile(r"___(.+?)___"),       "bold_italic"),
    # Bold **x** or __x__
    (_re.compile(r"\*\*(.+?)\*\*"),     "bold"),
    (_re.compile(r"__(.+?)__"),         "bold"),
    # Inline code `x`
    (_re.compile(r"`([^`]+?)`"),        "code"),
    # Link [text](url) — kind tracks the URL via the regex itself
    (_re.compile(r"\[([^\]]+?)\]\(([^)]+?)\)"), "link"),
    # Italic *x* or _x_  (after bold so ** isn't matched as italic-italic)
    (_re.compile(r"(?<![*\w])\*([^*\n]+?)\*(?!\*)"), "italic"),
    (_re.compile(r"(?<![_\w])_([^_\n]+?)_(?!_)"),   "italic"),
]


# Parenthetical citation like (Smith, 2025), (Smith et al., 2025), (Smith & Jones, 2024)
# Vietnamese diacritics included in the surname character class. Lookahead allows
# trailing page refs: (Smith, 2025, p. 14).
_CITATION_PATTERN = _re.compile(
    r"\(([^()\n]{2,150})\)"
)


def _split_citation_payload(payload: str) -> list[tuple[str, str]]:
    """For 'Smith, 2025; Jones et al., 2024' → [(surname_year_key, full_text)]."""
    out = []
    parts = [p.strip() for p in payload.split(";") if p.strip()]
    for part in parts:
        # First word is the surname (greedy). The rest can be any author-list
        # variation; we only require it ends with `, YYYY`.
        m = _re.match(
            r"^([\wÀ-ɏḀ-ỿ][\wÀ-ɏḀ-ỿ\.\-']*)"   # surname
            r"[^()0-9]*?"                       # author list, et al., &, and, initials
            r",\s*(\d{4})[a-z]?",               # year (optionally suffixed: 2025a)
            part,
        )
        if not m:
            continue
        surname = m.group(1).strip()
        year = m.group(2)
        if not surname:
            continue
        key = f"{surname.lower()}_{year}"
        out.append((key, part))
    return out


def _emit_inline_runs(
    para,
    text: str,
    *,
    base_bold: bool = False,
    citation_map: dict | None = None,
) -> None:
    """Walk a single line of markdown text and append python-docx Runs with formatting.

    If citation_map is provided, `(Author, year)` parenthetical references are
    rendered as hyperlinks to bookmarks placed at the matching bibliography entries.
    """
    if not text:
        return
    spans = []
    for pattern, kind in _INLINE_PATTERNS:
        for m in pattern.finditer(text):
            spans.append((m.start(), m.end(), kind, m))
    # Citation spans (only when a non-empty map is provided)
    if citation_map:
        for m in _CITATION_PATTERN.finditer(text):
            payload = m.group(1)
            entries = _split_citation_payload(payload)
            if any(k in citation_map for k, _ in entries):
                spans.append((m.start(), m.end(), "citation", m))

    spans.sort(key=lambda s: (s[0], -(s[1] - s[0])))
    chosen = []
    last_end = -1
    for s in spans:
        if s[0] >= last_end:
            chosen.append(s)
            last_end = s[1]

    cursor = 0
    for start, end, kind, m in chosen:
        if start > cursor:
            _add_plain_run(para, text[cursor:start], base_bold=base_bold)
        if kind == "link":
            label, url = m.group(1), m.group(2)
            run = para.add_run(label)
            run.font.color.rgb = _docx_color("#1c2eff")
            run.underline = True
            if base_bold:
                run.bold = True
        elif kind == "code":
            run = para.add_run(m.group(1))
            run.font.name = "Consolas"
            run.font.size = _docx_pt(10)
            if base_bold:
                run.bold = True
        elif kind == "bold_italic":
            run = para.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif kind == "bold":
            run = para.add_run(m.group(1))
            run.bold = True
        elif kind == "italic":
            run = para.add_run(m.group(1))
            run.italic = True
            if base_bold:
                run.bold = True
        elif kind == "citation":
            # Render the parens around a single hyperlink targeting the FIRST
            # matched bibliography entry inside the citation payload. This is a
            # pragmatic choice for `(A, 2020; B, 2021)` — only one link is
            # supported per cite group, jumping to whichever entry comes first.
            payload = m.group(1)
            entries = _split_citation_payload(payload)
            target_anchor = None
            for key, _part in entries:
                if key in citation_map:
                    target_anchor = citation_map[key]
                    break
            if target_anchor:
                _add_plain_run(para, "(", base_bold=base_bold)
                _emit_hyperlink_run(para, payload, target_anchor, bold=base_bold)
                _add_plain_run(para, ")", base_bold=base_bold)
            else:
                _add_plain_run(para, text[start:end], base_bold=base_bold)
        cursor = end
    if cursor < len(text):
        _add_plain_run(para, text[cursor:], base_bold=base_bold)


def _emit_hyperlink_run(para, text: str, anchor: str, *, bold: bool = False) -> None:
    """Append a hyperlink-styled run inside a paragraph, anchored to an in-doc bookmark."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1c2eff")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    if bold:
        b = OxmlElement("w:b")
        rpr.append(b)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)


def _add_plain_run(para, text: str, *, base_bold: bool = False) -> None:
    if not text:
        return
    run = para.add_run(text)
    if base_bold:
        run.bold = True


def _docx_color(hex_str: str):
    from docx.shared import RGBColor
    hex_str = hex_str.lstrip("#")
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def _docx_pt(n):
    from docx.shared import Pt
    return Pt(n)


def _add_bookmark(para, name: str, bid: int) -> None:
    """Insert a Word bookmark at the start of a paragraph so hyperlinks can jump to it."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    # bookmarkStart goes before the paragraph's content; bookmarkEnd at the end.
    para._p.insert(0, start)
    para._p.append(end)


_REFERENCE_HEADING_RE = _re.compile(
    r"^#{1,3}\s+("
    r"references|bibliography|works\s+cited|"
    r"tài\s+liệu\s+tham\s+khảo|"
    r"literaturverzeichnis|literatur|"
    r"bibliographie|"
    r"referencias|"
    r"参考文献"
    r")\s*$",
    _re.IGNORECASE,
)

_REF_ENTRY_KEY_RE = _re.compile(
    r"^\s*[-*]?\s*"
    r"([\wÀ-ɏḀ-ỿ][\wÀ-ɏḀ-ỿ\.\-']*)"  # surname (first whole word; greedy)
    r"[^()]*?"                          # author list, initials, et al., …
    r"\(?(\d{4})\)?"                    # year
)


def _build_citation_map(lines: list[str]) -> tuple[dict[str, str], dict[int, str]]:
    """Pre-scan the markdown for the references section.

    Returns (citation_map, line_to_anchor):
      citation_map  : "surname_year" -> "_Ref_NNNN"  (for inline lookups)
      line_to_anchor: line index of a reference entry -> "_Ref_NNNN" (so the body
                      walker can drop a bookmark when it emits that line)
    """
    citation_map: dict[str, str] = {}
    line_to_anchor: dict[int, str] = {}
    in_refs = False
    counter = 0
    in_fence = False
    for idx, raw in enumerate(lines):
        line = raw.rstrip()
        if line.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        if _REFERENCE_HEADING_RE.match(line.strip()):
            in_refs = True
            continue
        if not in_refs:
            continue
        # A new top-level heading (#) ends the references section
        if line.startswith("# ") or line.startswith("## "):
            in_refs = False
            continue
        if not line.strip():
            continue
        m = _REF_ENTRY_KEY_RE.match(line)
        if not m:
            continue
        surname = _re.split(r"[\s,]+", m.group(1).strip())[0].strip()
        year = m.group(2)
        if not surname:
            continue
        key = f"{surname.lower()}_{year}"
        # Don't overwrite — first entry with this key wins.
        if key in citation_map:
            continue
        counter += 1
        anchor = f"_Ref_{counter:04d}"
        citation_map[key] = anchor
        line_to_anchor[idx] = anchor
    return citation_map, line_to_anchor


def _extract_toc_entries(lines: list[str]) -> list[dict]:
    """Scan the (frontmatter-stripped) markdown for headings.

    Skips headings inside fenced code blocks. Returns a list of
    {"level": 1..3, "text": str, "anchor": "_Toc_NNN"} so the body
    walker and the TOC renderer can share the same bookmark names.
    """
    entries: list[dict] = []
    in_fence = False
    counter = 0
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        m = _re.match(r"^(#{1,3})\s+(.*)$", line)
        if not m:
            continue
        level = len(m.group(1))
        text = _strip_inline_markers(m.group(2).strip())
        if not text:
            continue
        counter += 1
        entries.append({"level": level, "text": text, "anchor": f"_Toc_{counter:04d}"})
    return entries


def _strip_inline_markers(text: str) -> str:
    """Strip the markers used by the inline parser so TOC lines stay clean."""
    out = text
    out = _re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", out)
    out = _re.sub(r"___(.+?)___", r"\1", out)
    out = _re.sub(r"\*\*(.+?)\*\*", r"\1", out)
    out = _re.sub(r"__(.+?)__", r"\1", out)
    out = _re.sub(r"`([^`]+?)`", r"\1", out)
    out = _re.sub(r"\[([^\]]+?)\]\(([^)]+?)\)", r"\1", out)
    out = _re.sub(r"(?<!\w)\*([^*\n]+?)\*(?!\*)", r"\1", out)
    out = _re.sub(r"(?<!\w)_([^_\n]+?)_(?!_)", r"\1", out)
    return out.strip()


def _render_static_toc(doc, entries: list[dict], *, body_font: str) -> None:
    """Render a TOC list with clickable hyperlinks to in-document bookmarks.

    The engine already numbers its own headings (1, 1.1, 1.1.1 …) so we do
    NOT prepend our own — that would cause "1.1  1.1 Title" double indicators.
    Each entry is wrapped in `<w:hyperlink w:anchor="...">` so Ctrl/Cmd-click
    jumps to the bookmark we emit alongside the matching heading.
    """
    from docx.shared import Pt, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    title = doc.add_heading("Table of Contents", level=1)
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    for entry in entries:
        level = entry["level"]
        text = entry["text"]
        anchor = entry["anchor"]

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25 * (level - 1))
        para.paragraph_format.space_after = Pt(4)

        # Build a hyperlink element pointing at the bookmark.
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        hyperlink.set(qn("w:history"), "1")

        run = OxmlElement("w:r")

        rpr = OxmlElement("w:rPr")
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "Hyperlink")
        rpr.append(rstyle)
        # Inline font for fallback when Hyperlink style is absent
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), body_font)
        rfonts.set(qn("w:hAnsi"), body_font)
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24" if level == 1 else "22")
        rpr.append(sz)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1c2eff")
        rpr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
        if level == 1:
            b = OxmlElement("w:b")
            rpr.append(b)
        run.append(rpr)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        run.append(t)
        hyperlink.append(run)
        para._p.append(hyperlink)


def _insert_word_toc_field(doc) -> None:
    """Insert a Word TOC field that auto-populates with page numbers.

    Users see the static list above; if they right-click → Update Field
    Word will replace this with a real linked TOC with page numbers.
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_before = _docx_pt(8)
    run = para.add_run()
    # Hint label
    hint = doc.add_paragraph()
    hint_run = hint.add_run(
        "(For page numbers, right-click the TOC below and choose Update Field in Microsoft Word.)"
    )
    hint_run.italic = True
    hint_run.font.size = _docx_pt(10)

    fld_para = doc.add_paragraph()
    fld_run = fld_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "[Right-click → Update Field to populate]"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_run._r.append(fld_begin)
    fld_run._r.append(instr)
    fld_run._r.append(fld_sep)
    fld_run._r.append(placeholder)
    fld_run._r.append(fld_end)


def _strip_yaml_frontmatter(md: str) -> str:
    """Drop a leading `---\\n...\\n---\\n` block if present."""
    if not md.startswith("---"):
        return md
    parts = md.split("\n", 1)
    if len(parts) < 2:
        return md
    end = md.find("\n---", 3)
    if end < 0:
        return md
    after = md[end + 4:]
    return after.lstrip("\n")


def export_docx_basic(md_file: Path, output_docx: Path) -> bool:
    """
    Export markdown to DOCX format using basic python-docx parsing.

    Note: This is a legacy/fallback function with limited formatting support.
    Does NOT support tables, complex formatting, or citations.

    For production use, prefer export_docx() which uses Pandoc with full
    markdown support and proper table formatting.

    Args:
        md_file: Path to input markdown file
        output_docx: Path for output DOCX file

    Returns:
        bool: True if DOCX generation succeeded, False otherwise
    """
    md_file = Path(md_file)
    output_docx = Path(output_docx)

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed")
        logger.error("Run: pip install python-docx")
        return False

    try:
        # Read markdown
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Drop the YAML frontmatter outright — basic exporter doesn't render a
        # cover page from it, so leaving it in produces a wall of key:value lines.
        md_content = _strip_yaml_frontmatter(md_content)
        lines = md_content.splitlines()

        # Create document with 1" margins
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Centered footer with a REAL auto-updating Word PAGE field, so the
        # document shows actual page numbers out of the box (no "Insert > Page
        # Number" step). Falls back to a static "1" if the OXML API is missing.
        footer_para = doc.sections[0].footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = "PAGE"
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            fld_num = OxmlElement("w:t")
            fld_num.text = "1"
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run._r.append(fld_begin)
            run._r.append(instr)
            run._r.append(fld_sep)
            run._r.append(fld_num)
            run._r.append(fld_end)
        except Exception:
            run.text = "1"

        BODY_FONT = "Times New Roman"
        BODY_SIZE = Pt(12)

        # --- Build Table of Contents from headings before walking the body ---
        toc_entries = _extract_toc_entries(lines)
        if toc_entries:
            _render_static_toc(doc, toc_entries, body_font=BODY_FONT)
            # Word's auto-updating TOC field. Hidden behind the static list so users
            # who Update Field get the real page numbers; everyone else sees the
            # static one we just wrote.
            _insert_word_toc_field(doc)
            doc.add_page_break()

        # Queue of anchors to consume when the body walker emits each heading.
        heading_anchors = [e["anchor"] for e in toc_entries]
        bookmark_id_counter = [0]  # mutable so the nested helper can bump it

        # Citation linkage: bibliography pre-scan
        citation_map, line_to_anchor = _build_citation_map(lines)

        def inline(p, txt, *, base_bold=False):
            _emit_inline_runs(p, txt, base_bold=base_bold, citation_map=citation_map)

        def style_body(para):
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
            for r in para.runs:
                if not r.font.name:
                    r.font.name = BODY_FONT
                if not r.font.size:
                    r.font.size = BODY_SIZE

        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # Blank
            if not line.strip():
                i += 1
                continue

            # Fenced code block ```lang ... ```
            if line.startswith("```"):
                lang = line[3:].strip()
                code_buf = []
                i += 1
                while i < len(lines) and not lines[i].rstrip().startswith("```"):
                    code_buf.append(lines[i].rstrip("\n"))
                    i += 1
                if i < len(lines):
                    i += 1  # skip closing ```
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                # Light gray shade via paragraph; python-docx doesn't easily set bg, so we just
                # indent and use the monospace font.
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.space_after = Pt(8)
                continue

            # Pipe table: header + separator + body
            if line.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
                table_rows = [_split_table_row(line)]
                i += 2
                while i < len(lines) and lines[i].rstrip().startswith("|"):
                    table_rows.append(_split_table_row(lines[i]))
                    i += 1
                _add_docx_table(doc, table_rows)
                continue

            # ATX headings #, ##, ###, ####, ##### (centered for top level only)
            if line.startswith("#"):
                m = _re.match(r"^(#{1,6})\s+(.*)$", line)
                if m:
                    level = len(m.group(1))
                    heading_text = m.group(2).strip()
                    para = doc.add_heading(level=min(level, 4))
                    # Anchor for TOC hyperlink: only levels 1-3 are in the TOC.
                    if 1 <= level <= 3 and heading_anchors:
                        anchor_name = heading_anchors.pop(0)
                        bookmark_id_counter[0] += 1
                        _add_bookmark(para, anchor_name, bookmark_id_counter[0])
                    inline(para, heading_text)
                    if level == 1:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    i += 1
                    continue

            # Horizontal rule (--- on its own line, not a table separator)
            if line.strip() in {"---", "***", "___"}:
                para = doc.add_paragraph()
                run = para.add_run("_" * 60)
                run.font.size = Pt(12)
                i += 1
                continue

            # Blockquote (one or more consecutive > lines)
            if line.startswith(">"):
                quote_lines = []
                while i < len(lines) and lines[i].lstrip().startswith(">"):
                    quote_lines.append(_re.sub(r"^>\s?", "", lines[i].rstrip()))
                    i += 1
                joined = " ".join(quote_lines).strip()
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.4)
                para.paragraph_format.right_indent = Inches(0.4)
                inline(para, joined)
                for r in para.runs:
                    r.italic = True
                    if not r.font.name:
                        r.font.name = BODY_FONT
                    if not r.font.size:
                        r.font.size = BODY_SIZE
                continue

            # Unordered list (- or * item)
            if _re.match(r"^[\-\*\+]\s+", line):
                while i < len(lines) and _re.match(r"^[\-\*\+]\s+", lines[i].rstrip()):
                    item_text = _re.sub(r"^[\-\*\+]\s+", "", lines[i].rstrip())
                    para = doc.add_paragraph(style="List Bullet")
                    # Reference list entries get a bookmark so inline citations can jump here.
                    if i in line_to_anchor:
                        bookmark_id_counter[0] += 1
                        _add_bookmark(para, line_to_anchor[i], bookmark_id_counter[0])
                    inline(para, item_text)
                    style_body(para)
                    i += 1
                continue

            # Ordered list (1. item, 2. item, ...)
            if _re.match(r"^\d+\.\s+", line):
                while i < len(lines) and _re.match(r"^\d+\.\s+", lines[i].rstrip()):
                    item_text = _re.sub(r"^\d+\.\s+", "", lines[i].rstrip())
                    para = doc.add_paragraph(style="List Number")
                    if i in line_to_anchor:
                        bookmark_id_counter[0] += 1
                        _add_bookmark(para, line_to_anchor[i], bookmark_id_counter[0])
                    inline(para, item_text)
                    style_body(para)
                    i += 1
                continue

            # Regular paragraph — collapse soft-wrapped lines until blank or block break
            start_line = i
            para_buf = [line]
            i += 1
            while i < len(lines):
                nxt = lines[i].rstrip()
                if not nxt.strip():
                    break
                if nxt.startswith(("#", ">", "|", "```")):
                    break
                if _re.match(r"^[\-\*\+]\s+", nxt) or _re.match(r"^\d+\.\s+", nxt):
                    break
                if nxt in {"---", "***", "___"}:
                    break
                para_buf.append(nxt)
                i += 1
            paragraph_text = " ".join(para_buf)
            para = doc.add_paragraph()
            # If this paragraph is a non-list-style reference entry, anchor it.
            if start_line in line_to_anchor:
                bookmark_id_counter[0] += 1
                _add_bookmark(para, line_to_anchor[start_line], bookmark_id_counter[0])
            inline(para, paragraph_text)
            style_body(para)

        doc.save(output_docx)
        logger.info(f"DOCX created successfully: {output_docx}")
        return True

    except Exception as e:
        logger.error(f"DOCX generation failed: {str(e)}", exc_info=True)
        return False


def _style_docx_tables(output_docx) -> None:
    """Give pandoc's tables a clean academic look.

    Pandoc + the reference doc render tables BORDERLESS with:
      - `w:tblLayout w:type="fixed"` (Word ignores autofit),
      - short `w:trHeight` on every row (numbers like "212" wrap to "21" / "2"),
      - guessed narrow column widths (headers "Tần số (n)" break to 3 lines).
    So we can't just flip `table.autofit = True` — Word ignores it while the
    fixed layout is on. For every table: apply a bordered grid style, FORCE
    `tblLayout=autofit`, allocate PROPORTIONAL widths (wide content columns,
    narrower number columns), clear the fixed cell widths + row heights, add
    consistent cell margins, vertically center every cell, and bold the header
    row. Best-effort — any failure leaves the (working) docx untouched.
    """
    try:
        import re as _re
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.shared import Pt
    except Exception:
        return
    try:
        doc = Document(str(output_docx))
    except Exception:
        logger.warning("table-styling: could not open docx", exc_info=True)
        return
    if not doc.tables:
        return

    def _set(elem, tag_name, **attrs):
        """Get-or-create a child element with the given w: attributes."""
        child = elem.find(qn(tag_name))
        if child is None:
            child = OxmlElement(tag_name)
            elem.append(child)
        for k, v in attrs.items():
            child.set(qn(k), str(v))
        return child

    for table in doc.tables:
        # 1) Bordered style. Try a few names since availability depends on the
        # reference doc; if none work, borders get set explicitly below.
        for style_name in ("Table Grid", "TableGrid", "Light Grid Accent 1"):
            try:
                table.style = style_name
                break
            except Exception:
                continue

        tblPr = table._tbl.tblPr
        # 2) Force AUTOFIT layout — Word obeys column widths as HINTS, not
        # constraints; long content grows the column instead of wrapping cruelly.
        _set(tblPr, "w:tblLayout", **{"w:type": "autofit"})
        # 3) Table width = 100% of the available space (5000 fiftieths).
        _set(tblPr, "w:tblW", **{"w:w": "5000", "w:type": "pct"})
        # 4) Consistent inner margins — tight enough that narrow columns
        # (VIF / t / Sig) can render their header on one line, generous enough
        # that text never touches the borders.
        cellMar = _set(tblPr, "w:tblCellMar")
        for side, val in (("top", 40), ("bottom", 40), ("left", 60), ("right", 60)):
            _set(cellMar, f"w:{side}", **{"w:w": str(val), "w:type": "dxa"})
        # 5) Explicit borders in case the built-in style didn't apply.
        borders = _set(tblPr, "w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            _set(borders, f"w:{side}", **{"w:val": "single", "w:sz": "4", "w:color": "auto"})

        # 6) Column widths sized to CONTENT, not a flat text/number split.
        # Previous heuristic gave every numeric column the same weight → in a
        # wide table like "Bảng 4.5" (8 cols: Giả thuyết / Biến / Hệ số B /
        # Beta / t / Sig / VIF / Kết luận) a header like "Giả thuyết" or "Kết
        # luận" broke into "Gi / ả thuyết" while single-letter headers "K", "IF"
        # were still cramped by margins. Fix: measure the longest word in the
        # HEADER of each column and use that as a MIN width, then split the
        # remaining space using the text-vs-numeric heuristic.
        n_cols = len(table.columns)
        header_cells = table.rows[0].cells if table.rows else []
        sample_rows = table.rows[1:4]

        def _longest_word_len(txt: str) -> int:
            # Word Vietnamese-friendly: split on whitespace only. A cell like
            # "Hệ số B (Unstandardized)" → 15 chars ("(Unstandardized)").
            return max((len(w) for w in (txt or "").split()), default=1)

        header_len: list[int] = []
        for c in range(n_cols):
            txt = header_cells[c].text.strip() if c < len(header_cells) else ""
            header_len.append(max(3, _longest_word_len(txt)))

        is_numeric_col: list[bool] = []
        for c in range(n_cols):
            numeric_hits = 0
            for r in sample_rows:
                if c < len(r.cells):
                    txt = r.cells[c].text.strip().replace(",", ".")
                    if _re.fullmatch(r"[\d.\s%\-]+", txt or ""):
                        numeric_hits += 1
            is_numeric_col.append(numeric_hits >= 1)

        # Word measures widths in dxa; a typical usable page width is ~9000.
        page_dxa = 9000
        # Min width per column ≈ header longest-word × ~120 dxa/char + margins,
        # so "Kết" (3 chars) reserves ~360 + 240 margins ≈ 600 dxa, and
        # "(Unstandardized)" (15 chars) reserves ~1800 dxa. Prevents header
        # breakage but leaves numeric single-line cells room to shrink.
        MIN_PER_CHAR = 120
        MARGINS = 240
        min_dxa = [max(500, header_len[c] * MIN_PER_CHAR + MARGINS) for c in range(n_cols)]
        # Cap the sum of minimums so we don't overshoot the page in extreme
        # cases; scale down proportionally if needed.
        min_sum = sum(min_dxa)
        if min_sum > page_dxa * 0.8:
            factor = (page_dxa * 0.8) / min_sum
            min_dxa = [max(500, int(v * factor)) for v in min_dxa]
        # Remaining space is split by weight: numeric = 1, text = 2.
        weights = [1 if is_numeric_col[c] else 2 for c in range(n_cols)]
        remaining = max(0, page_dxa - sum(min_dxa))
        total_w = sum(weights) or n_cols
        col_dxa = [min_dxa[c] + int(remaining * weights[c] / total_w) for c in range(n_cols)]
        # Replace tblGrid.
        old_grid = table._tbl.find(qn("w:tblGrid"))
        if old_grid is not None:
            table._tbl.remove(old_grid)
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(list(table._tbl).index(tblPr) + 1, grid)
        for w in col_dxa:
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)

        # 7) Per-cell: set width, vertical-center, clear fixed row heights,
        # slightly smaller header font on very narrow number columns is skipped
        # in favor of just giving numeric columns their fair share of space.
        for r_idx, row in enumerate(table.rows):
            # Clear row height rule so cells grow with content instead of
            # clipping into a two-line height.
            trPr = row._tr.get_or_add_trPr()
            trHeight = trPr.find(qn("w:trHeight"))
            if trHeight is not None:
                trPr.remove(trHeight)
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                # Fixed cell width override (in dxa).
                tcW = _set(tcPr, "w:tcW",
                           **{"w:w": str(col_dxa[c_idx] if c_idx < len(col_dxa) else 1500),
                              "w:type": "dxa"})
                # Vertical center — headers with "Tần số (n)" were bottom-aligned
                # and read as broken.
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Cell paragraphs inherit the reference doc's Normal style
                # (double line-spacing + 0.5" first-line indent) → tall, oddly
                # indented cells. Flatten them to compact single-spaced text.
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = Pt(0)
                    pf.line_spacing = 1.0
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)

        # 8) Bold the header row.
        if table.rows:
            for cell in table.rows[0].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

    try:
        doc.save(str(output_docx))
    except Exception:
        logger.warning("table-styling: could not save docx", exc_info=True)


def export_docx(
    md_file: Path,
    output_docx: Path,
    options: Optional[PDFGenerationOptions] = None,
    bibliography: Optional[Path] = None,
) -> bool:
    """
    Export markdown to DOCX with full formatting support (tables, citations, styles).

    Uses Pandoc with custom reference document for professional formatting that matches
    PDF output quality. Automatically falls back to basic export if Pandoc unavailable.

    Features:
    - ✅ Tables rendered as Word tables (not pipe text)
    - ✅ Proper heading styles (APA 7th edition compatible)
    - ✅ Citations and footnotes
    - ✅ Table of contents
    - ✅ Custom fonts and spacing from reference document

    Args:
        md_file: Path to input markdown file
        output_docx: Path for output DOCX file
        options: PDF generation options (uses academic defaults if None)

    Returns:
        bool: True if DOCX generation succeeded, False otherwise

    Examples:
        >>> export_docx(Path('draft.md'), Path('draft.docx'))
        >>> options = PDFGenerationOptions(enable_toc=True, toc_depth=3)
        >>> export_docx(Path('draft.md'), Path('draft.docx'), options)
    """
    md_file = Path(md_file)
    output_docx = Path(output_docx)

    import subprocess
    import shutil

    # Extract metadata from YAML frontmatter
    metadata = extract_metadata_from_yaml(md_file)

    # Create options with metadata if not provided
    if options is None:
        options = PDFGenerationOptions(
            title=metadata.get('title'),
            subtitle=metadata.get('subtitle'),
            author=metadata.get('author'),
            date=metadata.get('date'),
            institution=metadata.get('institution'),
            department=metadata.get('department'),
            faculty=metadata.get('faculty'),
            course=metadata.get('degree'),  # Map degree to course field
            instructor=metadata.get('advisor'),  # Map advisor to instructor field
            second_examiner=metadata.get('second_examiner'),
            student_id=metadata.get('student_id'),
            project_type=metadata.get('project_type'),
            system_credit=metadata.get('system_credit'),
            location=metadata.get('location')
        )

    # Try Pandoc method first (best quality)
    if not shutil.which('pandoc'):
        logger.warning("Pandoc not found - falling back to basic DOCX export")
        logger.warning("Tables and advanced formatting will be limited")
        logger.info("Install Pandoc for better results: sudo apt install pandoc")
        return export_docx_basic(md_file, output_docx)

    try:
        # Read and normalize YAML field names for Pandoc compatibility
        # (Pandoc only recognizes English field names like 'title', 'author', 'date')
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        md_content = _normalize_yaml_for_pandoc(md_content)

        # Write normalized content to temporary file for Pandoc
        import tempfile
        temp_md = None
        try:
            temp_fd, temp_path = tempfile.mkstemp(suffix='.md', text=True)
            temp_md = Path(temp_path)
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(md_content)
        finally:
            if temp_fd:
                import os
                os.close(temp_fd)

        # Locate reference document
        reference_doc = Path(__file__).parent.parent / "examples" / "custom-reference.docx"

        if not reference_doc.exists():
            logger.warning(f"Reference document not found: {reference_doc}")
            logger.warning("Continuing without reference document (may lose some formatting)")
            reference_doc = None

        # Build pandoc command (use normalized temp file)
        cmd = [
            'pandoc',
            str(temp_md),  # Use normalized markdown instead of original
            '-o', str(output_docx),
            '--from', 'markdown',
            '--to', 'docx',
        ]

        # Add reference document for styling
        if reference_doc:
            cmd.extend(['--reference-doc', str(reference_doc)])

        # Citation processing: when a CSL-JSON bibliography is supplied, turn on
        # pandoc citeproc so inline [@key] citations render as "(Author, Year)"
        # AND, with link-citations, become clickable hyperlinks to the
        # auto-generated References section. Harmless if there are no citations.
        if bibliography is not None and Path(bibliography).exists():
            cmd.extend([
                '--citeproc',
                '--bibliography', str(bibliography),
                '-M', 'link-citations=true',
            ])

        # Add table of contents (Pandoc generates a proper Word TOC field)
        if options.enable_toc:
            cmd.append('--toc')
            cmd.extend(['--toc-depth', str(options.toc_depth)])

        # Add metadata if provided
        if options.title:
            cmd.extend(['--metadata', f'title={options.title}'])
        if options.author:
            cmd.extend(['--metadata', f'author={options.author}'])

        logger.info("="*70)
        logger.info(f"Generating DOCX with Pandoc: {output_docx.name}")
        logger.info(f"Input: {md_file}")
        if reference_doc:
            logger.info(f"Reference doc: {reference_doc.name}")
        logger.info("="*70)

        # Run pandoc
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            logger.error(f"Pandoc failed with return code {result.returncode}")
            if result.stderr:
                logger.error(f"Error: {result.stderr}")
            return False

        logger.info(f"DOCX created successfully: {output_docx}")
        logger.info("Tables, formatting, and styling preserved from markdown")

        # Post-process DOCX to add academic structure (title page + TOC + page breaks)
        # Fixes Pandoc's inline title block by inserting professional page breaks
        from utils.docx_post_processor import insert_academic_structure

        # Build options dict from PDFGenerationOptions for cover page enhancement
        post_options = {}
        if options:
            if options.institution:
                post_options['institution'] = options.institution
            if hasattr(options, 'faculty') and options.faculty:
                post_options['faculty'] = options.faculty
            if options.department:
                post_options['department'] = options.department
            if options.course:
                post_options['course'] = options.course
            if options.instructor:
                post_options['instructor'] = options.instructor
            if hasattr(options, 'second_examiner') and options.second_examiner:
                post_options['second_examiner'] = options.second_examiner
            if options.student_id:
                post_options['student_id'] = options.student_id
            if hasattr(options, 'project_type') and options.project_type:
                post_options['project_type'] = options.project_type
            if hasattr(options, 'system_credit') and options.system_credit:
                post_options['system_credit'] = options.system_credit
            if hasattr(options, 'location') and options.location:
                post_options['location'] = options.location

        struct_ok = insert_academic_structure(output_docx, verbose=True, options=post_options if post_options else None)
        if not struct_ok:
            logger.warning("Post-processing failed - DOCX created but may lack page structure")
            logger.warning("DOCX will have inline title block instead of standalone pages")

        # Style tables LAST so it survives the structure pass. Borderless pandoc
        # tables → bordered, autofit, bold-header. Best-effort.
        _style_docx_tables(output_docx)

        return True

    except subprocess.TimeoutExpired:
        logger.error("DOCX generation timed out (>60s)")
        return False
    except Exception as e:
        logger.error(f"DOCX generation failed: {str(e)}")
        return False
    finally:
        # Clean up temporary markdown file
        if temp_md and temp_md.exists():
            temp_md.unlink()


def show_available_engines() -> None:
    """Display available PDF engines and their status."""
    logger.info("="*70)
    logger.info("Available PDF Engines")
    logger.info("="*70)

    engines = get_available_engines()
    if not engines:
        logger.warning("No PDF engines available")
        logger.info("Install at least one of:")
        logger.info("  - LibreOffice: sudo apt install libreoffice-writer libreoffice-core-nogui")
        logger.info("  - Pandoc: sudo apt install pandoc texlive-latex-base texlive-latex-recommended")
        logger.info("  - WeasyPrint: pip install weasyprint")
        return

    recommended = get_recommended_engine()

    for engine in engines:
        marker = "[Recommended]" if engine == recommended else ""
        logger.info(f"  {engine} {marker}")

    if recommended:
        logger.info(f"Recommended: {recommended}")


def main():
    """CLI entry point."""
    # Logging is automatically configured by utils.logging_config on import

    parser = argparse.ArgumentParser(
        description='Export markdown to professional academic PDF/DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  # Auto-select best engine
  python export_professional.py draft.md --pdf draft.pdf

  # Use specific engine
  python export_professional.py draft.md --pdf draft.pdf --engine libreoffice

  # Generate both PDF and DOCX
  python export_professional.py draft.md --pdf draft.pdf --docx draft.docx

  # List available engines
  python export_professional.py --list-engines

Engine Options:
  auto         - Automatically select best available engine (recommended)
  libreoffice  - Use LibreOffice (best quality, requires installation)
  pandoc       - Use Pandoc/LaTeX (highest quality, academic standard)
  weasyprint   - Use WeasyPrint (fallback, known font rendering issues)
        '''
    )
    parser.add_argument('input', nargs='?', help='Input markdown file')
    parser.add_argument('--pdf', help='Output PDF file')
    parser.add_argument('--docx', help='Output DOCX file')
    parser.add_argument(
        '--engine',
        choices=['auto', 'libreoffice', 'pandoc', 'weasyprint'],
        default='auto',
        help='PDF engine to use (default: auto)'
    )
    parser.add_argument(
        '--list-engines',
        action='store_true',
        help='List available PDF engines and exit'
    )

    args = parser.parse_args()

    # List engines mode
    if args.list_engines:
        show_available_engines()
        sys.exit(0)

    # Validate arguments
    if not args.input:
        parser.error("Input file required (or use --list-engines)")

    input_file = Path(args.input)

    if not input_file.exists():
        logger.error(f"Input file not found: {input_file}")
        sys.exit(1)

    if not args.pdf and not args.docx:
        parser.error("Specify --pdf and/or --docx output file")

    # Generate outputs
    success = True

    if args.pdf:
        pdf_success = export_pdf(input_file, Path(args.pdf), engine=args.engine)
        success = success and pdf_success

    if args.docx:
        docx_success = export_docx(input_file, Path(args.docx))
        success = success and docx_success

    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
