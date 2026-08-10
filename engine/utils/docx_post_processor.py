#!/usr/bin/env python3
"""
ABOUTME: Post-processor for DOCX files to add academic structure (title page, TOC)
ABOUTME: Fixes Pandoc's inline title block by inserting professional page breaks

Production-grade DOCX post-processing following SOLID principles.
"""

from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH


def insert_academic_structure(
    docx_path: Path,
    verbose: bool = False,
    options: Optional[Dict[str, Any]] = None
) -> bool:
    """
    Insert academic paper structure into Pandoc-generated DOCX.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    try:
        if verbose:
            print(f"📄 Post-processing DOCX: {docx_path.name}")

        doc = Document(docx_path)

        # Step 1: Find title block elements
        title_idx, date_idx = _find_title_block(doc)

        if title_idx is None:
            if verbose:
                print("   ⚠️  No title block found - skipping post-processing")
            return True

        if verbose:
            print(f"   ✓ Found title block (Title at {title_idx}, Date at {date_idx})")

        # Step 2: Insert institution info BEFORE title
        # ANY of the three, not institution alone: a project that knows only
        # its department (derived from m1.field, which every project has) had
        # the whole block dropped and rendered a cover with nothing above the
        # title.
        if options and any(options.get(k) for k in ('institution', 'faculty', 'department')):
            _insert_institution_block(doc, title_idx, options, verbose)
            # Recalculate positions after insertion
            title_idx, date_idx = _find_title_block(doc)

        # Step 2b: Center the title block (Title, Subtitle, Author, Date)
        _center_title_block(doc, title_idx, date_idx)

        # Step 3: Insert additional metadata AFTER date (supervisor, etc.)
        if options:
            _insert_metadata_after_date(doc, date_idx, options, verbose)
            # Recalculate date position
            _, date_idx = _find_title_block(doc)

        # Step 3b: make the cover look like a cover
        _polish_cover(doc, options or {}, verbose)

        # Step 4: Find end of cover page and insert page break
        cover_end_idx = _find_cover_end(doc)
        if cover_end_idx is not None:
            _insert_page_break_after(doc, cover_end_idx)
            if verbose:
                print(f"   ✓ Inserted page break after cover page")

        # Step 5: Pandoc generates TOC with --toc flag, so we don't insert manual TOC
        # Just need to insert page break after Abstract (before first chapter)
        abstract_end_idx = _find_abstract_end(doc)
        if abstract_end_idx is not None:
            _insert_page_break_after(doc, abstract_end_idx)
            if verbose:
                print(f"   ✓ Inserted page break after Abstract")

        # Step 6: Fix table widths to fit page
        _fix_table_widths(doc, verbose)

        doc.save(docx_path)

        if verbose:
            print(f"   ✅ Post-processing complete!")

        return True

    except Exception as e:
        if verbose:
            print(f"   ❌ Post-processing failed: {e}")
            import traceback
            traceback.print_exc()
        return False


def _find_title_block(doc: Document):
    """Find Title and Date paragraph indices.

    The Date style is not reliable — _center_title_block below says so in its
    own docstring and works around it, but this returned None and
    _insert_metadata_after_date gives up on None. So on every export where
    pandoc emitted the date as Normal (which is what the reference document
    produces here), the entire lower half of the cover page — degree, project
    type, supervisor, student id — was silently dropped while the institution
    block above the title rendered fine. Same fallback as the centring pass: the
    last non-empty paragraph of the cover, i.e. before the first Heading.
    """
    title_idx = None
    date_idx = None

    for i, para in enumerate(doc.paragraphs[:20]):
        style = para.style.name if para.style else ""
        if style == 'Title' and title_idx is None:
            title_idx = i
        elif style == 'Date':
            date_idx = i
            break

    if date_idx is None and title_idx is not None:
        last = None
        for j in range(title_idx + 1, min(len(doc.paragraphs), title_idx + 10)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if style.startswith("Heading"):
                break
            if doc.paragraphs[j].text.strip():
                last = j
        # Only when something actually follows the title — inserting the degree
        # block directly under a bare title would read as part of the title.
        date_idx = last

    return title_idx, date_idx


def _center_title_block(doc: Document, title_idx, date_idx):
    """Center the title block. Ends at the Date paragraph when one is styled;
    otherwise (pandoc doesn't always style the date) center from the Title down
    to the last non-empty cover paragraph before the first Heading — so a
    title + year cover is centred instead of stranded left-aligned."""
    if title_idx is None:
        return

    end = date_idx
    if end is None:
        end = title_idx
        for j in range(title_idx + 1, min(len(doc.paragraphs), title_idx + 10)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if style.startswith("Heading"):
                break
            if doc.paragraphs[j].text.strip():
                end = j

    defined = {st.style_id for st in doc.styles}
    for i in range(title_idx, end + 1):
        if i < len(doc.paragraphs):
            _drop_dangling_style(doc.paragraphs[i], defined)
            _center_hard(doc.paragraphs[i])


def _drop_dangling_style(para, defined_ids) -> None:
    """Remove a `w:pStyle` that names a style the document does not define.

    Pandoc tags the author and the date `Author` and `Date`. The reference
    document defines neither — grep styles.xml, they are simply absent — so both
    paragraphs carried a reference to nothing. LibreOffice does not lay those
    out like the plain centred paragraphs around them: the author and the year
    sat well left of the title while every line above and below them centred,
    which is the crooked cover a student sees. Nothing is lost by dropping the
    reference; it never resolved to any formatting in the first place.
    """
    from docx.oxml.ns import qn
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return
    for st in pPr.findall(qn("w:pStyle")):
        if st.get(qn("w:val")) not in defined_ids:
            pPr.remove(st)


def _center_hard(para) -> None:
    """Centre a paragraph AND clear the indents that decide what it centres in.

    `jc=center` alone is not enough: the paragraph is centred inside whatever box
    its indents leave, so an inherited indent quietly shifts it off the page's
    centre line while the XML insists it is centred.
    """
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.left_indent = 0
    pf.right_indent = 0
    pf.first_line_indent = 0


# Blank lines above the cover block. Pandoc starts the title at the top margin,
# which on a page whose only content is a cover leaves the title jammed against
# the header and two thirds of the sheet empty underneath.
_COVER_TOP_BLANKS = 5


def _polish_cover(doc: Document, options: Dict, verbose: bool) -> None:
    """Cosmetic pass over the finished cover. Never raises — a cover that looks
    plain is a far better outcome than an export that fails."""
    try:
        title_idx, _ = _find_title_block(doc)
        if title_idx is None:
            return
        title = doc.paragraphs[title_idx]

        # The reference document's Title style carries a bottom border, so every
        # cover had a blue rule ruled across it under the title — fine for a
        # section heading in the body, wrong on a title page. Override it on
        # this paragraph rather than editing the shared style.
        _clear_bottom_border(title)
        # …and single-space it. At the document's 2.0 academic line spacing a
        # three-line Vietnamese title sprawls down the page.
        title.paragraph_format.line_spacing = 1.15
        title.paragraph_format.space_after = Pt(18)

        # "by" / "Học viên thực hiện" above the author, as on a real cover.
        # The paragraph right after the title, if it is not the date. Looking
        # for style "Author" found nothing: the style is undefined, so
        # python-docx reports it as Normal and _drop_dangling_style has by now
        # removed the reference entirely.
        author = None
        for cand in doc.paragraphs[title_idx + 1:title_idx + 3]:
            txt = cand.text.strip()
            if txt and not _looks_like_date(txt):
                author = cand
                break
        if author is not None and author.text.strip():
            lead = author.insert_paragraph_before(_words(options)["by"])
            _center_hard(lead)
            for run in lead.runs:
                run.font.size = Pt(11)
            title_idx, _ = _find_title_block(doc)

        _move_date_to_foot(doc)

        # Breathing room above everything.
        first = doc.paragraphs[0]
        for _ in range(_COVER_TOP_BLANKS):
            first.insert_paragraph_before("")
        if verbose:
            print("   ✓ Polished cover (rule removed, spacing, by-line, date at foot)")
    except Exception as e:  # noqa: BLE001 — cosmetic only
        if verbose:
            print(f"   ⚠️  Cover polish skipped: {e}")


def _move_date_to_foot(doc: Document) -> None:
    """Put the date last on the cover, under the degree and supervisor lines.

    Pandoc emits it directly beneath the author because that is where a title
    BLOCK puts it, and the degree/supervisor lines are then inserted after it —
    so the year ended up wedged between the author and "LUẬN VĂN THẠC SĨ". On
    every thesis cover, the date is the last line.
    """
    end = _find_cover_end(doc)
    if end is None:
        return
    paras = doc.paragraphs[:end + 1]
    date = next((p for p in paras if _looks_like_date(p.text)), None)
    if date is None or date is paras[-1]:
        return
    last = paras[-1]
    # lxml moves an element that is already in the tree, so this is a move.
    last._p.addnext(date._p)
    spacer = date.insert_paragraph_before("")
    _center_hard(spacer)


def _looks_like_date(text: str) -> bool:
    """A cover date is a bare year, or a year with a word in front of it."""
    import re as _re
    return bool(_re.fullmatch(r"(?:Năm\s+)?\d{4}", text.strip(), _re.IGNORECASE))


def _clear_bottom_border(para) -> None:
    """Remove any bottom border the paragraph inherits from its style."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "none")
    bottom.set(qn("w:sz"), "0")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _find_cover_end(doc: Document):
    """Index of the LAST cover (title-block) paragraph — the last non-empty para
    before the first real Heading. Break AFTER it so the cover is its own page.

    The old heuristic (`style.startswith('Heading') and i > 5`) silently skipped
    the actual first heading when the title block is short: with no Abstract the
    first Heading is "Chương 1" at index ~2, the `i > 5` guard jumped past it AND
    "1.1", and the break landed in the MIDDLE of Chapter 1 — leaving the cover
    with no page break at all (title + year + TOC + chapter all crammed onto
    page 1). Pandoc's `--toc` lives in a `w:sdt` that python-docx doesn't surface
    as paragraphs, so we can't anchor on "Mục lục"; we anchor on the first VISIBLE
    Heading (the first chapter) instead and break right before it — which, in XML
    order, falls before the TOC sdt too, so the cover stands alone."""
    first_heading = None
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        if style.startswith('Heading'):
            first_heading = i
            break
    if not first_heading:  # None or 0 → no title block to separate
        return None
    for j in range(first_heading - 1, -1, -1):
        if doc.paragraphs[j].text.strip():
            return j
    return None


def _find_abstract_heading(doc: Document):
    """Find the Abstract heading paragraph index."""
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip().lower()

        if 'abstract' in text and 'Heading' in style:
            return i

    return None


def _insert_toc(doc: Document, insert_before_idx: int, verbose: bool = False):
    """Insert Table of Contents before the specified index."""
    # Collect all headings from the document
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        if style == 'Heading 1':
            headings.append((1, text))
        elif style == 'Heading 2':
            headings.append((2, text))
        elif style == 'Heading 3':
            headings.append((3, text))

    if not headings:
        return None

    # Get the paragraph before which to insert
    target_para = doc.paragraphs[insert_before_idx]

    # Insert TOC heading
    toc_heading = target_para.insert_paragraph_before("Table of Contents")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_heading.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    # Insert empty line after heading
    target_para.insert_paragraph_before('')

    # Insert TOC entries
    entry_count = 0
    for level, heading_text in headings:
        indent = "    " * (level - 1)
        entry_para = target_para.insert_paragraph_before(f"{indent}{heading_text}")
        entry_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in entry_para.runs:
            run.font.size = Pt(11)
        entry_count += 1

    # Insert empty line after TOC
    empty_para = target_para.insert_paragraph_before('')

    if verbose:
        print(f"   ✓ Inserted Table of Contents ({entry_count} entries)")

    # Return index of the last TOC paragraph (empty line after entries)
    # We need to find it by searching
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'abstract' and 'Heading' in (para.style.name if para.style else ''):
            return i - 1  # Return the paragraph before Abstract

    return insert_before_idx + entry_count + 2  # Approximate


def _find_abstract_end(doc: Document):
    """Find end of Abstract section (before first chapter heading)."""
    in_abstract = False
    last_abstract_para = None

    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip()

        # Start of Abstract
        if 'abstract' in text.lower() and 'Heading' in style:
            in_abstract = True
            continue

        if in_abstract:
            # Check for chapter heading (Heading 1 starting with number)
            if style == 'Heading 1' and text and (text[0].isdigit() or text.startswith('1')):
                return last_abstract_para
            last_abstract_para = i

    return None


def _insert_institution_block(doc: Document, title_idx: int, options: Dict, verbose: bool):
    """Insert institution, faculty, department before title."""
    insert_count = 0

    # Insert in reverse order so indices stay correct
    items = []

    if options.get('department'):
        items.append(('department', options['department'], True, 11))  # italic
    if options.get('faculty'):
        items.append(('faculty', options['faculty'], False, 11))
    if options.get('institution'):
        items.append(('institution', options['institution'].upper(), False, 14))  # uppercase

    # Insert each item at title_idx (they'll stack up correctly)
    for name, text, italic, size in items:
        para = doc.paragraphs[title_idx]
        new_para = para.insert_paragraph_before(text)
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in new_para.runs:
            run.font.size = Pt(size)
            if italic:
                run.font.italic = True
            if name == 'institution':
                run.font.small_caps = True
        insert_count += 1

    # Add empty line after institution block
    if insert_count > 0:
        para = doc.paragraphs[title_idx + insert_count]
        para.insert_paragraph_before('')

    if verbose and insert_count > 0:
        print(f"   ✓ Added institution block ({insert_count} lines)")


# The fixed words on a cover page, per language. These were English literals
# inside the builder, so a Vietnamese thesis came back with a Vietnamese title,
# a Vietnamese degree name and "submitted in partial fulfillment of the
# requirements for the degree of" between them. The Vietnamese wording is the
# standard on a VN thesis cover ("Người hướng dẫn khoa học" in particular).
_COVER_WORDS = {
    "en": {"degree_intro": "submitted in partial fulfillment of the requirements for the degree of",
           "student_id": "Matriculation No.: {v}",
           "supervisor": "First Supervisor: {v}",
           "examiner": "Second Examiner: {v}",
           "by": "by"},
    "vi": {"degree_intro": "Nộp để đáp ứng yêu cầu cấp bằng",
           "student_id": "Mã số học viên: {v}",
           "supervisor": "Người hướng dẫn khoa học: {v}",
           "examiner": "Người phản biện: {v}",
           "by": "Học viên thực hiện"},
}


def _words(options: Dict) -> Dict:
    lang = str((options or {}).get("language") or "en").lower()
    return _COVER_WORDS["vi" if lang.startswith("vi") else "en"]


def _insert_metadata_after_date(doc: Document, date_idx: int, options: Dict, verbose: bool):
    """Insert supervisor info, student ID, etc. after the date paragraph."""
    W = _words(options)
    if date_idx is None:
        return

    # Find the Date paragraph
    date_para = doc.paragraphs[date_idx]
    insert_after = date_para
    additions = 0

    # Project type
    if options.get('project_type'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['project_type'].upper(), size=11, small_caps=True)
        insert_after = new_para
        additions += 2

    # Degree info
    if options.get('course'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, W['degree_intro'], size=10)
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['course'], size=12, bold=True)
        insert_after = new_para
        additions += 3

    # Student ID
    if options.get('student_id'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, W['student_id'].format(v=options['student_id']), size=10)
        insert_after = new_para
        additions += 2

    # Supervisors
    if options.get('instructor'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, W['supervisor'].format(v=options['instructor']), size=10)
        insert_after = new_para
        additions += 2

    if options.get('second_examiner'):
        new_para = _insert_para_after(insert_after, W['examiner'].format(v=options['second_examiner']), size=10)
        insert_after = new_para
        additions += 1

    # System credit
    if options.get('system_credit'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['system_credit'], size=9, italic=True)
        insert_after = new_para
        additions += 2

    # Location
    if options.get('location'):
        new_para = _insert_para_after(insert_after, '')
        insert_after = new_para
        new_para = _insert_para_after(insert_after, options['location'], size=10)
        insert_after = new_para
        additions += 2

    if verbose and additions > 0:
        print(f"   ✓ Added metadata ({additions} lines)")


def _insert_para_after(after_para, text: str, size: int = 11, bold: bool = False,
                       italic: bool = False, small_caps: bool = False):
    """Insert a new centered paragraph after the given paragraph."""
    # Get the parent element and find position
    parent = after_para._element.getparent()
    index = list(parent).index(after_para._element)

    # Create new paragraph
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    new_p = OxmlElement('w:p')

    # Add paragraph properties for centering
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)

    # Add run with text
    if text:
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Font size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size * 2))  # Half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(size * 2))
        rPr.append(szCs)

        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if italic:
            i = OxmlElement('w:i')
            rPr.append(i)
        if small_caps:
            sc = OxmlElement('w:smallCaps')
            rPr.append(sc)

        run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        run.append(t)
        new_p.append(run)

    # Insert after the target paragraph
    parent.insert(index + 1, new_p)

    # Return a paragraph object for the new element
    # Find it in the document
    doc = after_para._element.getparent().getparent()
    for para in after_para._element.getparent().iterchildren(qn('w:p')):
        if para is new_p:
            from docx.text.paragraph import Paragraph
            return Paragraph(new_p, after_para._parent)

    return after_para  # Fallback


def _insert_page_break_after(doc: Document, para_index: int) -> None:
    """Insert a page break after the specified paragraph."""
    if para_index is None or para_index >= len(doc.paragraphs):
        return

    target_para = doc.paragraphs[para_index]
    run = target_para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _fix_table_widths(doc: Document, verbose: bool = False):
    """
    Fix table widths to fit within page margins.

    Pandoc-generated tables often overflow. This function:
    - Sets tables to auto-fit contents
    - Reduces font size for wide tables
    - Prevents mid-word breaks by setting noWrap on first column
    - Uses autofit layout to distribute column widths naturally
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tables_fixed = 0

    for table in doc.tables:
        num_cols = len(table.columns)

        # Get table properties
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Set table width to 100% (5000 = 100% in fifths of a percent)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')

        # Set layout to autofit (not fixed) - allows columns to resize
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'autofit')

        # Enable autofit
        table.autofit = True

        # Set font size based on column count
        if num_cols >= 5:
            font_size = Pt(8)
        elif num_cols >= 4:
            font_size = Pt(9)
        else:
            font_size = Pt(10)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)

                # Remove fixed cell width constraints
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)

                # For the first column (often labels like "Domain"), prevent word wrap
                # This stops mid-word breaks like "Dermat-ology"
                if col_idx == 0:
                    noWrap = tcPr.find(qn('w:noWrap'))
                    if noWrap is None:
                        noWrap = OxmlElement('w:noWrap')
                        tcPr.append(noWrap)

                # Set font size for all cells
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = font_size
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)

        tables_fixed += 1

    if verbose and tables_fixed > 0:
        print(f"   ✓ Fixed {tables_fixed} tables (auto-fit, {font_size.pt}pt font)")


# ============================================================================
# STANDALONE TESTING
# ============================================================================

def main():
    """Test DOCX post-processor."""
    import sys

    if len(sys.argv) > 1:
        docx_path = Path(sys.argv[1])
    else:
        docx_path = Path('examples/opensource_draft.docx')

    if not docx_path.exists():
        print(f"❌ File not found: {docx_path}")
        sys.exit(1)

    # Backup
    backup_path = docx_path.with_suffix('.docx.backup')
    import shutil
    shutil.copy2(docx_path, backup_path)
    print(f"📋 Created backup: {backup_path}")

    # Test options
    test_options = {
        'institution': 'Technical University of Berlin',
        'faculty': 'Faculty of Electrical Engineering and Computer Science',
        'department': 'Department of Security in Telecommunications',
        'course': 'Master of Science in Computer Science',
        'instructor': 'Prof. Dr. Maria Schmidt',
        'second_examiner': 'Prof. Dr. Hans Weber',
        'student_id': '123456',
        'project_type': 'Master Draft',
        'system_credit': 'Generated with DoThesis AI',
        'location': 'Berlin',
    }

    success = insert_academic_structure(docx_path, verbose=True, options=test_options)

    if success:
        print("✅ Test completed!")
    else:
        print("❌ Test failed")
        shutil.copy2(backup_path, docx_path)


if __name__ == '__main__':
    main()
